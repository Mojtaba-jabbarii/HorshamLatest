# -*- coding: utf-8 -*-
"""
Created on Thu Dec 14 12:07:03 2023

@author: ESCO

FUNCTIONALYTY:
The script will generate a steady state report using the results from steady state analysis

COMMENTS:
        The script reads inputs from common excel spreadsheet located in folder: test_scenario_definitions. 
        
@NOTE: 
        if the script is located on sharepoint folder, it will create an equivalent folder path locally for storing results -> reduce syncing burden
        
"""

# import sys
import os, sys
import pandas as pd
import numpy as np
from datetime import datetime
import time
from contextlib import contextmanager
from win32com.client import Dispatch
timestr=time.strftime("%Y%m%d-%H%M%S")

try:
    from StringIO import StringIO
except ImportError:
    from io import StringIO
from io import BytesIO

###############################################################################
#USER INPUTS
###############################################################################
TestDefinitionSheet=r'20230828_SUM_TESTINFO_V2.xlsx'
raw_SS_result_folder = '20240424-104251_S52512'
simulation_batches=['S52512_SS']
simulation_batch_label = simulation_batches[0]

# allign the input model name between the main analysis and reporting scripts
input_models={
                'HighLoad':{'on':'HighLoad_genon\\SUMSF_high_genon.sav',
                           'off':'HighLoad_genoff\\SUMSF_high_genoff.sav'},
                'LowLoad': {'on':'LowLoad_genon\\SUMSF_low_genon.sav',
                            'off':'LowLoad_genoff\\SUMSF_low_genoff.sav'},
             }


# allign the name excel sheet output with data frame - set the same sht name created from Steady state script to make sure the data is loaded properly
df_to_sheet = {'volt_levels':{'df':'df_volt_levels', 'sht':'Volt Levels'},
                'line_loadings':{'df':'df_line_loadings', 'sht':'Line Loadings'},
                'volt_fluc_gen_chng':{'df':'df_volt_fluc_gen_chng', 'sht':'Volt Fluc GenChg'},
                'volt_fluc_lol':{'df':'df_volt_fluc_lol', 'sht':'Volt Fluc Lol'},
                'fault_levels':{'df':'df_fault_levels', 'sht':'Fault Levels'},
               }

#summary_dfs={'HighLoad':{}, 'LowLoad':{},} # use input model names
summary_dfs = {}
for name in input_models.keys():
    summary_dfs[name] = {}
probs={df_to_sheet['volt_levels']['df']:0, df_to_sheet['line_loadings']['df']:0, df_to_sheet['volt_fluc_gen_chng']['df']:0, df_to_sheet['volt_fluc_lol']['df']:0} 
cases=[]
for case in summary_dfs.keys():
    cases.append('_'+case)

genoff = '_off' 
genon = '_on'

# For assessment
vol_nom_high = 1.1
vol_nom_low = 0.9
vol_fluc_gen_HV = 0.03
vol_fluc_gen_LV = 0.04
vol_fluc_lol = 0.05
###############################################################################
# Supporting functions
###############################################################################

def make_dir(dir_path, dir_name=""):
    dir_make = os.path.join(dir_path, dir_name)
    try:
        os.mkdir(dir_make)
    except OSError:
        pass
    return dir_make
        
def createPath(main_folder_out):
    path = os.path.normpath(main_folder_out)
    path_splits = path.split(os.sep) # Get the components of the path
    child_folder = r""+ path_splits[0] # Build up the output path from base drive
    for i in range(len(path_splits)-1):
        child_folder = child_folder + "\\" + path_splits[i+1]
        make_dir(child_folder)
    return child_folder

# read inputs: Iterate through the raw steady state analysis results file and create dataframes require for analysis based on high,low,genon and genoff scenarios.
def read_result_sheet(): 
#    genoff = '_off' 
#    genon = '_on'
    sheets_dict  = pd.read_excel(result_sheet_path, sheet_name = None)
    result_dfs={}

    for case in cases:
        for name,sheet in sheets_dict.items():
            if name == df_to_sheet['volt_levels']['sht']+ case + genoff:
                result_dfs[df_to_sheet['volt_levels']['df']+case+genoff] = pd.DataFrame(data=sheet, columns=['bus_name','PU']).round(3)
                result_dfs[df_to_sheet['volt_levels']['df']+case+genoff].rename(columns = {'bus_name':'Bus Name','PU':'Voltage Level(pu) GenOFF'}, inplace = True)
            elif name == df_to_sheet['volt_levels']['sht'] + case + genon:
                result_dfs[df_to_sheet['volt_levels']['df']+case+genon] = pd.DataFrame(data=sheet, columns=['bus_name','PU']).round(3)
                result_dfs[df_to_sheet['volt_levels']['df']+case+genon].rename(columns = {'bus_name':'Bus Name','PU':'Voltage Level(pu) GenON'}, inplace = True)
            elif name == df_to_sheet['line_loadings']['sht']+case+genoff:
                result_dfs[df_to_sheet['line_loadings']['df'] + case+genoff] = pd.DataFrame(data=sheet, columns = ['CaseNr','Case_Code','brch_name','Loading (%)']).round(3)
                result_dfs[df_to_sheet['line_loadings']['df'] + case+genoff].rename(columns = {'CaseNr':'Case Number','Case_Code': 'Case Reference','brch_name':'Branch Name','Loading (%)': 'Loading(%) GenOFF'}, inplace = True)
            elif name == df_to_sheet['line_loadings']['sht']+case+genon:
                 result_dfs[df_to_sheet['line_loadings']['df'] + case+genon] = pd.DataFrame(data=sheet, columns = ['CaseNr','Case_Code','brch_name','Loading (%)']).round(3)
                 result_dfs[df_to_sheet['line_loadings']['df'] + case+genon].rename(columns = {'CaseNr':'Case Number','Case_Code': 'Case Reference','brch_name':'Branch Name','Loading (%)': 'Loading(%) GenON'}, inplace = True)
            elif name == df_to_sheet['volt_fluc_gen_chng']['sht']+case+genon:
                result_dfs[df_to_sheet['volt_fluc_gen_chng']['df']+case+genon] = pd.DataFrame(data=sheet, columns = ['CaseNr','Case_Code','bus_name','PU','PU_final','VolDev (%)']).round(3)
                result_dfs[df_to_sheet['volt_fluc_gen_chng']['df']+case+genon].rename(columns = {'CaseNr':'Case Number','Case_Code': 'Case Reference','bus_name':'Bus Name','PU':'Voltage Level(pu) Prior','PU_final': 'Voltage Level(pu) After','VolDev (%)': 'Volt Fluc(%)'}, inplace = True)
            elif name == df_to_sheet['volt_fluc_lol']['sht']+case+genon:
                result_dfs[df_to_sheet['volt_fluc_lol']['df']+case+genon] = pd.DataFrame(data=sheet, columns = ['CaseNr','Case_Code','bus_name','PU','VolDev (%)']).round(3)
                result_dfs[df_to_sheet['volt_fluc_lol']['df']+case+genon].rename(columns = {'CaseNr':'Case Number','Case_Code': 'Case Reference','bus_name':'Bus Name','PU':'Voltage Level(pu) GenON','VolDev (%)': 'Volt Fluc(%) GenON'}, inplace = True)
            elif name == df_to_sheet['volt_fluc_lol']['sht']+case+genoff:
                result_dfs[df_to_sheet['volt_fluc_lol']['df']+case+genoff] = pd.DataFrame(data=sheet, columns = ['CaseNr','Case_Code','bus_name','PU','VolDev (%)']).round(3)
                result_dfs[df_to_sheet['volt_fluc_lol']['df']+case+genoff].rename(columns = {'CaseNr':'Case Number','Case_Code': 'Case Reference','bus_name':'Bus Name','PU':'Voltage Level(pu) GenOFF','VolDev (%)': 'Volt Fluc(%) GenOFF'}, inplace = True)
            elif name == df_to_sheet['fault_levels']['sht'] + case + genon:
                result_dfs[df_to_sheet['fault_levels']['df']+case+genon] = pd.DataFrame(data=sheet, columns=['bus_name','bus_no','fault_0_genoff','fault_0_genon', 'fault_1_genoff','fault_1_genon']).round(3)
                result_dfs[df_to_sheet['fault_levels']['df']+case+genon].rename(columns = {'bus_name':'Bus Name','bus_no':'Bus No.','fault_0_genoff':'1ph Fault (MVA) GenOFF','fault_0_genon':'1ph Fault (MVA) GenON', 'fault_1_genoff':'3ph Fault (MVA) GenOFF','fault_1_genon':'3ph Fault (MVA) GenON'}, inplace = True)
            else:
                pass
        

    return result_dfs

# Generate plots
def generate_plot(x_axis, y_axes, legends, label_x, label_y, title):
#    x_axis = df_vlt_lvl['Bus Name']
#    y_axes = [df_vlt_lvl['Voltage Level(pu) GenON'], df_vlt_lvl['Voltage Level(pu) GenOFF']]
#    legends = ['VL(pu) GenON', 'VL(pu) GenOFF']
#    label_x = 'Voltage(pu)'
#    label_y = 'Bus Names'
#    title = 'Voltage Levels'
    # size and positions
    fig = plt.figure(figsize=(7,5))
    fig.add_axes([0.1,0.1,0.8,0.8])
    # axis names and ticks
    plt.scatter(x_axis, y_axes[0], label = legends[0], color = 'k', marker = '*', s = 120)
    plt.scatter(x_axis, y_axes[1], label = legends[1])
    plt.title(title, fontsize =12, color = 'black' )
    plt.ylabel(label_y, fontsize = 10)
    plt.xlabel(label_x, fontsize = 10)
    plt.legend()
    plt.xticks(rotation = 90)
    plt.margins( tight = True)
    plt.grid()
    plt.minorticks_off()
    plt.savefig(case +'voltage_levels'+ 'plot.png', bbox_inches = 'tight')
#    imgdata= StringIO.StringIO()
#    imgdata = StringIO()
    imgdata= BytesIO() # version issues
    plt.savefig(imgdata, bbox_inches = 'tight', dpi =200)    
    return imgdata

def hl_vltg_lvls_violation(val): #check the violation of voltage levels
   color = 'red' if (val < 0.9 or val >1.1) else 'black'
   return 'color: %s' % color

def hl_line_ldng_violation(val): # check the violation of line loadings
    color = 'red' if (val >99.99) else 'black'
    return 'color: %s' % color

def hl_vltg_fluc_violation(val): # check the violation of voltage fluctuations
    color = 'red' if (val > 0.0299) else 'black'
    return 'color: %s' % color

# function to check pass fail criterion
def vl_check(row):
    if row['Voltage Level(pu) GenOFF'] == row['Voltage Level(pu) GenON']:
        val = 'yes'
    elif (row['Voltage Level(pu) GenOFF'] <0.9) and (row['Voltage Level(pu) GenON'] > row['Voltage Level(pu) GenOFF']):
        val = 'yes'
    elif (row['Voltage Level(pu) GenOFF'] >1.1) and (row['Voltage Level(pu) GenON'] < row['Voltage Level(pu) GenOFF']):
        val = 'yes'
    else:
        val = 'no'
    return val

 # function to check pass fail criterion
def ll_check(row):
    if (row['Loading(%) GenOFF'] == row['Loading(%) GenON']):
        val = 'yes'
    elif (row['Loading(%) GenOFF'] > row['Loading(%) GenON']):
        val = 'yes'
    else:
        val = 'no'
    return val


 # function to check pass fail criterion
def vf_gc_check(row):
    if (abs(row['Volt Fluc(%)']) < 3.00):
        val = 'yes'
    else:
        val = 'no'
    return val
                
# function to check pass fail criterion
def vl_lol_check(row):
    if (abs(row['Volt Fluc(%) GenOFF']) > 5.00) and (abs(row['Volt Fluc(%) GenOFF']) > abs(row['Volt Fluc(%) GenON'])):
        val = 'yes'
    else:
        val = 'no'
    return val
                
# prepare sumaries into data frames
def sumarise_results(result_dfs, summary_dfs):

    for case in summary_dfs.keys():
        summary_dfs[case][df_to_sheet['volt_levels']['df']]={}
        summary_dfs[case][df_to_sheet['line_loadings']['df']]={}
        summary_dfs[case][df_to_sheet['volt_fluc_gen_chng']['df']]={}
        summary_dfs[case][df_to_sheet['volt_fluc_lol']['df']]={}
        summary_dfs[case][df_to_sheet['fault_levels']['df']]={}
        for key,value in result_dfs.items():
            if df_to_sheet['volt_levels']['df']+'_'+case+genoff in key: # summarise the voltage level results with and without plant
                #Summary table
                df_name = df_to_sheet['volt_levels']['df']+'_'+case
                df_vlt_lvl = pd.merge(result_dfs[df_name+'_off'],result_dfs[df_name+'_on'], how = 'outer',on = 'Bus Name')
                pvt_df = pd.pivot_table(data = df_vlt_lvl,index = ['Bus Name'],values = ['Voltage Level(pu) GenON','Voltage Level(pu) GenOFF'] )
                pvt_df.style.applymap(hl_vltg_lvls_violation,subset = ['Voltage Level(pu) GenOFF','Voltage Level(pu) GenON']).format({'Voltage Level(pu) GenOFF':'{0:,.3f}','Voltage Level(pu) GenON':'{0:,.3f}'})
                summary_dfs[case][df_to_sheet['volt_levels']['df']]['summary'] = []
                summary_dfs[case][df_to_sheet['volt_levels']['df']]['summary'].append(df_vlt_lvl)
                
                # Voltage levels plot
                imgdata = generate_plot(df_vlt_lvl['Bus Name'], [df_vlt_lvl['Voltage Level(pu) GenON'], df_vlt_lvl['Voltage Level(pu) GenOFF']], ['VL(pu) GenON', 'VL(pu) GenOFF'], 'Bus Name', 'Voltage(pu)', 'Voltage Levels')
                summary_dfs[case][df_to_sheet['volt_levels']['df']]['plot'] = []
                summary_dfs[case][df_to_sheet['volt_levels']['df']]['plot'].append(imgdata)
                
                #Violation voltage level table
                vl_min_violation = df_vlt_lvl.loc[df_vlt_lvl['Voltage Level(pu) GenOFF']<0.9]
                vl_max_violation = df_vlt_lvl.loc[df_vlt_lvl['Voltage Level(pu) GenOFF']>1.1]
                vl_violation = pd.concat([vl_max_violation,vl_min_violation], axis =0)
                if not vl_violation.empty: vl_violation['Pass'] = vl_violation.apply(vl_check, axis = 1)
                summary_dfs[case][df_to_sheet['volt_levels']['df']]['violations'] = []
                summary_dfs[case][df_to_sheet['volt_levels']['df']]['violations'].append(vl_violation)
                
                #Empty apendix
                summary_dfs[case][df_to_sheet['volt_levels']['df']]['appendix'] = []
                
            elif df_to_sheet['line_loadings']['df']+'_'+case+'_off' in key: # summarise the line loading results for network normal and contingencies
                # line loading summary table
                df_name = df_to_sheet['line_loadings']['df']+'_'+case
                df_high = pd.concat([result_dfs[df_name+'_off'],result_dfs[df_name+'_on']],axis = 1)
                df_high = df_high.T.drop_duplicates().T
                pvt_df_high = pd.pivot_table(data=df_high, index=['Case Reference', 'Branch Name'], values = ['Loading(%) GenON','Loading(%) GenOFF'], aggfunc = ['max'])
                #pvt_df_high.style.applymap(hl_line_ldng_violation,subset = [('max','Loading(%) GenON'),('max','Loading(%) GenOFF')]).format({'Loading(%) GenON':'{0:,.3f}','Loading(%) GenOFF':'{0:,.3f}'})
                summary_dfs[case][df_to_sheet['line_loadings']['df']]['appendix'] = []
                summary_dfs[case][df_to_sheet['line_loadings']['df']]['appendix'].append(pvt_df_high)
                
                # line loading plotvltg_fluc_gen_chng.pop
                l_profile_plot = df_high.loc[(df_high['Case Number'] == 'case0 (base)')]
                imgdata = generate_plot(l_profile_plot['Branch Name'], [l_profile_plot['Loading(%) GenON'], l_profile_plot['Loading(%) GenOFF']], ['Loading(%) GenON', 'Loading(%) GenOFF'], 'Branch Name', 'Loading(%)', 'Line Loadings(%)')
                summary_dfs[case][df_to_sheet['line_loadings']['df']]['plot'] = []
                summary_dfs[case][df_to_sheet['line_loadings']['df']]['plot'].append(imgdata)
                
                # line loading violation table
                line_max_violation = df_high.loc[(df_high['Loading(%) GenOFF'] > 99.99) | (df_high['Loading(%) GenON'] > 99.99)]
                if not line_max_violation.empty: line_max_violation['Pass'] = line_max_violation.apply(ll_check, axis = 1)
                #line_max_violation.style.applymap(hl_line_ldng_violation,subset = ['Loading(%) GenON','Loading(%) GenOFF']).format({'Loading(%) GenON':'{0:,.3f}','Loading(%) GenOFF':'{0:,.3f}'})
                summary_dfs[case][df_to_sheet['line_loadings']['df']]['violations'] = []
                summary_dfs[case][df_to_sheet['line_loadings']['df']]['violations'].append(line_max_violation)
                
                #Line Loading Summary
                df_high.pop('Case Number')
                df_high['Loading(%) GenON'] = pd.to_numeric(df_high['Loading(%) GenON'])
                line_ldng_max = df_high.loc[df_high.groupby('Case Reference')['Loading(%) GenON'].idxmax()]
                line_ldng_min = df_high.loc[df_high.groupby('Case Reference')['Loading(%) GenON'].idxmin()]
                #line_ldng_smry = pd.merge(line_ldng_max,line_ldng_min, how = 'outer', on = ['Case Reference'])
                line_ldng_smry = line_ldng_max  #ignoring minimum loadings
                #vltg_fluc_gen_chng_smry = vltg_fluc_gen_chng_smry.T.drop_duplicates().T
                summary_dfs[case][df_to_sheet['line_loadings']['df']]['summary'] = []
                summary_dfs[case][df_to_sheet['line_loadings']['df']]['summary'].append(line_ldng_smry)
                

            elif df_to_sheet['volt_fluc_gen_chng']['df']+'_'+case+'_on' in key: # summary tables for voltage fluctuation due to change in generation
                df_name = df_to_sheet['volt_fluc_gen_chng']['df']+'_'+case
                vltg_fluc_gen_chng = result_dfs[df_name+'_on']
                vltg_fluc_gen_chng.iloc[:,5] = vltg_fluc_gen_chng.iloc[:,5].abs() #get absolute value of all fluctuations
                # adppendix of voltage flcutuation due to gen chnage
                vltg_fluc_gen_chng.pop('Case Number')
                vltg_fluc_gen_chng_pivot = pd.pivot_table(data = vltg_fluc_gen_chng, index = ['Case Reference','Bus Name'], aggfunc = ['max'])
                vltg_fluc_gen_chng_pivot.style.applymap( hl_vltg_fluc_violation,subset = [('max','Volt Fluc(%)')]).format({'Volt Fluc(%)':'{0:,.3f}'})
                summary_dfs[case][df_to_sheet['volt_fluc_gen_chng']['df']]['appendix'] = []
                summary_dfs[case][df_to_sheet['volt_fluc_gen_chng']['df']]['appendix'].append(vltg_fluc_gen_chng_pivot)
                
                # Plot for voltage flcutuation due to gen change
                imgdata = generate_plot(vltg_fluc_gen_chng['Bus Name'], [vltg_fluc_gen_chng['Voltage Level(pu) Prior'], vltg_fluc_gen_chng['Voltage Level(pu) After']], ['Voltage Level(pu) Prior', 'Voltage Level(pu) After'], 'Bus Name', 'Voltage(pu)', 'Voltage Levels Gen Change')
                summary_dfs[case][df_to_sheet['volt_fluc_gen_chng']['df']]['plot'] = []
                summary_dfs[case][df_to_sheet['volt_fluc_gen_chng']['df']]['plot'].append(imgdata)
                
                #Violation results
                vltg_fluc_gen_chng_violation = vltg_fluc_gen_chng.loc[vltg_fluc_gen_chng['Volt Fluc(%)']>3.00]
                if not vltg_fluc_gen_chng_violation.empty: vltg_fluc_gen_chng_violation['Pass'] = vltg_fluc_gen_chng_violation.apply( vf_gc_check, axis = 1)
                summary_dfs[case][df_to_sheet['volt_fluc_gen_chng']['df']]['violations'] = []
                summary_dfs[case][df_to_sheet['volt_fluc_gen_chng']['df']]['violations'].append(vltg_fluc_gen_chng_violation)
                
                #Summary results for voltage fluctuation due to gen change
                vltg_fluc_gen_chng.pop('Voltage Level(pu) Prior')
                vltg_fluc_gen_chng.pop('Voltage Level(pu) After')
                vltg_fluc_gen_chng_max = vltg_fluc_gen_chng.loc[vltg_fluc_gen_chng.groupby('Case Reference')['Volt Fluc(%)'].idxmax()]
                #vltg_fluc_gen_chng_min = vltg_fluc_gen_chng.loc[vltg_fluc_gen_chng.groupby('Case Reference')['Volt Fluc(%)'].idxmin()]
                #vltg_fluc_gen_chng_smry = pd.merge(vltg_fluc_gen_chng_max,vltg_fluc_gen_chng_min, how = 'outer', on = ['Case Reference'])
                #vltg_fluc_gen_chng_smry = vltg_fluc_gen_chng_smry.T.drop_duplicates().T
                vltg_fluc_gen_chng_smry = vltg_fluc_gen_chng_max.T.drop_duplicates().T
                summary_dfs[case][df_to_sheet['volt_fluc_gen_chng']['df']]['summary'] = []
                summary_dfs[case][df_to_sheet['volt_fluc_gen_chng']['df']]['summary'].append(vltg_fluc_gen_chng_smry)
                
            elif df_to_sheet['volt_fluc_lol']['df']+'_'+case+'_on' in key: # summary tables for voltage fluctuation due to loss of a line
                df_name = df_to_sheet['volt_fluc_lol']['df']+'_'+case
                vltfluc_lol = pd.concat([result_dfs[df_name+'_on'],result_dfs[df_name+'_off']],axis = 1)
                #vltfluc_lol = vltfluc_lol.T.drop_duplicates().T not sure why not working took lot of my time
                vltfluc_lol = vltfluc_lol.loc[:,~vltfluc_lol.columns.duplicated()]
                vltfluc_lol.iloc[:,[4,6]] = vltfluc_lol.iloc[:,[4,6]].abs()
                # adppendix of voltage flcutuation due to loss of line
                vltfluc_lol.pop('Case Number')# adppendix of voltage fluctuation due to loss of line'], aggfunc = ['max'])
                vltfluc_lol_pivot = pd.pivot_table(data =  vltfluc_lol, index = ['Case Reference','Bus Name'], aggfunc = ['max'])
                vltfluc_lol_pivot.style.applymap( hl_vltg_fluc_violation,subset = [('max','Volt Fluc(%) GenON'),('max','Volt Fluc(%) GenOFF')]).format({'Volt Fluc(%) GenOFF':'{0:,.3f}','Volt Fluc(%) GenON':'{0:,.3f}'})
                summary_dfs[case][df_to_sheet['volt_fluc_lol']['df']]['appendix'] = []
                summary_dfs[case][df_to_sheet['volt_fluc_lol']['df']]['appendix'].append(vltfluc_lol_pivot)
                
                # Plot for voltage flcutuation due to lol
                imgdata = generate_plot(vltfluc_lol['Bus Name'], [vltfluc_lol['Volt Fluc(%) GenON'], vltfluc_lol['Volt Fluc(%) GenOFF']], ['Volt Fluc(%) GenON', 'Volt Fluc(%) GenOFF'], 'Bus Name', 'Volt Fluc(%)', 'Voltage Fluctuation Contingencies')
                summary_dfs[case][df_to_sheet['volt_fluc_lol']['df']]['plot'] = []
                summary_dfs[case][df_to_sheet['volt_fluc_lol']['df']]['plot'].append(imgdata)
                
                
                #Violation results
                vltfluc_lol_violation = vltfluc_lol.loc[abs(vltfluc_lol['Volt Fluc(%) GenON'])>5.00]
                if not vltfluc_lol_violation.empty: vltfluc_lol_violation['Pass'] = vltfluc_lol_violation.apply(vl_lol_check, axis = 1)
                summary_dfs[case][df_to_sheet['volt_fluc_lol']['df']]['violations'] = []
                summary_dfs[case][df_to_sheet['volt_fluc_lol']['df']]['violations'].append(vltfluc_lol_violation)
                
                #Summary results for voltage fluctuation due to gen change
                vltfluc_lol.pop('Voltage Level(pu) GenOFF')
                vltfluc_lol.pop('Voltage Level(pu) GenON')
                vltfluc_lol_smry = vltfluc_lol.loc[vltfluc_lol.groupby('Case Reference')['Volt Fluc(%) GenON'].idxmax()]
                #vltfluc_lol_max = vltfluc_lol.loc[vltfluc_lol.groupby('Case Reference')['Volt Fluc(%) GenON'].idxmax()]
                #vltfluc_lol_min = vltfluc_lol.loc[vltfluc_lol.groupby('Case Reference')['Volt Fluc(%) GenON'].idxmin()]
                #vltfluc_lol_smry = pd.merge( vltfluc_lol_max,vltfluc_lol_min, how = 'outer', on = ['Case Reference'])
                #vltfluc_lol_smry = vltfluc_lol_smry.T.drop_duplicates().T
                summary_dfs[case][df_to_sheet['volt_fluc_lol']['df']]['summary'] = []
                summary_dfs[case][df_to_sheet['volt_fluc_lol']['df']]['summary'].append(vltfluc_lol_smry)
                
                
            elif df_to_sheet['fault_levels']['df']+'_' + case + '_on' in key: #  populate the summary for fault level results
                #summary table
                df_fault_lvl = pd.DataFrame(result_dfs[df_to_sheet['fault_levels']['df']+'_'+case+'_on'])
                #df_fault_lvl = df_fault_lvl.T.drop_duplicates().T
                summary_dfs[case][df_to_sheet['fault_levels']['df']]['summary'] = []
                summary_dfs[case][df_to_sheet['fault_levels']['df']]['summary'].append(df_fault_lvl)
                
                #appendix
                summary_dfs[case][df_to_sheet['fault_levels']['df']]['appendix'] = []
                
                #plots
                summary_dfs[case][df_to_sheet['fault_levels']['df']]['plot'] = []
                
                #violations

                summary_dfs[case][df_to_sheet['fault_levels']['df']]['violations'] = []    



def initialise_report():
    #read report template 
    report=Document(main_folder_path+"\\Plots\\ReportTemplate_steadystate.docx")
    return report


def replace_placeholders(report):
    replace_dict = {'[Project Name]':str(ProjectDetailsDict['Name']), '[Project Name Short]':str(ProjectDetailsDict['NameShrt']), '[Total Plant MW at POC]': str(ProjectDetailsDict['PlantMW']), 
                    '[Developer]': str(ProjectDetailsDict['Dev']), '[Network Service Provider]':str(ProjectDetailsDict['NSP']), '[Town]': str(ProjectDetailsDict['Town']), 
                    '[State]': str(ProjectDetailsDict['State']), '[Connection type]': str(ProjectDetailsDict['contyp']),'[POC Feeder]': str(ProjectDetailsDict['poc_fdr']),
                    '[Nominal POC voltage (kV)]': str(ProjectDetailsDict['VPOCkv']), '[PSSEversion]': str(PSSEmodelDict['PSSEversion']),'[Lot/DP]': str(ProjectDetailsDict['lot_dp']),
                    '[Address]': str(ProjectDetailsDict['addrs']), '[LGA]': str(ProjectDetailsDict['lga']), '[POC Substation]': str(ProjectDetailsDict['Sub']), '[Plant Model]': str(ProjectDetailsDict['plnt_mdl'])
                    }
    for key,value in replace_dict.items():
        for p in report.paragraphs:
            if key in p.text:
                p.text = p.text.replace(key,value)
    #for p in report.paragraphs:
     #   inline = p.runs
      #  for j in range(0,len(inline)):
       #     for k,v in replace_dict.items():
        #        if k in inline[j].text:
         #           inline[j].text = inline[j].text.replace(k,v)
    for key,value in replace_dict.items():
        for table in report.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if key in p.text:
                            p.text = p.text.replace(key,value)
                        
                    #inline = p.runs
                    #for j in range(0,len(inline)):
                    #    for k,v in replace_dict.items():
                    #       if k in inline[j].text:
                    #           inline[j].text = inline[j].text.replace(k,v)
    
    return report    

def add_report_intro(report):
    plant_rating=ProjectDetailsDict['PlantMW']
    POC_name=ProjectDetailsDict['Sub']
    location=ProjectDetailsDict['Town']+", "+ProjectDetailsDict['State']
    #generate general description and intro based on ProjectDetailsDict
    #Headline
    '''
    p=report.add_paragraph('')
    run=p.add_run()
    run.add_break(WD_BREAK.PAGE)
    report.add_heading("Introduction", level=1 )
    intro_text="This report summarises the findings of a steady state study carried out for a "+str(plant_rating)+" MW generator connected to "+str(POC_name)+" in "+str(location)+"."
    intro_text+="\n\nThe Report consists of five separate sections, each of which illustrate the findings of the respective part of the study:"
    p=report.add_paragraph(intro_text)
    #add description of each subsection of the report    
    temp_text="1) The first part of the steady state analysis looks at Bus voltages under system normal conditions (i.e. no line outages or other contingencies. The results of the analysis are presented in section 2. Adding generation to a bus or a line may shift the voltage levels at that bus or surrounding buses due to the impact it has on on power flows in the area. "
    temp_text+="\nThe normal operating band of the NEM is between 0.9 p.u. voltage and 1.1 p.u. voltage. The results presented in that section will highlight the changes in relevant bus voltages due to the addition of the new generator."
    p=report.add_paragraph('')
    for style in report.styles:
        print("style.name == %s" % style.name)
    report.add_paragraph(temp_text, style='List Paragraph')
    report.add_paragraph('', style='List Paragraph')
    temp_text="2) The second part of the Steady State analysis investigates line loading and transformer loading under N-1 conditions, with and without the proposed generator. This reveals pre-existing issues as well as issues caused by the inclusion of the proposed generator. "
    report.add_paragraph(temp_text, style='List Paragraph')
    report.add_paragraph('', style='List Paragraph')
    temp_text="3) The third part of the Steady State analysis explores voltage stability for a change in generation output. If there is a sudden disconnect the plant output can drop from 100% to 0%, which will instantly change the voltage at surrounding buses. This is quantified and compared against applicable thresholds. "
    report.add_paragraph(temp_text, style='List Paragraph')
    report.add_paragraph('', style='List Paragraph')
    temp_text="4) The fourth part of the Steady State analysis focusses on voltage stability under credible contingencies. The voltage magnitude of the voltage fluctuations "
    report.add_paragraph(temp_text, style='List Paragraph')
    report.add_paragraph('', style='List Paragraph')
    temp_text="5) The last part of the Steady State analysis investigates Fault current at buses of interest. These fault levels are required to not exceed planning levels. This item is unlikely to be problematic in most instances as the contribution from inverter-based resources behind transformer and cable impedances is generally normally small compared against existing headroom."
    report.add_paragraph(temp_text, style='List Paragraph')
    report.add_paragraph('', style='List Paragraph')
    '''
    p=report.add_paragraph('')
    run=p.add_run()
    run.add_break(WD_BREAK.PAGE)
    report.add_heading("Simulation Results", level=1 )
    return 0

#adds descritpion of the voltage level checks to the summary report and will add the table for high and low case, or just one if only one included.
#It will also add a graph showing the voltage levels.
def volt_lvl_report(report, summary_dfs, probs):
    report.add_heading("Voltage Levels", level=2 )
    temp_text="The voltage levels at buses of interest have been analysed for system normal conditions (no contingencies) with "+str(ProjectDetailsDict['Name'])+" in service and compared to the voltage levels prior to adding the new generator. "
    temp_text+="The normal operating band is defined as 0.9 p.u. to 1.1 p.u. voltage. If the plant pushes the voltage outside these boundaries this indicates a problem. Normally the voltage would be expected to remain within +/-5% of 1 p.u. "
    temp_text+="Any cases where the voltage goes outside this range should be carefully looked at and may indicate a need for further reactive support or voltage control."
    report.add_paragraph(temp_text)
    report.add_paragraph('')
    #Add Summary tables
    report.add_heading("Summary of findings", level=3 )
    temp_text="The results for each bus before and after addition of "+str(ProjectDetailsDict['Name'])+" are listed in the tables below."
    report.add_paragraph(temp_text)
    for case in summary_dfs.keys():
        if(summary_dfs[case][df_to_sheet['volt_levels']['df']]!={}):
            report.add_heading(ident_case_name(case), level=4)#maybe make this a label of the table instead of making it a  heading.
            for frame_id in range(0, len(summary_dfs[case][df_to_sheet['volt_levels']['df']]['summary'])):#add summary tables of results to word doc.
                data_frame_to_docx_table(summary_dfs[case][df_to_sheet['volt_levels']['df']]['summary'][frame_id], report)
                report.add_paragraph('')
    #Add plots
    plots_present=False
    for case in summary_dfs.keys():
        if('plot' in summary_dfs[case][df_to_sheet['volt_levels']['df']].keys()):
            if(summary_dfs[case][df_to_sheet['volt_levels']['df']]['plot']!=[]):
                plots_present=True
    if(plots_present):
        report.add_heading("Plots", level=3)
        temp_text="The scatter plots show the absolute voltage level in the reference case(s) analysed in this study."
        report.add_paragraph(temp_text)
        for case in summary_dfs.keys():
            if(summary_dfs[case][df_to_sheet['volt_levels']['df']]!={}):
                report.add_heading(ident_case_name(case), level=4)#maybe make this a label of the table instead of making it a  heading.
                for frame_id in range(0, len(summary_dfs[case][df_to_sheet['volt_levels']['df']]['plot'])):#add summary tables of results to word doc.
                    report.add_picture(summary_dfs[case][df_to_sheet['volt_levels']['df']]['plot'][frame_id], Inches(6))
                    report.add_paragraph('')
        
    #Add overview of violations
        #add check whether any violations exist
    report.add_heading("Violations", level=3 )
    temp_text="A voltage violation exists wherever the voltage is outside 0.9 to 1.1 p.u. If a given voltage violation is not exacerbated by the addition of the plant, it is considered unproblematic. "
    temp_text+="Where a violation occurs at dedicated buses, such as within other plant (SVCs etc.) these can often easily be mitigated by adjusting local transformer taps or control settings and may not be prohibitive."
    report.add_paragraph(temp_text)
    causer_flag=0
    for case in summary_dfs.keys():
        case_results_present=False
        viol_in_act_case=False
        if(summary_dfs[case][df_to_sheet['volt_levels']['df']]!={}):
            if('violations' in summary_dfs[case][df_to_sheet['volt_levels']['df']].keys()):
                case_results_present=True
                if(summary_dfs[case][df_to_sheet['volt_levels']['df']]['violations']!=[]):
                    viol_in_act_case=True
                    probs[df_to_sheet['volt_levels']['df']]=1
                    report.add_heading(ident_case_name(case), level=4)#maybe make this a label of the table instead of making it a  heading.
                    for frame_id in range(0, len(summary_dfs[case][df_to_sheet['volt_levels']['df']]['violations'])):#add summary tables of results to word doc.
                        problem_flag=data_frame_to_docx_table(summary_dfs[case][df_to_sheet['volt_levels']['df']]['violations'][frame_id], report)
                        if(problem_flag>0):
                            causer_flag=1#indicates whether the project is creating or exacerbating the issue.
                        report.add_paragraph('')
        if(not viol_in_act_case and case_results_present):
            report.add_paragraph('No voltage violations observed for '+ident_case_name(case)+'.') 
            
    if (causer_flag>0):
       probs[df_to_sheet['volt_levels']['df']]=2 #Add conclusion based on whether violations exist and provide some generic advice.
    return 0
#Adds description of line loading analysis along with summary table (comparing with and withou gen). It will also include a graph showing the highest loading of each line (this may occur for a different contingency per line. this should be marked in the graph)
#The function should be able - based on the result data - to point out pre-existing overloadings and new overloadings and distinguish between them.
def line_loadings_report(report, summary_dfs, probs):
    p=report.add_paragraph('')
    run=p.add_run()
    run.add_break(WD_BREAK.PAGE)
    report.add_heading("Line Loading", level=2 )
    temp_text="Line and transformer loading in the area around "+ProjectDetailsDict['Town']+" has been analysed under various conditions, including relevant N-1 scenarios. "
    temp_text+="Where overloading occurs, it is important to understand whether the overloading is caused and/or exacerbated by the proposed generator. Relevant information is listed under 'violations' in below. "
    report.add_paragraph(temp_text)
    report.add_paragraph('')
    #Add Summary tables
    report.add_heading("Summary of findings", level=3 )
    temp_text="The results for each bus before and after addition of "+str(ProjectDetailsDict['Name'])+" are listed in the tables below."
    report.add_paragraph(temp_text)
    for case in summary_dfs.keys():
        if(summary_dfs[case][df_to_sheet['line_loadings']['df']]!={}):
            report.add_heading(ident_case_name(case), level=4)#maybe make this a label of the table instead of making it a  heading.
            for frame_id in range(0, len(summary_dfs[case][df_to_sheet['line_loadings']['df']]['summary'])):#add summary tables of results to word doc.
                #data_frame_to_docx_table(summary_dfs[case][df_to_sheet['line_loadings']['df']]['summary'][frame_id], report)
                data_frame_to_docx_table(summary_dfs[case][df_to_sheet['line_loadings']['df']]['summary'][frame_id].iloc[:,[0,1,2,3]], report)
                report.add_paragraph('')
    #Add plots
    plots_present=False
    for case in summary_dfs.keys():
        if('plot' in summary_dfs[case][df_to_sheet['line_loadings']['df']].keys()):
            if(summary_dfs[case][df_to_sheet['line_loadings']['df']]['plot']!=[]):
                plots_present=True
    if(plots_present):
        report.add_heading("Plots", level=3 )
        temp_text="The scatter plots show the result for the worst case loading of each element (line or transformer). "
        temp_text+="The worst case scenario may differ between different network elements. If one particular outage lead to overloading of line X, it does not mean that the same outage constitutes the worst case for line Y. (line X and Y serving as a generic example in this explanation)."
        report.add_paragraph(temp_text)
        for case in summary_dfs.keys():
            if(summary_dfs[case][df_to_sheet['line_loadings']['df']]!={}):
                report.add_heading(ident_case_name(case), level=4)#maybe make this a label of the table instead of making it a  heading.
                for frame_id in range(0, len(summary_dfs[case][df_to_sheet['line_loadings']['df']]['plot'])):#add summary tables of results to word doc.
                    report.add_picture(summary_dfs[case][df_to_sheet['line_loadings']['df']]['plot'][frame_id], Inches(6))
                    report.add_paragraph('')
        
    #Add overview of violations
        #add check whether any violations exist
    report.add_heading("Violations", level=3 )
    temp_text="An overloading exists wherever the line or transformer loading exceeds 100% for any credible N-1 scenario. If a given overloading is not exacerbated by the addition of the plant, it is not generally a threat to the project, but may me an indication of emerging problems in the area and needs to be given careful consideration. "
    temp_text+="Overloading that is caused or exacerbated by the propsoed project may be addressed with a runback scheme. If the overloading occurs under System normal conditions, the "
    report.add_paragraph(temp_text)
    causer_flag=0
    for case in summary_dfs.keys():
        case_results_present=False
        viol_in_act_case=False
        if(summary_dfs[case][df_to_sheet['line_loadings']['df']]!={}):
            if('violations' in summary_dfs[case][df_to_sheet['line_loadings']['df']].keys()):
                case_results_present=True
                if(summary_dfs[case][df_to_sheet['volt_levels']['df']]['violations']!=[]):
                    viol_in_act_case=True
                    probs[df_to_sheet['line_loadings']['df']]=1
                    report.add_heading(ident_case_name(case), level=4)#maybe make this a label of the table instead of making it a  heading.
                    for frame_id in range(0, len(summary_dfs[case][df_to_sheet['line_loadings']['df']]['violations'])):#add summary tables of results to word doc.
                        problem_flag=data_frame_to_docx_table(summary_dfs[case][df_to_sheet['line_loadings']['df']]['violations'][frame_id], report)
                        if(problem_flag>0):
                            causer_flag=1#indicates whether the project is creating or exacerbating the issue.
                        report.add_paragraph('')
        if(not viol_in_act_case and case_results_present):
            report.add_paragraph('No overloading observed for '+ident_case_name(case)+'.') 
    if (causer_flag>0):
       probs[df_to_sheet['line_loadings']['df']]=2             #overloadings caused by proposed generator       
    return 0

#Add description of test type and context. include summary table of voltage fluctuations and critical fluctuations. 
#Additional comment on the table with fluctuations that excceed the limits and conclude whether or not there are problematic cases.
def bus_volt_fluct_gen_chng_report(report, summary_dfs, probs):
    p=report.add_paragraph('')
    run=p.add_run()
    run.add_break(WD_BREAK.PAGE)
    report.add_heading("Voltage Fluctuations for Change of Generation", level=2 )
    temp_text="This section explores the impact of a sudden loss of generation on the bus voltages at the buses of interest. The output of "+str(ProjectDetailsDict['Name'])+" can suddenly change due to unexpected tripping of the plant or weather-related circumstances such as a change in cloud cover. "
    temp_text+="A sudden change in the output of the plant must not cause a voltage disturbance larger than 3% due to loss of generation and larger than 5% due to trip of the plant. This voltage flcutation is acceptable within 5% at distribution level. If a project causes steps greater than the allowed margin at surrounding buses, mitigation measures can be considered, such as a Static Var Compensator (SVC) located at a nearby bus. "
    report.add_paragraph(temp_text)
    report.add_paragraph('')
    #Add Summary tables
    report.add_heading("Summary of findings", level=3 )
    #temp_text="The maximum and minimum voltage fluctuation for each bus under the given scenarios are listed in the tables below. "
    temp_text="The maximum voltage fluctuation for each bus under the given scenarios are listed in the tables below. "
    report.add_paragraph(temp_text)
    for case in summary_dfs.keys():
        if(summary_dfs[case][df_to_sheet['volt_fluc_gen_chng']['df']]!={}):
            report.add_heading(ident_case_name(case), level=4)#maybe make this a label of the table instead of making it a  heading.
            for frame_id in range(0, len(summary_dfs[case][df_to_sheet['volt_fluc_gen_chng']['df']]['summary'])):#add summary tables of results to word doc.
                data_frame_to_docx_table(summary_dfs[case][df_to_sheet['volt_fluc_gen_chng']['df']]['summary'][frame_id], report)
                report.add_paragraph('')
    #Add plots
    plots_present=False
    for case in summary_dfs.keys():
        if('plot' in summary_dfs[case][df_to_sheet['volt_fluc_gen_chng']['df']].keys()):
            if(summary_dfs[case][df_to_sheet['volt_fluc_gen_chng']['df']]['plot']!=[]):
                plots_present=True
    if(plots_present):
        report.add_heading("Plots", level=3 )
        temp_text="The scatter plots shows the voltage profile of the network during loss of generation events prior and after connecting the proposed plant. "
        report.add_paragraph(temp_text)
        for case in summary_dfs.keys():
            if(summary_dfs[case][df_to_sheet['volt_fluc_gen_chng']['df']]!={}):
                report.add_heading(ident_case_name(case), level=4)#maybe make this a label of the table instead of making it a  heading.
                for frame_id in range(0, len(summary_dfs[case][df_to_sheet['volt_fluc_gen_chng']['df']]['plot'])):#add summary tables of results to word doc.
                    report.add_picture(summary_dfs[case][df_to_sheet['volt_fluc_gen_chng']['df']]['plot'][frame_id], Inches(6))
                    report.add_paragraph('')
        
    #Add overview of violations
        #add check whether any violations exist
    report.add_heading("Violations", level=3 )
    temp_text="All voltage variations exceeding 3% are listed in this section. Where the variations occur at transmission level, up to 5% may be acceptable. "
    report.add_paragraph(temp_text)
    causer_flag=0 #flag contains information whether or not a given violation of limits is due to the addition of the proposed generator
    for case in summary_dfs.keys():
        case_results_present=False
        viol_in_act_case=False
        if(summary_dfs[case][df_to_sheet['volt_fluc_gen_chng']['df']]!={}):
            if('violations' in summary_dfs[case][df_to_sheet['volt_fluc_gen_chng']['df']].keys()):
                case_results_present=True
                if(summary_dfs[case][df_to_sheet['volt_fluc_gen_chng']['df']]['violations']!=[]):
                    viol_in_act_case=True
                    probs[df_to_sheet['volt_fluc_gen_chng']['df']]=1
                    report.add_heading(ident_case_name(case), level=4)#maybe make this a label of the table instead of making it a  heading.
                    for frame_id in range(0, len(summary_dfs[case][df_to_sheet['volt_fluc_gen_chng']['df']]['violations'])):#add summary tables of results to word doc.
                        problem_flag=data_frame_to_docx_table(summary_dfs[case][df_to_sheet['volt_fluc_gen_chng']['df']]['violations'][frame_id], report)
                        if(problem_flag>0):
                            causer_flag=1#indicates whether the project is creating or exacerbating the issue.
                        report.add_paragraph('')
        if(not viol_in_act_case and case_results_present):
            report.add_paragraph('No critical voltage fluctuatation observed for '+ident_case_name(case)+'.')
    if (causer_flag>0):
       probs[df_to_sheet['volt_fluc_gen_chng']['df']]=2 
    return 0
#Add description of test type and context. include summary table of voltage fluctuations and critical fluctuations. 
#Additional comment on the table with fluctuations and state whether or not the proposes project is causing the fluctuations to exceed limits.
def bus_volt_fluct_cont_report(report, summary_dfs, probs):
    p=report.add_paragraph('')
    run=p.add_run()
    run.add_break(WD_BREAK.PAGE)
    report.add_heading("Voltage fluctuations due to contingencies", level=2 )
    temp_text="This section explores voltage fluctuations at the buses of interest due to loss of line, transformer, and generator contingencies."
    temp_text+=" These voltage fluctions are than compared with the existing voltage fluctuations i.e. prior connecting the proposed plant. The voltage fluctuations should not be worse than the existing voltage fluctuations. "
    report.add_paragraph(temp_text)
    report.add_paragraph('')
    #Add Summary tables
    report.add_heading("Summary of findings", level=3 )
    temp_text="The maximum voltage fluctuation for each bus under the given scenarios are listed in the tables below. "
    report.add_paragraph(temp_text)
    for case in summary_dfs.keys():
        if(summary_dfs[case][df_to_sheet['volt_fluc_lol']['df']]!={}):
            report.add_heading(ident_case_name(case), level=4)#maybe make this a label of the table instead of making it a  heading.
            for frame_id in range(0, len(summary_dfs[case][df_to_sheet['volt_fluc_lol']['df']]['summary'])):#add summary tables of results to word doc.
                data_frame_to_docx_table(summary_dfs[case][df_to_sheet['volt_fluc_lol']['df']]['summary'][frame_id], report)
                report.add_paragraph('')
    #Add plots
    plots_present=False
    for case in summary_dfs.keys():
        if('plot' in summary_dfs[case][df_to_sheet['volt_fluc_lol']['df']].keys()):
            if(summary_dfs[case][df_to_sheet['volt_fluc_lol']['df']]['plot']!=[]):
                plots_present=True
    if(plots_present):
        report.add_heading("Plots", level=3 )
        temp_text="The scatter plots shows the voltage fluctuations due to contingencies. "
        report.add_paragraph(temp_text)
        for case in summary_dfs.keys():
            if(summary_dfs[case][df_to_sheet['volt_fluc_lol']['df']]!={}):
                report.add_heading(ident_case_name(case), level=4)#maybe make this a label of the table instead of making it a  heading.
                for frame_id in range(0, len(summary_dfs[case][df_to_sheet['volt_fluc_lol']['df']]['plot'])):#add summary tables of results to word doc.
                    report.add_picture(summary_dfs[case][df_to_sheet['volt_fluc_lol']['df']]['plot'][frame_id], Inches(6))
                    report.add_paragraph('')
        
    #Add overview of violations
        #add check whether any violations exist
    report.add_heading("Violations", level=3 )
    temp_text="All voltage fluctuations in this assessment are "
    report.add_paragraph(temp_text)
    causer_flag=0 #flag contains information whether or not a given violation of limits is due to the addition of the proposed generator
    for case in summary_dfs.keys():
        case_results_present=False
        viol_in_act_case=False
        if(summary_dfs[case][df_to_sheet['volt_fluc_lol']['df']]!={}):
            if('violations' in summary_dfs[case][df_to_sheet['volt_fluc_lol']['df']].keys()):
                case_results_present=True
                if(summary_dfs[case][df_to_sheet['volt_fluc_lol']['df']]['violations']!=[]):
                    viol_in_act_case=True
                    probs[df_to_sheet['volt_fluc_lol']['df']]=1
                    report.add_heading(ident_case_name(case), level=4)#maybe make this a label of the table instead of making it a  heading.
                    for frame_id in range(0, len(summary_dfs[case][df_to_sheet['volt_fluc_lol']['df']]['violations'])):#add summary tables of results to word doc.
                        problem_flag=data_frame_to_docx_table(summary_dfs[case][df_to_sheet['volt_fluc_lol']['df']]['violations'][frame_id], report)
                        if(problem_flag>0):
                            causer_flag=1#indicates whether the project is creating or exacerbating the issue.
                        report.add_paragraph('')
        if(not viol_in_act_case and case_results_present):
            report.add_paragraph('No critical voltage fluctuatation observed for '+ident_case_name(case)+'.')
    if (causer_flag>0):
       probs[df_to_sheet['volt_fluc_lol']['df']]=2 
    return 0
#Descritpion of the test type and context. Result table with fault levels. This shoudl include not only IEEE methodology but also N-something methodology.
#Maybe add logic  to check against some pre-defined levels.
def fautl_levels_report(report, summary_dfs, probs):
    p=report.add_paragraph('')
    run=p.add_run()
    run.add_break(WD_BREAK.PAGE)
    report.add_heading("Fault Level Analysis", level=2 )
#    temp_text=''
    temp_text="The fault levels were calculated using the NCSFCC (for current source generators) and ASCC function in PSSE) for both 3 phase and phase to ground faults. The short circuit current studies have been performed on the maximum load case."
    temp_text+=" Following table shows the fault levels on the monitored buses. "
    report.add_paragraph(temp_text)
    report.add_paragraph('')
    #Add Summary tables
    #report.add_heading("Summary of findings", level=3 )
    #temp_text="The maximum and minimum voltage fluctuation for each bus under the given scenarios are listed in the tables below. "
    #report.add_paragraph(temp_text)
    for case in summary_dfs.keys():
        if(summary_dfs[case][df_to_sheet['fault_levels']['df']]!={}):
            report.add_heading(ident_case_name(case), level=3)#maybe make this a label of the table instead of making it a  heading.
            for frame_id in range(0, len(summary_dfs[case][df_to_sheet['fault_levels']['df']]['summary'])):#add summary tables of results to word doc.
                data_frame_to_docx_table(summary_dfs[case][df_to_sheet['fault_levels']['df']]['summary'][frame_id], report)
                report.add_paragraph('')
    return 0

def add_conclusion(report, probs):
    p=report.add_paragraph('')
    run=p.add_run()
    run.add_break(WD_BREAK.PAGE)
    report.add_heading("Conclusion", level=1 )   
    temp_text=''
    flawless=True
    for test_type in probs.keys():
        if(probs[test_type])>0:
            flawless=False
    if not flawless:
        temp_text+="Some potential issues have been detected during the Steady State Analysis. "
        if(probs[df_to_sheet['volt_levels']['df']]>0):
            temp_text+="\nThe studies have shown issues with the normal voltage levels at one or more buses. "
            if(probs[df_to_sheet['volt_levels']['df']]>1):
                temp_text+="These issues are caused or exacerbated by the proposed project, indicating that there is a need for additional reactive power support like capacitor banks or an SVC, or a downsizing of the project. Please refer to section 2.1 for details. "
            else:
                temp_text+="These issues are pre-existing and are not exacerbated by the plant, hence this is not a reason against connecting the project, however it may indicate an issue with the network in the area and further investigation is recommended. "
        if(probs[df_to_sheet['line_loadings']['df']]>0):
            temp_text+="\nOne or more network elements exhibit overloading, the details are provided in section 2.2."
            if(probs[df_to_sheet['line_loadings']['df']]>1):
                temp_text+="The overloading is caused or exacerbated by "+str(ProjectDetailsDict['Name'])+" and must be addressed using runback schemes, a network augmentation or downsizing the project. "
            else:
                temp_text+="The overloading is pre-existing and not exacerbated by the project and is not a reason against the development, but an indicator for a nearby network issue. "
        if(probs[df_to_sheet['volt_fluc_gen_chng']['df']]>0):            
            temp_text+="\nVoltage violations have been observed for a change in generation output. Similar to voltage violations under system normal conditions this can be addressed using an SVC in a nearby location that will not trip for a trip of the proposed generator. The violations are detailed in section 2.3. "            
        if(probs[df_to_sheet['volt_fluc_lol']['df']]>0):
            temp_text+="\nSome contingency events have been observed to lead to unacceptable voltage fluctuations on the network. "
            if(probs[df_to_sheet['volt_fluc_lol']['df']]>1):
                temp_text+="These fluctuations are caused or exacerbated by the proposed generator and need to be mitigated in order for the project to procees.  Similar to voltage violations under system normal conditions this can also be addressed with the help of an SVC. Further details about the relevant cases are provided in the tables in section 2.4. "
            else:
                temp_text+="The observed voltage fluctuations under N-1 are not caused by the project and shoudl not negatively impact the development. Further details about the relevant cases are provided in the tables in section 2.4. "       
    else:
        temp_text+="No issues have been detected in the Steady State analysis, which may indicate that the proposed location on the network is suitable to host a" + str(ProjectDetailsDict['PlantMW']) + "MW generator without any further limitations from a grid perspective. It should however be kept in mind that the the network is subject to constant change and new incoming generation can change this situation. "
        temp_text+="The steady state study should be repeated from time to time as the project progresses and new information becomes available."
    
    p=report.add_paragraph(temp_text)
    #report.add_paragraph("test", color='red')

def add_appendices(report,summary_dfs):
    p=report.add_paragraph('')
    run=p.add_run()
    run.add_break(WD_BREAK.PAGE)
    report.add_heading("Appendices", level=1)
    temp_text ="This section provide the detailed results where applicable for the assessments performed in this steady state analysis."
    report.add_paragraph(temp_text)
    report.add_paragraph('')
    #Add Appendix tables
    #Voltage levels
    report.add_heading("Appendix 1", level=3 )
    temp_text="The voltage levels for all the monitored buses as per section[]"
    report.add_paragraph(temp_text)
    for case in summary_dfs.keys():
        if(summary_dfs[case][df_to_sheet['volt_levels']['df']]!={}):
            report.add_heading(ident_case_name(case), level=4)#maybe make this a label of the table instead of making it a  heading.
            for frame_id in range(0, len(summary_dfs[case][df_to_sheet['volt_levels']['df']]['appendix'])):#add summary tables of results to word doc.
                data_frame_to_docx_table(summary_dfs[case][df_to_sheet['volt_levels']['df']]['appendix'][frame_id], report)
                report.add_paragraph('')
    
    #Line Loadings
    report.add_heading("Appendix 2", level=3 )
    temp_text="The thermal loadings for all the monitored branches in network normal and N-1 contingencies as per section[]"
    report.add_paragraph(temp_text)
    for case in summary_dfs.keys():
        if(summary_dfs[case][df_to_sheet['line_loadings']['df']]!={}):
            report.add_heading(ident_case_name(case), level=4)#maybe make this a label of the table instead of making it a  heading.
            for frame_id in range(0, len(summary_dfs[case][df_to_sheet['line_loadings']['df']]['appendix'])):#add summary tables of results to word doc.
                data_frame_to_docx_table(summary_dfs[case][df_to_sheet['line_loadings']['df']]['appendix'][frame_id], report)
                report.add_paragraph('')
    
    #Voltage fluctuations GenChange
    report.add_heading("Appendix 3", level=3 )
    temp_text="The voltage fluctuations due to change in generation output on all the monitored buses as per section[]"
    report.add_paragraph(temp_text)
    for case in summary_dfs.keys():
        if(summary_dfs[case][df_to_sheet['volt_fluc_gen_chng']['df']]!={}):
            report.add_heading(ident_case_name(case), level=4)#maybe make this a label of the table instead of making it a  heading.
            for frame_id in range(0, len(summary_dfs[case][df_to_sheet['volt_fluc_gen_chng']['df']]['appendix'])):#add summary tables of results to word doc.
                data_frame_to_docx_table(summary_dfs[case][df_to_sheet['volt_fluc_gen_chng']['df']]['appendix'][frame_id], report)
                report.add_paragraph('')
                
    #Line Loadings
    report.add_heading("Appendix4", level=3 )
    temp_text="The voltage fluctuations due to contingencies on all the monitored buses as per section[]"
    report.add_paragraph(temp_text)
    for case in summary_dfs.keys():
        if(summary_dfs[case][df_to_sheet['volt_fluc_lol']['df']]!={}):
            report.add_heading(ident_case_name(case), level=4)#maybe make this a label of the table instead of making it a  heading.
            for frame_id in range(0, len(summary_dfs[case][df_to_sheet['volt_fluc_lol']['df']]['appendix'])):#add summary tables of results to word doc.
                data_frame_to_docx_table(summary_dfs[case][df_to_sheet['volt_fluc_lol']['df']]['appendix'][frame_id], report)
                report.add_paragraph('')
    
    return 0
                
            
    
            
#maps the shrot case names against a more detailed version and returns it as a string. This can be expanded to cover more network cases in the future.    
def ident_case_name(case):
    if('high' in case):
        return 'High demand Network Scenario'
    elif('low' in case):
        return "Low demand Network Scenario"
    else:
        return case

def data_frame_to_docx_table(df, report, skiprows=0, skipcolumns=0 ):
    from docx.shared import RGBColor
    grey='A5A5A5'
    white="FFFFFF"
    # add a table to the end and create a reference variable
    # extra row is so we can add the header row
    problem_flag=0
    # add the header rows
    #identify whether pivot table or regular table
    #if pivot table
    if(df.index.names[0]!=None):
        #add header rows
        t = report.add_table(df.shape[0]+2, df.shape[1]+len(df.index.names))
        #t.style='ESCO Data'
        t.style='ListTable3-Accent3'
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="A5A5A5"/>'.format(nsdecls('w')))
        t.rows[1].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)

        for lvl_id in range(0, len(df.columns.levels[0])): #create first header line, skipping cells on the left equal to lengt of df.index.names, then label remaining cells with df.columns.levels[0]. merging cells equal to amount of length of df.columns.levels[1]
            col_id= len(df.index.names)+lvl_id*len(df.columns.levels[1])
            add_cell_text(t.cell(0, col_id), str(df.columns.levels[0][lvl_id]), text_color=white, cell_color=grey, bolt=True)
            #merge cell(s) right of the entry, equivalent to the amount of len(df.columns.levels[1])
        for col_id in range(0, len(df.index.names)):#add first part of second header row
            add_cell_text(t.cell(1, col_id), str(df.index.names[col_id]), text_color=white, cell_color=grey, bolt=True)
        for col_id in range (0, len(df.columns.levels[1])): #add secodn part of second header row
            add_cell_text(t.cell(1, col_id+len(df.index.names)), str(df.columns.levels[1][col_id]), text_color=white, cell_color=grey, bolt=True)
        #add body of table
        for row_id in range(0, df.shape[0]):
            #first couple of cells from index table
            for index_id in range(0, len(df.index.names)):
                if(hasattr(df.index, 'levels')):
                    entry=df.index.levels[index_id][df.index.codes[index_id][row_id]]
                    if(row_id>0):
                        prev_entry=df.index.levels[index_id][df.index.codes[index_id][row_id-1]]
                    else:
                        prev_entry=''
                    if(prev_entry!=entry):
                        t.cell(row_id+2, index_id).text=str(entry)
                    #elif(entry == prev_entry):
                     #   t.cell(row_id-1, index_id).merge(t.cell(row_id+2, index_id))
                        
                else:
                    entry=df.index[row_id] #only single index avaialble
                    if(row_id>0):
                        prev_entry=df.index[row_id-1] #only single index available
                    else:
                        prev_entry=''
                    if(prev_entry!=entry):
                        t.cell(row_id+2, index_id).text=str(entry)
            
                    
                    
                #if entry same as previous entry: merge cells and only make one entry
                
            #remaining cells from values table
            for col_id in range(0, len(df.values[0])):
                t.cell(row_id+2, col_id+len(df.index.names)).text=str(df.values[row_id][col_id])
                
            
        
    else: #regular table
        t = report.add_table(df.shape[0]+1, df.shape[1])
        #t.style='ESCO Data'
        t.style='ListTable3-Accent3'
        for j in range(df.shape[-1]):
            t.cell(0,j).text = df.columns[j]
        
        # add the rest of the data frame
        
        for i in range(0, df.shape[0]):
            for j in range(df.shape[-1]):
                if(str(df.values[i,j])=='no'):
                    cell=t.cell(i+1,j)
                    paragraph=cell.paragraphs[0]
                    run1=paragraph.add_run('no')
                    red = RGBColor(255, 0, 0)
                    run1.font.color.rgb = red                
                    #cell.add_paragraph(str(df.values[i,j]), color="red")
                    problem_flag=1
                else:
                    t.cell(i+1,j).text = str(df.values[i,j])
    
    #df.values
   
    return problem_flag

def add_cell_text(cell, content, text_color, cell_color, bolt):
    from docx.shared import RGBColor
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="A5A5A5"/>'.format(nsdecls('w'))) #shading element for table cell
    cell._tc.get_or_add_tcPr().append(shading_elm_1)
        
    paragraph=cell.paragraphs[0]
    run1=paragraph.add_run(content)
    color = RGBColor(int('0x'+str(text_color[0:2]),0), int('0x'+str(text_color[2:4]),0), int('0x'+str(text_color[4:6]),0))
    run1.font.color.rgb = color    

def createShortcut(target, path):
    # target = ModelCopyDir # directory to which the shortcut is created
    # path = main_folder + "\\model_copies.lnk"  #This is where the shortcut will be created
    shell = Dispatch('WScript.Shell')
    shortcut = shell.CreateShortCut(path)
    shortcut.Targetpath = target
    shortcut.save()


###############################################################################
# Define Project Paths
###############################################################################

#main_folder_path=os.path.dirname(os.path.dirname(os.path.dirname(__file__)))
script_dir=os.getcwd()
script_dir_up=os.path.abspath(os.path.join(script_dir, os.pardir))
main_folder_path=os.path.abspath(os.path.join(script_dir_up, os.pardir))
sys.path.append(main_folder_path+"\\PSSE_sim\\scripts\\Libs")
# Create directory for storing the results
if "OneDrive - OX2" in main_folder_path: # if the current folder is online (under OneDrive - OX2), create a new directory to store the result
    user = os.path.expanduser('~')
    main_path_out = main_folder_path.replace(user + "\OneDrive - OX2","C:\work") # Change the path from Onedrive to Local in c drive
    main_folder_out = createPath(main_path_out)
else: # if the main folder is not in Onedrive, then store the results in the same location with the model
    main_folder_out = main_folder_path
    
dir_path =  main_folder_out +"\\Plots\\steady_state"
createPath(dir_path)

# Create shortcut linking to result folder if it is not stored in the main folder path
if main_folder_out != main_folder_path:
    createShortcut(main_folder_out, main_folder_path + "\\Plots\\steady_state.lnk")
else: # if the output location is same as input location, then delete the links
    try:os.remove(main_folder_path + "\\Plots\\steady_state.lnk")
    except: pass

result_sheet_path = main_folder_out +"\\PSSE_sim\\result_data\\steady_state\\" + raw_SS_result_folder + "\\Steady State Analysis Results.xlsx"


###############################################################################
# Import additional functions
###############################################################################
import matplotlib.pyplot as plt
sys.path.append(r"C:\ProgramData\Anaconda2\Lib\site-packages")
sys.path.append(r"C:\Python27\Lib\site-packages")
import docx
import openpyxl
import re
from docx import Document, shape
from docx.oxml import OxmlElement, parse_xml
from docx.shared import Inches, Pt
from docx.oxml.ns import qn, nsdecls
from docx.enum.text import WD_BREAK
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.dml import MSO_THEME_COLOR_INDEX
import readtestinfo


###############################################################################
# Main function
###############################################################################

return_dict =  readtestinfo.readTestdef(main_folder_path+"\\test_scenario_definitions\\"+TestDefinitionSheet, ['ProjectDetails','ModelDetailsPSSE'])
ProjectDetailsDict = return_dict['ProjectDetails']
PSSEmodelDict = return_dict['ModelDetailsPSSE']

def main():
    
    # Read data input
    result_dfs=read_result_sheet()
    reportname_prefix= timestr+"-"+str(ProjectDetailsDict['NameShrt']+ str(simulation_batch_label))
    
    # Prepare summary tables and plots
    sumarise_results(result_dfs, summary_dfs)
    
    # save to excel summary.
    writer = pd.ExcelWriter(main_folder_out+"\\Plots\\steady_state\\"+reportname_prefix+"_SteadyStateResultsSummaryTable.xlsx",engine = 'xlsxwriter')
    for key1 in summary_dfs.keys():
        for key2 in summary_dfs[key1].keys():
            if summary_dfs[key1][key2] == {}:
                pass
            else:
                for result in summary_dfs[key1][key2]['summary']:
                    if not result.empty: result.to_excel(writer,sheet_name = str(key1)+str(key2),startrow=0)
                for result in summary_dfs[key1][key2]['appendix']:
                    if not result.empty: result.to_excel(writer,sheet_name = str(key1)+str(key2),startrow=20)
                for result in summary_dfs[key1][key2]['violations']:
                    if not result.empty: result.to_excel(writer,sheet_name = str(key1)+str(key2),startcol=12)
                for result in summary_dfs[key1][key2]['plot']:
                    worksheet = writer.sheets[str(key1)+str(key2)]
                    worksheet.insert_image('AA1', str(key1)+str(key2) + 'plot.png' )
    writer.save()  
    
    #Generate summary Word doc.
    report=initialise_report()
    replace_placeholders(report)
    add_report_intro(report)
    volt_lvl_report(report, summary_dfs, probs) #add description of the results along with table(s) and plots(s)
    line_loadings_report(report, summary_dfs, probs) #add description of the results along with table(s) and plots(s)
    bus_volt_fluct_gen_chng_report(report, summary_dfs, probs)
    bus_volt_fluct_cont_report(report, summary_dfs, probs)
    fautl_levels_report(report, summary_dfs, probs)
    add_conclusion(report, probs)
    add_appendices(report,summary_dfs)
    
    
    report.save(main_folder_out+"\\Plots\\steady_state\\"+reportname_prefix+"_SteadyStateResultsSummaryReport.docx")

if __name__ == '__main__':
    main()
