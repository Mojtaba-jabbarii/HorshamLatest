# -*- coding: utf-8 -*-
"""
Created on Mon Jan 22 16:04:32 2024

@author: 341510davu
"""

# import sys
import os, sys
import pandas as pd
import numpy as np
from datetime import datetime
import time
from contextlib import contextmanager
from win32com.client import Dispatch
timestr=time.strftime("%Y%m%d-%H%M%S")

try:
    from StringIO import StringIO
except ImportError:
    from io import StringIO
from io import BytesIO

###############################################################################
#USER INPUTS
###############################################################################
TestDefinitionSheet=r'20230828_SUM_TESTINFO_V1.xlsx'
raw_SS_result_folder = '20240123-112638_S5251'
simulation_batches=['S5251_PQcurve']
simulation_batch_label = simulation_batches[0]

cases = ["35degC_BESS_PV_0.9", "35degC_BESS_PV_1.0"]

Vcp = ["0.9","1.0"]
temp_cases = ["35degC","50degC"] #"35degC" # "50degC" #"40degC"
    
summary_dfs={'35degC':{}, '50degC':{},} # use input model names

probs={'35degC':0, '50degC':0}

###############################################################################
# Supporting functions
###############################################################################

def make_dir(dir_path, dir_name=""):
    dir_make = os.path.join(dir_path, dir_name)
    try:
        os.mkdir(dir_make)
    except OSError:
        pass
    return dir_make
        
def createPath(main_folder_out):
    path = os.path.normpath(main_folder_out)
    path_splits = path.split(os.sep) # Get the components of the path
    child_folder = r""+ path_splits[0] # Build up the output path from base drive
    for i in range(len(path_splits)-1):
        child_folder = child_folder + "\\" + path_splits[i+1]
        make_dir(child_folder)
    return child_folder

# Generate plots
def generate_plot(x_axis, y_axes, legends, label_x, label_y, title):
#    x_axis = df_vlt_lvl['Bus Name']
#    y_axes = [df_vlt_lvl['Voltage Level(pu) GenON'], df_vlt_lvl['Voltage Level(pu) GenOFF']]
#    legends = ['VL(pu) GenON', 'VL(pu) GenOFF']
#    label_x = 'Voltage(pu)'
#    label_y = 'Bus Names'
#    title = 'Voltage Levels'
    # size and positions
    fig = plt.figure(figsize=(7,5))
    fig.add_axes([0.1,0.1,0.8,0.8])
    colors = ['k', 'r', 'b', 'g']
    markers = ['-', '--', '^','*', '+']
    # axis names and ticks
    for i in range(len(y_axes)):
        ydata = y_axes[i]
        ldata = legends[i]
        if len(x_axis) > 1: xdata = x_axis[i]
        else: xdata = x_axis[0]
        plt.scatter(xdata, ydata, label = ldata, color = colors[i], marker = markers[i], s = 100)
        
#    plt.scatter(x_axis, y_axes[0], label = legends[0], color = 'k', marker = '*', s = 120)
#    plt.scatter(x_axis, y_axes[1], label = legends[1])
    plt.title(title, fontsize =12, color = 'black' )
    plt.ylabel(label_y, fontsize = 10)
    plt.xlabel(label_x, fontsize = 10)
    plt.legend()
#    plt.xticks(rotation = 45)
    plt.margins( tight = True)
    plt.grid()
    plt.minorticks_off()
    #plt.savefig(case +'voltage_levels'+ 'plot.png', bbox_inches = 'tight')
    #imgdata= StringIO.StringIO()
    #imgdata = StringIO()
    imgdata= BytesIO() # version issues
    plt.savefig(imgdata, bbox_inches = 'tight', dpi =200)    
    return imgdata


def initialise_report():
    #read report template 
    report=Document(main_folder_path+"\\Plots\\ReportTemplate.docx")
    return report


def replace_placeholders(report):
    replace_dict = {'[Project Name]':str(ProjectDetailsDict['Name']), '[Project Name Short]':str(ProjectDetailsDict['NameShrt']), '[Total Plant MW at POC]': str(ProjectDetailsDict['PlantMW']), 
                    '[Developer]': str(ProjectDetailsDict['Dev']), '[Network Service Provider]':str(ProjectDetailsDict['NSP']), '[Town]': str(ProjectDetailsDict['Town']), 
                    '[State]': str(ProjectDetailsDict['State']), '[Connection type]': str(ProjectDetailsDict['contyp']), '[POC Feeder]': str(ProjectDetailsDict['poc_fdr']),
                    '[Nominal POC voltage (kV)]': str(ProjectDetailsDict['VPOCkv']), '[PSSEversion]': str(PSSEmodelDict['PSSEversion']), '[Lot/DP]': str(ProjectDetailsDict['lot_dp']),
                    '[Address]': str(ProjectDetailsDict['addrs']), '[LGA]': str(ProjectDetailsDict['lga']), '[POC Substation]': str(ProjectDetailsDict['Sub']),
                    '[Plant Model]': str(ProjectDetailsDict['plnt_mdl'])
                    }
    for key,value in replace_dict.items():
        for p in report.paragraphs:
            if key in p.text:
                p.text = p.text.replace(key,value)
    #for p in report.paragraphs:
     #   inline = p.runs
      #  for j in range(0,len(inline)):
       #     for k,v in replace_dict.items():
        #        if k in inline[j].text:
         #           inline[j].text = inline[j].text.replace(k,v)
    for key,value in replace_dict.items():
        for table in report.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if key in p.text:
                            p.text = p.text.replace(key,value)
                        
                    #inline = p.runs
                    #for j in range(0,len(inline)):
                    #    for k,v in replace_dict.items():
                    #       if k in inline[j].text:
                    #           inline[j].text = inline[j].text.replace(k,v)
    
    return report    

def add_report_intro(report):
    plant_rating=ProjectDetailsDict['PlantMW']
    POC_name=ProjectDetailsDict['Sub']
    location=ProjectDetailsDict['Town']+", "+ProjectDetailsDict['State']
    #generate general description and intro based on ProjectDetailsDict
    #Headline
    '''
    p=report.add_paragraph('')
    run=p.add_run()
    run.add_break(WD_BREAK.PAGE)
    report.add_heading("Introduction", level=1 )
    intro_text="This report summarises the findings of a steady state study carried out for a "+str(plant_rating)+" MW generator connected to "+str(POC_name)+" in "+str(location)+"."
    intro_text+="\n\nThe Report consists of five separate sections, each of which illustrate the findings of the respective part of the study:"
    p=report.add_paragraph(intro_text)
    #add description of each subsection of the report    
    temp_text="1) The first part of the steady state analysis looks at Bus voltages under system normal conditions (i.e. no line outages or other contingencies. The results of the analysis are presented in section 2. Adding generation to a bus or a line may shift the voltage levels at that bus or surrounding buses due to the impact it has on on power flows in the area. "
    temp_text+="\nThe normal operating band of the NEM is between 0.9 p.u. voltage and 1.1 p.u. voltage. The results presented in that section will highlight the changes in relevant bus voltages due to the addition of the new generator."
    p=report.add_paragraph('')
    for style in report.styles:
        print("style.name == %s" % style.name)
    report.add_paragraph(temp_text, style='List Paragraph')
    report.add_paragraph('', style='List Paragraph')
    temp_text="2) The second part of the Steady State analysis investigates line loading and transformer loading under N-1 conditions, with and without the proposed generator. This reveals pre-existing issues as well as issues caused by the inclusion of the proposed generator. "
    report.add_paragraph(temp_text, style='List Paragraph')
    report.add_paragraph('', style='List Paragraph')
    temp_text="3) The third part of the Steady State analysis explores voltage stability for a change in generation output. If there is a sudden disconnect the plant output can drop from 100% to 0%, which will instantly change the voltage at surrounding buses. This is quantified and compared against applicable thresholds. "
    report.add_paragraph(temp_text, style='List Paragraph')
    report.add_paragraph('', style='List Paragraph')
    temp_text="4) The fourth part of the Steady State analysis focusses on voltage stability under credible contingencies. The voltage magnitude of the voltage fluctuations "
    report.add_paragraph(temp_text, style='List Paragraph')
    report.add_paragraph('', style='List Paragraph')
    temp_text="5) The last part of the Steady State analysis investigates Fault current at buses of interest. These fault levels are required to not exceed planning levels. This item is unlikely to be problematic in most instances as the contribution from inverter-based resources behind transformer and cable impedances is generally normally small compared against existing headroom."
    report.add_paragraph(temp_text, style='List Paragraph')
    report.add_paragraph('', style='List Paragraph')
    '''
    p=report.add_paragraph('')
    run=p.add_run()
    run.add_break(WD_BREAK.PAGE)
    report.add_heading("Simulation Results", level=1 )
    return 0

def pq_curve_report(report, summary_dfs, probs):
    report.add_heading("Reactive Power Capability", level=2 )
    temp_text="The reactive power capability of "+str(ProjectDetailsDict['Name'])+" is studied by changing the active power from minimum to maximum value and get the plant to response of maximum and minimum level of Q for each individual P value. "
    temp_text+="The studies is conducted for 0.9 p.u.; 1.0 p.u. and 1.1 p.u. voltage at the connection point. "
    temp_text+="The automatic assess standard require the plant to have capability of providing 0.395 time rated power of the power plant."
    report.add_paragraph(temp_text)
    report.add_paragraph('')
    #Add Summary tables
    report.add_heading("Summary of findings", level=3 )
    temp_text="The results for each temperature scenario of  "+str(ProjectDetailsDict['Name'])+" are listed in the tables below."
    report.add_paragraph(temp_text)
    for case in summary_dfs.keys():
        if(summary_dfs[case][Vcp[0]]!={}):
            report.add_heading(ident_case_name(case), level=4)#maybe make this a label of the table instead of making it a  heading.
            for frame_id in range(0, len(summary_dfs[case][Vcp[0]]['summary'])):#add summary tables of results to word doc.
                data_frame_to_docx_table(summary_dfs[case][Vcp[0]]['summary'][frame_id], report)
                report.add_paragraph('')
    #Add plots
    plots_present=False
    for case in summary_dfs.keys():
        if('plot' in summary_dfs[case][Vcp[0]].keys()):
            if(summary_dfs[case][Vcp[0]]['plot']!=[]):
                plots_present=True
    if(plots_present):
        report.add_heading("Plots", level=3)
        temp_text="The scatter plots show the PQ capability in the reference case(s) analysed in this study."
        report.add_paragraph(temp_text)
        for case in summary_dfs.keys():
            if(summary_dfs[case][Vcp[0]]!={}):
                report.add_heading(ident_case_name(case), level=4)#maybe make this a label of the table instead of making it a  heading.
                for frame_id in range(0, len(summary_dfs[case][Vcp[0]]['plot'])):#add summary tables of results to word doc.
                    report.add_picture(summary_dfs[case][Vcp[0]]['plot'][frame_id], Inches(6))
                    report.add_paragraph('')
        
    #Add overview of violations
        #add check whether any violations exist
    report.add_heading("Violations", level=3 )
    temp_text="If the plant cannot be able to provide an amount of reactive power equal to 0.395 * Prated at any point in normal operation, it is considered to be not meeting the AAS. "
    temp_text+="If the plant cannot meet the automatic access standard, it is required to provide the reason and mitigation."
    report.add_paragraph(temp_text)
    causer_flag=0
    for case in summary_dfs.keys():
        case_results_present=False
        viol_in_act_case=False
        if(summary_dfs[case][Vcp[0]]!={}):
            if('violations' in summary_dfs[case][Vcp[0]].keys()):
                case_results_present=True
                if(summary_dfs[case][Vcp[0]]['violations']!=[]):
                    viol_in_act_case=True
                    probs[Vcp[0]]=1
                    report.add_heading(ident_case_name(case), level=4)#maybe make this a label of the table instead of making it a  heading.
                    for frame_id in range(0, len(summary_dfs[case][Vcp[0]]['violations'])):#add summary tables of results to word doc.
                        problem_flag=data_frame_to_docx_table(summary_dfs[case][Vcp[0]]['violations'][frame_id], report)
                        if(problem_flag>0):
                            causer_flag=1#indicates whether the project is creating or exacerbating the issue.
                        report.add_paragraph('')
        if(not viol_in_act_case and case_results_present):
            report.add_paragraph('No voltage violations observed for '+ident_case_name(case)+'.') 
            
    if (causer_flag>0):
       probs[df_to_sheet['volt_levels']['df']]=2 #Add conclusion based on whether violations exist and provide some generic advice.
    return 0


def data_frame_to_docx_table(df, report, skiprows=0, skipcolumns=0 ):
    from docx.shared import RGBColor
    grey='A5A5A5'
    white="FFFFFF"
    # add a table to the end and create a reference variable
    # extra row is so we can add the header row
    problem_flag=0
    # add the header rows
    #identify whether pivot table or regular table
    #if pivot table
    if(df.index.names[0]!=None):
        #add header rows
        t = report.add_table(df.shape[0]+2, df.shape[1]+len(df.index.names))
        #t.style='ESCO Data'
        t.style='ListTable3-Accent3'
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="A5A5A5"/>'.format(nsdecls('w')))
        t.rows[1].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)

        for lvl_id in range(0, len(df.columns.levels[0])): #create first header line, skipping cells on the left equal to lengt of df.index.names, then label remaining cells with df.columns.levels[0]. merging cells equal to amount of length of df.columns.levels[1]
            col_id= len(df.index.names)+lvl_id*len(df.columns.levels[1])
            add_cell_text(t.cell(0, col_id), str(df.columns.levels[0][lvl_id]), text_color=white, cell_color=grey, bolt=True)
            #merge cell(s) right of the entry, equivalent to the amount of len(df.columns.levels[1])
        for col_id in range(0, len(df.index.names)):#add first part of second header row
            add_cell_text(t.cell(1, col_id), str(df.index.names[col_id]), text_color=white, cell_color=grey, bolt=True)
        for col_id in range (0, len(df.columns.levels[1])): #add secodn part of second header row
            add_cell_text(t.cell(1, col_id+len(df.index.names)), str(df.columns.levels[1][col_id]), text_color=white, cell_color=grey, bolt=True)
        #add body of table
        for row_id in range(0, df.shape[0]):
            #first couple of cells from index table
            for index_id in range(0, len(df.index.names)):
                if(hasattr(df.index, 'levels')):
                    entry=df.index.levels[index_id][df.index.codes[index_id][row_id]]
                    if(row_id>0):
                        prev_entry=df.index.levels[index_id][df.index.codes[index_id][row_id-1]]
                    else:
                        prev_entry=''
                    if(prev_entry!=entry):
                        t.cell(row_id+2, index_id).text=str(entry)
                    #elif(entry == prev_entry):
                     #   t.cell(row_id-1, index_id).merge(t.cell(row_id+2, index_id))
                        
                else:
                    entry=df.index[row_id] #only single index avaialble
                    if(row_id>0):
                        prev_entry=df.index[row_id-1] #only single index available
                    else:
                        prev_entry=''
                    if(prev_entry!=entry):
                        t.cell(row_id+2, index_id).text=str(entry)
            
                    
                    
                #if entry same as previous entry: merge cells and only make one entry
                
            #remaining cells from values table
            for col_id in range(0, len(df.values[0])):
                t.cell(row_id+2, col_id+len(df.index.names)).text=str(df.values[row_id][col_id])
                
            
        
    else: #regular table
        t = report.add_table(df.shape[0]+1, df.shape[1])
        #t.style='ESCO Data'
        t.style='ListTable3-Accent3'
        for j in range(df.shape[-1]):
            t.cell(0,j).text = df.columns[j]
        
        # add the rest of the data frame
        
        for i in range(0, df.shape[0]):
            for j in range(df.shape[-1]):
                if(str(df.values[i,j])=='no'):
                    cell=t.cell(i+1,j)
                    paragraph=cell.paragraphs[0]
                    run1=paragraph.add_run('no')
                    red = RGBColor(255, 0, 0)
                    run1.font.color.rgb = red                
                    #cell.add_paragraph(str(df.values[i,j]), color="red")
                    problem_flag=1
                else:
                    t.cell(i+1,j).text = str(df.values[i,j])
    
    #df.values
   
    return problem_flag

def add_cell_text(cell, content, text_color, cell_color, bolt):
    from docx.shared import RGBColor
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="A5A5A5"/>'.format(nsdecls('w'))) #shading element for table cell
    cell._tc.get_or_add_tcPr().append(shading_elm_1)
        
    paragraph=cell.paragraphs[0]
    run1=paragraph.add_run(content)
    color = RGBColor(int('0x'+str(text_color[0:2]),0), int('0x'+str(text_color[2:4]),0), int('0x'+str(text_color[4:6]),0))
    run1.font.color.rgb = color    

def createShortcut(target, path):
    # target = ModelCopyDir # directory to which the shortcut is created
    # path = main_folder + "\\model_copies.lnk"  #This is where the shortcut will be created
    shell = Dispatch('WScript.Shell')
    shortcut = shell.CreateShortCut(path)
    shortcut.Targetpath = target
    shortcut.save()
    
def read_result_sheet(): 
#    genoff = '_off' 
#    genon = '_on'
    sheets_dict  = pd.read_excel(result_sheet_path, sheet_name = None)
    result_dfs={}
    
    
    for case in temp_cases:
#        result_dfs[case] = pd.DataFrame()
        result_dfs[case] = pd.DataFrame()
    
        for name,sheet in sheets_dict.items():
            if case in name:
                df_out = pd.DataFrame(data=sheet, columns=['1 V_Inv','1 V_Inv2','2 P_Inv','2 P_Inv2','3 Q_Inv','3 Q_Inv2','7 Q_Poc','6 P_Poc']).round(3)
#                result_dfs[case] = pd.concat([result_dfs[case],df_out],axis = 1)
                
                vol_txt = name[-3:]
                if vol_txt == '0.9': vol_txt = '90'
                if vol_txt == '1.0': vol_txt = '100'
                if vol_txt == '1.1': vol_txt = '110'
#                result_dfs[case][vol_txt] = pd.DataFrame(data=sheet, columns=['1 V_Inv','1 V_Inv2','2 P_Inv','2 P_Inv2','3 Q_Inv','3 Q_Inv2','7 Q_Poc','6 P_Poc']).round(3)
                result_dfs[case][vol_txt] = df_out
        #        result_dfs[case].rename(columns = {'1 V_Inv':"Vinv1"+vol_txt,'1 V_Inv2':"Vinv2"+vol_txt,'2 P_Inv':"Pinv1"+vol_txt,'2 P_Inv2':"Pinv2"+vol_txt,'3 Q_Inv':"Qinv1"+vol_txt,'3 Q_Inv2':"Qinv2"+vol_txt,'7 Q_Poc':"Q_Poc",'6 P_Poc':"P_Poc"}, inplace = True)
        #        else:
        #            pass
        

    return result_dfs


#maps the shrot case names against a more detailed version and returns it as a string. This can be expanded to cover more network cases in the future.    
def ident_case_name(case):
    if('35deg' in case):
        return '35 degree scenario'
    elif('50deg' in case):
        return "50 degree scenario"
    else:
        return case



# prepare sumaries into data frames
def sumarise_results(result_dfs, summary_dfs):

    for case in temp_cases:
        summary_dfs[case] = pd.DataFrame()

        for key,value in result_dfs.items():
            if case in key: # summarise the voltage level results with and without plant
                #Summary table
#                df_name = case + "_BESS_PV"
                vol_txt = key[-3:]
                df_pq_curve = result_dfs[key]
#                pvt_df = pd.pivot_table(data = df_pq_curve,index = ['Q_Poc [MVAr]'],values = ['Q_Poc [MVAr]','P_Poc [MW]'] )
#                pvt_df.style.applymap(hl_vltg_lvls_violation,subset = ['Voltage Level(pu) GenOFF','Voltage Level(pu) GenON']).format({'Voltage Level(pu) GenOFF':'{0:,.3f}','Voltage Level(pu) GenON':'{0:,.3f}'})
                summary_dfs[case][vol_txt]['summary'] = []
                summary_dfs[case][vol_txt]['summary'].append(df_pq_curve)
                
                # Voltage levels plot
#                imgdata = generate_plot(df_pq_curve['Q_Poc [MVAr]'], [df_pq_curve['P_Poc [MW]']], ['PQ curve'], 'Bus Names', 'Voltage(pu)', 'Voltage Levels')
                imgdata = generate_plot([df_pq_curve['Q_Poc']], [df_pq_curve["P_Poc"]], ['PQ curve'], 'Q_POC [MVAr]', 'P_Poc [MW]', 'Reactive Power Capability')
                summary_dfs[case][vol_txt]['plot'] = []
                summary_dfs[case][vol_txt]['plot'].append(imgdata)
                
                #Violation voltage level table
#                vl_min_violation = df_pq_curve.loc[df_pq_curve['Voltage Level(pu) GenOFF']<0.9]
#                vl_max_violation = df_pq_curve.loc[df_vlt_lvl['Voltage Level(pu) GenOFF']>1.1]
#                vl_violation = pd.concat([vl_max_violation,vl_min_violation], axis =0)
#                if not vl_violation.empty: vl_violation['Pass'] = vl_violation.apply(vl_check, axis = 1)
                summary_dfs[case][vol_txt]['violations'] = []
#                summary_dfs[case][Vcp[0]]['violations'].append(vl_violation)
                
                #Empty apendix
                summary_dfs[case][vol_txt]['appendix'] = []
                

def save_output(reportname_prefix):
    from openpyxl import load_workbook
    
    wb_in = xl.load_workbook(filename=result_sheet_path)# Read data input
    
#    reportname_prefix= timestr+"-"+str(ProjectDetailsDict['NameShrt']+ str(simulation_batch_label))
    writer = pd.ExcelWriter(main_folder_out+"\\Plots\\PQ_curve\\"+reportname_prefix+"_PQcurve.xlsx",engine = 'xlsxwriter') # Preparing for exporting the result
    
    for case in temp_cases:
        df_out = pd.DataFrame()
        for i in range(len(wb_in.worksheets)):
            if case in str(wb_in.worksheets[i]):
                ws_in1 = wb_in.worksheets[i]
                ws_in1.delete_cols(idx=8, amount =4) # current
                ws_in1.delete_cols(idx=1) # index
                data1 = ws_in1.values
                df_out1 = pd.DataFrame.from_dict(data = data1)  
#                imgdata = generate_plot([df_out1['7 Q_Poc']], [df_out1["6 P_Poc"]], ['PQ curve'], 'Q_POC [MVAr]', 'P_Poc [MW]', 'Reactive Power Capability')
                
                df_out = pd.concat([df_out,df_out1],axis = 1)
        df_out.to_excel(writer, sheet_name = str(case))

    writer.close() 

        
        
#    sheets_dicts  = pd.read_excel(result_sheet_path, sheet_name = None)
#    df_out={}
#    for case in temperature:
#        for name,sheet in sheets_dicts.items():
##            vol_txt = "_90"
#            vol_txt = name[-3:]
#            if case in name:
#                df_out[case] = pd.DataFrame(data=sheet, columns=['1 V_Inv','1 V_Inv2','2 P_Inv','2 P_Inv2','3 Q_Inv','3 Q_Inv2','7 Q_Poc','6 P_Poc']).round(3)
#                df_out[case].rename(columns = {'1 V_Inv':"Vinv1"+vol_txt,'1 V_Inv2':"Vinv2"+vol_txt,'2 P_Inv':"Pinv1"+vol_txt,'2 P_Inv2':"Pinv2"+vol_txt,'3 Q_Inv':"Qinv1"+vol_txt,'3 Q_Inv2':"Qinv2"+vol_txt,'7 Q_Poc':"Q_Poc"+vol_txt,'6 P_Poc':"P_Poc"+vol_txt}, inplace = True)
#            else:
#                pass
#    writer = pd.ExcelWriter(main_folder_out+"\\Plots\\PQ_curve\\"+reportname_prefix+"_PQcurve.xlsx",engine = 'xlsxwriter') # Preparing for exporting the result
#    df_out.to_excel(writer,sheet_name = "_PQcurve_cnv" ,startrow=1)
#    writer.close() 
    
###############################################################################
# Define Project Paths
###############################################################################

#main_folder_path=os.path.dirname(os.path.dirname(os.path.dirname(__file__)))
script_dir=os.getcwd()
script_dir_up=os.path.abspath(os.path.join(script_dir, os.pardir))
main_folder_path=os.path.abspath(os.path.join(script_dir_up, os.pardir))
sys.path.append(main_folder_path+"\\PSSE_sim\\scripts\\Libs")
# Create directory for storing the results
if "OneDrive - OX2" in main_folder_path: # if the current folder is online (under OneDrive - OX2), create a new directory to store the result
    user = os.path.expanduser('~')
    main_path_out = main_folder_path.replace(user + "\OneDrive - OX2","C:\work") # Change the path from Onedrive to Local in c drive
    main_folder_out = createPath(main_path_out)
else: # if the main folder is not in Onedrive, then store the results in the same location with the model
    main_folder_out = main_folder_path
    
dir_path =  main_folder_out +"\\Plots\\PQ_curve"
make_dir(dir_path)

# Create shortcut linking to result folder if it is not stored in the main folder path
if main_folder_out != main_folder_path:
    createShortcut(main_folder_out, main_folder_path + "\\Plots\\PQ_curve.lnk")
else: # if the output location is same as input location, then delete the links
    try:os.remove(main_folder_path + "\\Plots\\PQ_curve.lnk")
    except: pass

result_sheet_path = main_folder_out +"\\PSSE_sim\\result_data\\PQ_curve\\" + raw_SS_result_folder + "\\S5251_PQ curve results.xlsx"

###############################################################################
# Import additional functions
###############################################################################
import matplotlib.pyplot as plt
sys.path.append(r"C:\ProgramData\Anaconda2\Lib\site-packages")
sys.path.append(r"C:\Python27\Lib\site-packages")
import docx
import openpyxl as xl
import re
from docx import Document, shape
from docx.oxml import OxmlElement, parse_xml
from docx.shared import Inches, Pt
from docx.oxml.ns import qn, nsdecls
from docx.enum.text import WD_BREAK
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.dml import MSO_THEME_COLOR_INDEX
import readtestinfo

###############################################################################
# Main function
###############################################################################

return_dict =  readtestinfo.readTestdef(main_folder_path+"\\test_scenario_definitions\\"+TestDefinitionSheet, ['ProjectDetails','ModelDetailsPSSE'])
ProjectDetailsDict = return_dict['ProjectDetails']
PSSEmodelDict = return_dict['ModelDetailsPSSE']

def main():
    
    # Read data input
    result_dfs=read_result_sheet()
    reportname_prefix= timestr+"-"+str(ProjectDetailsDict['NameShrt']+ str(simulation_batch_label))
    
    # Prepare summary tables and plots
    sumarise_results(result_dfs, summary_dfs)
    
    # save to an excel summary based on the temperature cases.
    save_output(reportname_prefix)
  
    
    #Generate summary Word doc.
    report=initialise_report()
    replace_placeholders(report)
    add_report_intro(report)
    pq_curve_report(report, summary_dfs, probs) #add description of the results along with table(s) and plots(s)
    add_conclusion(report, probs)
    report.save(main_folder_out+"\\Plots\\PQ_curve\\"+reportname_prefix+"_PQ_curveReport.docx")

if __name__ == '__main__':
    main()
