# -*- coding: utf-8 -*-
"""
Created on Mon Nov  9 15:49:16 2020

@author: Mervin Kall

SMIB PLOT tool:

Script allows to generate plots of test results generated with the SMIB test tools, The plots are defined in the variable "reports"

plot range: for Benchmarking xrange is by default from 0.0 to the point where the first signal ends (e.g. if PSS/E simulation ran 5 seconds longer, those last 5 seconds will not be included in plot)

    
Ideas:
    For the benchmarking plots and also 5.2.5.13 and 5.2.5.5 settling time checks, the simulation time at which the step is applied, as well as the time at which the simulation reaches steady state again are required
    For any test involving a profile of any kind, this script could detect from the test metadata where the steps are located in the profile and take this into account for plotting settling bands, tolerance bands etc.

    Settling (GSMG) always based on first step --> try and limit tests to only one step (and back) per scenario when using GSMG bands
        --> detect first step from profile gradient: 
                set startwindow and endwiondow to 100 ms
                set starttime to >100ms before step
                set endtime to >100ms before either second step or end of scenario (whichever comes first)
                                
    include test profile in metadata, because the test profile may change in the config sheet after the test batch has been run
    
    
Versions
    V3.4: update the chanel plot for LSF
    #11/8/2022: correcting the fault list in MFRT test: remove the repeaation of first fault
    #V3.4.b: included option to export the rise, settle and recovery time: 'rise_t':0, 'set_t':0, 'rec_t':0
    #V3.4.c: Include DMAT id in the summary table
    #V3.4.d: include Generator data channels
    #V3.4.e: only consider one INV data as they are idendical
    #V4_LL08: update the summary table
    
"""

#------------------------------------------------------------------------------
# USER INPUT
#------------------------------------------------------------------------------
TestDefinitionSheet=r'20240403_HSFBESS_TESTINFO_V1.xlsx'
"""
# report_types = ["BENCH", #--> Will expect plots of type overlay. Will only include plots for test cases that are availabel in two or more of the specified datasets (to be able to create overlay)
                    "DMAT", #--> Will plot everything that is available, Either one or two sets of result data can be provided. Should include both single dataset plots and overlays. may contain different chapters (to allow for plots to change depending on test type)
                    "GPS"] #--> Various types of chapters and plots. 
# chapter_types = ["S5.2.5.5_inj", "S5.2.5.1"]
    """

# different datasets for SMIB and Network tests
datasets_PSSE = {'label':'PSSE_Data', 'path':r"PSSE_sim\result_data\dynamic_smib\20241007-1622_shallow_fault_dbg", 'ID': 0, 'timeID':'Time(s)', 'timeoffset':-3.0,#-3.0,
                  'calcCurrents':[{"P":"P_POC1", "Q":"Q_POC1", "V":"U_POC1", "nameLabel":"PLANT", "scaling":-0.007815126, }, # #negative due to the reversed Q measurement
                                  {"P":"P_LV1", "Q":"Q_LV1", "V":"U_LV1", "nameLabel":"PV", "scaling":0.00661375661, }, #1/Sbase_POC for Iq at POC or 1/Sbase_INV for Iq at INV -> convert QMVAr to Qpu for calculation
                                  {"P":"P_LV2", "Q":"Q_LV2", "V":"U_LV2", "nameLabel":"BESS", "scaling":0.0056818, },],  
                  'calPFs':[{"P":"P_POC1", "Q":"Q_POC1", "nameLabel":"PLANT", "scaling":-1.0, } ], } #calculate power factor from P and Q results scaling -1 due to reversed power measure
datasets_PSCAD = {'label':'PSCAD_Data', 'path':r"PSCAD_sim\result_data\dynamic_smib\20241007-1624_shallow_fault_dbg", 'ID': 4, 'timeID':'time(s)', 'timeoffset':-3.0,#-3.0, 
                  'calcCurrents':[{"P":"PLANT_P_HV", "Q":"PLANT_Q_HV", "V":"PLANT_V_HV_pu", "nameLabel":"PLANT", "scaling":0.007815126, }, 
                                  {"P":"PCU1_P_LV", "Q":"PCU1_Q_LV", "V":"PCU1_V_LV_pu", "nameLabel":"PV", "scaling":36.0, },
                                   {"P":"PCU2_P_LV", "Q":"PCU2_Q_LV", "V":"PCU2_V_LV_pu", "nameLabel":"BESS", "scaling":40.0, }],
                  'calPFs':[{"P":"PLANT_P_HV", "Q":"PLANT_Q_HV", "nameLabel":"PLANT", "scaling":1.0, } ],} #inverter apaprent power rating
datasets_nw_genon = {'label':'PSSE_genon', 'path':r"PSSE_sim\result_data\dynamic_network\Final Results GPS NEM genon", 'ID': 2, 'timeID':'Time(s)', 'timeoffset':-2.0, #20240305-101744_S5255_NW\HighLoad_genon
                  'calcCurrents':[{"P":"P_POC1", "Q":"Q_POC1", "V":"U_POC1", "nameLabel":"PLANT", "scaling":-0.007815126, }, # #negative due to the reversed Q measurement
                                  {"P":"P_LV1", "Q":"Q_LV1", "V":"U_LV1", "nameLabel":"PV", "scaling":0.00661375661, }, #1/Sbase_POC for Iq at POC or 1/Sbase_INV for Iq at INV -> convert QMVAr to Qpu for calculation
                                  {"P":"P_LV2", "Q":"Q_LV2", "V":"U_LV2", "nameLabel":"BESS", "scaling":0.0056818, },], }
datasets_nw_genoff = {'label':'PSSE_genoff', 'path':r"PSSE_sim\result_data\dynamic_network\Final Results GPS NEM genoff", 'ID': 1, 'timeID':'Time(s)', 'timeoffset':-2.0,
                  'calPFs':[{"P":"P_POC1", "Q":"Q_POC1", "nameLabel":"PLANT", "scaling":1.0, } ],}

data_PSCAD_Flatrun1 = {'label':'PSCAD_Data1', 'path':r"PSCAD_sim\result_data\dynamic_smib\Final Results DMAT SMIB\01_FlatRun1", 'ID': 5, 'timeID':'time(s)', 'timeoffset':-3.0, }
data_PSCAD_Flatrun2 = {'label':'PSCAD_Data2', 'path':r"PSCAD_sim\result_data\dynamic_smib\Final Results DMAT SMIB\01_FlatRun2", 'ID': 6, 'timeID':'time(s)', 'timeoffset':-3.0, }
data_PSCAD_Flatrun3 = {'label':'PSCAD_Data3', 'path':r"PSCAD_sim\result_data\dynamic_smib\Final Results DMAT SMIB\01_FlatRun3", 'ID': 7, 'timeID':'time(s)', 'timeoffset':-3.0, }

# selected datasets for the corresponding report
datasets = [datasets_PSCAD, datasets_PSSE]#, datasets_PSCAD, data_PSCAD_Flatrun1, data_PSCAD_Flatrun2, data_PSCAD_Flatrun3] # [datasets_PSSE, datasets_PSCAD,datasets_nw_genon,datasets_nw_genoff]
report_types = ["DMAT"] #["GPS", "BENCH", "DMAT", 'NetworkEvent']
chapter_types = ["general"]#["01_FlatRun","02_UFault","03_MFRT","04_TOV","05_SptChange","06_FreqChange","07_VolChange","08_ORT","09_AngChange","10_SCR1","11_FRT","12_IrrChange"]
 #"S5257", "S52511",GPS["general","S5253","S5254","S5255","S5255_HighLoad","S5255_LowLoad","S52512_HighLoad","S52512_LowLoad","S52513_HighLoad","S52513_LowLoad",] if chapter_types == [], then include all chapters into the report
                            #DMAT["general","01_FlatRun","02_UFault","03_MFRT","04_TOV","05_SptChange","06_FreqChange","07_VolChange","08_ORT","09_AngChange","10_SCR1","11_FRT","12_IrrChange"] following Table 22, Appendix A2 DMAT guideline 2021




channels_lib = { #to be used in the ploting each graph below
                'PSCAD':    {'V_POC':{'dataset':4, 'name':"PLANT_V_HV_pu", 'leg':'Voltage HV', 'offset':0.0, 'scale':1.0},
                             'P_POC':{'dataset':4, 'name':"PLANT_P_HV", 'leg':'PLANT_P_HV', 'offset':0.0, 'scale':1.0},
                             'Q_POC':{'dataset':4, 'name':"PLANT_Q_HV", 'leg':'PLANT_Q_HV', 'offset':0.0, 'scale':1.0},
                             'PF_POC':{'dataset':4, 'name':"PF_PLANT", 'leg':'Powerfactor', 'offset':0.0, 'scale':1.0},
                             'F_POC':{'dataset':4, 'name':"Hz_POI", 'leg':'PLANT_Freq', 'offset':0.0, 'scale':1.0},
                             'Iq_POC':{'dataset':4, 'name':"Iq_PLANT", 'leg':'Iq_PLANT', 'offset':0.0, 'scale':1.0},
                             'Ang_POC':{'dataset':4, 'name':"phAngSource", 'leg':'Angle_POC', 'offset':0.0, 'scale':1.0},
                             'FRT_PPC':{'dataset':4, 'name':"FrtActive", 'leg':'FRT_PPC', 'offset':0.0, 'scale':1.0},
                             'Vspt_POC':{'dataset':4, 'name':"PoiVolSpt_pu", 'leg':'Voltage stp', 'offset':0.0, 'scale':1.0},
                             'Pspt_POC':{'dataset':4, 'name':"Pspt_MW", 'leg':'Active Power stp', 'offset':0.0, 'scale':1.0},
                             'Qspt_POC':{'dataset':4, 'name':"Qspt_MVAr", 'leg':'Reactive Power stp', 'offset':0.0, 'scale':1.0},
                             'PFspt_POC':{'dataset':4, 'name':"PoiPfSpt", 'leg':'PF stp', 'offset':0.0, 'scale':1.0},
                             'V_INV1':{'dataset':4, 'name':"PCU1_V_LV_pu", 'leg':'Voltage PV', 'offset':0.0, 'scale':1.0},
                             'V_INV2':{'dataset':4, 'name':"PCU2_V_LV_pu", 'leg':'Voltage BESS', 'offset':0.0, 'scale':1.0},
                             'P_INV1':{'dataset':4, 'name':"PCU1_P_LV", 'leg':'P_PV', 'offset':0.0, 'scale':36.0},
                             'P_INV2':{'dataset':4, 'name':"PCU2_P_LV", 'leg':'P_BESS', 'offset':0.0, 'scale':40.0},
                             'Q_INV1':{'dataset':4, 'name':"PCU1_Q_LV", 'leg':'Q_PV', 'offset':0.0, 'scale':36.0},
                             'Q_INV2':{'dataset':4, 'name':"PCU2_Q_LV", 'leg':'Q_BESS', 'offset':0.0, 'scale':40.0},
                             'FRT_INV1':{'dataset':4, 'name':"FRT_flag_PV", 'leg':'PV FRT', 'offset':0.0, 'scale':1.0},
                             'FRT_INV2':{'dataset':4, 'name':"FRT_flag_BESS", 'leg':'BESS FRT', 'offset':0.0, 'scale':1.0},
#                             'FRT_INV1':{'dataset':4, 'name':"FRT_flag_PV", 'leg':'PV FRT', 'offset':-97.0, 'scale':1.0},
#                             'FRT_INV2':{'dataset':4, 'name':"FRT_flag_BESS", 'leg':'BESS FRT', 'offset':0.0, 'scale':0.01},
                             'Pcmd_PPC_INV1':{'dataset':4, 'name':"Pcmd_PPC_to_INV_PV", 'leg':'P_cmd to PV', 'offset':0.0, 'scale':151.2}, #S_PV
                             'Pcmd_PPC_INV2':{'dataset':4, 'name':"Pcmd_PPC_to_INV_BESS", 'leg':'P_cmd to BESS', 'offset':0.0, 'scale':168}, #S_BESS
                             'Qcmd_PPC_INV1':{'dataset':4, 'name':"Qcmd_PPC_to_INV_PV", 'leg':'Q_cmd to PV', 'offset':0.0, 'scale':90.72}, #S_PV*0.6
                             'Qcmd_PPC_INV2':{'dataset':4, 'name':"Qcmd_PPC_to_INV_BESS", 'leg':'Q_cmd to BESS', 'offset':0.0, 'scale':100.8}, #S_BESS*0.6 
#                             'V_MV':{'dataset':4, 'name':"PLANT_V_MV", 'leg':'Voltage MV', 'offset':0.0, 'scale':1.0},
                             'tap position':{'dataset':4, 'name':"Tap_position", 'leg':'Tap position', 'offset':0.0, 'scale':1.0},
                             'tap ratio':{'dataset':4, 'name':"maintap", 'leg':'tap ratio', 'offset':0.0, 'scale':1.0},
                             
                             'V_MV':{'dataset':4, 'name':"PLANT_V_MV", 'leg':'Voltage MV', 'offset':0.0, 'scale':1.0},
                             'P_MV':{'dataset':4, 'name':"PLANT_P_MV", 'leg':'PLANT_P_MV', 'offset':0.0, 'scale':1.0},
                             'Q_MV':{'dataset':4, 'name':"PLANT_Q_MV", 'leg':'PLANT_Q_MV', 'offset':0.0, 'scale':1.0},
                             
                             },
                'PSSE':     {'V_POC':{'dataset':0, 'name':"U_POC1", 'leg':'Voltage HV', 'offset':0.0, 'scale':1.0},
                             'P_POC':{'dataset':0, 'name':"P_POC1", 'leg':'PLANT_P_HV', 'offset':0.0, 'scale':-1.0},
                             'Q_POC':{'dataset':0, 'name':"Q_POC1", 'leg':'PLANT_Q_HV', 'offset':0.0, 'scale':-1.0},
                             'PF_POC':{'dataset':0, 'name':"PF_PLANT", 'leg':'Powerfactor', 'offset':0.0, 'scale':1.0},
                             'F_POC':{'dataset':0, 'name':"F_POC1", 'leg':'PLANT_Freq', 'offset':50.0, 'scale':50.0},
                             'Iq_POC':{'dataset':0, 'name':"Iq_PLANT", 'leg':'Iq_PLANT', 'offset':0.0, 'scale':1.0},
                             'Ang_POC':{'dataset':0, 'name':"ANG_POC1", 'leg':'Angle_POC', 'offset':0.0, 'scale':1.0},
                             'FRT_PPC':{'dataset':0, 'name':"FRTACTIVE", 'leg':'FRT_PPC', 'offset':0.0, 'scale':1.0},
                             'Vspt_POC':{'dataset':0, 'name':"VREF_POC", 'leg':'Voltage stp', 'offset':0.0, 'scale':1.0},
                             'Pspt_POC':{'dataset':0, 'name':"PREF_POC", 'leg':'Active Power stp', 'offset':0.0, 'scale':0.001},
                             'Qspt_POC':{'dataset':0, 'name':"QREF_POC", 'leg':'Reactive Power stp', 'offset':0.0, 'scale':0.001},
                             'PFspt_POC':{'dataset':0, 'name':"PFREF", 'leg':'PF stp', 'offset':0.0, 'scale':1.0},
                             'V_INV1':{'dataset':0, 'name':"U_LV1", 'leg':'Voltage PV', 'offset':0.0, 'scale':1.0},
                             'V_INV2':{'dataset':0, 'name':"U_LV2", 'leg':'Voltage BESS', 'offset':0.0, 'scale':1.0},
                             'P_INV1':{'dataset':0, 'name':"P_LV1", 'leg':'P_PV', 'offset':0.0, 'scale':1.0},
                             'P_INV2':{'dataset':0, 'name':"P_LV2", 'leg':'P_BESS', 'offset':0.0, 'scale':1.0},
                             'Q_INV1':{'dataset':0, 'name':"Q_LV1", 'leg':'Q_PV', 'offset':0.0, 'scale':1.0},
                             'Q_INV2':{'dataset':0, 'name':"Q_LV2", 'leg':'Q_BESS', 'offset':0.0, 'scale':1.0},
                             'FRT_INV1':{'dataset':0, 'name':"INV1_FRT_STATE", 'leg':'PV FRT', 'offset':0.0, 'scale':1.0},
                             'FRT_INV2':{'dataset':0, 'name':"INV2_FRT_FLAG", 'leg':'BESS FRT', 'offset':0.0, 'scale':1.0},
                             'Pcmd_PPC_INV1':{'dataset':0, 'name':"P_CMD_PV", 'leg':'P_cmd to PV', 'offset':0.0, 'scale':151.2}, #S_PV
                             'Pcmd_PPC_INV2':{'dataset':0, 'name':"P_CMD_BESS", 'leg':'P_cmd to BESS', 'offset':0.0, 'scale':168}, #S_BESS
                             'Qcmd_PPC_INV1':{'dataset':0, 'name':"Q_CMD_PV", 'leg':'Q_cmd to PV', 'offset':0.0, 'scale':90.72}, #S_PV*0.6
                             'Qcmd_PPC_INV2':{'dataset':0, 'name':"Q_CMD_BESS", 'leg':'Q_cmd to BESS', 'offset':0.0, 'scale':100.8}, #S_BESS*0.6
                             
                             'V_MV':{'dataset':0, 'name':"U_334091", 'leg':'Voltage MV', 'offset':0.0, 'scale':1.0},
                             'P_MV':{'dataset':0, 'name':"P_334091", 'leg':'PLANT_P_MV', 'offset':0.0, 'scale':1.0},
                             'Q_MV':{'dataset':0, 'name':"Q_334091", 'leg':'PLANT_Q_MV', 'offset':0.0, 'scale':1.0},
                             
                             },
                'NW_post':   {'V_POC':{'dataset':2, 'name':"U_POC1", 'leg':'after-SUMSF', 'offset':0.0, 'scale':1.0},
                             'P_POC':{'dataset':2, 'name':"P_POC1", 'leg':'PLANT_P_HV', 'offset':0.0, 'scale':-1.0},
                             'Q_POC':{'dataset':2, 'name':"Q_POC1", 'leg':'PLANT_Q_HV', 'offset':0.0, 'scale':-1.0},
                             'PF_POC':{'dataset':2, 'name':"PF_PLANT", 'leg':'Powerfactor', 'offset':0.0, 'scale':1.0},
                             'F_POC':{'dataset':2, 'name':"F_POC1", 'leg':'PLANT_Freq', 'offset':50.0, 'scale':50.0},
                             'Iq_POC':{'dataset':2, 'name':"Iq_PLANT", 'leg':'Iq_PLANT', 'offset':0.0, 'scale':1.0},
                             'Ang_POC':{'dataset':2, 'name':"ANG_POC1", 'leg':'Angle_POC', 'offset':0.0, 'scale':1.0},
                             'FRT_PPC':{'dataset':2, 'name':"FRTACTIVE", 'leg':'FRT_PPC', 'offset':0.0, 'scale':1.0},
                             'Vspt_POC':{'dataset':2, 'name':"VREF_POC", 'leg':'Voltage stp', 'offset':0.0, 'scale':1.0},
                             'Pspt_POC':{'dataset':2, 'name':"PREF_POC", 'leg':'Active Power stp', 'offset':0.0, 'scale':0.001},
                             'Qspt_POC':{'dataset':2, 'name':"QREF_POC", 'leg':'Reactive Power stp', 'offset':0.0, 'scale':0.001},
                             'PFspt_POC':{'dataset':2, 'name':"PFREF", 'leg':'PF stp', 'offset':0.0, 'scale':1.0},
                             'V_INV1':{'dataset':2, 'name':"U_LV1", 'leg':'Voltage PV', 'offset':0.0, 'scale':1.0},
                             'V_INV2':{'dataset':2, 'name':"U_LV2", 'leg':'Voltage BESS', 'offset':0.0, 'scale':1.0},
                             'P_INV1':{'dataset':2, 'name':"P_LV1", 'leg':'P_PV', 'offset':0.0, 'scale':1.0},
                             'P_INV2':{'dataset':2, 'name':"P_LV2", 'leg':'P_BESS', 'offset':0.0, 'scale':1.0},
                             'Q_INV1':{'dataset':2, 'name':"Q_LV1", 'leg':'Q_PV', 'offset':0.0, 'scale':1.0},
                             'Q_INV2':{'dataset':2, 'name':"Q_LV2", 'leg':'Q_BESS', 'offset':0.0, 'scale':1.0},
                             'FRT_INV1':{'dataset':2, 'name':"INV1_FRT_STATE", 'leg':'PV FRT', 'offset':0.0, 'scale':1.0},
                             'FRT_INV2':{'dataset':2, 'name':"INV2_FRT_FLAG", 'leg':'BESS FRT', 'offset':0.0, 'scale':1.0},
                             'Pcmd_PPC_INV1':{'dataset':2, 'name':"P_CMD_PV", 'leg':'P_cmd to PV', 'offset':0.0, 'scale':151.2}, #S_PV
                             'Pcmd_PPC_INV2':{'dataset':2, 'name':"P_CMD_BESS", 'leg':'P_cmd to BESS', 'offset':0.0, 'scale':168}, #S_BESS
                             'Qcmd_PPC_INV1':{'dataset':2, 'name':"Q_CMD_PV", 'leg':'Q_cmd to PV', 'offset':0.0, 'scale':90.72}, #S_PV*0.6
                             'Qcmd_PPC_INV2':{'dataset':2, 'name':"Q_CMD_BESS", 'leg':'Q_cmd to BESS', 'offset':0.0, 'scale':100.8}, #S_BESS*0.6
                             'V_Lismore330':{'dataset':2, 'name':"V_LSM_330", 'leg':'after-SUMSF', 'offset':0.0, 'scale':1.0},
                             'V_Coffs_Harbour330':{'dataset':2, 'name':"V_COFF_330", 'leg':'after-SUMSF', 'offset':0.0, 'scale':1.0},
                             'V_Lismore132':{'dataset':2, 'name':"V_LSM_132", 'leg':'after-SUMSF', 'offset':0.0, 'scale':1.0},
                             'V_Koolkhan132':{'dataset':2, 'name':"V_KOLK_132", 'leg':'after-SUMSF', 'offset':0.0, 'scale':1.0},

                             'Ang_Lismore330':{'dataset':2, 'name':"ANG_LSM_330", 'leg':'after-SUMSF', 'offset':0.0, 'scale':1.0},
                             'Ang_Coffs_Harbour330':{'dataset':2, 'name':"ANG_COFF_330", 'leg':'after-SUMSF', 'offset':0.0, 'scale':1.0},
                             'Ang_Lismore132':{'dataset':2, 'name':"ANG_LSM_132", 'leg':'after-SUMSF', 'offset':0.0, 'scale':1.0},
                             'Ang_Koolkhan132':{'dataset':2, 'name':"ANG_KOLK_132", 'leg':'after-SUMSF', 'offset':0.0, 'scale':1.0},
                             
                             'P_COFF_330 - LSM_330':{'dataset':2, 'name':"P_COFF_330 - LSM_330", 'leg':'after-SUMSF', 'offset':0.0, 'scale':1.0},
                             'P_ARMIDL_330D - COFF_330':{'dataset':2, 'name':"P_ARMIDL_330D - COFF_330", 'leg':'after-SUMSF', 'offset':0.0, 'scale':1.0},
                             'P_SUM_POC - KOLK_132':{'dataset':2, 'name':"P_SUM_POC - KOLK_132", 'leg':'after-SUMSF', 'offset':0.0, 'scale':1.0},
                             'P_SUM_POC - LSM_132':{'dataset':2, 'name':"P_SUM_POC - LSM_132", 'leg':'after-SUMSF', 'offset':0.0, 'scale':1.0},

                             'Q_COFF_330 - LSM_330':{'dataset':2, 'name':"Q_COFF_330 - LSM_330", 'leg':'after-SUMSF', 'offset':0.0, 'scale':1.0},
                             'Q_ARMIDL_330D - COFF_330':{'dataset':2, 'name':"Q_ARMIDL_330D - COFF_330", 'leg':'after-SUMSF', 'offset':0.0, 'scale':1.0},
                             'Q_SUM_POC - KOLK_132':{'dataset':2, 'name':"Q_SUM_POC - KOLK_132", 'leg':'after-SUMSF', 'offset':0.0, 'scale':1.0},
                             'Q_SUM_POC - LSM_132':{'dataset':2, 'name':"Q_SUM_POC - LSM_132", 'leg':'after-SUMSF', 'offset':0.0, 'scale':1.0},
                             },
                'NW_pre':  {'V_POC':{'dataset':1, 'name':"U_POC1", 'leg':'pre-SUMSF', 'offset':0.0, 'scale':1.0},
                             'V_Lismore330':{'dataset':1, 'name':"V_LSM_330", 'leg':'pre-SUMSF', 'offset':0.0, 'scale':1.0},
                             'V_Coffs_Harbour330':{'dataset':1, 'name':"V_COFF_330", 'leg':'pre-SUMSF', 'offset':0.0, 'scale':1.0},
                             'V_Lismore132':{'dataset':1, 'name':"V_LSM_132", 'leg':'pre-SUMSF', 'offset':0.0, 'scale':1.0},
                             'V_Koolkhan132':{'dataset':1, 'name':"V_KOLK_132", 'leg':'pre-SUMSF', 'offset':0.0, 'scale':1.0},

                             'Ang_Lismore330':{'dataset':1, 'name':"ANG_LSM_330", 'leg':'pre-SUMSF', 'offset':0.0, 'scale':1.0},
                             'Ang_Coffs_Harbour330':{'dataset':1, 'name':"ANG_COFF_330", 'leg':'pre-SUMSF', 'offset':0.0, 'scale':1.0},
                             'Ang_Lismore132':{'dataset':1, 'name':"ANG_LSM_132", 'leg':'pre-SUMSF', 'offset':0.0, 'scale':1.0},
                             'Ang_Koolkhan132':{'dataset':1, 'name':"ANG_KOLK_132", 'leg':'pre-SUMSF', 'offset':0.0, 'scale':1.0},
                             
                             'P_COFF_330 - LSM_330':{'dataset':1, 'name':"P_COFF_330 - LSM_330", 'leg':'pre-SUMSF', 'offset':0.0, 'scale':1.0},
                             'P_ARMIDL_330D - COFF_330':{'dataset':1, 'name':"P_ARMIDL_330D - COFF_330", 'leg':'pre-SUMSF', 'offset':0.0, 'scale':1.0},
                             'P_SUM_POC - KOLK_132':{'dataset':1, 'name':"P_SUM_POC - KOLK_132", 'leg':'pre-SUMSF', 'offset':0.0, 'scale':1.0},
                             'P_SUM_POC - LSM_132':{'dataset':1, 'name':"P_SUM_POC - LSM_132", 'leg':'pre-SUMSF', 'offset':0.0, 'scale':1.0},

                             'Q_COFF_330 - LSM_330':{'dataset':1, 'name':"Q_COFF_330 - LSM_330", 'leg':'pre-SUMSF', 'offset':0.0, 'scale':1.0},
                             'Q_ARMIDL_330D - COFF_330':{'dataset':1, 'name':"Q_ARMIDL_330D - COFF_330", 'leg':'pre-SUMSF', 'offset':0.0, 'scale':1.0},
                             'Q_SUM_POC - KOLK_132':{'dataset':1, 'name':"Q_SUM_POC - KOLK_132", 'leg':'pre-SUMSF', 'offset':0.0, 'scale':1.0},
                             'Q_SUM_POC - LSM_132':{'dataset':1, 'name':"Q_SUM_POC - LSM_132", 'leg':'pre-SUMSF', 'offset':0.0, 'scale':1.0},
                             
                             },
                'PSCAD_run1':{'V_POC':{'dataset':5, 'name':"PLANT_V_HV_pu", 'leg':'Voltage HV', 'offset':0.0, 'scale':1.0},
                             'P_POC':{'dataset':5, 'name':"PLANT_P_HV", 'leg':'PLANT_P_HV', 'offset':0.0, 'scale':1.0},
                             'Q_POC':{'dataset':5, 'name':"PLANT_Q_HV", 'leg':'PLANT_Q_HV', 'offset':0.0, 'scale':1.0},
                             },
                'PSCAD_run2':{'V_POC':{'dataset':6, 'name':"PLANT_V_HV_pu", 'leg':'Voltage HV', 'offset':0.0, 'scale':1.0},
                             'P_POC':{'dataset':6, 'name':"PLANT_P_HV", 'leg':'PLANT_P_HV', 'offset':0.0, 'scale':1.0},
                             'Q_POC':{'dataset':6, 'name':"PLANT_Q_HV", 'leg':'PLANT_Q_HV', 'offset':0.0, 'scale':1.0},
                             },
                'PSCAD_run3':{'V_POC':{'dataset':7, 'name':"PLANT_V_HV_pu", 'leg':'Voltage HV', 'offset':0.0, 'scale':1.0},
                             'P_POC':{'dataset':7, 'name':"PLANT_P_HV", 'leg':'PLANT_P_HV', 'offset':0.0, 'scale':1.0},
                             'Q_POC':{'dataset':7, 'name':"PLANT_Q_HV", 'leg':'PLANT_Q_HV', 'offset':0.0, 'scale':1.0},
                             },
                }



#####################################################
# Start report format
#####################################################

reports = {                                                                          

            'BENCH':{
                    'batchname': 'Benchmarking', #'GPS S5255_AEMO_PSSE', #'GPS S5254_AEMO_PSSE', #'GPS S52514_AEMO_PSSE', #'DMATsl_PSCAD_PSSSE_Final2', #'S5254a', #'Benchmarking', #'DMAT', #in DMAT
                    'report_definition':[          
                                           {'chapter':'general', #Array of chapters. 
                                            'datasets': [ dict(datasets[did], **{'path':datasets[did]['path']+r''}) for did in range(len(datasets))], # update the dataset path to the chapter result location
                                            'cases':[], #if empty, all cases are considered
                                            'plots_for_report': ['Overlays','PSSE Results','PSCAD Results'],
                                            'summary_items': {}, #Additional description/analysis of the result data can bebe specified here. For example amount of current injection for 5.2.5.5 analysis in table format. Other items can be defined in the future.
                                            'report':True,
                                            'plots':{
                                                                
                                                        'PSSE Results':       {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSSE']['V_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSSE']['P_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSSE']['Q_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5},
                                                                                        'Reactive Current POC':             {'channels':[dict(channels_lib['PSSE']['Iq_POC'], **{'markers':['dIq','rise_t','set_t']}),], 'unit':'p.u', 'rank':7, 'yminspan':0.1, 'ymaxlim':2.0, 'yminlim':-2.0},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSSE']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSSE']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSSE']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, },
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSSE']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, },
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSSE']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSSE']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },

                                                        'PSCAD Results':      {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5},
                                                                                        'Reactive Current POC':             {'channels':[dict(channels_lib['PSCAD']['Iq_POC'], **{}),], 'unit':'p.u', 'rank':7, 'yminspan':0.1, 'ymaxlim':2.0, 'yminlim':-2.0},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSCAD']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSCAD']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSCAD']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, },
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSCAD']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, },
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSCAD']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSCAD']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },


                                                        'Overlays':             {
                                                                                        'Voltage - POC':                  {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['V_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},#'GSMG':0, }, #first channel in array links to file in first result location
                                                                                        'Active Power - POC':             {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['P_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MW', 'rank':2, 'yminspan':50, 'yminlim':-200.0, 'ymaxlim':200.0},#'GSMG':0},
                                                                                        'Reactive Power - POC':           {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['Q_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MVAr', 'rank':3, 'yminspan':10.0, 'yminlim':-200.0, 'ymaxlim':200.0},  
                                                                                },

#                                                        'Overlays 33kV':        {
#                                                                                        'Voltage - 33kV':                  {'channels':[dict(channels_lib['PSCAD']['V_MV'], **{'leg':'PSCAD'}),
#                                                                                                                                 dict(channels_lib['PSSE']['V_MV'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},#'GSMG':0, }, #first channel in array links to file in first result location
#                                                                                        'Active Power - 33kV':             {'channels':[dict(channels_lib['PSCAD']['P_MV'], **{'leg':'PSCAD'}),
#                                                                                                                                 dict(channels_lib['PSSE']['P_MV'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MW', 'rank':2, 'yminspan':50, 'yminlim':-200.0, 'ymaxlim':200.0},#'GSMG':0},
#                                                                                        'Reactive Power - 33kV':           {'channels':[dict(channels_lib['PSCAD']['Q_MV'], **{'leg':'PSCAD'}),
#                                                                                                                                 dict(channels_lib['PSSE']['Q_MV'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MVAr', 'rank':3, 'yminspan':10.0, 'yminlim':-200.0, 'ymaxlim':200.0},  
#                                                                                },

                                                        'Overlays INV1':        {
                                                                                        'Voltage - INV1':                  {'channels':[dict(channels_lib['PSCAD']['V_INV1'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['V_INV1'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},#'GSMG':0, }, #first channel in array links to file in first result location
                                                                                        'Active Power - INV1':             {'channels':[dict(channels_lib['PSCAD']['P_INV1'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['P_INV1'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MW', 'rank':2, 'yminspan':50, 'yminlim':-200.0, 'ymaxlim':200.0},#'GSMG':0},
                                                                                        'Reactive Power - INV1':           {'channels':[dict(channels_lib['PSCAD']['Q_INV1'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['Q_INV1'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MVAr', 'rank':3, 'yminspan':10.0, 'yminlim':-200.0, 'ymaxlim':200.0},  
                                                                                },

                                                        'Overlays INV2':        {
                                                                                        'Voltage - INV2':                  {'channels':[dict(channels_lib['PSCAD']['V_INV2'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['V_INV2'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},#'GSMG':0, }, #first channel in array links to file in first result location
                                                                                        'Active Power - INV2':             {'channels':[dict(channels_lib['PSCAD']['P_INV2'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['P_INV2'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MW', 'rank':2, 'yminspan':50, 'yminlim':-200.0, 'ymaxlim':200.0},#'GSMG':0},
                                                                                        'Reactive Power - INV2':           {'channels':[dict(channels_lib['PSCAD']['Q_INV2'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['Q_INV2'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MVAr', 'rank':3, 'yminspan':10.0, 'yminlim':-200.0, 'ymaxlim':200.0},  
                                                                                },
                                                    },# end of plots
                                            }, # end of chapters
                                         ] #end of report_definition
                    }, # end of BENCH
                                                                                        
#####################################################
# DMAT
#####################################################
            'DMAT':{
                    'batchname': 'DMAT', #'GPS S5255_AEMO_PSSE', #'GPS S5254_AEMO_PSSE', #'GPS S52514_AEMO_PSSE', #'DMATsl_PSCAD_PSSSE_Final2', #'S5254a', #'Benchmarking', #'DMAT', #in DMAT
                    'report_definition':[          
                                           {'chapter':'general', #Array of chapters. 
                                            'datasets': [ dict(datasets[did], **{'path':datasets[did]['path']+r''}) for did in range(len(datasets))], # update the dataset path to the chapter result location
                                            'cases':[], #if empty, all cases are considered
                                            'plots_for_report': ['PSSE Results','PSCAD Results', 'Overlays',],
                                            'summary_items': {}, #Additional description/analysis of the result data can bebe specified here. For example amount of current injection for 5.2.5.5 analysis in table format. Other items can be defined in the future.
                                            'report':True,
                                            'plots':{
                                                                
                                                        'PSSE Results':       {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSSE']['V_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSSE']['P_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3, 'yminspan':5},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSSE']['Q_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5, 'yminspan':5},
                                                                                        'Frequency':                        {'channels':[dict(channels_lib['PSSE']['F_POC'], **{})], 'unit':'Hz', 'rank':7, 'yminspan':0.1},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSSE']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSSE']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSSE']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSSE']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSSE']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSSE']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },

                                                        'PSCAD Results':      {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3, 'yminspan':5},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5, 'yminspan':5},
                                                                                        'Frequency':                        {'channels':[dict(channels_lib['PSCAD']['F_POC'], **{})], 'unit':'Hz', 'rank':7, 'yminspan':0.1},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSCAD']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSCAD']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['V_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['V_MV'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSCAD']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSCAD']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSCAD']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSCAD']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },

                                                        'Overlays':             {
                                                                                        'Voltage':                  {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['V_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},#'GSMG':0, }, #first channel in array links to file in first result location
                                                                                        'Active Power':             {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['P_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MW', 'rank':2, 'yminspan':50, 'yminlim':-200.0, 'ymaxlim':200.0},#'GSMG':0},
                                                                                        'Reactive Power':           {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['Q_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MVAr', 'rank':3, 'yminspan':10.0, 'yminlim':-200.0, 'ymaxlim':200.0},  
                                                                                },


#                                                        'Overlays PPC':             {
#                                                                                        'Voltage':                  {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{'leg':'PSCAD'}),
#                                                                                                                                 dict(channels_lib['PSSE']['V_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'p.u.', 'rank':1},#'GSMG':0, }, #first channel in array links to file in first result location
#                                                                                        'PPC FRT signal':           {'channels':[dict(channels_lib['PSCAD']['FRT_PPC'], **{'leg':'PSCAD'}),
#                                                                                                                                 dict(channels_lib['PSSE']['FRT_PPC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'p.u.', 'rank':2, 'yminlim':-5.0, 'ymaxlim':5.0},#'GSMG':0},
##                                                                                        'Active Power':             {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{'leg':'PSCAD'}),
##                                                                                                                                 dict(channels_lib['PSSE']['P_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MW', 'rank':2, 'yminspan':50, 'yminlim':-200.0, 'ymaxlim':200.0},#'GSMG':0},
#                                                                                        'Reactive Power':           {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{'leg':'PSCAD'}),
#                                                                                                                                 dict(channels_lib['PSSE']['Q_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MVAr', 'rank':3},  
#                                                                                },
#
#                                                        'Overlays PV INV':        {
#                                                                                        'Voltage':                  {'channels':[dict(channels_lib['PSCAD']['V_INV1'], **{'leg':'PSCAD'}),
#                                                                                                                                 dict(channels_lib['PSSE']['V_INV1'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'p.u.', 'rank':1},#'GSMG':0, }, #first channel in array links to file in first result location
#                                                                                        'PV FRT signal':            {'channels':[dict(channels_lib['PSCAD']['FRT_INV1'], **{'leg':'PSCAD'}),
#                                                                                                                                 dict(channels_lib['PSSE']['FRT_INV1'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'p.u.', 'rank':2, 'yminlim':-5.0, 'ymaxlim':5.0}, 
##                                                                                        'Active Power':             {'channels':[dict(channels_lib['PSCAD']['P_INV1'], **{'leg':'PSCAD'}),
##                                                                                                                                 dict(channels_lib['PSSE']['P_INV1'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MW', 'rank':3, 'yminspan':10.0, 'yminlim':-200.0, 'ymaxlim':200.0},#'GSMG':0},
#                                                                                        'Reactive Power':           {'channels':[dict(channels_lib['PSCAD']['Q_INV1'], **{'leg':'PSCAD'}),
#                                                                                                                                 dict(channels_lib['PSSE']['Q_INV1'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MVAr', 'rank':3},  
#                                                                                },

#                                                        'Overlays BESS INV':        {
#                                                                                        'Voltage':                  {'channels':[dict(channels_lib['PSCAD']['V_INV2'], **{'leg':'PSCAD'}),
#                                                                                                                                 dict(channels_lib['PSSE']['V_INV2'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'p.u.', 'rank':1},#'GSMG':0, }, #first channel in array links to file in first result location
#                                                                                        'BESS FRT signal':          {'channels':[dict(channels_lib['PSCAD']['FRT_INV2'], **{'leg':'PSCAD'}),
#                                                                                                                                 dict(channels_lib['PSSE']['FRT_INV2'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'p.u.', 'rank':2}, 
##                                                                                        'Active Power':             {'channels':[dict(channels_lib['PSCAD']['P_INV2'], **{'leg':'PSCAD'}),
##                                                                                                                                 dict(channels_lib['PSSE']['P_INV2'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MW', 'rank':3, 'yminspan':10.0, 'yminlim':-200.0, 'ymaxlim':200.0},#'GSMG':0},
#                                                                                        'Reactive Power':           {'channels':[dict(channels_lib['PSCAD']['Q_INV2'], **{'leg':'PSCAD'}),
#                                                                                                                                 dict(channels_lib['PSSE']['Q_INV2'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MVAr', 'rank':3},  
#                                                                                },
                                                    },# end of plots
                                            }, # end of chapters

                                           {'chapter':'01_FlatRun', #Array of chapters. 
                                            'datasets': [ dict(datasets[did], **{'path':datasets[did]['path']+r''}) for did in range(len(datasets))], # update the dataset path to the chapter result location
                                            'cases':[], #if empty, all cases are considered
                                            'plots_for_report': ['Overlays'],
                                            'summary_items': {}, #Additional description/analysis of the result data can bebe specified here. For example amount of current injection for 5.2.5.5 analysis in table format. Other items can be defined in the future.
                                            'report':True,
                                            'plots':{
                                                                



                                                        'Overlays':             {
                                                                                        'Voltage':                  {'channels':[dict(channels_lib['PSCAD_run1']['V_POC'], **{'leg':'PSCAD_run1'}),
                                                                                                                                 dict(channels_lib['PSCAD_run2']['V_POC'], **{'leg':'PSCAD_run2'}),
                                                                                                                                 dict(channels_lib['PSCAD_run3']['V_POC'], **{'leg':'PSCAD_run3'}),
                                                                                                                                 ], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},#'GSMG':0, }, #first channel in array links to file in first result location
                                                                                        'Active Power':             {'channels':[dict(channels_lib['PSCAD_run1']['P_POC'], **{'leg':'PSCAD_run1'}),
                                                                                                                                 dict(channels_lib['PSCAD_run2']['P_POC'], **{'leg':'PSCAD_run2'}),
                                                                                                                                 dict(channels_lib['PSCAD_run3']['P_POC'], **{'leg':'PSCAD_run3'}),
                                                                                                                                 ], 'unit':'MW', 'rank':2, 'yminspan':50, 'yminlim':-200.0, 'ymaxlim':200.0},#'GSMG':0},
                                                                                        'Reactive Power':           {'channels':[dict(channels_lib['PSCAD_run1']['Q_POC'], **{'leg':'PSCAD_run1'}),
                                                                                                                                 dict(channels_lib['PSCAD_run2']['Q_POC'], **{'leg':'PSCAD_run2'}),
                                                                                                                                 dict(channels_lib['PSCAD_run3']['Q_POC'], **{'leg':'PSCAD_run3'}),
                                                                                                                                 ], 'unit':'MVAr', 'rank':3, 'yminspan':10.0, 'yminlim':-200.0, 'ymaxlim':200.0},  
                                                                                },

                                                    },# end of plots
                                            }, # end of chapters

                                           {'chapter':'02_UFault', #Array of chapters. 
                                            'datasets': [ dict(datasets[did], **{'path':datasets[did]['path']+r'\02_UFault'}) for did in range(len(datasets))], # update the dataset path to the chapter result location
                                            'cases':[], #if empty, all cases are considered
                                            'plots_for_report': ['PSCAD Results'],
                                            'summary_items': {}, #Additional description/analysis of the result data can bebe specified here. For example amount of current injection for 5.2.5.5 analysis in table format. Other items can be defined in the future.
                                            'report':True,
                                            'plots':{

                                                        'PSCAD Results':      {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3, 'yminspan':5},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5, 'yminspan':5},
                                                                                        'Frequency':                        {'channels':[dict(channels_lib['PSCAD']['F_POC'], **{})], 'unit':'Hz', 'rank':7, 'yminspan':0.1},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSCAD']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSCAD']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSCAD']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSCAD']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSCAD']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSCAD']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },

                                                    },# end of plots
                                            }, # end of chapters

                                           {'chapter':'03_MFRT', #Array of chapters. 
                                            'datasets': [ dict(datasets[did], **{'path':datasets[did]['path']+r'\03_MFRT'}) for did in range(len(datasets))], # update the dataset path to the chapter result location
                                            'cases':[], #if empty, all cases are considered
                                            'plots_for_report': ['Overlays','PSSE Results','PSCAD Results'],
                                            'summary_items': {}, #Additional description/analysis of the result data can bebe specified here. For example amount of current injection for 5.2.5.5 analysis in table format. Other items can be defined in the future.
                                            'report':True,
                                            'plots':{
                                                                
                                                        'PSSE Results':       {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSSE']['V_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSSE']['P_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3, 'yminspan':5},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSSE']['Q_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5, 'yminspan':5},
                                                                                        'Frequency':                        {'channels':[dict(channels_lib['PSSE']['F_POC'], **{})], 'unit':'Hz', 'rank':7, 'yminspan':0.1},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSSE']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSSE']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSSE']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSSE']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSSE']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSSE']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },

                                                        'PSCAD Results':      {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3, 'yminspan':5},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5, 'yminspan':5},
                                                                                        'Frequency':                        {'channels':[dict(channels_lib['PSCAD']['F_POC'], **{})], 'unit':'Hz', 'rank':7, 'yminspan':0.1},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSCAD']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSCAD']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSCAD']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSCAD']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSCAD']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSCAD']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },

                                                        'Overlays':             {
                                                                                        'Voltage':                  {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['V_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},#'GSMG':0, }, #first channel in array links to file in first result location
                                                                                        'Active Power':             {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['P_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MW', 'rank':2, 'yminspan':50, 'yminlim':-200.0, 'ymaxlim':200.0},#'GSMG':0},
                                                                                        'Reactive Power':           {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['Q_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MVAr', 'rank':3, 'yminspan':10.0, 'yminlim':-200.0, 'ymaxlim':200.0},  
                                                                                },

                                                    },# end of plots
                                            }, # end of chapters

                                           {'chapter':'04_TOV', #Array of chapters. 
                                            'datasets': [ dict(datasets[did], **{'path':datasets[did]['path']+r'\04_TOV'}) for did in range(len(datasets))], # update the dataset path to the chapter result location
                                            'cases':[], #if empty, all cases are considered
                                            'plots_for_report': ['Overlays','PSSE Results','PSCAD Results'],
                                            'summary_items': {}, #Additional description/analysis of the result data can bebe specified here. For example amount of current injection for 5.2.5.5 analysis in table format. Other items can be defined in the future.
                                            'report':True,
                                            'plots':{
                                                                
                                                        'PSSE Results':       {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSSE']['V_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSSE']['P_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3, 'yminspan':5},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSSE']['Q_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5, 'yminspan':5},
                                                                                        'Frequency':                        {'channels':[dict(channels_lib['PSSE']['F_POC'], **{})], 'unit':'Hz', 'rank':7, 'yminspan':0.1},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSSE']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSSE']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSSE']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSSE']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSSE']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSSE']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },

                                                        'PSCAD Results':      {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3, 'yminspan':5},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5, 'yminspan':5},
                                                                                        'Frequency':                        {'channels':[dict(channels_lib['PSCAD']['F_POC'], **{})], 'unit':'Hz', 'rank':7, 'yminspan':0.1},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSCAD']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSCAD']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSCAD']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSCAD']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSCAD']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSCAD']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },

                                                        'Overlays':             {
                                                                                        'Voltage':                  {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['V_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},#'GSMG':0, }, #first channel in array links to file in first result location
                                                                                        'Active Power':             {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['P_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MW', 'rank':2, 'yminspan':50, 'yminlim':-200.0, 'ymaxlim':200.0},#'GSMG':0},
                                                                                        'Reactive Power':           {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['Q_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MVAr', 'rank':3, 'yminspan':10.0, 'yminlim':-200.0, 'ymaxlim':200.0},  
                                                                                },

                                                    },# end of plots
                                            }, # end of chapters

                                           {'chapter':'05_SptChange', #Array of chapters. 
                                            'datasets': [ dict(datasets[did], **{'path':datasets[did]['path']+r'\05_SptChange'}) for did in range(len(datasets))], # update the dataset path to the chapter result location
                                            'cases':[], #if empty, all cases are considered
                                            'plots_for_report': ['Overlays','PSSE Results','PSCAD Results'],
                                            'summary_items': {}, #Additional description/analysis of the result data can bebe specified here. For example amount of current injection for 5.2.5.5 analysis in table format. Other items can be defined in the future.
                                            'report':True,
                                            'plots':{
                                                                
                                                        'PSSE Results':       {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSSE']['V_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSSE']['P_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3, 'yminspan':5},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSSE']['Q_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5, 'yminspan':5},
                                                                                        'Frequency':                        {'channels':[dict(channels_lib['PSSE']['F_POC'], **{})], 'unit':'Hz', 'rank':7, 'yminspan':0.1},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSSE']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSSE']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSSE']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSSE']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSSE']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSSE']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },

                                                        'PSCAD Results':      {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3, 'yminspan':5},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5, 'yminspan':5},
                                                                                        'Frequency':                        {'channels':[dict(channels_lib['PSCAD']['F_POC'], **{})], 'unit':'Hz', 'rank':7, 'yminspan':0.1},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSCAD']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSCAD']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSCAD']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSCAD']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSCAD']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSCAD']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },

                                                        'Overlays':             {
                                                                                        'Voltage':                  {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['V_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},#'GSMG':0, }, #first channel in array links to file in first result location
                                                                                        'Active Power':             {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['P_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MW', 'rank':2, 'yminspan':50, 'yminlim':-200.0, 'ymaxlim':200.0},#'GSMG':0},
                                                                                        'Reactive Power':           {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['Q_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MVAr', 'rank':3, 'yminspan':10.0, 'yminlim':-200.0, 'ymaxlim':200.0},  
                                                                                },

                                                    },# end of plots
                                            }, # end of chapters

                                           {'chapter':'06_FreqChange', #Array of chapters. 
                                            'datasets': [ dict(datasets[did], **{'path':datasets[did]['path']+r'\06_FreqChange'}) for did in range(len(datasets))], # update the dataset path to the chapter result location
                                            'cases':[], #if empty, all cases are considered
                                            'plots_for_report': ['Overlays','PSSE Results','PSCAD Results'],
                                            'summary_items': {}, #Additional description/analysis of the result data can bebe specified here. For example amount of current injection for 5.2.5.5 analysis in table format. Other items can be defined in the future.
                                            'report':True,
                                            'plots':{
                                                                
                                                        'PSSE Results':       {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSSE']['V_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSSE']['P_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3, 'yminspan':5},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSSE']['Q_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5, 'yminspan':5},
                                                                                        'Frequency':                        {'channels':[dict(channels_lib['PSSE']['F_POC'], **{})], 'unit':'Hz', 'rank':7, 'yminspan':0.1},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSSE']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSSE']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSSE']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSSE']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSSE']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSSE']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },

                                                        'PSCAD Results':      {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3, 'yminspan':5},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5, 'yminspan':5},
                                                                                        'Frequency':                        {'channels':[dict(channels_lib['PSCAD']['F_POC'], **{})], 'unit':'Hz', 'rank':7, 'yminspan':0.1},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSCAD']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSCAD']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSCAD']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSCAD']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSCAD']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSCAD']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },

                                                        'Overlays':             {
                                                                                        'Voltage':                  {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['V_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},#'GSMG':0, }, #first channel in array links to file in first result location
                                                                                        'Active Power':             {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['P_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MW', 'rank':2, 'yminspan':50, 'yminlim':-200.0, 'ymaxlim':200.0},#'GSMG':0},
                                                                                        'Reactive Power':           {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['Q_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MVAr', 'rank':3, 'yminspan':10.0, 'yminlim':-200.0, 'ymaxlim':200.0},  
                                                                                },

                                                    },# end of plots
                                            }, # end of chapters

                                           {'chapter':'07_VolChange', #Array of chapters. 
                                            'datasets': [ dict(datasets[did], **{'path':datasets[did]['path']+r'\07_VolChange'}) for did in range(len(datasets))], # update the dataset path to the chapter result location
                                            'cases':[], #if empty, all cases are considered
                                            'plots_for_report': ['Overlays','PSSE Results','PSCAD Results'],
                                            'summary_items': {}, #Additional description/analysis of the result data can bebe specified here. For example amount of current injection for 5.2.5.5 analysis in table format. Other items can be defined in the future.
                                            'report':True,
                                            'plots':{
                                                                
                                                        'PSSE Results':       {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSSE']['V_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSSE']['P_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3, 'yminspan':5},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSSE']['Q_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5, 'yminspan':5},
                                                                                        'Frequency':                        {'channels':[dict(channels_lib['PSSE']['F_POC'], **{})], 'unit':'Hz', 'rank':7, 'yminspan':0.1},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSSE']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSSE']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSSE']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSSE']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSSE']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSSE']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },

                                                        'PSCAD Results':      {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3, 'yminspan':5},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5, 'yminspan':5},
                                                                                        'Frequency':                        {'channels':[dict(channels_lib['PSCAD']['F_POC'], **{})], 'unit':'Hz', 'rank':7, 'yminspan':0.1},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSCAD']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSCAD']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSCAD']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSCAD']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSCAD']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSCAD']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },

                                                        'Overlays':             {
                                                                                        'Voltage':                  {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['V_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},#'GSMG':0, }, #first channel in array links to file in first result location
                                                                                        'Active Power':             {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['P_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MW', 'rank':2, 'yminspan':50, 'yminlim':-200.0, 'ymaxlim':200.0},#'GSMG':0},
                                                                                        'Reactive Power':           {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['Q_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MVAr', 'rank':3, 'yminspan':10.0, 'yminlim':-200.0, 'ymaxlim':200.0},  
                                                                                },

                                                    },# end of plots
                                            }, # end of chapters

                                           {'chapter':'08_ORT', #Array of chapters. 
                                            'datasets': [ dict(datasets[did], **{'path':datasets[did]['path']+r'\08_ORT'}) for did in range(len(datasets))], # update the dataset path to the chapter result location
                                            'cases':[], #if empty, all cases are considered
                                            'plots_for_report': ['Overlays','PSSE Results','PSCAD Results'],
                                            'summary_items': {}, #Additional description/analysis of the result data can bebe specified here. For example amount of current injection for 5.2.5.5 analysis in table format. Other items can be defined in the future.
                                            'report':True,
                                            'plots':{
                                                                
                                                        'PSSE Results':       {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSSE']['V_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSSE']['P_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3, 'yminspan':5},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSSE']['Q_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5, 'yminspan':5},
                                                                                        'Frequency':                        {'channels':[dict(channels_lib['PSSE']['F_POC'], **{})], 'unit':'Hz', 'rank':7, 'yminspan':0.1},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSSE']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSSE']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSSE']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSSE']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSSE']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSSE']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },

                                                        'PSCAD Results':      {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3, 'yminspan':5},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5, 'yminspan':5},
                                                                                        'Frequency':                        {'channels':[dict(channels_lib['PSCAD']['F_POC'], **{})], 'unit':'Hz', 'rank':7, 'yminspan':0.1},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSCAD']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSCAD']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSCAD']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSCAD']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSCAD']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSCAD']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },

                                                        'Overlays':             {
                                                                                        'Voltage':                  {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['V_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},#'GSMG':0, }, #first channel in array links to file in first result location
                                                                                        'Active Power':             {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['P_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MW', 'rank':2, 'yminspan':50, 'yminlim':-200.0, 'ymaxlim':200.0},#'GSMG':0},
                                                                                        'Reactive Power':           {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['Q_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MVAr', 'rank':3, 'yminspan':10.0, 'yminlim':-200.0, 'ymaxlim':200.0},  
                                                                                },

                                                    },# end of plots
                                            }, # end of chapters

                                           {'chapter':'09_AngChange', #Array of chapters. 
                                            'datasets': [ dict(datasets[did], **{'path':datasets[did]['path']+r'\09_AngChange'}) for did in range(len(datasets))], # update the dataset path to the chapter result location
                                            'cases':[], #if empty, all cases are considered
                                            'plots_for_report': ['Overlays','PSSE Results','PSCAD Results'],
                                            'summary_items': {}, #Additional description/analysis of the result data can bebe specified here. For example amount of current injection for 5.2.5.5 analysis in table format. Other items can be defined in the future.
                                            'report':True,
                                            'plots':{
                                                                
                                                        'PSSE Results':       {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSSE']['V_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSSE']['P_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3, 'yminspan':5},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSSE']['Q_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5, 'yminspan':5},
#                                                                                        'Frequency':                        {'channels':[dict(channels_lib['PSSE']['F_POC'], **{})], 'unit':'Hz', 'rank':7, 'yminspan':0.1},
                                                                                        'Angle':                            {'channels':[dict(channels_lib['PSSE']['Ang_POC'], **{})], 'unit':'Deg', 'rank':7, 'yminspan':0.1},
#                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSSE']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSSE']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSSE']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSSE']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'Frequency':                        {'channels':[dict(channels_lib['PSSE']['F_POC'], **{})], 'unit':'Hz', 'rank':8, 'yminspan':0.1},
#                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSSE']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
#                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSSE']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },

                                                        'PSCAD Results':      {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3, 'yminspan':5},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5, 'yminspan':5},
#                                                                                        'Frequency':                        {'channels':[dict(channels_lib['PSCAD']['F_POC'], **{})], 'unit':'Hz', 'rank':7, 'yminspan':0.1},
                                                                                        'Angle':                            {'channels':[dict(channels_lib['PSCAD']['Ang_POC'], **{})], 'unit':'Deg', 'rank':7, 'yminspan':0.1},
#                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSCAD']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSCAD']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSCAD']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSCAD']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'Frequency':                        {'channels':[dict(channels_lib['PSCAD']['F_POC'], **{})], 'unit':'Hz', 'rank':8, 'yminspan':0.1},
#                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSCAD']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
#                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSCAD']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },

                                                        'Overlays':             {
                                                                                        'Voltage':                  {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['V_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},#'GSMG':0, }, #first channel in array links to file in first result location
                                                                                        'Active Power':             {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['P_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MW', 'rank':2, 'yminspan':50, 'yminlim':-200.0, 'ymaxlim':200.0},#'GSMG':0},
                                                                                        'Reactive Power':           {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['Q_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MVAr', 'rank':3, 'yminspan':10.0, 'yminlim':-200.0, 'ymaxlim':200.0},  
                                                                                },

                                                    },# end of plots
                                            }, # end of chapters

                                           {'chapter':'10_SCR1', #Array of chapters. 
                                            'datasets': [ dict(datasets[did], **{'path':datasets[did]['path']+r'\10_SCR1'}) for did in range(len(datasets))], # update the dataset path to the chapter result location
                                            'cases':[], #if empty, all cases are considered
                                            'plots_for_report': ['Overlays','PSSE Results','PSCAD Results'],
                                            'summary_items': {}, #Additional description/analysis of the result data can bebe specified here. For example amount of current injection for 5.2.5.5 analysis in table format. Other items can be defined in the future.
                                            'report':True,
                                            'plots':{
                                                                
                                                        'PSSE Results':       {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSSE']['V_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSSE']['P_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3, 'yminspan':5},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSSE']['Q_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5, 'yminspan':5},
                                                                                        'Frequency':                        {'channels':[dict(channels_lib['PSSE']['F_POC'], **{})], 'unit':'Hz', 'rank':7, 'yminspan':0.1},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSSE']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSSE']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSSE']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSSE']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSSE']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSSE']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },

                                                        'PSCAD Results':      {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3, 'yminspan':5},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5, 'yminspan':5},
                                                                                        'Frequency':                        {'channels':[dict(channels_lib['PSCAD']['F_POC'], **{})], 'unit':'Hz', 'rank':7, 'yminspan':0.1},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSCAD']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSCAD']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSCAD']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSCAD']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSCAD']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSCAD']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },

                                                        'Overlays':             {
                                                                                        'Voltage':                  {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['V_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},#'GSMG':0, }, #first channel in array links to file in first result location
                                                                                        'Active Power':             {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['P_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MW', 'rank':2, 'yminspan':50, 'yminlim':-200.0, 'ymaxlim':200.0},#'GSMG':0},
                                                                                        'Reactive Power':           {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['Q_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MVAr', 'rank':3, 'yminspan':10.0, 'yminlim':-200.0, 'ymaxlim':200.0},  
                                                                                },

                                                    },# end of plots
                                            }, # end of chapters

                                           {'chapter':'11_FRT', #Array of chapters. 
                                            'datasets': [ dict(datasets[did], **{'path':datasets[did]['path']+r'\11_FRT'}) for did in range(len(datasets))], # update the dataset path to the chapter result location
                                            'cases':[], #if empty, all cases are considered
                                            'plots_for_report': ['Overlays','PSSE Results','PSCAD Results'],
                                            'summary_items': {}, #Additional description/analysis of the result data can bebe specified here. For example amount of current injection for 5.2.5.5 analysis in table format. Other items can be defined in the future.
                                            'report':True,
                                            'plots':{
                                                                
                                                        'PSSE Results':       {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSSE']['V_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSSE']['P_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3, 'yminspan':5},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSSE']['Q_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5, 'yminspan':5},
                                                                                        'Frequency':                        {'channels':[dict(channels_lib['PSSE']['F_POC'], **{})], 'unit':'Hz', 'rank':7, 'yminspan':0.1},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSSE']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSSE']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSSE']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSSE']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSSE']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSSE']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },

                                                        'PSCAD Results':      {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3, 'yminspan':5},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5, 'yminspan':5},
                                                                                        'Frequency':                        {'channels':[dict(channels_lib['PSCAD']['F_POC'], **{})], 'unit':'Hz', 'rank':7, 'yminspan':0.1},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSCAD']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSCAD']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSCAD']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSCAD']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSCAD']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSCAD']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },

                                                        'Overlays':             {
                                                                                        'Voltage':                  {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['V_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},#'GSMG':0, }, #first channel in array links to file in first result location
                                                                                        'Active Power':             {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['P_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MW', 'rank':2, 'yminspan':50, 'yminlim':-200.0, 'ymaxlim':200.0},#'GSMG':0},
                                                                                        'Reactive Power':           {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['Q_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MVAr', 'rank':3, 'yminspan':10.0, 'yminlim':-200.0, 'ymaxlim':200.0},  
                                                                                },

                                                    },# end of plots
                                            }, # end of chapters

                                           {'chapter':'12_IrrChange', #Array of chapters. 
                                            'datasets': [ dict(datasets[did], **{'path':datasets[did]['path']+r'\12_IrrChange'}) for did in range(len(datasets))], # update the dataset path to the chapter result location
                                            'cases':[], #if empty, all cases are considered
                                            'plots_for_report': ['Overlays','PSSE Results','PSCAD Results'],
                                            'summary_items': {}, #Additional description/analysis of the result data can bebe specified here. For example amount of current injection for 5.2.5.5 analysis in table format. Other items can be defined in the future.
                                            'report':True,
                                            'plots':{
                                                                
                                                        'PSSE Results':       {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSSE']['V_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSSE']['P_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3, 'yminspan':5},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSSE']['Q_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5, 'yminspan':5},
                                                                                        'Frequency':                        {'channels':[dict(channels_lib['PSSE']['F_POC'], **{})], 'unit':'Hz', 'rank':7, 'yminspan':0.1},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSSE']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSSE']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSSE']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSSE']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSSE']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSSE']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },

                                                        'PSCAD Results':      {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3, 'yminspan':5},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5, 'yminspan':5},
                                                                                        'Frequency':                        {'channels':[dict(channels_lib['PSCAD']['F_POC'], **{})], 'unit':'Hz', 'rank':7, 'yminspan':0.1},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSCAD']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSCAD']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSCAD']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSCAD']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':5},
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSCAD']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSCAD']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },

                                                        'Overlays':             {
                                                                                        'Voltage':                  {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['V_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},#'GSMG':0, }, #first channel in array links to file in first result location
                                                                                        'Active Power':             {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['P_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MW', 'rank':2, 'yminspan':50, 'yminlim':-200.0, 'ymaxlim':200.0},#'GSMG':0},
                                                                                        'Reactive Power':           {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['Q_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MVAr', 'rank':3, 'yminspan':10.0, 'yminlim':-200.0, 'ymaxlim':200.0},  
                                                                                },

                                                    },# end of plots
                                            }, # end of chapters

                                         ] #end of report_definition
                    }, # end of DMAT
                                                                                        
#####################################################
# GPS
#####################################################
            'GPS':{ # in network contingency events
                    'batchname':'GPS', 
                    'report_definition':[          

                                           {'chapter':'general', #Array of chapters. 
                                            'datasets': [ dict(datasets[did], **{'path':datasets[did]['path']+r''}) for did in range(len(datasets))], # update the dataset path to the chapter result location
                                            'cases':[], #if empty, all cases are considered
                                            'plots_for_report': ['Overlays','PSSE Results','PSCAD Results'],
                                            'summary_items': {}, #Additional description/analysis of the result data can bebe specified here. For example amount of current injection for 5.2.5.5 analysis in table format. Other items can be defined in the future.
                                            'report':True,
                                            'plots':{
                                                                
                                                        'PSSE Results':       {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSSE']['V_POC'], **{'scale':1.0,'markers':['dV']}),
                                                                                                                                         dict(channels_lib['PSSE']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSSE']['P_POC'], **{'markers':['rec_t']}),
                                                                                                                                         dict(channels_lib['PSSE']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSSE']['Q_POC'], **{'markers':[]}),
                                                                                                                                         dict(channels_lib['PSSE']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5},
                                                                                        'Reactive Current POC':             {'channels':[dict(channels_lib['PSSE']['Iq_POC'], **{'markers':['dIq','rise_t','set_t']}),], 'unit':'p.u', 'rank':7, 'yminspan':0.1, 'ymaxlim':2.0, 'yminlim':-2.0},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSSE']['FRT_PPC'], **{'markers':[]}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSSE']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSSE']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, },
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSSE']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, },
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSSE']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSSE']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },

                                                        'PSCAD Results':      {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{'markers':['dV']}),
                                                                                                                                         dict(channels_lib['PSCAD']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{'markers':['rec_t']}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{'markers':[]}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5},
                                                                                        'Reactive Current POC':             {'channels':[dict(channels_lib['PSCAD']['Iq_POC'], **{'markers':['dIq','rise_t','set_t']}),], 'unit':'p.u', 'rank':7, 'yminspan':0.1, 'ymaxlim':2.0, 'yminlim':-2.0},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSCAD']['FRT_PPC'], **{'markers':[]}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSCAD']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSCAD']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, },
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSCAD']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, },
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSCAD']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSCAD']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },


                                                        'Overlays':             {
                                                                                        'Voltage':                  {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['V_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},#'GSMG':0, }, #first channel in array links to file in first result location
                                                                                        'Active Power':             {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['P_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MW', 'rank':2, 'yminspan':50, 'yminlim':-200.0, 'ymaxlim':200.0},#'GSMG':0},
                                                                                        'Reactive Power':           {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['Q_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MVAr', 'rank':3, 'yminspan':10.0, 'yminlim':-200.0, 'ymaxlim':200.0},  
                                                                                },
                                                    },# end of plots
                                            }, # end of chapters
                                                                                                                                 
                                           {'chapter':'S5253', #Array of chapters. 
                                            'datasets': [ dict(datasets[did], **{'path':datasets[did]['path']+r'\S5253'}) for did in range(len(datasets))], # update the dataset path to the chapter result location
                                            'cases':[], #if empty, all cases are considered
                                            'plots_for_report': ['Overlays','PSSE Results','PSCAD Results'],
                                            'summary_items': {}, #Additional description/analysis of the result data can bebe specified here. For example amount of current injection for 5.2.5.5 analysis in table format. Other items can be defined in the future.
                                            'report':True,
                                            'plots':{
                                                                
                                                        'PSSE Results':       {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSSE']['V_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSSE']['P_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3, 'yminspan':10},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSSE']['Q_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5, 'yminspan':10},
                                                                                        'Frequency':                        {'channels':[dict(channels_lib['PSSE']['F_POC'], **{'markers':['callout', [1,7,17,27]]})], 'unit':'Hz', 'rank':7, 'yminspan':0.1},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSSE']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSSE']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.1},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSSE']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':10},
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSSE']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':10},
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSSE']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSSE']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },

                                                        'PSCAD Results':      {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5},
                                                                                        'Frequency':                        {'channels':[dict(channels_lib['PSCAD']['F_POC'], **{'markers':['callout', [1,7,17,27]]})], 'unit':'Hz', 'rank':7, 'yminspan':0.1},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSCAD']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSCAD']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSCAD']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, },
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSCAD']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, },
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSCAD']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSCAD']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },


                                                        'Overlays':             {
                                                                                        'Voltage':                  {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['V_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},#'GSMG':0, }, #first channel in array links to file in first result location
                                                                                        'Active Power':             {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['P_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MW', 'rank':2, 'yminspan':50, 'yminlim':-200.0, 'ymaxlim':200.0},#'GSMG':0},
                                                                                        'Reactive Power':           {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['Q_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MVAr', 'rank':3, 'yminspan':10.0, 'yminlim':-200.0, 'ymaxlim':200.0},  
                                                                                },
                                                    },# end of plots
                                            }, # end of chapters


                                           {'chapter':'S5254', #Array of chapters. 
                                            'datasets': [ dict(datasets[did], **{'path':datasets[did]['path']+r'\S5254'}) for did in range(len(datasets))], # update the dataset path to the chapter result location
                                            'cases':[], #if empty, all cases are considered
                                            'plots_for_report': ['Overlays','PSSE Results','PSCAD Results'],
                                            'summary_items': {}, #Additional description/analysis of the result data can bebe specified here. For example amount of current injection for 5.2.5.5 analysis in table format. Other items can be defined in the future.
                                            'report':True,
                                            'plots':{
                                                                
                                                        'PSSE Results':       {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSSE']['V_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSSE']['P_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSSE']['Q_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5},
                                                                                        'Reactive Current POC':             {'channels':[dict(channels_lib['PSSE']['Iq_POC'], **{}),], 'unit':'p.u', 'rank':7, 'yminspan':0.1, 'ymaxlim':2.0, 'yminlim':-2.0},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSSE']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSSE']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSSE']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, },
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSSE']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, },
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSSE']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSSE']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },

                                                        'PSCAD Results':      {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5},
                                                                                        'Reactive Current POC':             {'channels':[dict(channels_lib['PSCAD']['Iq_POC'], **{}),], 'unit':'p.u', 'rank':7, 'yminspan':0.1, 'ymaxlim':2.0, 'yminlim':-2.0},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSCAD']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSCAD']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSCAD']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, },
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSCAD']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, },
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSCAD']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSCAD']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },


                                                        'Overlays':             {
                                                                                        'Voltage':                  {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['V_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},#'GSMG':0, }, #first channel in array links to file in first result location
                                                                                        'Active Power':             {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['P_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MW', 'rank':2, 'yminspan':50, 'yminlim':-200.0, 'ymaxlim':200.0},#'GSMG':0},
                                                                                        'Reactive Power':           {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['Q_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MVAr', 'rank':3, 'yminspan':10.0, 'yminlim':-200.0, 'ymaxlim':200.0},  
                                                                                },

                                                    },# end of plots
                                            }, # end of chapters

                                           {'chapter':'S5255', #Array of chapters. 
                                            'datasets': [ dict(datasets[did], **{'path':datasets[did]['path']+r'\S5255'}) for did in range(len(datasets))], # update the dataset path to the chapter result location
                                            'cases':[], #if empty, all cases are considered
                                            'plots_for_report': ['Overlays','PSSE Results','PSCAD Results'],
                                            'summary_items': {}, #Additional description/analysis of the result data can bebe specified here. For example amount of current injection for 5.2.5.5 analysis in table format. Other items can be defined in the future.
                                            'report':True,
                                            'plots':{
                                                                
                                                        'PSSE Results':       {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSSE']['V_POC'], **{'scale':1.0,'markers':['dV']}),
                                                                                                                                         dict(channels_lib['PSSE']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSSE']['P_POC'], **{'markers':['rec_t']}),
                                                                                                                                         dict(channels_lib['PSSE']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSSE']['Q_POC'], **{'markers':[]}),
                                                                                                                                         dict(channels_lib['PSSE']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5},
                                                                                        'Reactive Current POC':             {'channels':[dict(channels_lib['PSSE']['Iq_POC'], **{'markers':['dIq','rise_t','set_t']}),], 'unit':'p.u', 'rank':7, 'yminspan':0.1, 'ymaxlim':2.0, 'yminlim':-2.0},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSSE']['FRT_PPC'], **{'markers':[]}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSSE']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSSE']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, },
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSSE']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, },
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSSE']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSSE']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },

                                                        'PSCAD Results':      {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{'markers':['dV']}),
                                                                                                                                         dict(channels_lib['PSCAD']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{'markers':['rec_t']}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{'markers':[]}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5},
                                                                                        'Reactive Current POC':             {'channels':[dict(channels_lib['PSCAD']['Iq_POC'], **{'markers':['dIq','rise_t','set_t']}),], 'unit':'p.u', 'rank':7, 'yminspan':0.1, 'ymaxlim':2.0, 'yminlim':-2.0},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSCAD']['FRT_PPC'], **{'markers':[]}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSCAD']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSCAD']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, },
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSCAD']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, },
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSCAD']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSCAD']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },


                                                        'Overlays':             {
                                                                                        'Voltage':                  {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['V_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},#'GSMG':0, }, #first channel in array links to file in first result location
                                                                                        'Active Power':             {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['P_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MW', 'rank':2, 'yminspan':50, 'yminlim':-200.0, 'ymaxlim':200.0},#'GSMG':0},
                                                                                        'Reactive Power':           {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['Q_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MVAr', 'rank':3, 'yminspan':10.0, 'yminlim':-200.0, 'ymaxlim':200.0},  
                                                                                },

                                                    },# end of plots
                                            }, # end of chapters

                                           {'chapter':'S5257', #Array of chapters. 
                                            'datasets': [ dict(datasets[did], **{'path':datasets[did]['path']+r'\S5257'}) for did in range(len(datasets))], # update the dataset path to the chapter result location
                                            'cases':[], #if empty, all cases are considered
                                            'plots_for_report': ['Overlays','PSSE Results','PSCAD Results'],
                                            'summary_items': {}, #Additional description/analysis of the result data can bebe specified here. For example amount of current injection for 5.2.5.5 analysis in table format. Other items can be defined in the future.
                                            'report':True,
                                            'plots':{
                                                                
                                                        'PSSE Results':       {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSSE']['V_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSSE']['P_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3, 'yminspan':10},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSSE']['Q_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5, 'yminspan':10},
#                                                                                        'Reactive Current POC':             {'channels':[dict(channels_lib['PSSE']['Iq_POC'], **{}),], 'unit':'p.u', 'rank':7, 'yminspan':0.1, 'ymaxlim':2.0, 'yminlim':-2.0},
                                                                                        'Frequency':                        {'channels':[dict(channels_lib['PSSE']['F_POC'], **{'markers':['callout', [1,6.5]]})], 'unit':'Hz', 'rank':7, 'yminspan':0.1},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSSE']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSSE']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.1},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSSE']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':10 },
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSSE']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':10},
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSSE']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSSE']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },

                                                        'PSCAD Results':      {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5},
                                                                                        'Reactive Current POC':             {'channels':[dict(channels_lib['PSCAD']['Iq_POC'], **{}),], 'unit':'p.u', 'rank':7, 'yminspan':0.1, 'ymaxlim':2.0, 'yminlim':-2.0},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSCAD']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSCAD']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSCAD']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, },
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSCAD']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, },
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSCAD']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSCAD']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },


                                                        'Overlays':             {
                                                                                        'Voltage':                  {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['V_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},#'GSMG':0, }, #first channel in array links to file in first result location
                                                                                        'Active Power':             {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['P_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MW', 'rank':2, 'yminspan':50, 'yminlim':-200.0, 'ymaxlim':200.0},#'GSMG':0},
                                                                                        'Reactive Power':           {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['Q_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MVAr', 'rank':3, 'yminspan':10.0, 'yminlim':-200.0, 'ymaxlim':200.0},  
                                                                                },
                                                    },# end of plots
                                            }, # end of chapters

                                           {'chapter':'S5258', #Array of chapters. 
                                            'datasets': [ dict(datasets[did], **{'path':datasets[did]['path']+r'\S5258'}) for did in range(len(datasets))], # update the dataset path to the chapter result location
                                            'cases':[], #if empty, all cases are considered
                                            'plots_for_report': ['Overlays','PSSE Results','PSCAD Results'],
                                            'summary_items': {}, #Additional description/analysis of the result data can bebe specified here. For example amount of current injection for 5.2.5.5 analysis in table format. Other items can be defined in the future.
                                            'report':True,
                                            'plots':{
                                                                
                                                        'PSSE Results':       {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSSE']['V_POC'], **{'scale':1.0,'markers':['dV']}),
                                                                                                                                         dict(channels_lib['PSSE']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSSE']['P_POC'], **{'markers':['rec_t']}),
                                                                                                                                         dict(channels_lib['PSSE']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3, 'yminspan':10},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSSE']['Q_POC'], **{'markers':[]}),
                                                                                                                                         dict(channels_lib['PSSE']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5, 'yminspan':10},
                                                                                        'Frequency':                        {'channels':[dict(channels_lib['PSSE']['F_POC'], **{'markers':['callout', [1,6.5]]})], 'unit':'Hz', 'rank':7, 'yminspan':0.1},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSSE']['FRT_PPC'], **{'markers':[]}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSSE']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.1},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSSE']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':10},
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSSE']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':10},
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSSE']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSSE']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },

                                                        'PSCAD Results':      {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{'markers':['dV']}),
                                                                                                                                         dict(channels_lib['PSCAD']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{'markers':['rec_t']}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3, 'yminspan':10},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{'markers':[]}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5, 'yminspan':10},
                                                                                        'Frequency':                        {'channels':[dict(channels_lib['PSCAD']['F_POC'], **{'markers':['callout', [1,6.5]]})], 'unit':'Hz', 'rank':7, 'yminspan':0.1},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSCAD']['FRT_PPC'], **{'markers':[]}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSCAD']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.1},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSCAD']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':10},
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSCAD']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':10},
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSCAD']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSCAD']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },


                                                        'Overlays':             {
                                                                                        'Voltage':                  {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['V_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},#'GSMG':0, }, #first channel in array links to file in first result location
                                                                                        'Active Power':             {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['P_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MW', 'rank':2, 'yminspan':50, 'yminlim':-200.0, 'ymaxlim':200.0},#'GSMG':0},
                                                                                        'Reactive Power':           {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['Q_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MVAr', 'rank':3, 'yminspan':10.0, 'yminlim':-200.0, 'ymaxlim':200.0},  
                                                                                },
                                                    },# end of plots
                                            }, # end of chapters

                                           {'chapter':'S52511', #Array of chapters. 
                                            'datasets': [ dict(datasets[did], **{'path':datasets[did]['path']+r'\S52511'}) for did in range(len(datasets))], # update the dataset path to the chapter result location
                                            'cases':[], #if empty, all cases are considered
                                            'plots_for_report': ['Overlays','PSSE Results','PSCAD Results'],
                                            'summary_items': {}, #Additional description/analysis of the result data can bebe specified here. For example amount of current injection for 5.2.5.5 analysis in table format. Other items can be defined in the future.
                                            'report':True,
                                            'plots':{
                                                                
                                                        'PSSE Results':       {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSSE']['V_POC'], **{'scale':1.0}),
                                                                                                                                         dict(channels_lib['PSSE']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSSE']['P_POC'], **{'markers':['callout', [1,6.5]]}),
                                                                                                                                         dict(channels_lib['PSSE']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3, 'yminspan':10.0},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSSE']['Q_POC'], **{'markers':[]}),
                                                                                                                                         dict(channels_lib['PSSE']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5, 'yminspan':10.0},
                                                                                        'Frequency':                        {'channels':[dict(channels_lib['PSSE']['F_POC'], **{'markers':['callout', [1,6.5]]})], 'unit':'Hz', 'rank':7, 'yminspan':0.1},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSSE']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSSE']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.1},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSSE']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':10.0},
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSSE']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':10.0},
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSSE']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSSE']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },

                                                        'PSCAD Results':      {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{'markers':['callout', [1,6.5]]}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3, 'yminspan':2.0},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5, 'yminspan':2.0},
                                                                                        'Frequency':                        {'channels':[dict(channels_lib['PSCAD']['F_POC'], **{'markers':['callout', [1,6.5]]})], 'unit':'Hz', 'rank':7, 'yminspan':0.1},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSCAD']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSCAD']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSCAD']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, },
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSCAD']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, },
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSCAD']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSCAD']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },


                                                        'Overlays':             {
                                                                                        'Voltage':                  {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['V_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},#'GSMG':0, }, #first channel in array links to file in first result location
                                                                                        'Active Power':             {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['P_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MW', 'rank':2, 'yminspan':50, 'yminlim':-200.0, 'ymaxlim':200.0},#'GSMG':0},
                                                                                        'Reactive Power':           {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['Q_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MVAr', 'rank':3, 'yminspan':10.0, 'yminlim':-200.0, 'ymaxlim':200.0},  
                                                                                },
                                                    },# end of plots
                                            }, # end of chapters

                                           {'chapter':'S52513', #Array of chapters. 
                                            'datasets': [ dict(datasets[did], **{'path':datasets[did]['path']+r'\S52513'}) for did in range(len(datasets))], # update the dataset path to the chapter result location
                                            'cases':[], #if empty, all cases are considered
                                            'plots_for_report': ['Overlays','PSSE Results','PSCAD Results'],
                                            'summary_items': {}, #Additional description/analysis of the result data can bebe specified here. For example amount of current injection for 5.2.5.5 analysis in table format. Other items can be defined in the future.
                                            'report':True,
                                            'plots':{
                                                                
                                                        'PSSE Results':       {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSSE']['V_POC'], **{'markers':["rise_t","set_t",'callout', [7]]}),
                                                                                                                                         dict(channels_lib['PSSE']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5, 'markers':['callout', [7]]}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSSE']['P_POC'], ),#**{'markers':["set_t"]}
                                                                                                                                         dict(channels_lib['PSSE']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSSE']['Q_POC'], **{'markers':["rise_t","set_t",'callout', [7]]}),
                                                                                                                                         dict(channels_lib['PSSE']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5},
                                                                                        'Powerfactor POC':                  {'channels':[dict(channels_lib['PSSE']['PF_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['PFspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u', 'rank':7},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSSE']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSSE']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSSE']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, },
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSSE']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, },
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSSE']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSSE']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },

                                                        'PSCAD Results':      {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{'markers':["rise_t","set_t",'callout', [8]]}),
                                                                                                                                         dict(channels_lib['PSCAD']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5, 'markers':['callout', [8]]}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSCAD']['P_POC'], ),#**{'markers':["set_t"]}
                                                                                                                                         dict(channels_lib['PSCAD']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3, 'yminspan':10},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{'markers':["rise_t","set_t",'callout', [8]]}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5, 'yminspan':10},
                                                                                        'Powerfactor POC':                  {'channels':[dict(channels_lib['PSCAD']['PF_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['PFspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':7},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSCAD']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSCAD']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.1},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSCAD']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0,'yminspan':10 },
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSCAD']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, 'yminspan':10 },
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSCAD']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSCAD']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },


                                                        'Overlays':             {
                                                                                        'Voltage':                  {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['V_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},#'GSMG':0, }, #first channel in array links to file in first result location
                                                                                        'Active Power':             {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['P_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MW', 'rank':2, 'yminspan':50, 'yminlim':-200.0, 'ymaxlim':200.0},#'GSMG':0},
                                                                                        'Reactive Power':           {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['Q_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MVAr', 'rank':3, 'yminspan':10.0, 'yminlim':-200.0, 'ymaxlim':200.0},  
                                                                                },

                                                    },# end of plots
                                            }, # end of chapters



                                           {'chapter':'S52514', #Array of chapters. 
                                            'datasets': [ dict(datasets[did], **{'path':datasets[did]['path']+r'\S52514'}) for did in range(len(datasets))], # update the dataset path to the chapter result location
                                            'cases':[], #if empty, all cases are considered
                                            'plots_for_report': ['Overlays','PSSE Results','PSCAD Results'],
                                            'summary_items': {}, #Additional description/analysis of the result data can bebe specified here. For example amount of current injection for 5.2.5.5 analysis in table format. Other items can be defined in the future.
                                            'report':True,
                                            'plots':{
                                                                
                                                        'PSSE Results':       {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSSE']['V_POC'], **{'scale':1.0}),
                                                                                                                                         dict(channels_lib['PSSE']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSSE']['P_POC'], **{'markers':['callout', [1,15]]}),
                                                                                                                                         dict(channels_lib['PSSE']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3,'yminspan':10},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSSE']['Q_POC'], **{'markers':[]}),
                                                                                                                                         dict(channels_lib['PSSE']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5,'yminspan':10},
                                                                                        'Frequency':                        {'channels':[dict(channels_lib['PSSE']['F_POC'], **{})], 'unit':'Hz', 'rank':7, 'yminspan':0.1},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSSE']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSSE']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.1},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSSE']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0,'yminspan':10 },
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSSE']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0,'yminspan':10 },
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSSE']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSSE']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },

                                                        'PSCAD Results':      {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{'markers':['callout', [1,15]]}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5},
                                                                                        'Frequency':                        {'channels':[dict(channels_lib['PSCAD']['F_POC'], **{})], 'unit':'Hz', 'rank':7, 'yminspan':0.1},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSCAD']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSCAD']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSCAD']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, },
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSCAD']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, },
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSCAD']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSCAD']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },


                                                        'Overlays':             {
                                                                                        'Voltage':                  {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['V_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},#'GSMG':0, }, #first channel in array links to file in first result location
                                                                                        'Active Power':             {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['P_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MW', 'rank':2, 'yminspan':50, 'yminlim':-200.0, 'ymaxlim':200.0},#'GSMG':0},
                                                                                        'Reactive Power':           {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['Q_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MVAr', 'rank':3, 'yminspan':10.0, 'yminlim':-200.0, 'ymaxlim':200.0},  
                                                                                },

                                                    },# end of plots
                                            }, # end of chapters

                                           {'chapter':'S52515', #Array of chapters. 
                                            'datasets': [ dict(datasets[did], **{'path':datasets[did]['path']+r'\S52515'}) for did in range(len(datasets))], # update the dataset path to the chapter result location
                                            'cases':[], #if empty, all cases are considered
                                            'plots_for_report': ['Overlays','PSSE Results','PSCAD Results'],
                                            'summary_items': {}, #Additional description/analysis of the result data can bebe specified here. For example amount of current injection for 5.2.5.5 analysis in table format. Other items can be defined in the future.
                                            'report':True,
                                            'plots':{
                                                                
                                                        'PSSE Results':       {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSSE']['V_POC'], **{'scale':1.0}),
                                                                                                                                         dict(channels_lib['PSSE']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSSE']['P_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSSE']['Q_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5},
                                                                                        'Frequency':                        {'channels':[dict(channels_lib['PSSE']['F_POC'], **{})], 'unit':'Hz', 'rank':7, 'yminspan':0.1},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSSE']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSSE']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSSE']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, },
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSSE']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, },
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSSE']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSSE']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },

                                                        'PSCAD Results':      {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5},
                                                                                        'Frequency':                        {'channels':[dict(channels_lib['PSCAD']['F_POC'], **{})], 'unit':'Hz', 'rank':7, 'yminspan':0.1},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSCAD']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSCAD']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSCAD']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, },
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSCAD']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, },
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSCAD']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSCAD']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },


                                                        'Overlays':             {
                                                                                        'Voltage':                  {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['V_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},#'GSMG':0, }, #first channel in array links to file in first result location
                                                                                        'Active Power':             {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['P_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MW', 'rank':2, 'yminspan':50, 'yminlim':-200.0, 'ymaxlim':200.0},#'GSMG':0},
                                                                                        'Reactive Power':           {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['Q_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MVAr', 'rank':3, 'yminspan':10.0, 'yminlim':-200.0, 'ymaxlim':200.0},  
                                                                                },

                                                    },# end of plots
                                            }, # end of chapters

                                           {'chapter':'S52516', #Array of chapters. 
                                            'datasets': [ dict(datasets[did], **{'path':datasets[did]['path']+r'\S52516'}) for did in range(len(datasets))], # update the dataset path to the chapter result location
                                            'cases':[], #if empty, all cases are considered
                                            'plots_for_report': ['Overlays','PSSE Results','PSCAD Results'],
                                            'summary_items': {}, #Additional description/analysis of the result data can bebe specified here. For example amount of current injection for 5.2.5.5 analysis in table format. Other items can be defined in the future.
                                            'report':True,
                                            'plots':{
                                                                
                                                        'PSSE Results':       {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSSE']['V_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSSE']['P_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSSE']['Q_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5},
#                                                                                        'Frequency':                        {'channels':[dict(channels_lib['PSSE']['F_POC'], **{})], 'unit':'Hz', 'rank':7, 'yminspan':0.1},
                                                                                        'Angle':                            {'channels':[dict(channels_lib['PSSE']['Ang_POC'], **{})], 'unit':'Deg', 'rank':7, 'yminspan':0.1},

                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSSE']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSSE']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSSE']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, },
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSSE']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSSE']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, },
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSSE']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSSE']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },

                                                        'PSCAD Results':      {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5},
#                                                                                        'Frequency':                        {'channels':[dict(channels_lib['PSCAD']['F_POC'], **{})], 'unit':'Hz', 'rank':7, 'yminspan':0.1},
                                                                                        'Angle':                            {'channels':[dict(channels_lib['PSCAD']['Ang_POC'], **{})], 'unit':'Deg', 'rank':7, 'yminspan':0.1},

                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['PSCAD']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['PSCAD']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['PSCAD']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, },
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['PSCAD']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['PSCAD']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, },
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['PSCAD']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['PSCAD']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },


                                                        'Overlays':             {
                                                                                        'Voltage':                  {'channels':[dict(channels_lib['PSCAD']['V_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['V_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},#'GSMG':0, }, #first channel in array links to file in first result location
                                                                                        'Active Power':             {'channels':[dict(channels_lib['PSCAD']['P_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['P_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MW', 'rank':2, 'yminspan':50, 'yminlim':-200.0, 'ymaxlim':200.0},#'GSMG':0},
                                                                                        'Reactive Power':           {'channels':[dict(channels_lib['PSCAD']['Q_POC'], **{'leg':'PSCAD'}),
                                                                                                                                 dict(channels_lib['PSSE']['Q_POC'], **{'leg':'PSSE','markers':['GSMG']}),], 'unit':'MVAr', 'rank':3, 'yminspan':10.0, 'yminlim':-200.0, 'ymaxlim':200.0},  
                                                                                },

                                                    },# end of plots
                                            }, # end of chapters
                                         ] #end of report_definition
                    }, # end of GPS
                                                                                        
#####################################################
# NetworkEvent
#####################################################
            'NetworkEvent':{ # in network contingency events
                    'batchname':'Network Event', #'HighLoad1', 'HighLoad2', 'HighLoad3', 'LowLoad1','LowLoad2'
                    'report_definition':[          

                                           {'chapter':'general', #Array of chapters. 
                                            'datasets': [ dict(datasets[did], **{'path':datasets[did]['path']+r''}) for did in range(len(datasets))], # update the dataset path to the chapter result location
                                            'cases':[], #if empty, all cases are considered
                                            'plots_for_report': ['PSSE network general'],
                                            'summary_items': {}, #Additional description/analysis of the result data can bebe specified here. For example amount of current injection for 5.2.5.5 analysis in table format. Other items can be defined in the future.
                                            'report':True,
                                            'plots':{
                                                                
                                                        'PSSE network general':       {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['NW_post']['V_POC'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['NW_post']['P_POC'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['NW_post']['Q_POC'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5},
                                                                                        'Frequency':                        {'channels':[dict(channels_lib['NW_post']['F_POC'], **{})], 'unit':'Hz', 'rank':7, 'yminspan':0.1},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['NW_post']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['NW_post']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['NW_post']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['NW_post']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, },
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['NW_post']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['NW_post']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, },
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['NW_post']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['NW_post']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },
                                                    },# end of plots
                                            }, # end of chapters
                                                                                                                                 
                                                                                                                                 
                                           {'chapter':'S5255_HighLoad', #Array of chapters. 
                                            'datasets': [ dict(datasets[did], **{'path':datasets[did]['path']+r'\S5255_HL'}) for did in range(len(datasets))], # update the dataset path to the chapter result location
                                            'cases':[], #if empty, all cases are considered
                                            'plots_for_report': ['PSSE S5255 network'],
                                            'summary_items': {}, #Additional description/analysis of the result data can bebe specified here. For example amount of current injection for 5.2.5.5 analysis in table format. Other items can be defined in the future.
                                            'report':True,
                                            'plots':{
                                                                
                                                        'PSSE S5255 network':       {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['NW_post']['V_POC'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['NW_post']['P_POC'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['NW_post']['Q_POC'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5},
                                                                                        'Frequency':                        {'channels':[dict(channels_lib['NW_post']['F_POC'], **{})], 'unit':'Hz', 'rank':7, 'yminspan':0.1},
#                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['NW_post']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                        'Lismore 330kV bus voltage':        {'channels':[dict(channels_lib['NW_post']['V_Lismore330'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.1, 'ymaxlim':1.4, 'xmaxlim':3},
                                                                                        'Coffs Harbour 330kV bus voltage':  {'channels':[dict(channels_lib['NW_post']['V_Coffs_Harbour330'], **{}),], 'unit':'p.u.', 'rank':4, 'yminspan':0.1, 'ymaxlim':1.4, 'xmaxlim':3},
                                                                                        'Lismore 132kV bus voltage':        {'channels':[dict(channels_lib['NW_post']['V_Lismore132'], **{}),], 'unit':'p.u.', 'rank':6, 'yminspan':0.1, 'ymaxlim':1.4, 'xmaxlim':3},
                                                                                        'Koolkhan 132kV bus voltage':       {'channels':[dict(channels_lib['NW_post']['V_Koolkhan132'], **{}),], 'unit':'p.u.', 'rank':8, 'yminspan':0.1, 'ymaxlim':1.4, 'xmaxlim':3},


#                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['NW_post']['V_INV1'], **{}),
#                                                                                                                                         dict(channels_lib['NW_post']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
#                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['NW_post']['P_INV1'], **{}),
#                                                                                                                                         dict(channels_lib['NW_post']['P_INV2'], **{}),
#                                                                                                                                         dict(channels_lib['NW_post']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
#                                                                                                                                         dict(channels_lib['NW_post']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
#                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, },
#                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['NW_post']['Q_INV1'], **{}),
#                                                                                                                                         dict(channels_lib['NW_post']['Q_INV2'], **{}),
#                                                                                                                                         dict(channels_lib['NW_post']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
#                                                                                                                                         dict(channels_lib['NW_post']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
#                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, },
#                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['NW_post']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
#                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['NW_post']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },
                                                    },# end of plots
                                            }, # end of chapters

                                           {'chapter':'S5255_LowLoad', #Array of chapters. 
                                            'datasets': [ dict(datasets[did], **{'path':datasets[did]['path']+r'\S5255_LL'}) for did in range(len(datasets))], # update the dataset path to the chapter result location
                                            'cases':[], #if empty, all cases are considered
                                            'plots_for_report': ['PSSE S5255 network'],
                                            'summary_items': {}, #Additional description/analysis of the result data can bebe specified here. For example amount of current injection for 5.2.5.5 analysis in table format. Other items can be defined in the future.
                                            'report':True,
                                            'plots':{
                                                                
                                                        'PSSE S5255 network':       {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['NW_post']['V_POC'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['NW_post']['P_POC'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['NW_post']['Q_POC'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5},
                                                                                        'Frequency':                        {'channels':[dict(channels_lib['NW_post']['F_POC'], **{})], 'unit':'Hz', 'rank':7, 'yminspan':0.1},
#                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['NW_post']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                        'Lismore 330kV bus voltage':        {'channels':[dict(channels_lib['NW_post']['V_Lismore330'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.1, 'ymaxlim':1.4, 'xmaxlim':3},
                                                                                        'Coffs Harbour 330kV bus voltage':  {'channels':[dict(channels_lib['NW_post']['V_Coffs_Harbour330'], **{}),], 'unit':'p.u.', 'rank':4, 'yminspan':0.1, 'ymaxlim':1.4, 'xmaxlim':3},
                                                                                        'Lismore 132kV bus voltage':        {'channels':[dict(channels_lib['NW_post']['V_Lismore132'], **{}),], 'unit':'p.u.', 'rank':6, 'yminspan':0.1, 'ymaxlim':1.4, 'xmaxlim':3},
                                                                                        'Koolkhan 132kV bus voltage':       {'channels':[dict(channels_lib['NW_post']['V_Koolkhan132'], **{}),], 'unit':'p.u.', 'rank':8, 'yminspan':0.1, 'ymaxlim':1.4, 'xmaxlim':3},


#                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['NW_post']['V_INV1'], **{}),
#                                                                                                                                         dict(channels_lib['NW_post']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
#                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['NW_post']['P_INV1'], **{}),
#                                                                                                                                         dict(channels_lib['NW_post']['P_INV2'], **{}),
#                                                                                                                                         dict(channels_lib['NW_post']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
#                                                                                                                                         dict(channels_lib['NW_post']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
#                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, },
#                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['NW_post']['Q_INV1'], **{}),
#                                                                                                                                         dict(channels_lib['NW_post']['Q_INV2'], **{}),
#                                                                                                                                         dict(channels_lib['NW_post']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
#                                                                                                                                         dict(channels_lib['NW_post']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
#                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, },
#                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['NW_post']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
#                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['NW_post']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },
                                                    },# end of plots
                                            }, # end of chapters

                                           {'chapter':'S52512_HighLoad', #Array of chapters. 
                                            'datasets': [ dict(datasets[did], **{'path':datasets[did]['path']+r'\S52512_HL'}) for did in range(len(datasets))], # update the dataset path to the chapter result location
                                            'cases':[], #if empty, all cases are considered
                                            'plots_for_report': ['PSSE S52512 network'],
                                            'summary_items': {}, #Additional description/analysis of the result data can bebe specified here. For example amount of current injection for 5.2.5.5 analysis in table format. Other items can be defined in the future.
                                            'report':True,
                                            'plots':{

                                                        'PSSE S52512 network':       {

                                                                                        'COFF_330 - LSM_330 P':           {'channels':[dict(channels_lib['NW_pre']['P_COFF_330 - LSM_330'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['P_COFF_330 - LSM_330'], **{}),
                                                                                                                                         ], 'unit':'MW', 'rank':1, 'yminspan':10},
                                                                                        'ARMIDL_330D - COFF_330 P':       {'channels':[dict(channels_lib['NW_pre']['P_ARMIDL_330D - COFF_330'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['P_ARMIDL_330D - COFF_330'], **{}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'yminspan':10},
                                                                                        'SUM_POC - KOLK_132 P':           {'channels':[dict(channels_lib['NW_pre']['P_SUM_POC - KOLK_132'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['P_SUM_POC - KOLK_132'], **{}),
                                                                                                                                         ], 'unit':'MW', 'rank':7, 'yminspan':10},
                                                                                        'SUM_POC - LSM_132 P':            {'channels':[dict(channels_lib['NW_pre']['P_SUM_POC - LSM_132'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['P_SUM_POC - LSM_132'], **{}),
                                                                                                                                         ], 'unit':'MW', 'rank':10, 'yminspan':10},
                                                                                                                                       

                                                                                        'COFF_330 - LSM_330 Q':           {'channels':[dict(channels_lib['NW_pre']['Q_COFF_330 - LSM_330'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['Q_COFF_330 - LSM_330'], **{}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':2, 'yminspan':10},
                                                                                        'ARMIDL_330D - COFF_330 Q':       {'channels':[dict(channels_lib['NW_pre']['Q_ARMIDL_330D - COFF_330'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['Q_ARMIDL_330D - COFF_330'], **{}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':5, 'yminspan':10},
                                                                                        'SUM_POC - KOLK_132Q':           {'channels':[dict(channels_lib['NW_pre']['Q_SUM_POC - KOLK_132'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['Q_SUM_POC - KOLK_132'], **{}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':8, 'yminspan':10},
                                                                                        'SUM_POC - LSM_132 Q':            {'channels':[dict(channels_lib['NW_pre']['Q_SUM_POC - LSM_132'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['Q_SUM_POC - LSM_132'], **{}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':11, 'yminspan':10},

                                                                                        'Lismore 330kV bus voltage':        {'channels':[dict(channels_lib['NW_pre']['V_Lismore330'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['V_Lismore330'], **{}),], 'unit':'p.u.', 'rank':3, 'yminspan':0.1, 'ymaxlim':1.4, 'xmaxlim':3},
                                                                                        'Coffs Harbour 330kV bus voltage':  {'channels':[dict(channels_lib['NW_pre']['V_Coffs_Harbour330'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['V_Coffs_Harbour330'], **{}),], 'unit':'p.u.', 'rank':6, 'yminspan':0.1, 'ymaxlim':1.4, 'xmaxlim':3},
                                                                                        'Lismore 132kV bus voltage':        {'channels':[dict(channels_lib['NW_pre']['V_Lismore132'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['V_Lismore132'], **{}),], 'unit':'p.u.', 'rank':9, 'yminspan':0.1, 'ymaxlim':1.4, 'xmaxlim':3},
                                                                                        'Koolkhan 132kV bus voltage':       {'channels':[dict(channels_lib['NW_pre']['V_Koolkhan132'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['V_Koolkhan132'], **{}),], 'unit':'p.u.', 'rank':12, 'yminspan':0.1, 'ymaxlim':1.4, 'xmaxlim':3},


                                                                              },
                                                                                                                                       
#                                                        'PSSE S52512 network - voltage, Angle':       {
##                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['NW_pre']['V_POC'], **{}),
##                                                                                                                                         dict(channels_lib['NW_post']['V_POC'], **{}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1, 'ymaxlim':1.4, 'xmaxlim':3},
#                                                                                        'Lismore 330kV bus voltage':        {'channels':[dict(channels_lib['NW_pre']['V_Lismore330'], **{}),
#                                                                                                                                         dict(channels_lib['NW_post']['V_Lismore330'], **{}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1, 'ymaxlim':1.4, 'xmaxlim':3},
#                                                                                        'Coffs Harbour 330kV bus voltage':  {'channels':[dict(channels_lib['NW_pre']['V_Coffs_Harbour330'], **{}),
#                                                                                                                                         dict(channels_lib['NW_post']['V_Coffs_Harbour330'], **{}),], 'unit':'p.u.', 'rank':3, 'yminspan':0.1, 'ymaxlim':1.4, 'xmaxlim':3},
#                                                                                        'Lismore 132kV bus voltage':        {'channels':[dict(channels_lib['NW_pre']['V_Lismore132'], **{}),
#                                                                                                                                         dict(channels_lib['NW_post']['V_Lismore132'], **{}),], 'unit':'p.u.', 'rank':5, 'yminspan':0.1, 'ymaxlim':1.4, 'xmaxlim':3},
#                                                                                        'Koolkhan 132kV bus voltage':       {'channels':[dict(channels_lib['NW_pre']['V_Koolkhan132'], **{}),
#                                                                                                                                         dict(channels_lib['NW_post']['V_Koolkhan132'], **{}),], 'unit':'p.u.', 'rank':7, 'yminspan':0.1, 'ymaxlim':1.4, 'xmaxlim':3},
#
#                                                                                        'Lismore 330kV bus angle':        {'channels':[dict(channels_lib['NW_pre']['Ang_Lismore330'], **{}),
#                                                                                                                                         dict(channels_lib['NW_post']['Ang_Lismore330'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':5},
#                                                                                        'Coffs Harbour 330kV bus angle':  {'channels':[dict(channels_lib['NW_pre']['Ang_Coffs_Harbour330'], **{}),
#                                                                                                                                         dict(channels_lib['NW_post']['Ang_Coffs_Harbour330'], **{}),], 'unit':'p.u.', 'rank':4, 'yminspan':5},
#                                                                                        'Lismore 132kV bus angle':        {'channels':[dict(channels_lib['NW_pre']['Ang_Lismore132'], **{}),
#                                                                                                                                         dict(channels_lib['NW_post']['Ang_Lismore132'], **{}),], 'unit':'p.u.', 'rank':6, 'yminspan':5},
#                                                                                        'Koolkhan 132kV bus angle':       {'channels':[dict(channels_lib['NW_pre']['Ang_Koolkhan132'], **{}),
#                                                                                                                                         dict(channels_lib['NW_post']['Ang_Koolkhan132'], **{}),], 'unit':'p.u.', 'rank':8, 'yminspan':5},
#
#                                                                              },
#
#                                                        'PSSE S52512 network - P,Q':       {
##                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['NW_pre']['V_POC'], **{}),
##                                                                                                                                         dict(channels_lib['NW_post']['V_POC'], **{}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1, 'ymaxlim':1.4, 'xmaxlim':3},
##                                                                                        'COFF_330 - LSM_330 P,Q':           {'channels':[dict(channels_lib['NW_pre']['P_COFF_330 - LSM_330'], **{}),
##                                                                                                                                         dict(channels_lib['NW_post']['P_COFF_330 - LSM_330'], **{}),
##                                                                                                                                         dict(channels_lib['NW_pre']['Q_COFF_330 - LSM_330'], **{}),
##                                                                                                                                         dict(channels_lib['NW_post']['Q_COFF_330 - LSM_330'], **{}),
##                                                                                                                                         ], 'unit':'MW', 'rank':1, 'yminspan':10},
##                                                                                        'ARMIDL_330D - COFF_330 P,Q':       {'channels':[dict(channels_lib['NW_pre']['P_ARMIDL_330D - COFF_330'], **{}),
##                                                                                                                                         dict(channels_lib['NW_post']['P_ARMIDL_330D - COFF_330'], **{}),
##                                                                                                                                         dict(channels_lib['NW_pre']['Q_ARMIDL_330D - COFF_330'], **{}),
##                                                                                                                                         dict(channels_lib['NW_post']['Q_ARMIDL_330D - COFF_330'], **{}),
##                                                                                                                                         ], 'unit':'MW', 'rank':2, 'yminspan':10},
##                                                                                        'SUM_POC - KOLK_132 P,Q':           {'channels':[dict(channels_lib['NW_pre']['P_SUM_POC - KOLK_132'], **{}),
##                                                                                                                                         dict(channels_lib['NW_post']['P_SUM_POC - KOLK_132'], **{}),
##                                                                                                                                         dict(channels_lib['NW_pre']['Q_SUM_POC - KOLK_132'], **{}),
##                                                                                                                                         dict(channels_lib['NW_post']['Q_SUM_POC - KOLK_132'], **{}),
##                                                                                                                                         ], 'unit':'MW', 'rank':3, 'yminspan':10},
##                                                                                        'SUM_POC - LSM_132 P,Q':            {'channels':[dict(channels_lib['NW_pre']['P_SUM_POC - LSM_132'], **{}),
##                                                                                                                                         dict(channels_lib['NW_post']['P_SUM_POC - LSM_132'], **{}),
##                                                                                                                                         ], 'unit':'MW', 'rank':4, 'yminspan':10},
#
#                                                                                        'COFF_330 - LSM_330 P':           {'channels':[dict(channels_lib['NW_pre']['P_COFF_330 - LSM_330'], **{}),
#                                                                                                                                         dict(channels_lib['NW_post']['P_COFF_330 - LSM_330'], **{}),
#                                                                                                                                         ], 'unit':'MW', 'rank':1, 'yminspan':10},
#                                                                                        'ARMIDL_330D - COFF_330 P':       {'channels':[dict(channels_lib['NW_pre']['P_ARMIDL_330D - COFF_330'], **{}),
#                                                                                                                                         dict(channels_lib['NW_post']['P_ARMIDL_330D - COFF_330'], **{}),
#                                                                                                                                         ], 'unit':'MW', 'rank':3, 'yminspan':10},
#                                                                                        'SUM_POC - KOLK_132 P':           {'channels':[dict(channels_lib['NW_pre']['P_SUM_POC - KOLK_132'], **{}),
#                                                                                                                                         dict(channels_lib['NW_post']['P_SUM_POC - KOLK_132'], **{}),
#                                                                                                                                         ], 'unit':'MW', 'rank':5, 'yminspan':10},
#                                                                                        'SUM_POC - LSM_132 P':            {'channels':[dict(channels_lib['NW_pre']['P_SUM_POC - LSM_132'], **{}),
#                                                                                                                                         dict(channels_lib['NW_post']['P_SUM_POC - LSM_132'], **{}),
#                                                                                                                                         ], 'unit':'MW', 'rank':7, 'yminspan':10},
#                                                                                                                                       
#
#                                                                                        'COFF_330 - LSM_330 Q':           {'channels':[dict(channels_lib['NW_pre']['Q_COFF_330 - LSM_330'], **{}),
#                                                                                                                                         dict(channels_lib['NW_post']['Q_COFF_330 - LSM_330'], **{}),
#                                                                                                                                         ], 'unit':'MW', 'rank':2, 'yminspan':10},
#                                                                                        'ARMIDL_330D - COFF_330 Q':       {'channels':[dict(channels_lib['NW_pre']['Q_ARMIDL_330D - COFF_330'], **{}),
#                                                                                                                                         dict(channels_lib['NW_post']['Q_ARMIDL_330D - COFF_330'], **{}),
#                                                                                                                                         ], 'unit':'MW', 'rank':4, 'yminspan':10},
#                                                                                        'SUM_POC - KOLK_132Q':           {'channels':[dict(channels_lib['NW_pre']['Q_SUM_POC - KOLK_132'], **{}),
#                                                                                                                                         dict(channels_lib['NW_post']['Q_SUM_POC - KOLK_132'], **{}),
#                                                                                                                                         ], 'unit':'MW', 'rank':6, 'yminspan':10},
#                                                                                        'SUM_POC - LSM_132 Q':            {'channels':[dict(channels_lib['NW_pre']['Q_SUM_POC - LSM_132'], **{}),
#                                                                                                                                         dict(channels_lib['NW_post']['Q_SUM_POC - LSM_132'], **{}),
#                                                                                                                                         ], 'unit':'MW', 'rank':8, 'yminspan':10},
#                                                                                                                                       
#                                                                              },
                                                                                                                                         
                                                    },# end of plots
                                            }, # end of chapters


                                           {'chapter':'S52512_LowLoad', #Array of chapters. 
                                            'datasets': [ dict(datasets[did], **{'path':datasets[did]['path']+r'\S52512_LL'}) for did in range(len(datasets))], # update the dataset path to the chapter result location
                                            'cases':[], #if empty, all cases are considered
                                            'plots_for_report': ['PSSE S52512 network'],
                                            'summary_items': {}, #Additional description/analysis of the result data can bebe specified here. For example amount of current injection for 5.2.5.5 analysis in table format. Other items can be defined in the future.
                                            'report':True,
                                            'plots':{
                                                                
                                                        'PSSE S52512 network':       {

                                                                                        'COFF_330 - LSM_330 P':           {'channels':[dict(channels_lib['NW_pre']['P_COFF_330 - LSM_330'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['P_COFF_330 - LSM_330'], **{}),
                                                                                                                                         ], 'unit':'MW', 'rank':1, 'yminspan':10},
                                                                                        'ARMIDL_330D - COFF_330 P':       {'channels':[dict(channels_lib['NW_pre']['P_ARMIDL_330D - COFF_330'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['P_ARMIDL_330D - COFF_330'], **{}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'yminspan':10},
                                                                                        'SUM_POC - KOLK_132 P':           {'channels':[dict(channels_lib['NW_pre']['P_SUM_POC - KOLK_132'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['P_SUM_POC - KOLK_132'], **{}),
                                                                                                                                         ], 'unit':'MW', 'rank':7, 'yminspan':10},
                                                                                        'SUM_POC - LSM_132 P':            {'channels':[dict(channels_lib['NW_pre']['P_SUM_POC - LSM_132'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['P_SUM_POC - LSM_132'], **{}),
                                                                                                                                         ], 'unit':'MW', 'rank':10, 'yminspan':10},
                                                                                                                                       

                                                                                        'COFF_330 - LSM_330 Q':           {'channels':[dict(channels_lib['NW_pre']['Q_COFF_330 - LSM_330'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['Q_COFF_330 - LSM_330'], **{}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':2, 'yminspan':10},
                                                                                        'ARMIDL_330D - COFF_330 Q':       {'channels':[dict(channels_lib['NW_pre']['Q_ARMIDL_330D - COFF_330'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['Q_ARMIDL_330D - COFF_330'], **{}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':5, 'yminspan':10},
                                                                                        'SUM_POC - KOLK_132Q':           {'channels':[dict(channels_lib['NW_pre']['Q_SUM_POC - KOLK_132'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['Q_SUM_POC - KOLK_132'], **{}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':8, 'yminspan':10},
                                                                                        'SUM_POC - LSM_132 Q':            {'channels':[dict(channels_lib['NW_pre']['Q_SUM_POC - LSM_132'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['Q_SUM_POC - LSM_132'], **{}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':11, 'yminspan':10},

                                                                                        'Lismore 330kV bus voltage':        {'channels':[dict(channels_lib['NW_pre']['V_Lismore330'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['V_Lismore330'], **{}),], 'unit':'p.u.', 'rank':3, 'yminspan':0.1, 'ymaxlim':1.4, 'xmaxlim':3},
                                                                                        'Coffs Harbour 330kV bus voltage':  {'channels':[dict(channels_lib['NW_pre']['V_Coffs_Harbour330'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['V_Coffs_Harbour330'], **{}),], 'unit':'p.u.', 'rank':6, 'yminspan':0.1, 'ymaxlim':1.4, 'xmaxlim':3},
                                                                                        'Lismore 132kV bus voltage':        {'channels':[dict(channels_lib['NW_pre']['V_Lismore132'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['V_Lismore132'], **{}),], 'unit':'p.u.', 'rank':9, 'yminspan':0.1, 'ymaxlim':1.4, 'xmaxlim':3},
                                                                                        'Koolkhan 132kV bus voltage':       {'channels':[dict(channels_lib['NW_pre']['V_Koolkhan132'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['V_Koolkhan132'], **{}),], 'unit':'p.u.', 'rank':12, 'yminspan':0.1, 'ymaxlim':1.4, 'xmaxlim':3},


                                                                              },
                                                    },# end of plots
                                            }, # end of chapters



                                           {'chapter':'S52513_HighLoad', #Array of chapters. 
                                            'datasets': [ dict(datasets[did], **{'path':datasets[did]['path']+r'\S52513_HL'}) for did in range(len(datasets))], # update the dataset path to the chapter result location
                                            'cases':[], #if empty, all cases are considered
                                            'plots_for_report': ['PSSE S52513 network'],
                                            'summary_items': {}, #Additional description/analysis of the result data can bebe specified here. For example amount of current injection for 5.2.5.5 analysis in table format. Other items can be defined in the future.
                                            'report':True,
                                            'plots':{

                                                        'PSSE S52513 network':       {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['NW_post']['V_POC'], **{'markers':["rise_t","set_t"]}),
                                                                                                                                         dict(channels_lib['NW_post']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5, 'markers':[]}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['NW_post']['P_POC'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['NW_post']['Q_POC'], **{'markers':["rise_t","set_t"]}),
                                                                                                                                         dict(channels_lib['NW_post']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5},
                                                                                        'Frequency':                        {'channels':[dict(channels_lib['NW_post']['F_POC'], **{})], 'unit':'Hz', 'rank':7, 'yminspan':0.1},

#                                                                                        'Powerfactor POC':                  {'channels':[dict(channels_lib['PSSE']['PF_POC'], **{}),
#                                                                                                                                         dict(channels_lib['PSSE']['PFspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u', 'rank':7},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['NW_post']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['NW_post']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['NW_post']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['NW_post']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, },
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['NW_post']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['NW_post']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, },
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['NW_post']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['NW_post']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },
                                                    },# end of plots
                                            }, # end of chapters

                                           {'chapter':'S52513_LowLoad', #Array of chapters. 
                                            'datasets': [ dict(datasets[did], **{'path':datasets[did]['path']+r'\S52513_LL'}) for did in range(len(datasets))], # update the dataset path to the chapter result location
                                            'cases':[], #if empty, all cases are considered
                                            'plots_for_report': ['PSSE S52513 network'],
                                            'summary_items': {}, #Additional description/analysis of the result data can bebe specified here. For example amount of current injection for 5.2.5.5 analysis in table format. Other items can be defined in the future.
                                            'report':True,
                                            'plots':{
                                                                
                                                        'PSSE S52513 network':       {
                                                                                        'Voltage POC':                      {'channels':[dict(channels_lib['NW_post']['V_POC'], **{'markers':["rise_t","set_t"]}),
                                                                                                                                         dict(channels_lib['NW_post']['Vspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5, 'markers':[]}),], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
                                                                                        'Active Power POC':                 {'channels':[dict(channels_lib['NW_post']['P_POC'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['Pspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MW', 'rank':3},
                                                                                        'Reactive Power POC':               {'channels':[dict(channels_lib['NW_post']['Q_POC'], **{'markers':["rise_t","set_t"]}),
                                                                                                                                         dict(channels_lib['NW_post']['Qspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'MVAr', 'rank':5},
                                                                                        'Frequency':                        {'channels':[dict(channels_lib['NW_post']['F_POC'], **{})], 'unit':'Hz', 'rank':7, 'yminspan':0.1},

#                                                                                        'Powerfactor POC':                  {'channels':[dict(channels_lib['PSSE']['PF_POC'], **{}),
#                                                                                                                                         dict(channels_lib['PSSE']['PFspt_POC'], **{'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}),], 'unit':'p.u', 'rank':7},
                                                                                        'PPC FRT signal':                   {'channels':[dict(channels_lib['NW_post']['FRT_PPC'], **{}),], 'unit':'PPC code', 'rank':9},
    
                                                                                                                                            
                                                                                        'Voltage INV':                      {'channels':[dict(channels_lib['NW_post']['V_INV1'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['V_INV2'], **{}),], 'unit':'p.u.', 'rank':2, 'yminspan':0.01},                                                                                                                                                
                                                                                        'Active Power INV':                 {'channels':[dict(channels_lib['NW_post']['P_INV1'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['P_INV2'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['Pcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['NW_post']['Pcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MW', 'rank':4, 'ymaxlim':200.0, 'yminlim':-200.0, },
                                                                                        'Reactive Power INV':               {'channels':[dict(channels_lib['NW_post']['Q_INV1'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['Q_INV2'], **{}),
                                                                                                                                         dict(channels_lib['NW_post']['Qcmd_PPC_INV1'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         dict(channels_lib['NW_post']['Qcmd_PPC_INV2'], **{'linestyle':'--', 'linewidth': 1.5}),
                                                                                                                                         ], 'unit':'MVAr', 'rank':6, 'ymaxlim':200.0, 'yminlim':-200.0, },
                                                                                        'PV FRT signal':                    {'channels':[dict(channels_lib['NW_post']['FRT_INV1'], **{}),], 'unit':'inverter code', 'rank':8},
                                                                                        'BESS FRT signal':                  {'channels':[dict(channels_lib['NW_post']['FRT_INV2'], **{}),], 'unit':'inverter code', 'rank':10},
                                                                              },
                                                    },# end of plots
                                            }, # end of chapters


                                         ] #end of report_definition
                    }, # end of NetworkEvent
    
    

            
            } #specify which type of output document(s) shall be generated

#channels={} #specified which channels to be used 

#datasets={} #specify locations in which to look for data sets
#------------------------------------------------------------------------------
# IMPORTS
#------------------------------------------------------------------------------
import os
import sys
from win32com.client import Dispatch
#main_folder_path=os.path.dirname(os.path.dirname(os.path.dirname(__file__)))
#main_folder_path_out=os.path.dirname(os.path.dirname(os.path.dirname(__file__)))

def make_dir(dir_path, dir_name=""):
    dir_make = os.path.join(dir_path, dir_name)
    try:
        os.mkdir(dir_make)
    except OSError:
        pass
    
def createPath(main_folder_out):
    path = os.path.normpath(main_folder_out)
    path_splits = path.split(os.sep) # Get the components of the path
    child_folder = r""+ path_splits[0] # Build up the output path from base drive
    for i in range(len(path_splits)-1):
        child_folder = child_folder + "\\" + path_splits[i+1]
        make_dir(child_folder)
    return child_folder

def createShortcut(target, path):
    # target = ModelCopyDir # directory to which the shortcut is created
    # path = main_folder + "\\model_copies.lnk"  #This is where the shortcut will be created
    shell = Dispatch('WScript.Shell')
    shortcut = shell.CreateShortCut(path)
    shortcut.Targetpath = target
    shortcut.save()
    
#-----------------------------------------------------------------------------
# Define Project Paths
#-----------------------------------------------------------------------------
script_dir=os.getcwd()
main_folder=os.path.abspath(os.path.join(script_dir, os.pardir))
main_folder_path = os.path.dirname(main_folder)
if "OneDrive - OX2" in main_folder_path: # if the current folder is online (under OneDrive - OX2), create a new directory to store the result
    user = os.path.expanduser('~')
    main_path_out = main_folder_path.replace(user + "\OneDrive - OX2","C:\work") # Change the path from Onedrive to Local in c drive
    main_folder_path_out = createPath(main_path_out)
else:
    main_folder_path_out = main_folder_path

output_loc=main_folder_path_out+"\\Plots\\" + report_types[0]
createPath(output_loc)

# Create shortcut linking to result folder if it is not stored in the main folder path
if main_folder_path_out != main_folder_path:
    createShortcut(output_loc, main_folder_path + "\\Plots\\"+report_types[0]+".lnk")
else: # if the output location is same as input location, then delete the links
    try:os.remove(main_folder_path + "\\Plots\\"+report_types[0]+".lnk")
    except: pass
sys.path.append(main_folder_path+"\\PSSE_sim\\scripts\\Libs")
sys.path.append(r"C:\Python27\Lib\site-packages")
#sys.path.append(r"C:\Program Files\Python37\Lib\site-packages")
import EscoPlot as ep
import shelve

from io import StringIO

import docx
from docx import Document, shape
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK
from docx import Document

from docx.enum.dml import MSO_THEME_COLOR_INDEX

import datetime
import readtestinfo


#------------------------------------------------------------------------------
# GLOBAL VARIABLES
#------------------------------------------------------------------------------
# [ProjectDetailsDict, PSCADmodelDict, PSSEmodelDict, SetpointsDict, TestTypesDict, ScenariosDict, ProfilesDict] =  readtestinfo.readTestdef(main_folder_path+"\\test_scenario_definitions\\"+TestDefinitionSheet)

return_dict =  readtestinfo.readTestdef(main_folder_path+"\\test_scenario_definitions\\"+TestDefinitionSheet, ['ProjectDetails', 'ModelDetailsPSSE', 'ModelDetailsPSCAD', 'Setpoints', 'ScenariosSMIB', 'Profiles'])
ProjectDetailsDict = return_dict['ProjectDetails']
# SimulationSettingsDict = return_dict['SimulationSettings']
PSSEmodelDict = return_dict['ModelDetailsPSSE']
PSCADmodelDict = return_dict['ModelDetailsPSCAD']
SetpointsDict = return_dict['Setpoints']
ScenariosDict = return_dict['ScenariosSMIB']
ProfilesDict = return_dict['Profiles']

landscape_flag=0
#------------------------------------------------------------------------------
# FUNCTIONS
#------------------------------------------------------------------------------
#sort cases by type and by ID (which is located at the end of the strin identifier)
def sort_cases(cases):
    sorted_cases=[]
    order = [ 'small', 'large', 'ort', 'tov', 'con', 'Case_']
    for testType in order:
        cases_temp={}
        for case in cases:
            if(testType in case): #will ignore any folders that are do not belong to one of thet est types
                number=float(case.replace(testType, ''))
                if(number>0):## additional criterion added for debugging purposes. Should be removed for final version
                    cases_temp[number]=case
        keys = cases_temp.keys()
        # if(keys!=[]):
        if(cases_temp!={}): # avoid the return of keys=dict_keys([]) in Python3
            for id in range(0, int(max(keys))+1):
                if id in keys:
                    sorted_cases.append(cases_temp[id])    
    return sorted_cases

def check_dataset_pos(dataset_ID, dataset_info):
    return_value=-1
    for dataset_pos in range (0, len(dataset_info)): #iterate over datasets for plot
        if (dataset_info[dataset_pos]['ID']==dataset_ID):
            return_value=dataset_pos
    return return_value
#check if string contains number

#for cases where the project consists of multiple plants, include the information which plant contributes which amount of active power and return is as a string
def retrieve_P_split(setpoint):
    P_split=''
    for key in setpoint.keys():
        if ("P_") in key:
            if (is_number(setpoint[key])):
                if (P_split==''):
                    P_split='('+key[2:-1]+'='+str(setpoint[key])+'MW; '
                else:
                    P_split+=key[2:-1]+'='+str(setpoint[key])+'MW; '
    if(P_split!=''):
        P_split=P_split[0:-2]+')'
    return P_split

#check if string contains number
def is_number(s):
    try:
        float(s)
        return True
    except:
        return False
#function is given test profile (pointwise definition) and returns the start times and end times of the steps included in that profile. A step is considered to occur whenever the gradien is greater than 0, and to end whenever the gradient becomes 0 again
def detect_steps(offset, profile):
    steps=[] 
    magnitude=max(profile['y'])
    i=0
    prev_grad=0
    while (i < len(profile['x'])-1):
        gradient=float((profile['y'][i+1]-profile['y'][i]))/(magnitude*(profile['x'][i+1]-profile['x'][i]))
        if((abs(gradient)>0.01) and (abs(prev_grad)<=0.01)):
            stepStart=profile['x'][i]
        elif( (i>0) and ((abs(gradient)<=0.01) or ( (abs(gradient)>0.01) and (i+1==len(profile['x']-1)) ) ) ):
            stepEnd=profile['x'][i]
            steps.append([stepStart, stepEnd])
        i+=1        
    return steps

def add_bookmark(paragraph, bookmark_text, bookmark_name):
    run = paragraph.add_run()
    tag = run._r
    start = docx.oxml.shared.OxmlElement('w:bookmarkStart')
    start.set(docx.oxml.ns.qn('w:id'), '0')
    start.set(docx.oxml.ns.qn('w:name'), bookmark_name)
    tag.append(start)

    text = docx.oxml.OxmlElement('w:r')
    text.text = bookmark_text
    tag.append(text)

    end = docx.oxml.shared.OxmlElement('w:bookmarkEnd')
    end.set(docx.oxml.ns.qn('w:id'), '0')
    end.set(docx.oxml.ns.qn('w:name'), bookmark_name)
    tag.append(end)

def add_link(paragraph, link_to, text, tool_tip=None):
    # create hyperlink node
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')

    # set attribute for link to bookmark
    hyperlink.set(docx.oxml.shared.qn('w:anchor'), link_to,)

    if tool_tip is not None:
        # set attribute for link to bookmark
        hyperlink.set(docx.oxml.shared.qn('w:tooltip'), tool_tip,)

    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    r = paragraph.add_run()
    r._r.append(hyperlink)
    r.font.name = "Calibri"
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True
    

def initialise_report(report_type):
    #read report template 
    report=Document(main_folder_path+"\\Plots\\ReportTemplate.docx")
    return report

def add_report_intro(report, report_type, datasets, cases): #change it so that data location is an array of datasets. For overlay plots the inclusion of time step can then be decided based on if PSS/E results are present.
    #generate general description and intro based on ProjectDetailsDict
    p=report.add_paragraph('')
    run=p.add_run()
    run.add_break(WD_BREAK.PAGE)
    report.add_heading("Introduction", level=1 )
    if(report_type=="BENCHMARKING"):
        intro_text=str(ProjectDetailsDict['Dev'])+" is developing "+str(ProjectDetailsDict['Name'])+" in "+str(ProjectDetailsDict['Town'])+", "+str(ProjectDetailsDict['State'])+". "
        intro_text+="The project will connect into "+str(ProjectDetailsDict['Sub'])+" and feature a total active power rating of "+str(ProjectDetailsDict['PlantMW'])+" MW at the point of connection. " 
        pass
        p=report.add_paragraph(intro_text)
        
        software_types=[]
        for dataset in datasets:
            if (( ('PSS/E' in dataset['label']) or ("PSSE" in dataset['label']) ) and not ("PSS/E" in software_types)):
                software_types.append("PSS/E")
        if( ('PSCAD' in dataset['label']) and not ("PSCAD" in software_types)):
            software_types.append('PSCAD')
        intro_text="As part of the connection application, a PSS/E and a PSCAD model are submitted. These models are required to adequately represent the performance of the hardware proposed to be installed on site. As such, the two models need to show comparable behaviour and performance when subjected to the same test conditions in both PSS/E and PSCAD. "
        intro_text+="This report shows the results of benchmarking studies that have been carried out to demonstrate adequate alignment between the PSCAD and PSS/E model. "
        p=report.add_paragraph(intro_text)
        intro_text="The studies have been conducted in "+str(PSSEmodelDict['PSSEversion'])+" and "+str(PSCADmodelDict['pscad version'])+" using the compiler "+str(PSCADmodelDict['compiler'])+". "
        intro_text+="The tests included in this report are listed below." 
        summary_table=True

        
    elif(report_type=="DMAT"): 
        intro_text=str(ProjectDetailsDict['Dev'])+" is developing "+str(ProjectDetailsDict['Name'])+" in "+str(ProjectDetailsDict['Town'])+" ,"+str(ProjectDetailsDict['State'])+". "
        intro_text+="The project will connect into "+str(ProjectDetailsDict['Sub'])+" and feature a total active power rating of "+str(ProjectDetailsDict['PlantMW'])+" MW at the point of connection. " 
        pass
        p=report.add_paragraph(intro_text)
        
        software_types=[]
        for dataset in datasets:
            if (( ('PSS/E' in dataset['label']) or ("PSSE" in dataset['label']) ) and not ("PSS/E" in software_types)):
                software_types.append("PSS/E")
        if( ('PSCAD' in dataset['label']) and not ("PSCAD" in software_types)):
            software_types.append('PSCAD')
        
        intro_text="As part of the connection application "
        if(len(software_types)==1):
            intro_text+='a '+software_types[0]+" model is submitted. The model is required to adequately represent the performance of the hardware proposed to be installed on site, and the model is also expected to show acceptable performance in the test scenarios outlined in AEMO's Model Acceptance Test Guideline. "
            intro_text+="This report shows the results of Model Acceptance test studies that have been carried out to demonstrate the performance of the "+software_types[0]+" model. "
        elif(len(software_types)>1):
            intro_text+='models in '
            for software in range(0, len(software_types)-1):
                intro_text+=software_types[software]+", "
            intro_text+=' and '+software_types[-1]+" are submitted. The models are required to adequately represent the performance of the hardware proposed to be installed on site, and the models are also expected to show acceptable performance in the test scenarios outlined in AEMO's Model Acceptance Test Guideline. "
            intro_text+="This report shows the results of Model Acceptance test studies that have been carried out to demonstrate the performance of the models. "
        p=report.add_paragraph(intro_text)
        if('PSCAD' in software_types):
            intro_text="The PSCAD studies have been conducted in "+str(PSCADmodelDict['pscad version'])+" using the compiler "+str(PSCADmodelDict['compiler'])+". "
        if('PSS/E' in software_types):
            intro_text="The PSSE studies have been conducted in "+str(PSSEmodelDict['PSSEversion'])+". "
        intro_text+="The tests included in this report are listed below."   
        summary_table=True

    elif(report_type=="GPS"): 
        intro_text=str(ProjectDetailsDict['Dev'])+" is developing "+str(ProjectDetailsDict['Name'])+" in "+str(ProjectDetailsDict['Town'])+" ,"+str(ProjectDetailsDict['State'])+". "
        intro_text+="The project will connect into "+str(ProjectDetailsDict['Sub'])+" and feature a total active power rating of "+str(ProjectDetailsDict['PlantMW'])+" MW at the point of connection. " 
        pass
        p=report.add_paragraph(intro_text)
        
        software_types=[]
        for dataset in datasets:
            if (( ('PSS/E' in dataset['label']) or ("PSSE" in dataset['label']) ) and not ("PSS/E" in software_types)):
                software_types.append("PSS/E")
        if( ('PSCAD' in dataset['label']) and not ("PSCAD" in software_types)):
            software_types.append('PSCAD')
        
        intro_text="As part of the connection application "
        if(len(software_types)==1):
            intro_text+='a '+software_types[0]+" model is submitted. The model is required to adequately represent the performance of the hardware proposed to be installed on site, and the model is also expected to show acceptable performance in the test scenarios outlined in the NER S525. "
            intro_text+="This report shows the results of GPS test studies that have been carried out to demonstrate the performance of the "+software_types[0]+" model. "
        elif(len(software_types)>1):
            intro_text+='models in '
            for software in range(0, len(software_types)-1):
                intro_text+=software_types[software]+", "
            intro_text+=' and '+software_types[-1]+" are submitted. The models are required to adequately represent the performance of the hardware proposed to be installed on site, and the models are also expected to show acceptable performance in the test scenarios outlined in NER S525. "
            intro_text+="This report shows the results of GPS test studies that have been carried out to demonstrate the performance of the models. "
        p=report.add_paragraph(intro_text)
        if('PSCAD' in software_types):
            intro_text="The PSCAD studies have been conducted in "+str(PSCADmodelDict['pscad version'])+" using the compiler "+str(PSCADmodelDict['compiler'])+". "
        if('PSS/E' in software_types):
            intro_text="The PSSE studies have been conducted in "+str(PSSEmodelDict['PSSEversion'])+". "
        intro_text+="The tests included in this report are listed below."   
        summary_table=True
        
    elif(report_type=="NetworkEvent"): 
        intro_text=str(ProjectDetailsDict['Dev'])+" is developing "+str(ProjectDetailsDict['Name'])+" in "+str(ProjectDetailsDict['Town'])+" ,"+str(ProjectDetailsDict['State'])+". "
        intro_text+="The project will connect into "+str(ProjectDetailsDict['Sub'])+" and feature a total active power rating of "+str(ProjectDetailsDict['PlantMW'])+" MW at the point of connection. " 
        pass
        p=report.add_paragraph(intro_text)
        
        software_types=[]
        for dataset in datasets:
            if (( ('PSS/E' in dataset['label']) or ("PSSE" in dataset['label']) ) and not ("PSS/E" in software_types)):
                software_types.append("PSS/E")
        if( ('PSCAD' in dataset['label']) and not ("PSCAD" in software_types)):
            software_types.append('PSCAD')
        
        intro_text="As part of the connection application "
        if(len(software_types)==1):
            intro_text+='a '+software_types[0]+" model is submitted. The model is required to adequately represent the performance of the hardware proposed to be installed on site, and the model is also expected to show acceptable performance in the test scenarios outlined in AEMO's Model Acceptance Test Guideline. "
            intro_text+="This report shows the results of Model Acceptance test studies that have been carried out to demonstrate the performance of the "+software_types[0]+" model. "
        elif(len(software_types)>1):
            intro_text+='models in '
            for software in range(0, len(software_types)-1):
                intro_text+=software_types[software]+", "
            intro_text+=' and '+software_types[-1]+" are submitted. The models are required to adequately represent the performance of the hardware proposed to be installed on site, and the models are also expected to show acceptable performance in the test scenarios outlined in AEMO's Model Acceptance Test Guideline. "
            intro_text+="This report shows the results of Model Acceptance test studies that have been carried out to demonstrate the performance of the models. "
        p=report.add_paragraph(intro_text)
        if('PSCAD' in software_types):
            intro_text="The PSCAD studies have been conducted in "+str(PSCADmodelDict['pscad version'])+" using the compiler "+str(PSCADmodelDict['compiler'])+". "
        if('PSS/E' in software_types):
            intro_text="The PSSE studies have been conducted in "+str(PSSEmodelDict['PSSEversion'])+". "
        intro_text+="The tests included in this report are listed below."   
        summary_table=True

    p=report.add_paragraph('')        
    run=p.add_run()
    run.add_break(WD_BREAK.PAGE)
    
#    if(summary_table):     
#        p=report.add_paragraph(intro_text)        
#        change_orientation(report)
#        p=report.add_paragraph("")
#        tableCnt=1
#        if(any('small' in case for case in cases)):
#            p.add_run('Table '+str(tableCnt)+': Scenario list - Small Disturbance tests').bold=True        
#
#            if('PSS/E' in software_types):
##                headers=['Case Nr.', 'Test Type', 'Test profile', 'POC voltage', 'Grid MVA', 'Grid X/R', 'P at POC (MW)', 'Q at POC (MVAr)', 'PSSE time step', 'PSSE acc. factor', 'passed']
#                headers=['Case Nr.', 'Test Type', 'Test profile', 'POC voltage', 'Grid SCR', 'Grid X/R', 'P at POC (MW)', 'Q at POC (MVAr)', 'PSSE time step (ms)', 'PSSE acc. factor', 'DMAT id', 'passed']
#            else:
##                headers=['Case Nr.', 'Test Type', 'Test profile', 'POC voltage', 'Grid MVA', 'Grid X/R', 'P at POC (MW)', 'Q at POC (MVAr)', 'passed']
#                headers=['Case Nr.', 'Test Type', 'Test profile', 'POC voltage', 'Grid SCR', 'Grid X/R', 'P at POC (MW)', 'Q at POC (MVAr)', 'DMAT id', 'passed']
#            table=report.add_table(rows=1, cols=len(headers))
#            table.style='ListTable3-Accent3'
#            hdr_cells=table.rows[0].cells
#            for header_id in range(0, len(headers)): hdr_cells[header_id].text=headers[header_id]
#            # read test metadata to get test info
#            for case_id in range(0, len(cases)):
#                if ('small' in cases[case_id]):
#                    dataset_number=0
#                    test_details={}
#                    while (dataset_number < len(datasets)) and (test_details=={}):
#                        try:
#                            test_details=shelve.open(main_folder_path_out+"\\"+datasets[dataset_number]['path']+"\\"+cases[case_id]+"\\testInfo\\"+cases[case_id])
#                        except:
#                            dataset_number+=1
#                    row_cells=table.add_row().cells
#                    cell_paragraph=row_cells[0].paragraphs[0]
#                    add_link(paragraph=cell_paragraph, link_to=str(cases[case_id]), text=str(cases[case_id]), tool_tip="link to test results")
#                    row_cells[1].text=str(test_details['scenario_params']['Test Type'])
#                    row_cells[2].text=str(test_details['scenario_params']['Test profile'])
#                    row_cells[3].text=str(test_details['setpoint']['V_POC'])
##                    row_cells[4].text=str(test_details['setpoint']['GridMVA'])
##                    row_cells[4].text=str(test_details['setpoint']['SCR'])
#                    row_cells[4].text=str(round(test_details['setpoint']['SCR'],2))
#                    row_cells[5].text=str(test_details['setpoint']['X_R'])
#                    row_cells[6].text=str(test_details['setpoint']['P'])
#                    row_cells[7].text=str(test_details['setpoint']['Q'])
#                    if('PSS/E' in software_types):
#                        row_cells[8].text=str(test_details['scenario_params']['TimeStep'])
#                        row_cells[9].text=str(test_details['scenario_params']['AccFactor'])
#                    row_cells[10].text=str(test_details['scenario_params']['Comment/Corresponding DMAT case'])
##                    row_cells[10].text=str(test_details['scenario_params']['comment'])
#            for row in table.rows:
#                for cell in row.cells:
#                    paragraphs = cell.paragraphs
#                    for paragraph in paragraphs:
#                        for run in paragraph.runs:
#                            font = run.font
#                            font.size= Pt(9)
#        
#            tableCnt+=1
#            #run.add_break()
#            p=report.add_paragraph(" ")
#            p=report.add_paragraph('')
#        if(any('large' in case for case in cases)):
#            p.add_run('Table '+str(tableCnt)+': Scenario list - Large Disturbance tests').bold=True  
#            if('PSS/E' in software_types):
##                headers=['Case Nr.', 'Test Type', 'Fault Type', 'Fault time', 'Fault duration', 'Fault impedance', 'V residual', 'Fault X/R', 'POC voltage', 'Grid MVA', 'Grid X/R', 'Grid MVA post-fault','Grid X/R post-fault', 'P at POC (MW)', 'Q at POC (MVAr)', 'PSSE time step', 'PSSE acc. factor', 'passed']
#                headers=['Case Nr.', 'Test Type', 'Fault Type', 'Fault time', 'Fault duration', 'Fault impedance', 'V residual', 'Fault X/R', 'POC voltage', 'Grid SCR', 'Grid X/R', 'Grid SCR post-fault','Grid X/R post-fault', 'P at POC (MW)', 'Q at POC (MVAr)', 'PSSE time step (ms)', 'PSSE acc. factor', 'DMAT id', 'passed']
#            else:    
##                headers=['Case Nr.', 'Test Type', 'Fault Type', 'Fault time', 'Fault duration', 'Fault impedance', 'V residual', 'Fault X/R', 'POC voltage', 'Grid MVA', 'Grid X/R', 'Grid MVA post-fault','Grid X/R post-fault', 'P at POC (MW)', 'Q at POC (MVAr)', 'passed']
#                headers=['Case Nr.', 'Test Type', 'Fault Type', 'Fault time', 'Fault duration', 'Fault impedance', 'V residual', 'Fault X/R', 'POC voltage', 'Grid SCR', 'Grid X/R', 'Grid SCR post-fault','Grid X/R post-fault', 'P at POC (MW)', 'Q at POC (MVAr)', 'DMAT id', 'passed']
#            table1=report.add_table(rows=1, cols=len(headers))
#            table1.style='ListTable3-Accent3'
#            hdr_cells1=table1.rows[0].cells
#            for header_id in range(0, len(headers)): hdr_cells1[header_id].text=headers[header_id]
#            # read test metadata to get test info
#            for case_id in range(0, len(cases)):
#                if ('large' in cases[case_id]):
#                    dataset_number=0
#                    test_details={}
#                    while (dataset_number < len(datasets)) and (test_details=={}):
#                        try:
#                            test_details=shelve.open(main_folder_path_out+"\\"+datasets[dataset_number]['path']+"\\"+cases[case_id]+"\\testInfo\\"+cases[case_id])
#                        except:
#                            dataset_number+=1
#                    row_cells=table1.add_row().cells
#                    cell_paragraph=row_cells[0].paragraphs[0]
#                    add_link(paragraph=cell_paragraph, link_to=str(cases[case_id]), text=str(cases[case_id]), tool_tip="link to test results")
#                    row_cells[1].text=str(test_details['scenario_params']['Test Type'])
#
#                    if(not 'Multifault' in test_details['scenario_params']['Test Type'] ):#single fault test
#                        row_cells[2].text=str(test_details['scenario_params']['Ftype'])
#                        row_cells[3].text=str(test_details['scenario_params']['Ftime'])
#                        row_cells[4].text=str(test_details['scenario_params']['Fduration'])
#                        row_cells[5].text=str(round(test_details['scenario_params']['F_Impedance'],2))
#                        if(test_details['scenario_params']['Vresidual']!=''):
#                            row_cells[6].text=str(test_details['scenario_params']['Vresidual'])
#                        else:
#                            row_cells[6].text='-'
#                        row_cells[7].text=str(test_details['scenario_params']['Fault X_R'])
#                    else:
#                        row_cells[2].text='various'
#                        row_cells[3].text='various'
#                        row_cells[4].text='various'
#                        row_cells[5].text='various'
#                        row_cells[6].text='various'
#                        row_cells[7].text='various'
#                    row_cells[8].text=str(test_details['setpoint']['V_POC'])
##                    row_cells[9].text=str(test_details['setpoint']['GridMVA'])
#                    row_cells[9].text=str(round(test_details['setpoint']['SCR'],2))
#                    row_cells[10].text=str(test_details['setpoint']['X_R'])
#                    if('SCL_post' in test_details['scenario_params']):
#                        if(test_details['scenario_params']['SCL_post']!=''):
#                            row_cells[11].text=str(test_details['scenario_params']['SCL_post'])
#                        else:
##                            row_cells[11].text=str(test_details['setpoint']['GridMVA'])
#                            row_cells[11].text=str(round(test_details['setpoint']['SCR'],2))
#                    else:
##                        row_cells[11].text=str(test_details['setpoint']['GridMVA'])
#                        row_cells[11].text=str(round(test_details['setpoint']['SCR'],2))
#                    if('X_R_post' in test_details['scenario_params']):
#                        if(test_details['scenario_params']['X_R_post']!=''):
#                            row_cells[12].text=str(test_details['scenario_params']['X_R_post'])
#                        else:
#                            row_cells[12].text=str(test_details['setpoint']['X_R'])
#                    else:
#                        row_cells[12].text=str(test_details['setpoint']['X_R'])  
#                    row_cells[13].text=str(test_details['setpoint']['P'])
#                    row_cells[14].text=str(test_details['setpoint']['Q'])
#                    if('PSS/E' in software_types):
#                        row_cells[15].text=str(test_details['scenario_params']['TimeStep'])
#                        row_cells[16].text=str(test_details['scenario_params']['AccFactor'])
#                    row_cells[17].text=str(test_details['scenario_params']['Comment/Corresponding DMAT case'])
#            for row in table1.rows:
#                for cell in row.cells:
#                    paragraphs = cell.paragraphs
#                    for paragraph in paragraphs:
#                        for run in paragraph.runs:
#                            font = run.font
#                            font.size= Pt(9)
#                        
#            tableCnt+=1
#            #run.add_break()
#            p=report.add_paragraph(" ")
#            p=report.add_paragraph('')
#        if(any('ort' in case for case in cases)):
#            p.add_run('Table '+str(tableCnt)+': Scenario list - Oscillatory Rejection tests').bold=True        
#            if('PSS/E' in software_types):
##                headers=['Case Nr.', 'Test Type', 'Time', 'Disturbance Frequency', 'Disturbance Magnitude', 'Phase Oscillation Magnitude', 'POC voltage', 'Grid MVA', 'Grid X/R',  'P at POC (MW)', 'Q at POC (MVAr)', 'PSSE time step', 'PSSE acc. factor', 'passed']
#                headers=['Case Nr.', 'Test Type', 'Time', 'Disturbance Frequency', 'Disturbance Magnitude', 'Phase Oscillation Magnitude', 'POC voltage', 'Grid SCR', 'Grid X/R',  'P at POC (MW)', 'Q at POC (MVAr)', 'PSSE time step (ms)', 'PSSE acc. factor', 'DMAT id', 'passed']
#            else:
##                headers=['Case Nr.', 'Test Type', 'Time', 'Disturbance Frequency', 'Disturbance Magnitude', 'Phase Oscillation Magnitude', 'POC voltage', 'Grid MVA', 'Grid X/R',  'P at POC (MW)', 'Q at POC (MVAr)', 'passed']
#                headers=['Case Nr.', 'Test Type', 'Time', 'Disturbance Frequency', 'Disturbance Magnitude', 'Phase Oscillation Magnitude', 'POC voltage', 'Grid SCR', 'Grid X/R',  'P at POC (MW)', 'Q at POC (MVAr)', 'DMAT id', 'passed']
#            table1=report.add_table(rows=1, cols=len(headers))
#            table1.style='ListTable3-Accent3'
#            hdr_cells1=table1.rows[0].cells
#            for header_id in range(0, len(headers)): hdr_cells1[header_id].text=headers[header_id]
#            # read test metadata to get test info
#            for case_id in range(0, len(cases)):
#                if ('ort' in cases[case_id]):
#                    dataset_number=0
#                    test_details={}
#                    while (dataset_number < len(datasets)) and (test_details=={}):
#                        try:
#                            test_details=shelve.open(main_folder_path_out+"\\"+datasets[dataset_number]['path']+"\\"+cases[case_id]+"\\testInfo\\"+cases[case_id])
#                        except:
#                            dataset_number+=1
#                    row_cells=table1.add_row().cells
#                    cell_paragraph=row_cells[0].paragraphs[0]
#                    add_link(paragraph=cell_paragraph, link_to=str(cases[case_id]), text=str(cases[case_id]), tool_tip="link to test results")
#                    row_cells[1].text=str(test_details['scenario_params']['Test Type'])
#                    if('time' in test_details['scenario_params'].keys()):
#                        row_cells[2].text=str(test_details['scenario_params']['time'])
#                    else:
#                        row_cells[3].text='0.0'
#                    row_cells[3].text=str(test_details['scenario_params']['Disturbance Frequency'])
#                    row_cells[4].text=str(round(test_details['scenario_params']['Disturbance Magnitude'],2))
#                    row_cells[5].text=str(test_details['scenario_params']['PhaseOsc Magnitude'])
#                    row_cells[6].text=str(test_details['setpoint']['V_POC'])
##                    row_cells[7].text=str(test_details['setpoint']['GridMVA'])
#                    row_cells[7].text=str(round(test_details['setpoint']['SCR'],2))
#                    row_cells[8].text=str(test_details['setpoint']['X_R'])
#                    row_cells[9].text=str(test_details['setpoint']['P'])
#                    row_cells[10].text=str(test_details['setpoint']['Q'])
#                    if('PSS/E' in software_types):
#                        if ('TimeStep' in test_details['scenario_params'].keys()):
#                            row_cells[11].text=str(test_details['scenario_params']['TimeStep'])
#                        elif('time step' in test_details['scenario_params'].keys()):
#                            row_cells[11].text=str(test_details['scenario_params']['time step'])
#                            row_cells[12].text=str(test_details['scenario_params']['AccFactor'])
#                    row_cells[13].text=str(test_details['scenario_params']['Comment/Corresponding DMAT case'])
#            for row in table1.rows:
#                for cell in row.cells:
#                    paragraphs = cell.paragraphs
#                    for paragraph in paragraphs:
#                        for run in paragraph.runs:
#                            font = run.font
#                            font.size= Pt(9)
#                        
#            tableCnt+=1
#            #run.add_break()
#            p=report.add_paragraph(" ")
#            p=report.add_paragraph('')
#        if(any('tov' in case for case in cases)):
#            p.add_run('Table '+str(tableCnt)+': Scenario list - Temporary Over-Voltage tests').bold=True        
#            if('PSS/E' in software_types):
##                headers=['Case Nr.', 'Test Type', 'Time', 'TOV duration', 'Capacity(F)', 'V residual', 'POC voltage', 'Grid MVA', 'Grid X/R', 'Grid MVA post-TOV','Grid X/R post-TOV', 'P at POC (MW)', 'Q at POC (MVAr)', 'PSSE time step', 'PSSE acc. factor', 'passed']
#                headers=['Case Nr.', 'Test Type', 'Time', 'TOV duration', 'Capacity(F)', 'V residual', 'POC voltage', 'Grid SCR', 'Grid X/R', 'Grid SCR post-TOV','Grid X/R post-TOV', 'P at POC (MW)', 'Q at POC (MVAr)', 'PSSE time step (ms)', 'PSSE acc. factor', 'DMAT id', 'passed']
#            else:
##                headers=['Case Nr.', 'Test Type', 'Time', 'TOV duration', 'Capacity(F)', 'V residual', 'POC voltage', 'Grid MVA', 'Grid X/R', 'Grid MVA post-TOV','Grid X/R post-TOV', 'P at POC (MW)', 'Q at POC (MVAr)', 'passed']
#                headers=['Case Nr.', 'Test Type', 'Time', 'TOV duration', 'Capacity(F)', 'V residual', 'POC voltage', 'Grid SCR', 'Grid X/R', 'Grid SCR post-TOV','Grid X/R post-TOV', 'P at POC (MW)', 'Q at POC (MVAr)', 'DMAT id', 'passed']
#            table1=report.add_table(rows=1, cols=len(headers))
#            table1.style='ListTable3-Accent3'
#            hdr_cells1=table1.rows[0].cells
#            for header_id in range(0, len(headers)): hdr_cells1[header_id].text=headers[header_id]
#            # read test metadata to get test info
#            for case_id in range(0, len(cases)):
#                if ('tov' in cases[case_id]):
#                    dataset_number=0
#                    test_details={}
#                    while (dataset_number < len(datasets)) and (test_details=={}):
#                        try:
#                            test_details=shelve.open(main_folder_path_out+"\\"+datasets[dataset_number]['path']+"\\"+cases[case_id]+"\\testInfo\\"+cases[case_id])
#                        except:
#                            dataset_number+=1
#                    row_cells=table1.add_row().cells
#                    cell_paragraph=row_cells[0].paragraphs[0]
#                    add_link(paragraph=cell_paragraph, link_to=str(cases[case_id]), text=str(cases[case_id]), tool_tip="link to test results")
#                    row_cells[1].text=str(test_details['scenario_params']['Test Type'])
#                    row_cells[2].text=str(test_details['scenario_params']['time'])
#                    row_cells[3].text=str(test_details['scenario_params']['Fduration'])
#                    if('Capacity(F)' in test_details['scenario_params'].keys()):
#                        if(test_details['scenario_params']['Capacity(F)']!=''):
#                            row_cells[4].text=str(round(float(test_details['scenario_params']['Capacity(F)']),2))
#                        else:row_cells[4].text='-'
#                    else:row_cells[4].text='-'
#                    row_cells[5].text=str(test_details['scenario_params']['Vresidual'])
#                    row_cells[6].text=str(test_details['setpoint']['V_POC'])
##                    row_cells[7].text=str(test_details['setpoint']['GridMVA'])
#                    row_cells[7].text=str(round(test_details['setpoint']['SCR'],2))
#                    row_cells[8].text=str(test_details['setpoint']['X_R'])
#                    if('SCL_post' in test_details['scenario_params']):
#                        if(test_details['scenario_params']['SCL_post']!=''):
#                            row_cells[9].text=str(test_details['scenario_params']['SCL_post'])
#                        else:
##                            row_cells[9].text=str(test_details['setpoint']['GridMVA'])
#                            row_cells[9].text=str(round(test_details['setpoint']['SCR'],2))
#                    else:
##                        row_cells[9].text=str(test_details['setpoint']['GridMVA'])
#                        row_cells[9].text=str(round(test_details['setpoint']['SCR'],2))
#                    if('X_R_post' in test_details['scenario_params']):
#                        if(test_details['scenario_params']['X_R_post']!=''):
#                            row_cells[10].text=str(test_details['scenario_params']['X_R_post'])
#                        else:
#                            row_cells[10].text=str(test_details['setpoint']['X_R'])
#                    else:
#                        row_cells[10].text=str(test_details['setpoint']['X_R'])            
#                    row_cells[11].text=str(test_details['setpoint']['P'])
#                    row_cells[12].text=str(test_details['setpoint']['Q'])
#                    if('PSS/E' in software_types):
#                        row_cells[13].text=str(test_details['scenario_params']['TimeStep'])
#                        row_cells[14].text=str(test_details['scenario_params']['AccFactor'])
#                    row_cells[15].text=str(test_details['scenario_params']['Comment/Corresponding DMAT case'])
#            for row in table1.rows:
#                for cell in row.cells:
#                    paragraphs = cell.paragraphs
#                    for paragraph in paragraphs:
#                        for run in paragraph.runs:
#                            font = run.font
#                            font.size= Pt(9)
#                            
#            tableCnt+=1    
#            #run.add_break()
#            p=report.add_paragraph(" ")
#            p=report.add_paragraph('') 
#
#
#        if(any('con' in case for case in cases)):
#            p.add_run('Table '+str(tableCnt)+': Scenario list - Network Contingency tests').bold=True        
#
#            if('PSS/E' in software_types):
##                headers=['Case Nr.', 'Test Type', 'Test profile', 'POC voltage', 'Grid MVA', 'Grid X/R', 'P at POC (MW)', 'Q at POC (MVAr)', 'PSSE time step', 'PSSE acc. factor', 'passed']
#                headers=['Case Nr.', 'Test Type', 'Test profile', 'POC voltage', 'Grid SCR', 'Grid X/R', 'P at POC (MW)', 'Q at POC (MVAr)', 'PSSE time step (ms)', 'PSSE acc. factor', 'DMAT id', 'passed']
#            else:
##                headers=['Case Nr.', 'Test Type', 'Test profile', 'POC voltage', 'Grid MVA', 'Grid X/R', 'P at POC (MW)', 'Q at POC (MVAr)', 'passed']
#                headers=['Case Nr.', 'Test Type', 'Test profile', 'POC voltage', 'Grid SCR', 'Grid X/R', 'P at POC (MW)', 'Q at POC (MVAr)', 'DMAT id', 'passed']
#            table=report.add_table(rows=1, cols=len(headers))
#            table.style='ListTable3-Accent3'
#            hdr_cells=table.rows[0].cells
#            for header_id in range(0, len(headers)): hdr_cells[header_id].text=headers[header_id]
#            # read test metadata to get test info
#            for case_id in range(0, len(cases)):
#                if ('small' in cases[case_id]):
#                    dataset_number=0
#                    test_details={}
#                    while (dataset_number < len(datasets)) and (test_details=={}):
#                        try:
#                            test_details=shelve.open(main_folder_path_out+"\\"+datasets[dataset_number]['path']+"\\"+cases[case_id]+"\\testInfo\\"+cases[case_id])
#                        except:
#                            dataset_number+=1
#                    row_cells=table.add_row().cells
#                    cell_paragraph=row_cells[0].paragraphs[0]
#                    add_link(paragraph=cell_paragraph, link_to=str(cases[case_id]), text=str(cases[case_id]), tool_tip="link to test results")
#                    row_cells[1].text=str(test_details['scenario_params']['Test Type'])
#                    row_cells[2].text=str(test_details['scenario_params']['Test profile'])
#                    row_cells[3].text=str(test_details['setpoint']['V_POC'])
##                    row_cells[4].text=str(test_details['setpoint']['GridMVA'])
##                    row_cells[4].text=str(test_details['setpoint']['SCR'])
#                    row_cells[4].text=str(round(test_details['setpoint']['SCR'],2))
#                    row_cells[5].text=str(test_details['setpoint']['X_R'])
#                    row_cells[6].text=str(test_details['setpoint']['P'])
#                    row_cells[7].text=str(test_details['setpoint']['Q'])
#                    if('PSS/E' in software_types):
#                        row_cells[8].text=str(test_details['scenario_params']['TimeStep'])
#                        row_cells[9].text=str(test_details['scenario_params']['AccFactor'])
##                    row_cells[10].text=str(test_details['scenario_params']['Comment/Corresponding DMAT case'])
#                    row_cells[10].text=str(test_details['scenario_params']['comment'])
#            for row in table.rows:
#                for cell in row.cells:
#                    paragraphs = cell.paragraphs
#                    for paragraph in paragraphs:
#                        for run in paragraph.runs:
#                            font = run.font
#                            font.size= Pt(9)
#        
#
#        p=report.add_paragraph('')        
#        run=p.add_run()
#        run.add_break(WD_BREAK.PAGE)
#        change_orientation(report)
#            
##    report.add_heading("Simulation Results", level=1 )
#               
#                    
#        #generate table with scenarios and scenario details based on list of cases and metadata saved in 'data_location'
#    
#    return 0

def add_summary_table(report, report_type, datasets, cases): #change it so that data location is an array of datasets. For overlay plots the inclusion of time step can then be decided based on if PSS/E results are present.

#    p=report.add_paragraph(intro_text)        
    change_orientation(report)
    p=report.add_paragraph("")
    tableCnt=1
    if(any('small' in case for case in cases)):
        p.add_run('Table '+str(tableCnt)+': Scenario list - Small Disturbance tests').bold=True        

#            headers=['Case Nr.', 'Test Type', 'Test profile', 'POC voltage', 'Grid SCR', 'Grid X/R', 'P at POC (MW)', 'Q at POC (MVAr)', 'PSSE time step (ms)', 'PSSE acc. factor', 'DMAT id', 'passed']
        headers=['Case Nr.', 'Test Type', 'Test profile', 'POC voltage', 'Grid SCR', 'Grid X/R', 'P at POC (MW)', 'P BESS (MW)', 'Q at POC (MVAr)', 'DMAT id', 'Control mode']
        table=report.add_table(rows=1, cols=len(headers))
        table.style='ListTable3-Accent3'
        hdr_cells=table.rows[0].cells
        for header_id in range(0, len(headers)): hdr_cells[header_id].text=headers[header_id]
        # read test metadata to get test info
        for case_id in range(0, len(cases)):
            if ('small' in cases[case_id]):
                dataset_number=0
                test_details={}
                while (dataset_number < len(datasets)) and (test_details=={}):
                    try:
                        test_details=shelve.open(main_folder_path_out+"\\"+datasets[dataset_number]['path']+"\\"+cases[case_id]+"\\testInfo\\"+cases[case_id])
                    except:
                        dataset_number+=1
                row_cells=table.add_row().cells
                cell_paragraph=row_cells[0].paragraphs[0]
                add_link(paragraph=cell_paragraph, link_to=str(cases[case_id]), text=str(cases[case_id]), tool_tip="link to test results")
                row_cells[1].text=str(test_details['scenario_params']['Test Type'])
#                row_cells[1].text=str(test_details['scenario_params']['Test Type'])
                row_cells[2].text=str(test_details['scenario_params']['Test profile'])
                row_cells[3].text=str(round(test_details['setpoint']['V_POC'],3))
#                    row_cells[4].text=str(test_details['setpoint']['GridMVA'])
#                    row_cells[4].text=str(test_details['setpoint']['SCR'])
                row_cells[4].text=str(round(test_details['setpoint']['SCR'],2))
                row_cells[5].text=str(test_details['setpoint']['X_R'])
                row_cells[6].text=str(test_details['setpoint']['P'])
                row_cells[7].text=str(test_details['setpoint']['P_PV2'])
                row_cells[8].text=str(test_details['setpoint']['Q'])
                try: row_cells[9].text=str(test_details['scenario_params']['Comment/Corresponding DMAT case'])
                except: row_cells[9].text='N/A'
                try: row_cells[10].text=str(test_details['scenario_params']['test group'])
                except: row_cells[10].text='test group'
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size= Pt(9)
    
        tableCnt+=1
        #run.add_break()
        p=report.add_paragraph(" ")
        p=report.add_paragraph('')
    if(any('large' in case for case in cases)):
        p.add_run('Table '+str(tableCnt)+': Scenario list - Large Disturbance tests').bold=True  
        headers=['Case Nr.', 'Test Type', 'Fault Type', 'Fault time', 'Fault duration', 'Fault impedance', 'V residual', 'Fault X/R', 'POC voltage', 'Grid SCR', 'Grid X/R', 'Grid SCR post-fault','Grid X/R post-fault', 'P at POC (MW)', 'Q at POC (MVAr)', 'DMAT id', 'Control mode','passed']
        table1=report.add_table(rows=1, cols=len(headers))
        table1.style='ListTable3-Accent3'
        hdr_cells1=table1.rows[0].cells
        for header_id in range(0, len(headers)): hdr_cells1[header_id].text=headers[header_id]
        # read test metadata to get test info
        for case_id in range(0, len(cases)):
            if ('large' in cases[case_id]):
                dataset_number=0
                test_details={}
                while (dataset_number < len(datasets)) and (test_details=={}):
                    try:
                        test_details=shelve.open(main_folder_path_out+"\\"+datasets[dataset_number]['path']+"\\"+cases[case_id]+"\\testInfo\\"+cases[case_id])
                    except:
                        dataset_number+=1
                row_cells=table1.add_row().cells
                cell_paragraph=row_cells[0].paragraphs[0]
                add_link(paragraph=cell_paragraph, link_to=str(cases[case_id]), text=str(cases[case_id]), tool_tip="link to test results")
                row_cells[1].text=str(test_details['scenario_params']['Test Type'])

                if(not 'Multifault' in test_details['scenario_params']['Test Type'] ):#single fault test
                    row_cells[2].text=str(test_details['scenario_params']['Ftype'])
                    row_cells[3].text=str(test_details['scenario_params']['Ftime'])
                    row_cells[4].text=str(test_details['scenario_params']['Fduration'])
                    row_cells[5].text=str(round(test_details['scenario_params']['F_Impedance'],2))
                    if(test_details['scenario_params']['Vresidual']!=''):
                        row_cells[6].text=str(round(test_details['scenario_params']['Vresidual'],2))
                    else:
                        row_cells[6].text='-'
                    row_cells[7].text=str(test_details['scenario_params']['Fault X_R'])
                else:
                    row_cells[2].text='various'
                    row_cells[3].text='various'
                    row_cells[4].text='various'
                    row_cells[5].text='various'
                    row_cells[6].text='various'
                    row_cells[7].text='various'
                row_cells[8].text=str(test_details['setpoint']['V_POC'])
#                    row_cells[9].text=str(test_details['setpoint']['GridMVA'])
                row_cells[9].text=str(round(test_details['setpoint']['SCR'],2))
                row_cells[10].text=str(test_details['setpoint']['X_R'])
                if('SCL_post' in test_details['scenario_params']):
                    if(test_details['scenario_params']['SCL_post']!=''):
                        row_cells[11].text=str(test_details['scenario_params']['SCL_post'])
                    else:
#                            row_cells[11].text=str(test_details['setpoint']['GridMVA'])
                        row_cells[11].text=str(round(test_details['setpoint']['SCR'],2))
                else:
#                        row_cells[11].text=str(test_details['setpoint']['GridMVA'])
                    row_cells[11].text=str(round(test_details['setpoint']['SCR'],2))
                if('X_R_post' in test_details['scenario_params']):
                    if(test_details['scenario_params']['X_R_post']!=''):
                        row_cells[12].text=str(test_details['scenario_params']['X_R_post'])
                    else:
                        row_cells[12].text=str(test_details['setpoint']['X_R'])
                else:
                    row_cells[12].text=str(test_details['setpoint']['X_R'])  
                row_cells[13].text=str(test_details['setpoint']['P'])
                row_cells[14].text=str(test_details['setpoint']['Q'])
                try: row_cells[15].text=str(test_details['scenario_params']['Comment/Corresponding DMAT case'])
                except: row_cells[15].text=str('N/A')
                try: row_cells[16].text=str(test_details['scenario_params']['test group'])
                except: row_cells[16].text=str('test group')
                
        for row in table1.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size= Pt(9)
                    
        tableCnt+=1
        #run.add_break()
        p=report.add_paragraph(" ")
        p=report.add_paragraph('')
    if(any('ort' in case for case in cases)):
        p.add_run('Table '+str(tableCnt)+': Scenario list - Oscillatory Rejection tests').bold=True        
        headers=['Case Nr.', 'Test Type', 'Time', 'Disturbance Frequency', 'Disturbance Magnitude', 'Phase Oscillation Magnitude', 'POC voltage', 'Grid SCR', 'Grid X/R',  'P at POC (MW)', 'Q at POC (MVAr)', 'DMAT id', 'Control mode', 'passed']
        table1=report.add_table(rows=1, cols=len(headers))
        table1.style='ListTable3-Accent3'
        hdr_cells1=table1.rows[0].cells
        for header_id in range(0, len(headers)): hdr_cells1[header_id].text=headers[header_id]
        # read test metadata to get test info
        for case_id in range(0, len(cases)):
            if ('ort' in cases[case_id]):
                dataset_number=0
                test_details={}
                while (dataset_number < len(datasets)) and (test_details=={}):
                    try:
                        test_details=shelve.open(main_folder_path_out+"\\"+datasets[dataset_number]['path']+"\\"+cases[case_id]+"\\testInfo\\"+cases[case_id])
                    except:
                        dataset_number+=1
                row_cells=table1.add_row().cells
                cell_paragraph=row_cells[0].paragraphs[0]
                add_link(paragraph=cell_paragraph, link_to=str(cases[case_id]), text=str(cases[case_id]), tool_tip="link to test results")
                row_cells[1].text=str(test_details['scenario_params']['Test Type'])
                if('time' in test_details['scenario_params'].keys()):
                    row_cells[2].text=str(test_details['scenario_params']['time'])
                else:
                    row_cells[3].text='0.0'
                row_cells[3].text=str(test_details['scenario_params']['Disturbance Frequency'])
                row_cells[4].text=str(round(test_details['scenario_params']['Disturbance Magnitude'],2))
                row_cells[5].text=str(test_details['scenario_params']['PhaseOsc Magnitude'])
                row_cells[6].text=str(round(test_details['setpoint']['V_POC'],3))
#                    row_cells[7].text=str(test_details['setpoint']['GridMVA'])
                row_cells[7].text=str(round(test_details['setpoint']['SCR'],2))
                row_cells[8].text=str(test_details['setpoint']['X_R'])
                row_cells[9].text=str(test_details['setpoint']['P'])
                row_cells[10].text=str(test_details['setpoint']['Q'])
                row_cells[11].text=str(test_details['scenario_params']['Comment/Corresponding DMAT case'])
                row_cells[12].text=str(test_details['scenario_params']['test group'])

        for row in table1.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size= Pt(9)
                    
        tableCnt+=1
        #run.add_break()
        p=report.add_paragraph(" ")
        p=report.add_paragraph('')
    if(any('tov' in case for case in cases)):
        p.add_run('Table '+str(tableCnt)+': Scenario list - Temporary Over-Voltage tests').bold=True        
        headers=['Case Nr.', 'Test Type', 'Time', 'TOV duration', 'Capacity(uF)', 'V residual', 'POC voltage', 'Grid SCR', 'Grid X/R', 'Grid SCR post-TOV','Grid X/R post-TOV', 'P at POC (MW)', 'Q at POC (MVAr)', 'DMAT id', 'Control mode', 'passed']
        table1=report.add_table(rows=1, cols=len(headers))
        table1.style='ListTable3-Accent3'
        hdr_cells1=table1.rows[0].cells
        for header_id in range(0, len(headers)): hdr_cells1[header_id].text=headers[header_id]
        # read test metadata to get test info
        for case_id in range(0, len(cases)):
            if ('tov' in cases[case_id]):
                dataset_number=0
                test_details={}
                while (dataset_number < len(datasets)) and (test_details=={}):
                    try:
                        test_details=shelve.open(main_folder_path_out+"\\"+datasets[dataset_number]['path']+"\\"+cases[case_id]+"\\testInfo\\"+cases[case_id])
                    except:
                        dataset_number+=1
                row_cells=table1.add_row().cells
                cell_paragraph=row_cells[0].paragraphs[0]
                add_link(paragraph=cell_paragraph, link_to=str(cases[case_id]), text=str(cases[case_id]), tool_tip="link to test results")
                row_cells[1].text=str(test_details['scenario_params']['Test Type'])
                row_cells[2].text=str(test_details['scenario_params']['time'])
                row_cells[3].text=str(test_details['scenario_params']['Fduration'])
                if('Capacity(uF)' in test_details['scenario_params'].keys()):
                    if(test_details['scenario_params']['Capacity(uF)']!=''):
                        row_cells[4].text=str(round(float(test_details['scenario_params']['Capacity(uF)']),2))
                    else:row_cells[4].text='-'
                else:row_cells[4].text='-'
                row_cells[5].text=str(round(test_details['scenario_params']['Vresidual'],2))
                row_cells[6].text=str(round(test_details['setpoint']['V_POC'],3))
#                    row_cells[7].text=str(test_details['setpoint']['GridMVA'])
                row_cells[7].text=str(round(test_details['setpoint']['SCR'],2))
                row_cells[8].text=str(test_details['setpoint']['X_R'])
                if('SCL_post' in test_details['scenario_params']):
                    if(test_details['scenario_params']['SCL_post']!=''):
                        row_cells[9].text=str(test_details['scenario_params']['SCL_post'])
                    else:
#                            row_cells[9].text=str(test_details['setpoint']['GridMVA'])
                        row_cells[9].text=str(round(test_details['setpoint']['SCR'],2))
                else:
#                        row_cells[9].text=str(test_details['setpoint']['GridMVA'])
                    row_cells[9].text=str(round(test_details['setpoint']['SCR'],2))
                if('X_R_post' in test_details['scenario_params']):
                    if(test_details['scenario_params']['X_R_post']!=''):
                        row_cells[10].text=str(test_details['scenario_params']['X_R_post'])
                    else:
                        row_cells[10].text=str(test_details['setpoint']['X_R'])
                else:
                    row_cells[10].text=str(test_details['setpoint']['X_R'])            
                row_cells[11].text=str(test_details['setpoint']['P'])
                row_cells[12].text=str(test_details['setpoint']['Q'])
                row_cells[13].text=str(test_details['scenario_params']['Comment/Corresponding DMAT case'])
                try: row_cells[14].text=str(test_details['scenario_params']['test group'])
                except: row_cells[14].text='test group'

        for row in table1.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size= Pt(9)
                        
        tableCnt+=1    
        #run.add_break()
        p=report.add_paragraph(" ")
        p=report.add_paragraph('') 


    if(any('con' in case for case in cases)):
        p.add_run('Table '+str(tableCnt)+': Scenario list - Network Contingency tests').bold=True        
        headers=['Case Nr.', 'Test Type', 'Test profile', 'Event Element', 'Event Type', 'Event Description','passed']
        table=report.add_table(rows=1, cols=len(headers))
        table.style='ListTable3-Accent3'
        hdr_cells=table.rows[0].cells
        for header_id in range(0, len(headers)): hdr_cells[header_id].text=headers[header_id]
        # read test metadata to get test info
        for case_id in range(0, len(cases)):
            if ('con' in cases[case_id]):
                dataset_number=0
                test_details={}
                while (dataset_number < len(datasets)) and (test_details=={}):
                    try:
                        test_details=shelve.open(main_folder_path_out+"\\"+datasets[dataset_number]['path']+"\\"+cases[case_id]+"\\testInfo\\"+cases[case_id])
                    except:
                        dataset_number+=1
                row_cells=table.add_row().cells
                cell_paragraph=row_cells[0].paragraphs[0]
                add_link(paragraph=cell_paragraph, link_to=str(cases[case_id]), text=str(cases[case_id]), tool_tip="link to test results")
                row_cells[1].text=str(test_details['scenario_params']['Test Type'])
                row_cells[2].text=str(test_details['scenario_params']['Test profile'])
                row_cells[3].text=str(test_details['scenario_params']['Event_Element'])
                row_cells[4].text=str(test_details['scenario_params']['Event_Type'])
                row_cells[5].text=str(test_details['scenario_params']['CaseDescription'])

        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size= Pt(9)
    

    p=report.add_paragraph('')        
    run=p.add_run()
    run.add_break(WD_BREAK.PAGE)
    change_orientation(report)
            
#    report.add_heading("Simulation Results", level=1 )
               
                    
        #generate table with scenarios and scenario details based on list of cases and metadata saved in 'data_location'
    
    return 0
def add_plots_to_report(case, report, datasets, plots, plot_list, assessment):
    software_type=''
    if(assessment['PSSE_flag']>0):
        software_type='PSS/E'
    #retrieve test info
    dataset_number=0
    test_details={}
    while (dataset_number < len(datasets)) and (test_details=={}):
        try:
            test_details=shelve.open(main_folder_path_out+"\\"+datasets[dataset_number]['path']+"\\"+case+"\\testInfo\\"+case)
        except:
            dataset_number+=1
    if("Comment/Corresponding DMAT case" in test_details['scenario_params'].keys()):
        if(test_details['scenario_params']["Comment/Corresponding DMAT case"]!=''):
            paragraph_temp=report.add_heading('Case '+case+' (DMAT '+str(test_details['scenario_params']["Comment/Corresponding DMAT case"])+')', level=2 )
        else:
            paragraph_temp=report.add_heading('Case '+case, level=2 ) 
    elif("Comment/corresponding DMAT case" in test_details['scenario_params'].keys()):
        if(test_details['scenario_params']["Comment/corresponding DMAT case"]!=''):
            paragraph_temp=report.add_heading('Case '+case+' (DMAT '+str(test_details['scenario_params']["Comment/corresponding DMAT case"])+')', level=2 )
        else:
            paragraph_temp=report.add_heading('Case '+case, level=2 ) 
    elif("comment/corresponding DMAT case" in test_details['scenario_params'].keys()):
        if(test_details['scenario_params']["comment/corresponding DMAT case"]!=''):
            paragraph_temp=report.add_heading('Case '+case+' (DMAT '+str(test_details['scenario_params']["comment/corresponding DMAT case"])+')', level=2 )
        else:
            paragraph_temp=report.add_heading('Case '+case, level=2 ) 
    elif("comment" in test_details['scenario_params'].keys()):
        if(test_details['scenario_params']["comment"]!=''):
            paragraph_temp=report.add_heading('Case '+case+' (DMAT '+str(test_details['scenario_params']["comment"])+')', level=2 )
        else:
            paragraph_temp=report.add_heading('Case '+case, level=2 ) 
    else:        
    #add level 2 headline with test name
        paragraph_temp=report.add_heading('Case '+case, level=2 ) 
    add_bookmark(paragraph=paragraph_temp, bookmark_text='', bookmark_name=case )#add bookmark that link in table at start of document points to
    #add summary table with test details
#    table=report.add_table(rows=1, cols=2)
    table=report.add_table(rows=1, cols=4)
    table.style='ListTable3-Accent3-OX2'   
#    table.autofit = False
#    table.allow_autofit = False
#    table.columns[0].width = Inches(1.0)
#    table.columns[1].width = Inches(0.5)
#    table.columns[2].width = Inches(1.0)
#    table.columns[3].width = Inches(0.5)
    headers=['Parameter', 'Value','Parameter', 'Value']
    hdr_cells=table.rows[0].cells
    for header_id in range(0, len(headers)): hdr_cells[header_id].text=headers[header_id]    
    if('small') in case:
        dataset_number=0
        test_details={}
        while (dataset_number < len(datasets)) and (test_details=={}):
            try:
                test_details=shelve.open(main_folder_path_out+"\\"+datasets[dataset_number]['path']+"\\"+case+"\\testInfo\\"+case)
            except:
                dataset_number+=1
        row_cells=table.add_row().cells
        row_cells[0].text='Test Type'
        row_cells[1].text=str(test_details['scenario_params']['Test Type'])
        try: 
            row_cells[2].text='Control mode'
            row_cells[3].text=str(test_details['scenario_params']['test group'])
        except:
            row_cells[2].text='Test profile'
            row_cells[3].text=str(test_details['scenario_params']['Test profile'])
#        row_cells[2].text='Control mode'
#        row_cells[3].text=str(test_details['scenario_params']['test group'])
        row_cells=table.add_row().cells
        row_cells[0].text='Grid short circuit ratio'
#        row_cells[1].text=str(test_details['setpoint']['GridMVA'])
        row_cells[1].text=str(round(test_details['setpoint']['SCR'],2))
        row_cells[2].text='Grid X/R-ratio'
        row_cells[3].text=str(test_details['setpoint']['X_R'])
        row_cells=table.add_row().cells
        row_cells[0].text='P at POC (MW)'
        row_cells[1].text=str(test_details['setpoint']['P'])
        row_cells[2].text='Q at POC (MVAr)'
        row_cells[3].text=str(round(test_details['setpoint']['Q'],2))
        row_cells=table.add_row().cells
        row_cells[0].text='P from BESS (MW)'
        row_cells[1].text=str(test_details['setpoint']['P_PV2'])
        row_cells[2].text='POC voltage (p.u.)'
        row_cells[3].text=str(test_details['setpoint']['V_POC'])
        if(software_type=='PSS/E'):
            row_cells=table.add_row().cells
            row_cells[0].text='PSSE time step (ms)'
            row_cells[1].text=str(test_details['scenario_params']['TimeStep'])
#            row_cells=table.add_row().cells
            row_cells[2].text='PSSE acc. factor'
            row_cells[3].text=str(test_details['scenario_params']['AccFactor'])
    if('large') in case:
#        headers=['Case Nr.', 'Test Type', 'Fault Type', 'Fault time', ' Fault duration', 'Fault impedance', 'V residual', 'Fault X/R', 'POC voltage', 'Grid MVA', 'Grid X/R',  'P at POC (MW)', 'Q at POC (MVAr)', 'PSSE time step', 'PSSE acc. factor']
#        headers=['Case Nr.', 'Test Type', 'Fault Type', 'Fault time', ' Fault duration', 'Fault impedance', 'V residual', 'Fault X/R', 'POC voltage', 'Grid SCR', 'Grid X/R',  'P at POC (MW)', 'Q at POC (MVAr)', 'PSSE time step', 'PSSE acc. factor']
#        headers=['Case Nr.', 'Test Type', 'Fault Type', 'Fault time', ' Fault duration', 'Fault impedance', 'V residual', 'Fault X/R', 'POC voltage', 'Grid SCR', 'Grid X/R',  'P at POC (MW)', 'Q at POC (MVAr)', 'PSSE time step (ms)', 'PSSE acc. factor']
        dataset_number=0
        test_details={}
        while (dataset_number < len(datasets)) and (test_details=={}):
            try:
                test_details=shelve.open(main_folder_path_out+"\\"+datasets[dataset_number]['path']+"\\"+case+"\\testInfo\\"+case)
            except:
                dataset_number+=1
        row_cells=table.add_row().cells
        row_cells[0].text='Test Type'
        row_cells[1].text=str(test_details['scenario_params']['Test Type'])
        row_cells[2].text='Fault Type'
        row_cells[3].text=str(test_details['scenario_params']['Ftype'])
        row_cells=table.add_row().cells
#        row_cells[0].text='Grid Fault Level (MVA)'
        row_cells[0].text='Grid short circuit ratio'
#        row_cells[1].text=str(test_details['setpoint']['GridMVA'])
        row_cells[1].text=str(round(test_details['setpoint']['SCR'],2))
        row_cells[2].text='Grid X/R-ratio'
        row_cells[3].text=str(test_details['setpoint']['X_R'])
        row_cells=table.add_row().cells
        row_cells[0].text='P at POC (MW)'
        row_cells[1].text=str(test_details['setpoint']['P'])
        row_cells[2].text='Q at POC (MVAr)'
        row_cells[3].text=str(round(test_details['setpoint']['Q'],2))
        row_cells=table.add_row().cells
        row_cells[0].text='P from BESS (MW)'
        row_cells[1].text=str(test_details['setpoint']['P_PV2'])
        row_cells[2].text='POC voltage (p.u.)'
        row_cells[3].text=str(test_details['setpoint']['V_POC'])
        

        if(not 'Multifault' in test_details['scenario_params']['Test Type'] ):#single fault test
            row_cells=table.add_row().cells
            row_cells[0].text='Fault time (s)'
            row_cells[1].text=str(test_details['scenario_params']['Ftime'])
#            row_cells=table.add_row().cells
            row_cells[2].text='Fault duration (s)'
            row_cells[3].text=str(test_details['scenario_params']['Fduration'])
            row_cells=table.add_row().cells
            row_cells[0].text='Fault impedance (Ohm)'
            row_cells[1].text=str(round(test_details['scenario_params']['F_Impedance'],2))
            row_cells[2].text='Fault X/R-ratio'
            row_cells[3].text=str(test_details['scenario_params']['Fault X_R'])
            
#            row_cells=table.add_row().cells
#            row_cells[0].text='V residual (p.u.)'
#            if(test_details['scenario_params']['Vresidual']!=''):
#                row_cells[1].text=str(test_details['scenario_params']['Vresidual'])
#            else:
#                row_cells[1].text='-'
#            row_cells=table.add_row().cells

        else: #Multifault case
            if(len(test_details['scenario_params']['Ftype'])>0):
                tmp=str(test_details['scenario_params']['Ftype'][0])
#                for fault_id in range(0, len(test_details['scenario_params']['Ftype'])-1) :
                for fault_id in range(1, len(test_details['scenario_params']['Ftype'])) : # #11/8/2022: correcting the fault list
                    tmp+=', '+str(test_details['scenario_params']['Ftype'][fault_id])
                row_cells=table.add_row().cells
                row_cells[0].text='Fault Type'
                row_cells[1].text=tmp
            if(len(test_details['scenario_params']['Ftime'])>0):
                tmp=str(round(test_details['scenario_params']['Ftime'][0],3))
#                for fault_id in range(0, len(test_details['scenario_params']['Ftime'])-1) :
                for fault_id in range(1, len(test_details['scenario_params']['Ftime'])) :
                    tmp+=', '+str(round(test_details['scenario_params']['Ftime'][fault_id],3))
                row_cells[2].text='Fault time (s)'
                row_cells[3].text=tmp
            if(len(test_details['scenario_params']['Fduration'])>0):
                tmp=str(test_details['scenario_params']['Fduration'][0])
#                for fault_id in range(0, len(test_details['scenario_params']['Fduration'])-1) :
                for fault_id in range(1, len(test_details['scenario_params']['Fduration'])) :
                    tmp+=', '+str(test_details['scenario_params']['Fduration'][fault_id])
                row_cells=table.add_row().cells
                row_cells[0].text='Fault duration (s)'
                row_cells[1].text=tmp
            if(len(test_details['scenario_params']['F_Impedance'])>0):
                tmp=str(round(test_details['scenario_params']['F_Impedance'][0],2))
                for fault_id in range(0, len(test_details['scenario_params']['F_Impedance'])-1) :
#                for fault_id in range(1, len(test_details['scenario_params']['F_Impedance'])) :
                    tmp+=', '+str(round(test_details['scenario_params']['F_Impedance'][fault_id],2))
#                row_cells=table.add_row().cells
                row_cells[2].text='Fault impedance (Ohm)'
                row_cells[3].text=tmp
#            if(len(test_details['scenario_params']['Vresidual'])>0):
#                tmp=str(test_details['scenario_params']['Vresidual'][0])
##                for fault_id in range(0, len(test_details['scenario_params']['Vresidual'])-1) :
#                for fault_id in range(1, len(test_details['scenario_params']['Vresidual'])) :
#                    tmp+=', '+str(test_details['scenario_params']['Vresidual'][fault_id])
#                row_cells=table.add_row().cells
#                row_cells[0].text='V residual (p.u.)'
#                row_cells[1].text=tmp
    

        row_cells=table.add_row().cells
#        row_cells[0].text='Grid Fault Level post-fault (MVA)'
        row_cells[0].text='Grid short circuit ratio post-fault'
        if('SCL_post' in test_details['scenario_params']):
            if(test_details['scenario_params']['SCL_post']!=''):
                row_cells[1].text=str(float(test_details['scenario_params']['SCL_post'])/float(ProjectDetailsDict['PlantMW']))
#            else: row_cells[1].text=str(test_details['setpoint']['GridMVA'])
            else: row_cells[1].text=str(round(test_details['setpoint']['SCR'],2))
#        else: row_cells[1].text=str(test_details['setpoint']['GridMVA'])
#        else: row_cells[1].text=str(round(test_details['setpoint']['SCR'],2))
#        row_cells=table.add_row().cells
        row_cells[2].text='Grid X/R-ratio post-fault'
        if('X_R_post' in test_details['scenario_params']):
            if(test_details['scenario_params']['X_R_post']!=''): 
                row_cells[3].text=str(test_details['scenario_params']['X_R_post'])
            else:row_cells[3].text=str(test_details['setpoint']['X_R'])
#        else:row_cells[3].text=str(test_details['setpoint']['X_R'])


        if(software_type=='PSS/E'):
            row_cells=table.add_row().cells
            row_cells[0].text='PSSE time step'
            row_cells[1].text=str(test_details['scenario_params']['TimeStep'])
#            row_cells=table.add_row().cells
            row_cells[2].text='PSSE acc. factor'
            row_cells[3].text=str(test_details['scenario_params']['AccFactor'])
            
    if('ort') in case:
#        headers=['Case Nr.', 'Test Type', 'Time', 'Disturbance Frequency', 'Disturbance Magnitude', 'Phase Oscillation Magnitude', 'POC voltage', 'Grid MVA', 'Grid X/R',  'P at POC (MW)', 'Q at POC (MVAr)', 'PSSE time step', 'PSSE acc. factor']
#        headers=['Case Nr.', 'Test Type', 'Time', 'Disturbance Frequency', 'Disturbance Magnitude', 'Phase Oscillation Magnitude', 'POC voltage', 'Grid SCR', 'Grid X/R',  'P at POC (MW)', 'Q at POC (MVAr)', 'PSSE time step (ms)', 'PSSE acc. factor']
        dataset_number=0
        test_details={}
        while (dataset_number < len(datasets)) and (test_details=={}):
            try:
                test_details=shelve.open(main_folder_path_out+"\\"+datasets[dataset_number]['path']+"\\"+case+"\\testInfo\\"+case)
            except:
                dataset_number+=1
        row_cells=table.add_row().cells
        row_cells[0].text='Test Type'
        row_cells[1].text=str(test_details['scenario_params']['Test Type'])
        row_cells[2].text='Disturbance Frequency'
        row_cells[3].text=str(test_details['scenario_params']['Disturbance Frequency'])        
        row_cells=table.add_row().cells
        row_cells[0].text='Disturbance Magnitude'
        row_cells[1].text=str(round(test_details['scenario_params']['Disturbance Magnitude'],2))
        row_cells[2].text='Phase Oscillation Magnitude'
        row_cells[3].text=str(test_details['scenario_params']['PhaseOsc Magnitude'])
        row_cells=table.add_row().cells
#        row_cells[0].text='Grid Fault Level (MVA)'
        row_cells[0].text='Grid short circuit ratio'
#        row_cells[1].text=str(test_details['setpoint']['GridMVA'])
        row_cells[1].text=str(round(test_details['setpoint']['SCR'],2))
        row_cells[2].text='Grid X/R-ratio'
        row_cells[3].text=str(test_details['setpoint']['X_R'])
        row_cells=table.add_row().cells
        row_cells[0].text='P at POC (MW)'
        row_cells[1].text=str(test_details['setpoint']['P'])
        row_cells[2].text='Q at POC (MVAr)'
        row_cells[3].text=str(round(test_details['setpoint']['Q'],2))
        row_cells=table.add_row().cells
        row_cells[0].text='P from BESS (MW)'
        row_cells[1].text=str(test_details['setpoint']['P_PV2'])
        row_cells[2].text='POC voltage (p.u.)'
        row_cells[3].text=str(test_details['setpoint']['V_POC'])
        if(software_type=='PSS/E'):
            row_cells=table.add_row().cells
            row_cells[0].text='PSSE time step'
            row_cells[1].text=str(test_details['scenario_params']['TimeStep'])
            row_cells=table.add_row().cells
            row_cells[2].text='PSSE acc. factor'
            row_cells[3].text=str(test_details['scenario_params']['AccFactor'])
        
        
#        if('time' in test_details['scenario_params'].keys()):
#            row_cells=table.add_row().cells
#            row_cells[0].text='Fault time (s)'
#            row_cells[1].text=str(test_details['scenario_params']['time'])
#        row_cells=table.add_row().cells
#
#
#        row_cells=table.add_row().cells
#
#        row_cells=table.add_row().cells
#        row_cells[0].text='POC voltage (p.u.)'
#        row_cells[1].text=str(test_details['setpoint']['V_POC'])
#        row_cells=table.add_row().cells
##        row_cells[0].text='Grid Fault Level (MVA)'
#        row_cells[0].text='Grid short circuit ratio'
##        row_cells[1].text=str(test_details['setpoint']['GridMVA'])
#        row_cells[1].text=str(round(test_details['setpoint']['SCR'],2))
#        row_cells=table.add_row().cells
#        row_cells[0].text='Grid X/R-ratio'
#        row_cells[1].text=str(test_details['setpoint']['X_R'])
#        row_cells=table.add_row().cells
#        row_cells[0].text='P at POC (MW)'
#        row_cells[1].text=str(test_details['setpoint']['P'])
#        row_cells=table.add_row().cells
#        row_cells[0].text='Q at POC (MVAr)'
#        row_cells[1].text=str(test_details['setpoint']['Q'])

            
    if('tov') in case:
#        headers=['Case Nr.', 'Test Type', 'Time', 'TOV duration', 'Capacity(F)', 'V residual', 'POC voltage', 'Grid MVA', 'Grid X/R', 'Grid MVA post-TOV','Grid X/R post-TOV', 'P at POC (MW)', 'Q at POC (MVAr)', 'PSSE time step', 'PSSE acc. factor']
#        headers=['Case Nr.', 'Test Type', 'Time', 'TOV duration', 'Capacity(F)', 'V residual', 'POC voltage', 'Grid SCR', 'Grid X/R', 'Grid SCR post-TOV','Grid X/R post-TOV', 'P at POC (MW)', 'Q at POC (MVAr)', 'PSSE time step (ms)', 'PSSE acc. factor']
        dataset_number=0
        test_details={}
        while (dataset_number < len(datasets)) and (test_details=={}):
            try:
                test_details=shelve.open(main_folder_path_out+"\\"+datasets[dataset_number]['path']+"\\"+case+"\\testInfo\\"+case)
            except:
                dataset_number+=1
        row_cells=table.add_row().cells
        row_cells[0].text='Test Type'
        row_cells[1].text=str(test_details['scenario_params']['Test Type'])
        row_cells[2].text='V_tov (p.u.)'
        row_cells[3].text=str(round(test_details['scenario_params']['Vresidual'],2))
        row_cells=table.add_row().cells
#        row_cells[0].text='Grid Fault Level (MVA)'
        row_cells[0].text='Grid short circuit ratio'
#        row_cells[1].text=str(test_details['setpoint']['GridMVA'])
        row_cells[1].text=str(round(test_details['setpoint']['SCR'],2))
        row_cells[2].text='Grid X/R-ratio'
        row_cells[3].text=str(test_details['setpoint']['X_R'])
        row_cells=table.add_row().cells
        row_cells[0].text='P at POC (MW)'
        row_cells[1].text=str(test_details['setpoint']['P'])
        row_cells[2].text='Q at POC (MVAr)'
        row_cells[3].text=str(round(test_details['setpoint']['Q'],2))
        row_cells=table.add_row().cells
        row_cells[0].text='P from BESS (MW)'
        row_cells[1].text=str(test_details['setpoint']['P_PV2'])
        row_cells[2].text='POC voltage (p.u.)'
        row_cells[3].text=str(test_details['setpoint']['V_POC'])
        row_cells=table.add_row().cells
        row_cells[0].text='Time (s)'
        row_cells[1].text=str(test_details['scenario_params']['time'])
        row_cells[2].text='TOV duration (s)'
        row_cells[3].text=str(test_details['scenario_params']['Fduration'])

        if('Capacity(uF)' in test_details['scenario_params'].keys()):
            if(test_details['scenario_params']['Capacity(uF)']!=''):
                row_cells=table.add_row().cells
                row_cells[0].text='Capacity (uF)'
                row_cells[1].text=str(round(test_details['scenario_params']['Capacity(uF)'],2))
#                row_cells[1].text=str(round(1000000*test_details['scenario_params']['Capacity(F)'],2))

        row_cells=table.add_row().cells
        row_cells[0].text='Grid Fault Level post-TOV'
        if('SCL_post' in test_details['scenario_params']):
            if(test_details['scenario_params']['SCL_post']!=''):
                row_cells[1].text=str(test_details['scenario_params']['SCL_post'])
#            else: row_cells[1].text=str(test_details['setpoint']['GridMVA'])
            else: row_cells[1].text=str(round(test_details['setpoint']['SCR'],2))
#        else: row_cells[1].text=str(test_details['setpoint']['GridMVA'])
        else: row_cells[1].text=str(round(test_details['setpoint']['SCR'],2))
        row_cells[2].text='Grid X/R-ratio post-TOV'
        if('X_R_post' in test_details['scenario_params']):
            if(test_details['scenario_params']['X_R_post']!=''): 
                row_cells[3].text=str(test_details['scenario_params']['X_R_post'])
            else:row_cells[3].text=str(test_details['setpoint']['X_R'])
        else:row_cells[3].text=str(test_details['setpoint']['X_R'])
        
        if(software_type=='PSS/E'):
            row_cells=table.add_row().cells
            row_cells[0].text='PSSE time step'
            row_cells[1].text=str(test_details['scenario_params']['TimeStep'])
            row_cells=table.add_row().cells
            row_cells[2].text='PSSE acc. factor'
            row_cells[3].text=str(test_details['scenario_params']['AccFactor'])
            
#    table.set_fontsize(9)
#    table.scale(0.8, 0.8)
    #add comment generated from automated assessment of the results (optional)
    #add Graphs (with titles?)
    for plot_id in range(0, len(plot_list)): 
        if plot_list[plot_id] in plots.keys():
            report.add_heading(plot_list[plot_id], level=3 )
            report.add_picture(plots[plot_list[plot_id]], Inches(6)) #normal graphs smaller
    # add page break
    p=report.add_paragraph('')
    run=p.add_run()
    run.add_break(WD_BREAK.PAGE)     
    return 0

def add_plots_to_report_nw(case, report, datasets, plots, plot_list, assessment):
    software_type=''
    if(assessment['PSSE_flag']>0):
        software_type='PSS/E'
    #retrieve test info
    dataset_number=0
    test_details={}
    while (dataset_number < len(datasets)) and (test_details=={}):
        try:
            test_details=shelve.open(main_folder_path_out+"\\"+datasets[dataset_number]['path']+"\\"+case+"\\testInfo\\"+case)
        except:
            dataset_number+=1

  
    #add level 2 headline with test name
    paragraph_temp=report.add_heading('Case '+case, level=2 ) 
    add_bookmark(paragraph=paragraph_temp, bookmark_text='', bookmark_name=case )#add bookmark that link in table at start of document points to
    #add summary table with test details
    table=report.add_table(rows=1, cols=2)
    table.style='ListTable3-Accent3-OX2'   
    headers=['Parameter', 'Value']
    hdr_cells=table.rows[0].cells
    for header_id in range(0, len(headers)): hdr_cells[header_id].text=headers[header_id]    
    if('con') in case:
        dataset_number=0
        test_details={}
        while (dataset_number < len(datasets)) and (test_details=={}):
            try:
                test_details=shelve.open(main_folder_path_out+"\\"+datasets[dataset_number]['path']+"\\"+case+"\\testInfo\\"+case)
            except:
                dataset_number+=1
        if test_details['scenario_params']['Test Type'] in ['V_stp_profile']: # if the event is not a fault
            row_cells=table.add_row().cells
            row_cells[0].text='Test Type'
            row_cells[1].text=str(test_details['scenario_params']['Test Type'])
            row_cells=table.add_row().cells
            row_cells[0].text='Event Type'
            row_cells[1].text=str(test_details['scenario_params']['Event_Type'])
            row_cells=table.add_row().cells
            row_cells[0].text='Event Element'
            row_cells[1].text=str(test_details['scenario_params']['Event_Element'])
            row_cells=table.add_row().cells
            row_cells[0].text='Event Description'
            row_cells[1].text=str(test_details['scenario_params']['CaseDescription'])
            row_cells=table.add_row().cells
            row_cells[0].text='Test Profile'
            row_cells[1].text=str(test_details['scenario_params']['Test profile'])           
                                
        else:
            row_cells=table.add_row().cells
            row_cells[0].text='Test Type'
            row_cells[1].text=str(test_details['scenario_params']['Test Type'])
            row_cells=table.add_row().cells
            row_cells[0].text='Fault Type'
            row_cells[1].text=str(test_details['scenario_params']['Event_Type'])
            row_cells=table.add_row().cells
            row_cells[0].text='Fault Element'
            row_cells[1].text=str(test_details['scenario_params']['Event_Element'])
            row_cells=table.add_row().cells
            row_cells[0].text='Fault Description'
            row_cells[1].text=str(test_details['scenario_params']['CaseDescription'])
            row_cells=table.add_row().cells
            row_cells[0].text='Fault From Bus'
            row_cells[1].text=str(test_details['scenario_params']['i_bus'])           
            if test_details['scenario_params']['j_bus'] != '':
                row_cells=table.add_row().cells
                row_cells[0].text='Fault To Bus'
                row_cells[1].text=str(test_details['scenario_params']['j_bus'])   
    #        row_cells=table.add_row().cells
    #        row_cells[0].text='Fault Time'
    #        row_cells[1].text=str(test_details['scenario_params'][0]['Ftime'])  
            row_cells=table.add_row().cells
            row_cells[0].text='Local Clearing Time (s)'
            row_cells[1].text=str(test_details['scenario_params']['trip_near']) 
            if test_details['scenario_params']['trip_far'] != '':
                row_cells=table.add_row().cells
                row_cells[0].text='Remote Clearing Time (s)'
                row_cells[1].text=str(test_details['scenario_params']['trip_far'])         
            if test_details['scenario_params']['arc_success'] != '' and test_details['scenario_params']['arc_time'] > 0:
                row_cells=table.add_row().cells
                row_cells[0].text='Auto-reclose Time (s)'
                row_cells[1].text=str(test_details['scenario_params']['arc_time'])     
            if len(test_details['runback']) > 1:
                for rb_id in range(len(test_details['runback'])):
                    row_cells=table.add_row().cells
                    row_cells[0].text='Runback Action'
                    if test_details['runback'][rb_id]['j_bus'] != '':
                        row_cells[1].text=str(test_details['runback'][rb_id]['CaseDescription']) +' ('+ str(test_details['runback'][rb_id]['i_bus']) +'-'+ str(test_details['runback'][rb_id]['j_bus'])+')'
                    else:
                        row_cells[1].text=str(test_details['runback'][rb_id]['CaseDescription']) +' ('+ str(test_details['runback'][rb_id]['i_bus'])+')'

 
                    
    #add comment generated from automated assessment of the results (optional)
    #add Graphs (with titles?)
    for plot_id in range(0, len(plot_list)): 
        if plot_list[plot_id] in plots.keys():
            report.add_heading(plot_list[plot_id], level=3 )
            report.add_picture(plots[plot_list[plot_id]], Inches(6)) #normal graphs smaller
    # add page break
    p=report.add_paragraph('')
    run=p.add_run()
    run.add_break(WD_BREAK.PAGE)     
    return 0

    
def change_orientation(document):
    global landscape_flag
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    
    if(landscape_flag==0):
        landscape_flag = 1
    else:
        landscape_flag=0

    return new_section    

        
def DMAT_report(info, batchname, chapter_types):
    #checks which data is available and generates all defined plots for which all data is available. If 'cases' is not empty, the plots (and report) are generated only for the cases listed in 'cases'. All Plots for which only partial data is availabel are left out (e.g. an overlay plot for which only the PSS/E results data but no PSCAD results data is available would not be generated)
    #create output folder
    output_loc=main_folder_path_out+"\\plots\\DMAT"
    if(info[0]['report']==True):
        report=initialise_report('DMAT')
    report_intro = True
    for chapter_cnt in range(0,len(info)):
        chapter_info=info[chapter_cnt]
        if chapter_info['chapter'] in chapter_types or chapter_types == []: # if chapter_types is empty, then include all chapters into the report
            datasets=chapter_info['datasets']
            relevant_cases=chapter_info['cases']# make this more sopphisticated to allow for ranges of cases to be defined in a more elegant manner
            all_cases=[] #all cases for which any data is available (may just be from a single dataset and not across multiple)
            for dataset_id in range(0, len(datasets)):#create list of results included in each of the listed datasets. for each plot check if data specified in the plot is available in the datasets. Only produce plots for which all data is available, otherwise produce an error message.
                if os. path. exists(main_folder_path_out+"\\"+datasets[dataset_id]['path']): # only proceed if path is valid
                    cases_temp=next(os.walk(main_folder_path_out+"\\"+datasets[dataset_id]['path']))
                    cases_temp=cases_temp[1]
                    datasets[dataset_id]['cases']=sort_cases(cases_temp)
                    for case in datasets[dataset_id]['cases']:
                        if not case in all_cases:
                            all_cases.append(case)
                else:
                    print("Input dataset path is not correct. Please check again")
            cases=[]
            if(relevant_cases==[]):
                cases=sort_cases(all_cases)
            else:
                for case in all_cases:
                    if case in relevant_cases:
                        cases.append(case)
                cases=sort_cases(cases)

            if report_intro: # only include introduction once
                add_report_intro(report, 'DMAT', datasets , cases) #This should include summary table of the test cases the pass/no pass column can be left empty and filled in manually while reviewing the results. 
                report_intro = False
                
#            if(chapter_info['report']==True):
##                report=initialise_report('DMAT')
#                add_report_intro(report, 'DMAT', datasets , cases) #This should include summary table of the test cases the pass/no pass column can be left empty and filled in manually while reviewing the results. 
    
            if cases != []: # only add heading if data is available 
                report.add_heading("Simulation Results - "+chapter_info['chapter'], level=1 )   
            add_summary_table(report, 'DMAT', datasets, cases)
            summary = ""
            
            for case in cases:
                plots, assessment, summary = generate_plots(summary, case, output_loc, chapter_info, batchname+' - '+chapter_info['chapter']) #"info" contains dataset information as well
#                plots, assessment = generate_plots(case, output_loc, chapter_info, batchname) #"info" contains dataset information as well
                if(plots==-1):
                    print(assessment)
                if(chapter_info['report']==True):
                    add_plots_to_report(case, report, datasets, plots, chapter_info['plots_for_report'], assessment)
    if( chapter_info['report']==True):
       reportname= str(datetime.datetime.now().strftime("%Y%m%d-%H%M"))+"-"+str(ProjectDetailsDict['NameShrt'])+"_"+str(batchname)+"_report.docx"
       report.save(main_folder_path_out+"\\Plots\\DMAT\\"+reportname)
            
            
#Generate Benchmarking plots and report
def benchmarking_report(info, batchname, chapter_types):
    output_loc=main_folder_path_out+"\\plots\\BENCH"
    if(info[0]['report']==True):
        report=initialise_report('BENCHMARKING')
    report_intro = True
    for chapter_cnt in range(0,len(info)):
        chapter_info=info[chapter_cnt]
        if chapter_info['chapter'] in chapter_types or chapter_types == []: # if chapter_types is empty, then include all chapters into the report
            datasets=chapter_info['datasets']
            relevant_cases=chapter_info['cases']# make this more sopphisticated to allow for ranges of cases to be defined in a more elegant manner
            all_cases=[] #all cases for which any data is available (may just be from a single dataset and not across multiple)
    
            for dataset_id in range(0, len(datasets)):#create list of results included in each of the listed datasets. for each plot check if data specified in the plot is available in the datasets. Only produce plots for which all data is available, otherwise produce an error message.
                if os. path. exists(main_folder_path_out+"\\"+datasets[dataset_id]['path']): # only proceed if path is valid
                    cases_temp=next(os.walk(main_folder_path_out+"\\"+datasets[dataset_id]['path']))
                    cases_temp=cases_temp[1]
                    datasets[dataset_id]['cases']=sort_cases(cases_temp)
                    for case in datasets[dataset_id]['cases']:
                        if not case in all_cases:
                            all_cases.append(case)
                else:
                    print("Input dataset path is not correct. Please check again")
                
            cases=[]
            if(relevant_cases==[]):
                cases=sort_cases(all_cases)
            else:
                for case in all_cases:
                    if case in relevant_cases:
                        cases.append(case)
                cases=sort_cases(cases)
                
            if report_intro: # only include introduction once
                add_report_intro(report, 'GPS', datasets , cases) #This should include summary table of the test cases the pass/no pass column can be left empty and filled in manually while reviewing the results. 
                report_intro = False
                
            if cases != []: # only add heading if data is available 
                report.add_heading("Simulation Results - "+chapter_info['chapter'], level=1 ) 
                
            add_summary_table(report, 'GPS', datasets, cases)
            summary = ""
            for case in cases:
                plots, assessment, summary = generate_plots(summary, case, output_loc, chapter_info, batchname+' - '+chapter_info['chapter']) #"info" contains dataset information as well
                if(plots==-1):
                    print(assessment)
                if(chapter_info['report']==True):
                    add_plots_to_report(case, report, datasets, plots, chapter_info['plots_for_report'], assessment)
    if( chapter_info['report']==True):
       reportname= str(datetime.datetime.now().strftime("%Y%m%d-%H%M"))+"-"+str(ProjectDetailsDict['NameShrt'])+"_"+str(batchname)+"_report.docx"
       report.save(main_folder_path_out+"\\Plots\\BENCH\\"+reportname)
       
       
       
#    for chapter_cnt in range(0,len(info)):
#        chapter_info=info[chapter_cnt]
#        datasets=chapter_info['datasets']
#        relevant_cases=chapter_info['cases']# make this more sopphisticated to allow for ranges of cases to be defined in a more elegant manner
#        all_cases=[] #all cases for which any data is available (may just be from a single dataset and not across multiple)
#        overlap_cases=[]
#        for dataset_id in range(0, len(datasets)):#create list of results included in each of the listed datasets. for each plot check if data specified in the plot is available in the datasets. Only produce plots for which all data is available, otherwise produce an error message.
#            if os. path. exists(main_folder_path_out+"\\"+datasets[dataset_id]['path']): # only proceed if path is valid
#                datasets[dataset_id]['cases']=sort_cases(next(os.walk(main_folder_path_out+"\\"+datasets[dataset_id]['path']))[1])
#                for case in datasets[dataset_id]['cases']:
#                    if not case in all_cases:
#                        all_cases.append(case)
#            else:
#                print("Input dataset path is not correct. Please check again")
#            for case in all_cases:
#                n_datasets_including_case=0
#                for dataset_id in range(0, len(datasets)):
#                    if case in datasets[dataset_id]['cases']:
#                       n_datasets_including_case+=1
#                if(n_datasets_including_case>=1):
#                    overlap_cases.append(case)
#                    
#            if(relevant_cases==[]):
#                benchmark_cases=sort_cases(overlap_cases)
#            else:
#                benchmark_cases=[]
#                for case in all_cases:
#                    if case in relevant_cases:
#                        benchmark_cases.append(case)
#                benchmark_cases=sort_cases(benchmark_cases)
#            
#    
#            #create output folder
#            output_loc=main_folder_path_out+"\\plots\\Overlays"
#            if(chapter_info['report']==True):
#                report=initialise_report('BENCHMARKING')
#                add_report_intro(report, 'BENCHMARKING', datasets , benchmark_cases) #This should include summary table of the test cases the pass/no pass column can be left empty and filled in manually while reviewing the results. 
#            for case in benchmark_cases:
#                plots, assessment = generate_plots(case, output_loc, chapter_info, batchname, x_range='common') #"info" contains dataset information as well
#                if(plots==-1):
#                    print(assessment)
#                if(chapter_info['report']==True):
#                    add_plots_to_report(case, report, datasets, plots, chapter_info['plots_for_report'], assessment)
#    if(chapter_info['report']==True):       
#        reportname= str(datetime.datetime.now().strftime("%Y%m%d-%H%M"))+"-"+str(ProjectDetailsDict['NameShrt'])+"_benchmarking_report.docx" 
#        report.save(main_folder_path_out+"\\Plots\\Overlays\\"+reportname)

#Generates the partial GPS report, based on SMIB test data
def GPS_report(info, batchname, chapter_types):
    output_loc=main_folder_path_out+"\\plots\\GPS"
    if(info[0]['report']==True):
        report=initialise_report('GPS')
        
    report_intro = True
    for chapter_cnt in range(0,len(info)):
        chapter_info=info[chapter_cnt]
        if chapter_info['chapter'] in chapter_types or chapter_types == []: # if chapter_types is empty, then include all chapters into the report
            datasets=chapter_info['datasets']
            relevant_cases=chapter_info['cases']# make this more sopphisticated to allow for ranges of cases to be defined in a more elegant manner
            all_cases=[] #all cases for which any data is available (may just be from a single dataset and not across multiple)
    
            for dataset_id in range(0, len(datasets)):#create list of results included in each of the listed datasets. for each plot check if data specified in the plot is available in the datasets. Only produce plots for which all data is available, otherwise produce an error message.
                if os. path. exists(main_folder_path_out+"\\"+datasets[dataset_id]['path']): # only proceed if path is valid
                    cases_temp=next(os.walk(main_folder_path_out+"\\"+datasets[dataset_id]['path']))
                    cases_temp=cases_temp[1]
                    datasets[dataset_id]['cases']=sort_cases(cases_temp)
                    for case in datasets[dataset_id]['cases']:
                        if not case in all_cases:
                            all_cases.append(case)
                else:
                    print("Input dataset path is not correct. Please check again")
                
            cases=[]
            if(relevant_cases==[]):
                cases=sort_cases(all_cases)
            else:
                for case in all_cases:
                    if case in relevant_cases:
                        cases.append(case)
                cases=sort_cases(cases)
                
            if report_intro: # only include introduction once
                add_report_intro(report, 'GPS', datasets , cases) #This should include summary table of the test cases the pass/no pass column can be left empty and filled in manually while reviewing the results. 
                report_intro = False
            
            if cases != []: # only add introduction if data is available 
                report.add_heading("Simulation Results - "+chapter_info['chapter'], level=1 )   

            add_summary_table(report, 'GPS', datasets, cases)
            # store rise and settling time if needed
            filename = 'Chapter Summary '+chapter_info['chapter']+'.csv'#This needs to be adjusted to work for mutil-report setup
            csvfile = output_loc + '\\' + filename
            f = open(csvfile, 'w')
            f.write('Record rise, settling and recovery time when the markers are provided\n')
            summary = ""

            for case in cases:
                plots, assessment, summary = generate_plots(summary, case, output_loc, chapter_info, batchname+' - '+chapter_info['chapter']) #"info" contains dataset information as well
                if(plots==-1):
                    print(assessment)
                if(chapter_info['report']==True):
                    add_plots_to_report(case, report, datasets, plots, chapter_info['plots_for_report'], assessment)
            f.write(summary)
            f.close
    if( chapter_info['report']==True):
       reportname= str(datetime.datetime.now().strftime("%Y%m%d-%H%M"))+"-"+str(ProjectDetailsDict['NameShrt'])+"_"+str(batchname)+"_report.docx"
       report.save(main_folder_path_out+"\\Plots\\GPS\\"+reportname)

def NetworkEvent_report(info, batchname, chapter_types):
    #checks which data is available and generates all defined plots for which all data is available. If 'cases' is not empty, the plots (and report) are generated only for the cases listed in 'cases'. All Plots for which only partial data is availabel are left out (e.g. an overlay plot for which only the PSS/E results data but no PSCAD results data is available would not be generated)
    #create output folder
    output_loc=main_folder_path_out+"\\plots\\NetworkEvent"
    if(info[0]['report']==True):
        report=initialise_report('NetworkEvent')
    report_intro = True
    for chapter_cnt in range(0,len(info)):
        chapter_info=info[chapter_cnt]
        if chapter_info['chapter'] in chapter_types or chapter_types == []: # if chapter_types is empty, then include all chapters into the report
            datasets=chapter_info['datasets']
            relevant_cases=chapter_info['cases']# make this more sopphisticated to allow for ranges of cases to be defined in a more elegant manner
            all_cases=[] #all cases for which any data is available (may just be from a single dataset and not across multiple)
    
            for dataset_id in range(0, len(datasets)):#create list of results included in each of the listed datasets. for each plot check if data specified in the plot is available in the datasets. Only produce plots for which all data is available, otherwise produce an error message.
                if os. path. exists(main_folder_path_out+"\\"+datasets[dataset_id]['path']): # only proceed if path is valid
                    cases_temp=next(os.walk(main_folder_path_out+"\\"+datasets[dataset_id]['path']))
                    cases_temp=cases_temp[1]
                    datasets[dataset_id]['cases']=sort_cases(cases_temp)
                    for case in datasets[dataset_id]['cases']:
                        if not case in all_cases:
                            all_cases.append(case)
                else:
                    print("Input dataset path is not correct. Please check again")
                
            cases=[]
            if(relevant_cases==[]):
                cases=sort_cases(all_cases)
            else:
                for case in all_cases:
                    if case in relevant_cases:
                        cases.append(case)
                cases=sort_cases(cases)

            if report_intro: # only include introduction once
                add_report_intro(report, 'NetworkEvent', datasets , cases) #This should include summary table of the test cases the pass/no pass column can be left empty and filled in manually while reviewing the results. 
                report_intro = False
                
 
            if cases != []: # only add introduction if data is available 
                report.add_heading("Simulation Results - "+chapter_info['chapter'], level=1 )   
            
            add_summary_table(report, 'NetworkEvent', datasets, cases)
            summary = ""
            for case in cases:
#                plots, assessment = generate_plots(case, output_loc, chapter_info, batchname+' - '+chapter_info['chapter']) #"info" contains dataset information as well
                plots, assessment, summary = generate_plots(summary, case, output_loc, chapter_info, batchname+' - '+chapter_info['chapter']) #"info" contains dataset information as well
                if(plots==-1):
                    print(assessment)
                if(chapter_info['report']==True):
                    add_plots_to_report_nw(case, report, datasets, plots, chapter_info['plots_for_report'], assessment)
    if( chapter_info['report']==True):
       reportname= str(datetime.datetime.now().strftime("%Y%m%d-%H%M"))+"-"+str(ProjectDetailsDict['NameShrt'])+"_"+str(batchname)+"_report.docx"
       report.save(main_folder_path_out+"\\Plots\\NetworkEvent\\"+reportname)
            
def generate_plots(summary, case, output_loc, info, batchname, x_range='max'): #xrange can be 'max', 'common' or can be user-defined as a range
    #read all relevant datasets (possibly cross-check which ones are needed for specified plots, to reduce loading time
    assessment={}
    assessment['PSSE_flag']=0
    assessment['rise_t']={}
    plots={}
    plot = ep.ESCOPlot()
    dataset_names={}
    dataset_pos={} #this stores the position of the datasets within the report declaration in the header of this script (addressable by ID)
    dataset_pos_in_mem={} #this stores the position of the datasets within the plot object (addressable by ID)
    dataset_in_mem_cnt=0
    dataset_aux_cnt=0
    for dataset_id in range(0, len(info['datasets'])):
        dataset=info['datasets'][dataset_id]
        # check if data exists for the given case and if not, don't load.
        try:           
            plot.read_data(main_folder_path_out+"\\"+dataset['path']+"\\"+case+"\\"+case+'_results.csv', timeID=dataset['timeID'])
            dataset_pos_in_mem[info['datasets'][dataset_id]['ID']]=dataset_in_mem_cnt
            if('PSSE' in dataset['path']):
                assessment['PSSE_flag']=1
            dataset_in_mem_cnt+=1
            dataset_names[info['datasets'][dataset_id]['ID']]=dataset['label']
            dataset_pos[info['datasets'][dataset_id]['ID']]=dataset_id
            
    #        dataset_names[dataset_id]=dataset['label']
            plot.timeoffset[dataset_id]=dataset['timeoffset']
            
            if 'calcCurrents' in dataset.keys():
                    calcCurrents=dataset['calcCurrents']
                    for calcCurrentCnt in range (0, len(calcCurrents)):
                        calcCurrent=calcCurrents[calcCurrentCnt]
                        if('I' in calcCurrent.keys()):
                            currents=plot.calcCurrents(dataset_aux_cnt, P=calcCurrent['P'], Q=calcCurrent['Q'], I=calcCurrent['I'], nameLabel=calcCurrent['nameLabel'], scaling=calcCurrent['scaling'])
                        else:
                            currents=plot.calcCurrents(dataset_aux_cnt, P=calcCurrent['P'], Q=calcCurrent['Q'], V=calcCurrent['V'], nameLabel=calcCurrent['nameLabel'], scaling=calcCurrent['scaling'])

            if 'calPFs' in dataset.keys():
                    calPFs=dataset['calPFs']
                    for calPFCnt in range (0, len(calPFs)):
                        calPF=calPFs[calPFCnt]
                        PFs=plot.calPFs(dataset_aux_cnt, P=calPF['P'], Q=calPF['Q'], nameLabel=calPF['nameLabel'], scaling=calPF['scaling'])
            dataset_aux_cnt+=1
        except IOError:
            print(str(main_folder_path_out+"\\"+dataset['path']+"\\"+case+"\\"+case+'_results.csv')+' is not available.')
                        
        
    time_range=plot.check_min_max_time() #returns ---> by default limit x-axis of plots to this value. 
    common_range=[max(time_range[0]), min(time_range[1])]
    max_range=[min(time_range[0]), max(time_range[1])]
#    settling_time = {}
#    filename = 'Rise, settling and recovery times.csv'
#    csvfile = output_loc + '\\' + filename
#    f = open(csvfile, 'w')
##    f.write('Filename,Vsettle,Qsettle,Qrise,Vrise\n')
#    f.write('Record rise, settling and recovery time\n')
    
    for plot_name in info['plots'].keys():
        #determine datasets required for the plot and traces required in the datasets. Then send request to plot script so see if the signals are available in the dataset 
        #--> if not, generate error message and don't plot, Otherwise proceed to plot
        proceed=True
        for subplot in info['plots'][plot_name].keys():
            for channel in info['plots'][plot_name][subplot]['channels']:
#                if( (channel['dataset'] not in dataset_names.keys()) or not (case in info['datasets'][dataset_pos[channel['dataset']]]['cases']) ):
                if( (channel['dataset'] not in dataset_names.keys()) ):
                    proceed=False # "Data missing - Plots could not be generated."
        if(proceed):
            print(plot_name)
            test=1
            plot_xrange=[] #this is 
            for subplot_ID in range(0, len(info['plots'][plot_name].keys())):
                for subplot in info['plots'][plot_name].keys():
                    if(info['plots'][plot_name][subplot]['rank']-1==subplot_ID) :
                        subplot_info=info['plots'][plot_name][subplot]
                        print(subplot)
                        dataset_number=0
                        test_details={}
#                        test_details=shelve.open(main_folder_path_out+"\\"+info['datasets'][dataset_number]['path']+"\\"+case+"\\testInfo\\"+case)#take metadata from case of first dataset (shoudl be the same across all datasets for a given case)
                        while (dataset_number < len(info['datasets'])) and (test_details=={}):
                            try:
                                test_details=shelve.open(main_folder_path_out+"\\"+info['datasets'][dataset_number]['path']+"\\"+case+"\\testInfo\\"+case)#take metadata from case of first dataset (shoudl be the same across all datasets for a given case)
                            except:
                                dataset_number+=1
                        traces=info['plots'][plot_name][subplot]['channels']                        
                        subplot_legend=[]
                        for trace_ID in range(0, len(traces)):
                            print(subplot_info['channels'][trace_ID]['name'])
                            
                            if('linestyle' in subplot_info['channels'][trace_ID].keys() ):
                                linestyle=subplot_info['channels'][trace_ID]['linestyle']
                            else:
                                linestyle='-'
                            if('colour' in subplot_info['channels'][trace_ID].keys() ):
                                colour=subplot_info['channels'][trace_ID]['colour']
                            else:
                                colour=''
                            if('linewidth' in subplot_info['channels'][trace_ID].keys() ):
                                linewidth=subplot_info['channels'][trace_ID]['linewidth']
                            else:
                                linewidth=2.5
                            #check on which axis trace shouls be plotted
                            temp_subplot_properties={'markers':[], 'tolerance_band_offset':0.05, 'tolerance_band_base':-1}
                            if('twinX' in subplot_info['channels'][trace_ID].keys()):
                                plot.subplot_spec(subplot_ID, (dataset_pos_in_mem[subplot_info['channels'][trace_ID]['dataset']], subplot_info['channels'][trace_ID]['name']), title=subplot,y2label=subplot_info['unit2'], scale=subplot_info['channels'][trace_ID]['scale'], offset=subplot_info['channels'][trace_ID]['offset'], colour=colour, linestyle=linestyle, linewidth=linewidth, twinX=subplot_info['channels'][trace_ID]['twinX'])
                            else:
                                if('markers' in subplot_info['channels'][trace_ID]):
                                    if not ('show_markers' in subplot_info['channels'][trace_ID]): #unless showing the markers is suppressed, include them
                                        temp_subplot_properties['markers']=subplot_info['channels'][trace_ID]['markers']
                                    elif(subplot_info['channels'][trace_ID]['show_markers']!=False):#unless showing the markers is suppressed, include them
                                        temp_subplot_properties['markers']=subplot_info['channels'][trace_ID]['markers']
                                if('tolerance_bands' in subplot_info.keys()): #Only tolerance_bands is a special marker that is not included in the trace (channel) definition but instead in the subplot definition. Might change that in the future
                                    if (trace_ID==subplot_info['tolerance_bands']['trace']):
                                        #plot.subplot_spec(subplot_ID, (dataset_pos_in_mem[subplot_info['channels'][trace_ID]['dataset']], subplot_info['channels'][trace_ID]['name']), title=subplot,ylabel=subplot_info['unit'], scale=subplot_info['channels'][trace_ID]['scale'], offset=subplot_info['channels'][trace_ID]['offset'], colour=colour, linestyle=linestyle, linewidth=linewidth, markers=['tolerance_bands'],tolerance_band_offset=subplot_info['tolerance_bands']['percent']/100.0, tolerance_band_base=subplot_info['tolerance_bands']['base'])
                                        temp_subplot_properties['markers'].append('tolerance_bands')
                                        temp_subplot_properties['tolerance_band_offset']=subplot_info['tolerance_bands']['percent']/100.0
                                        temp_subplot_properties['tolerance_band_base']=subplot_info['tolerance_bands']['base']
                                plot.subplot_spec(subplot_ID, (dataset_pos_in_mem[subplot_info['channels'][trace_ID]['dataset']], subplot_info['channels'][trace_ID]['name']), title=subplot,ylabel=subplot_info['unit'], scale=subplot_info['channels'][trace_ID]['scale'], offset=subplot_info['channels'][trace_ID]['offset'], colour=colour, linestyle=linestyle, linewidth=linewidth, markers=temp_subplot_properties['markers'], tolerance_band_offset=temp_subplot_properties['tolerance_band_offset'], tolerance_band_base=temp_subplot_properties['tolerance_band_base']) 
                            #add Legend: If not specified: only add legend to first subplot only, using the label of the respective dataset (in case there are multiple datasets). Otherwise add no legend.
    #                        if( (trace_ID==0) and not ('leg' in (subplot_info['channels'][trace_ID].keys())) ):
    #                            subplot_legend.append(info['datasets'][subplot_info['channels'][trace_ID]['dataset']]['label']) #this should return label of dataset to which trace belongs and add it as lagend for the first trace.
                            #if legend is explicitly specified for a trace, add legend for that trace.
                            if('leg' in subplot_info['channels'][trace_ID].keys() ): 
                                if(subplot_info['channels'][trace_ID]['leg']!=''):
                                    subplot_legend.append(subplot_info['channels'][trace_ID]['leg'])
                        
                        for trace_ID in range(0, len(traces)):
                            dataset_nr=check_dataset_pos(subplot_info['channels'][trace_ID]['dataset'], info['datasets'])
                            dataset=info['datasets'][dataset_nr] #retrieves dataset position to which trace belongs for which rise-T to be calculated
                            timeoffset_0=dataset['timeoffset']
                            if('markers' in subplot_info['channels'][trace_ID]):#calculate the parameters per the markers per trace, no matter if the markers are shown or not. Can be suppressed with 'show_markers' variable
                                markers=subplot_info['channels'][trace_ID]['markers']
                                #include settling bands if marker is set                   
                                if('GSMG' in markers ): #only if marker is set and there is either a profile defined or a a fault test applied
                                    if (test_details['scenario_params']['Test Type']=='Fault'):
                                        starttime=test_details['scenario_params']['Ftime']+timeoffset_0 # starting at begining of fault
        #                                endtime=starttime+0.8*test_details['scenario_params']['Fduration']-0.05+timeoffset_0 #leave 5 ms buffer. use as "endTime" the last 20% of the time where the fault is applied
                                        endtime=starttime+0.9*test_details['scenario_params']['Fduration'] #use as "endTime" the last 10% of the time where the fault is applied
                                        startWindow=0.01
                                        endWindow=0.05*test_details['scenario_params']['Fduration']
                                    elif (test_details['scenario_params']['Test Type']=='TOV'):
                                        starttime=test_details['scenario_params']['time']+timeoffset_0 # starting at begining of fault
                                        endtime=starttime+0.9*test_details['scenario_params']['Fduration'] #use as "endTime" the last 10% of the time where the fault is applied
                                        startWindow=0.01
                                        endWindow=0.05*test_details['scenario_params']['Fduration']
        #                            elif (test_details['scenario_params']['Test Type']=='V_stp_profile'):
                                    else:
                                        starttime=5+timeoffset_0 # starting at 5 seconds
                                        endtime=-1 #If endtime = -1, consider all simulation period in finding the base for the error band
                                        startWindow=0.01
                                        endWindow=0.05 # Use 0.5sec for endWindow
                                    GSMG_error_bands = plot.GSMG_bands(dataset_nr, subplot_info['channels'][trace_ID]['name'], starttime=starttime, endtime=endtime, startWindow=startWindow, endWindow=endWindow)

                                if('rise_t' in markers ): #only if marker is set and there is either a profile defined or a a fault test applied
                                    starttime=5+timeoffset_0 # starting at 5 seconds
                                    endtime=starttime+7 #use 8 seconds after applying voltage changes for assessement
                                    if (test_details['scenario_params']['Test Type']=='Fault'):
                                        starttime=test_details['scenario_params']['Ftime']+timeoffset_0 # starting at begining of fault
                                        endtime=starttime+0.8*test_details['scenario_params']['Fduration'] #use as "endTime" the last 20% of the time where the fault is applied
                                    elif (test_details['scenario_params']['Test Type']=='TOV'):
                                        starttime=test_details['scenario_params']['time']+timeoffset_0 # starting at begining of fault
                                        endtime=starttime+0.8*test_details['scenario_params']['Fduration'] #use as "endTime" the last 20% of the time where the fault is applied
        #                            elif (test_details['scenario_params']['Test Type']=='V_stp_profile'):
                                    else:
                                        try: starttime=test_details['scenario_params']['Etime']+timeoffset_0 # starting at begining of fault,
                                        except: pass
                                        endtime=starttime+5
                                    risetime = plot.qrise(dataset_nr, subplot_info['channels'][trace_ID]['name'], starttime=starttime, endtime=endtime)
                                    # print " rise time = {:2.4f}".format(risetime)
#                                        f.write('Case ID: {},Test Type: {},Channel: {},Rising Time:{},{:1.3f}\n'.format(case, test_details['scenario_params']['Test Type'], subplot_info['channels'][trace_ID]['name'], "", risetime))
                                    summary += ('Case ID: {},Test Type: {},Channel: {},Rising Time:{},{:1.3f}\n'.format(case, test_details['scenario_params']['Test Type'], subplot_info['channels'][trace_ID]['name'], "", risetime))
                                    assessment['rise_t'][subplot_info['channels'][trace_ID]['name']]=risetime # for assessing rise time in next step
#                                    assessment['rise_t']=risetime
                                    
                                if('set_t' in markers): #only if marker is set and there is either a profile defined or a a fault test applied
                                    starttime=5+timeoffset_0 # starting at 5 seconds
                                    endtime=starttime+7 #use 8 seconds after applying voltage changes for assessement
                                    if (test_details['scenario_params']['Test Type']=='Fault'):
                                        starttime=test_details['scenario_params']['Ftime']+timeoffset_0 # starting at begining of fault
        #                                endtime=starttime+0.8*test_details['scenario_params']['Fduration']-0.05+timeoffset_0 #leave 5 ms buffer. use as "endTime" the last 20% of the time where the fault is applied
                                        endtime=starttime+0.9*test_details['scenario_params']['Fduration'] #use as "endTime" the last 10% of the time where the fault is applied
                                        startWindow=0.01
                                        endWindow=0.05*test_details['scenario_params']['Fduration']
                                    elif (test_details['scenario_params']['Test Type']=='TOV'):
                                        starttime=test_details['scenario_params']['time']+timeoffset_0 # starting at begining of fault
                                        endtime=starttime+0.9*test_details['scenario_params']['Fduration'] #use as "endTime" the last 10% of the time where the fault is applied
                                        startWindow=0.01
                                        endWindow=0.05*test_details['scenario_params']['Fduration']
                                    else:
                                        try: starttime=test_details['scenario_params']['Etime']+timeoffset_0 # starting at begining of fault
                                        except: pass
                                        endtime=starttime+7 #use as "endTime" the last 10% of the time where the fault is applied
                                        startWindow=0.01
                                        endWindow=0.05 # Use 0.5sec for endWindow
                                    tempsettime = plot.settleTime(dataset_nr, subplot_info['channels'][trace_ID]['name'], starttime=starttime, endtime=endtime, startWindow=startWindow, endWindow=endWindow)
                                    settletime = tempsettime - starttime
#                                        f.write('Case ID: {},Test Type: {},Channel: {},Settling Time:{},{:1.3f}\n'.format(case, test_details['scenario_params']['Test Type'], subplot_info['channels'][trace_ID]['name'], "", settletime))
                                    summary += ('Case ID: {},Test Type: {},Channel: {},Settling Time:{},{:1.3f}\n'.format(case, test_details['scenario_params']['Test Type'], subplot_info['channels'][trace_ID]['name'], "", settletime))


                                if('rec_t' in markers): #only if marker is set and there is either a profile defined or a a fault test applied
                                    starttime=5
                                    endtime=starttime+5
                                    if (test_details['scenario_params']['Test Type']=='Fault'):
                                        starttime=test_details['scenario_params']['Ftime']+timeoffset_0 # starting at beggining of fault
                                        endtime=test_details['scenario_params']['Ftime']+test_details['scenario_params']['Fduration']+timeoffset_0 #Disturbance ended
                                    elif (test_details['scenario_params']['Test Type']=='TOV'):
                                        starttime=test_details['scenario_params']['time']+timeoffset_0 # starting at beggining of fault
                                        endtime=test_details['scenario_params']['time']+test_details['scenario_params']['Fduration']+timeoffset_0 #Disturbance ended
                                    else:
                                        try: starttime=test_details['scenario_params']['Etime']+timeoffset_0 # starting at begining of fault
                                        except: pass
                                        endtime=starttime+5
                                    p_recovery = plot.prise(dataset_nr, subplot_info['channels'][trace_ID]['name'],100, 'U_POC1',1, distStartTime=starttime, distEndTime=endtime)
                                    p_recovery = p_recovery - endtime
#                                        f.write('Case ID: {},Test Type: {},Channel: {},p_recovery Time:{},{:1.3f}\n'.format(case, test_details['scenario_params']['Test Type'], subplot_info['channels'][trace_ID]['name'], "", p_recovery))
                                    summary += ('Case ID: {},Test Type: {},Channel: {},p_recovery Time:{},{:1.3f}\n'.format(case, test_details['scenario_params']['Test Type'], subplot_info['channels'][trace_ID]['name'], "", p_recovery))

        
                                if('dIq' in markers):
                                    starttime=5
                                    endtime=starttime+5
                                    if (test_details['scenario_params']['Test Type']=='Fault'):
                                        starttime=test_details['scenario_params']['Ftime']+timeoffset_0 # starting at begining of fault
                                        endtime=starttime+test_details['scenario_params']['Fduration'] #use as "endTime" at fault ended

                                    elif (test_details['scenario_params']['Test Type']=='TOV'):
                                        starttime=test_details['scenario_params']['time']+timeoffset_0 # starting at begining of fault
                                        endtime=starttime+test_details['scenario_params']['Fduration'] #use as "endTime" at fault ended
                                    else:
                                        try: starttime=test_details['scenario_params']['Etime']+timeoffset_0 # starting at begining of fault
                                        except: pass
                                        endtime=starttime+5
                                    delta_Iq = plot.deltaIq(dataset_nr, subplot_info['channels'][trace_ID]['name'], Vchan=-1, Iqbase=1,Vbase=1.0, distStartTime=starttime, distEndTime=endtime, endoffset=20)
#                                        print " delta_Iq = {:2.4f}".format(delta_Iq)
#                                        f.write('Case ID: {},Test Type: {},Channel: {},delta_Iq:{},{:1.3f}\n'.format(case, test_details['scenario_params']['Test Type'], subplot_info['channels'][trace_ID]['name'], "", delta_Iq))
                                    summary += ('Case ID: {},Test Type: {},Channel: {},delta_Iq:{},{:1.3f}\n'.format(case, test_details['scenario_params']['Test Type'], subplot_info['channels'][trace_ID]['name'], "", delta_Iq))
                                   
        
                                    
                                if('dV' in markers):
                                    HV_calc_threshold=1.20
                                    LV_calc_threshold=0.8
                                    starttime=5
                                    endtime=starttime+5
                                    if (test_details['scenario_params']['Test Type']=='Fault'):
                                        starttime=test_details['scenario_params']['Ftime']+timeoffset_0 # starting at begining of fault
                                        endtime=starttime+test_details['scenario_params']['Fduration'] #use as "endTime" at fault ended
                                    elif (test_details['scenario_params']['Test Type']=='TOV'):
                                        starttime=test_details['scenario_params']['time']+timeoffset_0 # starting at begining of fault
                                        endtime=starttime+test_details['scenario_params']['Fduration'] #use as "endTime" at fault ended
                                    else:
                                        try: starttime=test_details['scenario_params']['Etime']+timeoffset_0 # starting at begining of fault
                                        except: pass
                                        endtime=starttime+5
                                    status, delta_V, V_fault = plot.deltaV(dataset_nr, subplot_info['channels'][trace_ID]['name'], distStartTime=starttime, distEndTime=endtime, HV_calc_threshold=HV_calc_threshold, LV_calc_threshold=LV_calc_threshold, endoffset=20)
                                    # print " delta_V = {:2.4f}".format(delta_V)
#                                        f.write('Case ID: {},Test Type: {},Channel: {},delta_V:{},{:1.3f}\n'.format(case, test_details['scenario_params']['Test Type'], subplot_info['channels'][trace_ID]['name'], "", delta_V))
#                                        f.write('Case ID: {},Test Type: {},Channel: {},V_fault:{},{:1.3f}\n'.format(case, test_details['scenario_params']['Test Type'], subplot_info['channels'][trace_ID]['name'], "", V_fault))
                                    summary += ('Case ID: {},Test Type: {},Channel: {},delta_V:{},{:1.3f}\n'.format(case, test_details['scenario_params']['Test Type'], subplot_info['channels'][trace_ID]['name'], "", delta_V))
                                    summary += ('Case ID: {},Test Type: {},Channel: {},V_fault:{},{:1.3f}\n'.format(case, test_details['scenario_params']['Test Type'], subplot_info['channels'][trace_ID]['name'], "", V_fault))


                                if('callout' in markers):
                                    callout_times = [1.5, 5.5] #time stamp on the graph would like to call the data out (by default at 1.5 and 5.5secs)
                                    if type(markers[-1]) != str:
                                        callout_times = markers[-1]
                                    dstamp = plot.calloutd(dataset_nr, subplot_info['channels'][trace_ID]['name'], callout_times=callout_times) #interpolate the callout values
                                    for cID in range(len(callout_times)):
                                        summary += ('Case ID: {},Test Type: {},Channel: {},Callout Time:{},{:1.3f}\n'.format(case, test_details['scenario_params']['Test Type'], subplot_info['channels'][trace_ID]['name'], "", callout_times[cID]))
                                        summary += ('Case ID: {},Test Type: {},Channel: {},Callout Value:{},{:1.3f}\n'.format(case, test_details['scenario_params']['Test Type'], subplot_info['channels'][trace_ID]['name'], "", dstamp[cID]))


#                            plot.prise(self, entry, pchan, pmax=100, vchan=-1, vbase = 1.0, distStartTime=0, distEndTime=-1)
#                            GSMG_bands(self, entry, channel, starttime = 0.0, endtime = 10.0, startWindow = 1.0, endWindow = 5.0):
#                                plot.prise(subplot_info['rec_t'], subplot_info['channels'][subplot_info['rec_t']]['name'],100, subplot_info['channels']['U_POC1'],1, distStartTime=starttime, distEndTime=endtime)
#                                plot.GSMG_bands(subplot_info['GSMG'], subplot_info['channels'][subplot_info['GSMG']]['name'], starttime=starttime, endtime=endtime, startWindow=startWindow, endWindow=endWindow)
                            
#                        if('set_band' in subplot_info.keys() ): #only if marker is set and there is either a profile defined or a a fault test applied
                            
    #                    if('dV' in subplot_info.keys() ):
    #                        
    #                    if('dIq' in subplot_info.keys() ):
                                    
                        # if plot contains overlays, include labels of datasets --> include legend using the label of the datasets for that subplot
                        legendAll=''
                        for i in range(0, len(subplot_legend)  ):
                            legendAll+=subplot_legend[i]
                        if(subplot_ID==0):
                            if (legendAll=='') and len(subplot_info['channels'])>1:
                                if(subplot_info['channels'][0]['dataset'] != subplot_info['channels'][1]['dataset']): #if first two traces belong to different datasets it is likely that it is an overlay plot
                                    for i in range(0, len(subplot_info['channels'])):
                                        subplot_legend.append(dataset_names[subplot_info['channels'][i]['dataset']])
                                
                        plot.legends[subplot_ID] = subplot_legend#try and set legend per subplot if possible. If not possible, modify EscoPlot lib to be able to do that.
                        if('xrange' in subplot_info.keys()):
                            plot_xrange=subplot_info['xrange']
                        if('yrange' in subplot_info.keys()):
                            plot.ylimit[subplot_ID]=subplot_info['yrange']
                        elif('yminspan' in subplot_info.keys()):
                            plot.yspan[subplot_ID]=subplot_info['yminspan']
                        if('y2range' in subplot_info.keys()):
                            plot.y2limit[subplot_ID]=subplot_info['y2range'] 
                        elif('y2minspan' in subplot_info.keys()):
                            plot.y2span[subplot_ID]=subplot_info['y2minspan']
                        if('ymaxlim' in subplot_info.keys()):
                            plot.ymaxlim[subplot_ID]=subplot_info['ymaxlim']
                        if('yminlim' in subplot_info.keys()):
                            plot.yminlim[subplot_ID]=subplot_info['yminlim']
            output_folder=output_loc+"\\"+batchname
            try:
                os.mkdir(output_folder)
            except:
                print("plot folder already exists")
            else:
                print("plot directory created")
                
            if(plot_xrange!=[]):
                plot.xlimit=plot_xrange
            else:
                if(x_range=='max'):
                    plot.xlimit=max_range
                elif(x_range=='common'):
                    plot.xlimit=common_range
            
            imgdata=plot.plot(figname = output_folder+"\\"+case+'_'+plot_name, show = 0, legloc = 'best')#, markers={'settleTime':['Qout_MV (MW)']} ) #legloc likely =legend location
            plots[plot_name]=imgdata
            plot.clear_subplot_spec()
            plot.clear_ylabels()
            plot.clear_ylimits()
#    f.close()

    return plots, assessment, summary
            
        
        
    #iterate over all plot definitions
        #iterale over all subplot definitions
        #generate plots and metadata, add lables and legen, etc.
        #reset subplot definition
    
    

def generate_plot(case, output_loc, result_data, info, plot_type):
    if(plot_type=='BENCHMARKING'): #expecting result_data to contain two file paths for the two overlays
        assessment={}        
        plots={}        
        for plot_name in info['overlay_plots'].keys():
            print(plot_name)
            plot = ep.ESCOPlot()            
            dataset_names={}
            for i in [0,1]:
                for dataset in info['datasets'].keys():
                    if(info['datasets'][dataset]['ID']==i):
                        dataset_names[i]=dataset
                        plot.read_data(result_data[i], timeID=info['datasets'][dataset]['timeID'])    
                        plot.timeoffset[i]=info['datasets'][dataset]['timeoffset']
            test=1
            for subplot_id in range(0, len(info['overlay_plots'][plot_name].keys()) ) : 
                for subplot in info['overlay_plots'][plot_name].keys():
                    if(info['overlay_plots'][plot_name][subplot]['rank']-1==subplot_id):
                        subplot_info=info['overlay_plots'][plot_name][subplot]
                        print(subplot)                
                        duration_0=plot.subplot_spec(subplot_id, (0, subplot_info['channels'][0]['name']), title = subplot,  ylabel = subplot_info['unit'], scale=subplot_info['channels'][0]['scale'], offset = subplot_info['channels'][0]['offset'], markers=['GSMG'])
                        duration_1=plot.subplot_spec(subplot_id, (1, subplot_info['channels'][1]['name']), title = subplot,  ylabel = subplot_info['unit'], scale=subplot_info['channels'][1]['scale'], offset = subplot_info['channels'][1]['offset'], markers=['GSMG'])
                        #read test scenario metadata
                        test_details=shelve.open(os.path.dirname(result_data[0])+"\\testInfo\\"+case)
                        timeoffset_0=info['datasets'][dataset]['timeoffset']
                        if('GSMG' in subplot_info.keys() ): #only if marker is set and there is either a profile defined or a a fault test applied
                            if ('Test profile' in test_details['scenario_params'].keys()): #if scenario contains a test profile
                                test_profile=ProfilesDict[test_details['scenario_params']['Test profile']]
                                #test_profile=test_details['Test profile']
                                steps=detect_steps(subplot_info['channels'][subplot_info['GSMG']]['offset'], {'x':test_profile['x_data'], 'y':test_profile['y_data']})#whatever test profile is given 
                                if(len(steps)==1):
#                                    endtime=#last entry in profile -0.105
                                    starttime=steps[0][0]-0.105+timeoffset_0
                                    endtime=test_profile['x_data'][-1]-0.105+timeoffset_0 #endtime is end of scenario -105 ms
                                    plot.GSMG_bands(subplot_info['GSMG'], subplot_info['channels'][subplot_info['GSMG']]['name'], starttime=starttime, endtime=endtime, startWindow=0.1, endWindow=0.1)
                                elif(len(steps)>1):
                                    starttime=steps[0][0]-0.105+timeoffset_0
                                    endtime=steps[1][0]-0.105+timeoffset_0 #endtime right before next step - 105 ms
                                    plot.GSMG_bands(subplot_info['GSMG'], subplot_info['channels'][subplot_info['GSMG']]['name'], starttime=starttime, endtime=endtime, startWindow=0.1, endWindow=0.1)
                            elif (test_details['scenario_params']['Test Type']=='Fault'):
                                starttime=test_details['scenario_params']['Ftime']-0.105+timeoffset_0
                                endtime=starttime+0.8*test_details['scenario_params']['Fduration']-0.05+timeoffset_0 #leave 5 ms buffer. use as "endTime" the last 20% of the time where the fault is applied
                                startWindow=0.1
                                endWindow=0.2*test_details['scenario_params']['Fduration']
                                plot.GSMG_bands(subplot_info['GSMG'], subplot_info['channels'][subplot_info['GSMG']]['name'], starttime=starttime, endtime=endtime, startWindow=startWindow, endWindow=endWindow)
                                pass
                        common_range=max([duration_0, duration_1])
                        plot.xlimit=[0,common_range]        
                                
                            
            plot.legends[0] = [dataset_names[0], dataset_names[1]]  
                
            #plot.xlimit=[0,10] #cut x-axis at the point where the shorter dataset ends.
            
            #plot.ylimit = [[0.8,1.2],[], [],[]] #scale y-axis automatically, with option to scale manually in plot_info
            #modify plot routine to return hander that can be used to copy plot into word report.
            # create output folder
            if('batchname' in info.keys()):
                if(info['batchname']!=''):
                    output_folder=output_loc+"\\"+info['batchname']
            else:
                output_folder=output_loc+"\\"+dataset_names[0]+'_vs_'+dataset_names[1]
            try:
                os.mkdir(output_folder)
            except:
                print("plot folder already exists")
            else:
                print("plot directory created")
                        
            imgdata=plot.plot(figname = output_folder+"\\"+case+'_'+plot_name, show = 1, legloc = 'best')#, markers={'settleTime':['Qout_MV (MW)']} ) #legloc likely =legend location
            plots[plot_name]=imgdata        
        return plots, assessment
    
    if(plot_type=='MAT'):
        assessment={}
        plots={}     
        for plot_name in info['plots']:
            print(plot_name)
            plot = ep.ESCOPlot()
            dataset_names={}
            dataset=info['datasets'].keys()[0]
            plot.read_data(result_data, timeID=info['datasets'][dataset]['timeID'])
            plot.timeoffset[0]=info['datasets'][dataset]['timeoffset']
        
            plot_xrange=[]
            subplots=info['plots'][plot_name].keys()
            for subplot_ID in range(0, len(subplots) ):
                for subplot in subplots:
                    if(info['plots'][plot_name][subplot]['rank']==subplot_ID+1):
                        subplot_info=info['plots'][plot_name][subplot]
                        traces=info['plots'][plot_name][subplot]['channels']
                        
                        subplot_legend=[]
                        for trace_ID in range(0, len(traces)):
                            print(subplot_info['channels'][trace_ID]['name'])
                            plot.subplot_spec(subplot_ID, (0, subplot_info['channels'][trace_ID]['name']), title=subplot,ylabel=subplot_info['unit'], scale=subplot_info['channels'][trace_ID]['scale'], offset=subplot_info['channels'][trace_ID]['offset'])
                            subplot_legend.append(subplot_info['channels'][trace_ID]['leg'])
                        plot.legends[subplot_ID] = subplot_legend#try and set legend per subplot if possible. If not possible, modify EscoPlot lib to be able to do that.
                        if('xrange' in subplot_info.keys()):
                            plot_xrange=subplot_info['xrange']
            if('batchname' in info.keys()):
                if(info['batchname']!=''):
                    output_folder=output_loc+"\\"+info['batchname']
                else:
                    output_folder=output_loc+"\\"+info['datasets'].keys()[0] #just one dataset included for MAT
            try:
                os.mkdir(output_folder)
            except:
                print("plot folder already exists")
            else:
                print("plot directory created")
                
            if(plot_xrange!=[]):
                plot.xlimit=plot_xrange
                            
            imgdata=plot.plot(figname = output_folder+"\\"+case+'_'+plot_name, show = 1, legloc = 'best')#, markers={'settleTime':['Qout_MV (MW)']} ) #legloc likely =legend location
            plots[plot_name]=imgdata
        return plots, assessment
                 

def main():
#    if('BENCH' in reports.keys()):
#        benchmarking_report(reports['BENCH']['report_definition'], reports['BENCH']['batchname'])
#    if('MAT' in reports.keys()):
#        for batch in reports['MAT'].keys():
#            reports['MAT'][batch]['batchname']=batch
#            MAT_report(reports['MAT'][batch], batch)
#    if('DMAT' in reports.keys()):
    if 'BENCH' in report_types:
        benchmarking_report(reports['BENCH']['report_definition'], reports['BENCH']['batchname'], chapter_types)
    
    if 'DMAT' in report_types:
        DMAT_report(reports['DMAT']['report_definition'], reports['DMAT']['batchname'], chapter_types)

    if 'GPS' in report_types:
        GPS_report(reports['GPS']['report_definition'], reports['GPS']['batchname'], chapter_types)
        
#    elif('NetworkEvent' in reports.keys()):
    if 'NetworkEvent' in report_types:
        NetworkEvent_report(reports['NetworkEvent']['report_definition'], reports['NetworkEvent']['batchname'], chapter_types)
        
    pass        
        
        
if __name__ == "__main__":
    main()
        






