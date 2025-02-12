# -*- coding: utf-8 -*-
"""
Created on Mon Nov  9 15:49:16 2020

@author: Mervin Kall

SMIB PLOT tool:

Script allows to generate plots of test results generated with the SMIB test tools, The plots are defined in the variable "reports"

plot range: for Benchmarking xrange is by default from 0.0 to the point where the first signal ends (e.g. if PSS/E simulation ran 5 seconds longer, those last 5 seconds will not be included in plot)

    
Ideas:
    For the benchmarking plots and also 5.2.5.13 and 5.2.5.5 settling time checks, the simulation time at which the step is applied, as well as the time at which the simulation reaches steady state again are required
    For any test involving a profile of any kind, this script could detect from the test metadata where the steps are located in the profile and take this into account for plotting settling bands, tolerance bands etc.

    Settling (GSMG) always based on first step --> try and limit tests to only one step (and back) per scenario when using GSMG bands
        --> detect first step from profile gradient: 
                set startwindow and endwiondow to 100 ms
                set starttime to >100ms before step
                set endtime to >100ms before either second step or end of scenario (whichever comes first)
                                
    include test profile in metadata, because the test profile may change in the config sheet after the test batch has been run
    
    
Versions
    V3.4: update the chanel plot for LSF
    #11/8/2022: correcting the fault list in MFRT test: remove the repeaation of first fault
    #V3.4.b: included option to export the rise, settle and recovery time: 'rise_t':0, 'set_t':0, 'rec_t':0
    #V3.4.c: Include DMAT id in the summary table
    #V3.4.d: include Generator data channels
    #V3.4.e: only consider one INV data as they are idendical
    #V4_LL08: update the summary table
    
"""

#------------------------------------------------------------------------------
# USER INPUT
#------------------------------------------------------------------------------
#main_folder_path=r'C:\work\202O'
#main_folder_path_out = r'C:\work\20220803_LSF_I2_1_DDupdate_Vspt_CUO'20803_LSF_I2_1_DDupdate_Vspt_CU
TestDefinitionSheet=r'20230828_SUM_TESTINFO_V1.xlsx'
"""
# report_types = ["BENCH", #--> Will expect plots of type overlay. Will only include plots for test cases that are availabel in two or more of the specified datasets (to be able to create overlay)
                    "DMAT", #--> Will plot everything that is available, Either one or two sets of result data can be provided. Should include both single dataset plots and overlays. may contain different chapters (to allow for plots to change depending on test type)
                    "GPS"] #--> Various types of chapters and plots. 
# chapter_types = ["S5.2.5.5_inj", "S5.2.5.1"]
    """

reports = {                                                                             
#            'DMAT':{'batchname': 'Typical Cases', #in DMAT
#            'DMAT':{'batchname': 'S5255', #'DMAT', #in DMAT
            'NetworkFault':{'batchname': 'HighLoad_Overlay', #'HighLoad1', 'HighLoad2', 'HighLoad3', 'LowLoad1','LowLoad2'
#            'DMAT':{'batchname': 'GPS Clauses', #in GPS Compliance Report
                    'report_definition':[          {'chapter':'general', #Array of chapters. In this case there is only one chapter.
                                                    'datasets': [
                                                                     {'label':'PSSE_DMAT', 'path':r"PSSE_sim\result_data\dynamic_network\20230914-1219_S5255_NW\HL_dyn_genon", 'ID': 0, 'timeID':'Time(s)', 'timeoffset':-2.0, },
                                                                     {'label':'PSSE_DMAT', 'path':r"PSSE_sim\result_data\dynamic_network\20230914-1219_S5255_NW\HL_dyn_genoff", 'ID': 1, 'timeID':'Time(s)', 'timeoffset':-2.0,  },

#                                                                     {'label':'PSSE_DMAT', 'path':r"PSSE_sim\result_data\dynamic_network\20230914-1219_S5255_NW\LL_dyn_genon", 'ID': 0, 'timeID':'Time(s)', 'timeoffset':-2.0, },
#                                                                     {'label':'PSSE_DMAT', 'path':r"PSSE_sim\result_data\dynamic_network\20230914-1219_S5255_NW\LL_dyn_genoff", 'ID': 1, 'timeID':'Time(s)', 'timeoffset':-2.0,  },


#                                                                     {'label':'PSSE_DMAT', 'path':r"PSSE_sim\result_data\dynamic_network\20230523-2234_NetworkFault_final\HighLoad_dyn", 'ID': 0, 'timeID':'Time(s)', 'timeoffset':-2.0, 
#                                                                      'calcCurrents':[{"P":"P_POC1", "Q":"Q_POC1", "V":"U_POC1", "nameLabel":"PLANT", "scaling":0.010334129, }, # 1/Sbase_POC for Iq at POC or 1/Sbase_INV for Iq at INV -> convert QMVAr to Qpu for calculation
#                                                                                      {"P":"P_LV1", "Q":"Q_LV1", "V":"U_LV1", "nameLabel":"PV", "scaling":0.009157509, },
#                                                                                      {"P":"P_LV2", "Q":"Q_LV2", "V":"U_LV2", "nameLabel":"BESS", "scaling":0.007716049, },]  },
#                                                                     {'label':'PSSE_DMAT', 'path':r"PSSE_sim\result_data\dynamic_network\20230523-2234_NetworkFault_final\HighLoad_dyn_org", 'ID': 1, 'timeID':'Time(s)', 'timeoffset':-2.0, 
#                                                                      'calcCurrents':[{"P":"P_POC1", "Q":"Q_POC1", "V":"U_POC1", "nameLabel":"PLANT", "scaling":0.010334129, }, # 1/Sbase_POC for Iq at POC or 1/Sbase_INV for Iq at INV -> convert QMVAr to Qpu for calculation
#                                                                                      {"P":"P_LV1", "Q":"Q_LV1", "V":"U_LV1", "nameLabel":"PV", "scaling":0.009157509, },
#                                                                                      {"P":"P_LV2", "Q":"Q_LV2", "V":"U_LV2", "nameLabel":"BESS", "scaling":0.007716049, },]  },
#                                                                     {'label':'PSCAD_DMAT', 'path':r"PSCAD_sim\result_data\dynamic_smib\20220513_DMAT3", 'ID': 4, 'timeID':'time(s)', 'timeoffset':-3.0},
#                                                                     {'label':'PSCAD_DMAT', 'path':r"PSCAD_sim\result_data\dynamic_smib\20220519_DMAT", 'ID': 4, 'timeID':'time(s)', 'timeoffset':-3.0, 'calcCurrents':[
#                                                                             {"P":"PPOC", "Q":"QPOC", "V":"VPOC", "nameLabel":"PLANT", "scaling":0.010532968, }, ]},
#                                                                    {'label':'PSSE_INV1', 'path':r"PSSE_sim\result_data\dynamic_studies\20220421_DMAT_2", 'ID': 0, 'timeID':'Time(s)', 'timeoffset':0.0}, #settling bands are added relating to the first dataset
                                                                    # {'label':'PSSE_INV2', 'path':r"PSSE_sim\result_data\20220329_DMAT", 'ID': 1, 'timeID':'Time(s)', 'timeoffset':0.0},
                                                                    # {'label':'PSSE_POC', 'path':r"PSSE_sim\result_data\20220329_DMAT", 'ID': 2, 'timeID':'Time(s)', 'timeoffset':0.0},
#                                                                    {'label':'PSCAD_INIT', 'path':r"PSCAD_sim\result_data\20211111_INIT3", 'ID': 0, 'timeID':'time(s)', 'timeoffset':0.0, },
                                                                #   {'label':'PSCAD_new_meter', 'path':r"PSCAD_sim\result_data\20220127_f_m_last_test_1", 'ID': 0, 'timeID':'time(s)', 'timeoffset':-4.0, },
                                                                #   {'label':'PSCAD_old_meter', 'path':r"PSCAD_sim\result_data\20220127_f_m_last_test_prev1", 'ID': 1, 'timeID':'time(s)', 'timeoffset':-4.0,},  
                                                                  #{'label':'PSSE_BM_new', 'path':r"PSSE_sim\result_data\20211207_PSSE_msm_4200UP", 'ID': 2, 'timeID':'Time(s)', 'timeoffset':0.0, 'calcCurrents':[
#                                                                            {"P":"P_POC1", "Q":"Q_POC1", "V":"U_POC1", "nameLabel":"PLANT", "scaling":1.0, }, ]},
#                                                                  {'label':'PSSE_BM_prev', 'path':r"PSSE_sim\result_data\20211207_PSSE_msm_2750EV", 'ID': 3, 'timeID':'Time(s)', 'timeoffset':0.0, 'calcCurrents':[
#                                                                            {"P":"P_POC1", "Q":"Q_POC1", "V":"U_POC1", "nameLabel":"PLANT", "scaling":1.0, }, ]},                                                                 
                                                                  #{'label':'PSCAD_DMAT_new', 'path':r"PSCAD_sim\result_data\20220121_f_lim_test3", 'ID': 4, 'timeID':'time(s)', 'timeoffset':-4.0, },
#                                                                  {'label':'PSSE_DMAT_new', 'path':r"PSSE_sim\result_data\20211217_addendum4", 'ID': 5, 'timeID':'Time(s)', 'timeoffset':-4.0, 'calcCurrents':[
#                                                                            {"P":"P_POC1", "Q":"Q_POC1", "V":"U_POC1", "nameLabel":"PLANT", "scaling":1.0, }, ]},

                                                                     
                                                                ],
                                                    'cases':[], #if empty, all cases are considered
#                                                    'cases':['con39'], #if empty, all cases are considered
#                                                    'cases':['small130','small131','small132','small133','small134','small135','small136','small137','small138','small139'],
#                                                    'cases':['con1'],
#                                                    'cases':['small65','small66','small67','small68','small69'],
#                                                    'cases':['ort9'],
                                                    # 'plots_for_report': ['PSCAD overlays'], #list pecifying which of the below plots shoudl be added to the report. Allows to omit plots without deleting the definition.
#                                                   'plots_for_report': [ 'Flat Run test PSCAD', 'Init test PSCAD'], #list pecifying which of the below plots shoudl be added to the report. Allows to omit plots without deleting the definition.
                                                    # 'plots_for_report': ['PSSE simple INV1','PSSE simple INV2','PSSE simple POC'], #list pecifying which of the below plots shoudl be added to the report. Allows to omit plots without deleting the definition.
#                                                    'plots_for_report': ['PSCAD simple','PSSE simple','Overlays(POC)'], #list pecifying which of the below plots shoudl be added to the report. Allows to omit plots without deleting the definition.
                                                     'plots_for_report': ['Overlays(NEM)'],
#                                                     'plots_for_report': ['Overlays(NEM)'],
                                                    
                                                    'summary_items': {}, #Additional description/analysis of the result data can bebe specified here. For example amount of current injection for 5.2.5.5 analysis in table format. Other items can be defined in the future.
                                                    'report':True,
                                                    'plots':{
#                                                             
#                                                            'PSCAD simple':                {
#                                                                                              'Voltage POC':                    {'channels':[{'dataset':4, 'name':"VPOC", 'leg':'Voltage HV', 'offset':0.0, 'scale':1.0}, 
#                                                                                                                                          {'dataset':4, 'name':"Vref_POC (pu)", 'leg':'Voltage stp', 'offset':0.0, 'scale':1.0, 'colour':'grey', 'linestyle':'--', 'linewidth': 1.5},  ], 'unit':'p.u.', 'rank':1, 'ymaxlim':1.3},
#                                                                                              'Active Power POC':               {'channels':[{'dataset':4, 'name':"PPOC", 'leg':'Active Power HV', 'offset':0.0, 'scale':1.0}, 
#                                                                                                                                          {'dataset':4, 'name':"Pref_POC (%)", 'leg':'Active Power stp', 'offset':0.0, 'scale':0.8, 'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}, ], 'unit':'MW', 'rank':3, 'yminlim':-200.0, 'ymaxlim':200.0}, #add active power setpoint here
#                                                                                              'Reactive Power POC':             {'channels':[{'dataset':4, 'name':"QPOC", 'leg':'Reactive Power HV', 'offset':0.0, 'scale':1.0}, 
#                                                                                                                                          {'dataset':4, 'name':"Qref_POC (%)", 'leg':'Reactive Power stp', 'offset':0.0, 'scale':0.8, 'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}, ], 'unit':'MVAr', 'rank':5, 'yminlim':-200.0, 'ymaxlim':200.0}, #add reactive power setpoint here
#                                                                                              'Reactive Current POC':           {'channels':[{'dataset':4, 'name':"Iq_PLANT", 'leg':'Reactive Current HV', 'offset':0.0, 'scale':1.0},], 'unit':'p.u', 'rank':7, 'yminspan':0.1, 'ymaxlim':2.0, 'yminlim':-2.0 },
##                                                                                              'Frequency':                      {'channels':[{'dataset':4, 'name':"fpoc", 'leg':'Frequency', 'offset':0.0, 'scale':1.0},], 'unit':'Hz', 'rank':7, 'yminspan':5.0},
##                                                                                             'Angle':                           {'channels':[{'dataset':4, 'name':"phAngSource", 'leg':'Angle at POC', 'offset':0.0, 'scale':1.0},], 'unit':'degrees', 'rank':7, 'yminspan':10.0},
#
#                                                                                             'Voltage INV':                     {'channels':[{'dataset':4, 'name':"Vrms_LV_inv1_pu", 'leg':'Voltage LV', 'offset':0.0, 'scale':1.0},], 'unit':'p.u.', 'rank':2, 'yminspan':0.1},
#                                                                                             'Active Power INV':                {'channels':[{'dataset':4, 'name':"Pout_LV (MW)", 'leg':'P_INV1', 'offset':0.0, 'scale':1.0},], 'unit':'MW', 'rank':4, 'yminspan':5, 'ymaxlim':200.0, 'yminlim':-200.0, },
#                                                                                             'Reactive Power INV':              {'channels':[{'dataset':4, 'name':"Qout_LV (MVar)", 'leg':'Q_INV1', 'offset':0.0, 'scale':1.0},], 'unit':'MVAr', 'rank':6, 'yminspan':5, 'ymaxlim':200.0, 'yminlim':-200.0, },
#                                                                                             'Inverter State (FRT)':            {'channels':[{'dataset':4, 'name':"LVRT_ON", 'leg':'Inverter LVRT state', 'offset':0.0, 'scale':1.0},], 'unit':'inverter code', 'rank':8},
#
#                                                                                             
#                                                                                              },
#                                                            'PSSE simple':                {
#                                                                                            'Voltage POC':                      {'channels':[{'dataset':0, 'name':"V_POC", 'leg':'Voltage HV', 'offset':0.0, 'scale':1.0},], 'unit':'p.u.', 'rank':1, 'yminspan':0.1, 'ymaxlim':1.3},
#                                                                                            'Active Power POC':                 {'channels':[{'dataset':0, 'name':"P_POC", 'leg':'P - POC', 'offset':0.0, 'scale':1.0},], 'unit':'MW', 'rank':3, 'yminspan':10, 'rec_t':0},
#                                                                                            'Reactive Power POC':               {'channels':[{'dataset':0, 'name':"Q_POC", 'leg':'Q - POC', 'offset':0.0, 'scale':1.0},], 'unit':'MVAr', 'rank':5, 'yminspan':5},
#                                                                                            'Reactive Current POC':             {'channels':[{'dataset':0, 'name':"Iq_PLANT", 'leg':'Reactive Current HV', 'offset':0.0, 'scale':1.0},], 'unit':'p.u', 'rank':7, 'yminspan':0.1, 'ymaxlim':2.0, 'yminlim':-2.0, 'rise_t':0, 'set_t':0},
##                                                                                            'Frequency':                        {'channels':[{'dataset':0, 'name':"F_POC1", 'leg':'Frequency', 'offset':50.0, 'scale':50.0},], 'unit':'Hz', 'rank':7, 'yminspan':5.0},
##                                                                                            'Angle':                           {'channels':[{'dataset':0, 'name':"ANG_POC1", 'leg':'Angle at POC', 'offset':0.0, 'scale':1.0},], 'unit':'MVAr', 'rank':7, 'yminspan':10.0},
#
#                                                                                                                                                
#                                                                                            'Voltage INV':                      {'channels':[{'dataset':0, 'name':"V_INV1", 'leg':'Voltage INV1', 'offset':0.0, 'scale':1.0},], 'unit':'p.u.', 'rank':2, 'yminspan':0.1},                                                                                                                                                
#                                                                                            'Active Power INV':                 {'channels':[{'dataset':0, 'name':"P_INV1", 'leg':'P_INV1', 'offset':0.0, 'scale':1.0},], 'unit':'MW', 'rank':4, 'yminspan':5, 'ymaxlim':200.0, 'yminlim':-200.0, },
#                                                                                            'Reactive Power INV':               {'channels':[{'dataset':0, 'name':"Q_INV1", 'leg':'Q_INV1', 'offset':0.0, 'scale':1.0},], 'unit':'MVAr', 'rank':6, 'yminspan':5, 'ymaxlim':200.0, 'yminlim':-200.0, },
#                                                                                            'Inverter State (FRT)':             {'channels':[{'dataset':0, 'name':"INV1_VDFLAG", 'leg':'INV1 state', 'offset':0.0, 'scale':1.0},], 'unit':'inverter code', 'rank':8},
#                                                                                            
#                                                                                            },

#                                                            'PSSE Network':                {
#                                                                                            'Voltage':                      {'channels':[{'dataset':0, 'name':"U_POC1", 'leg':'Voltage POC', 'offset':0.0, 'scale':1.0},
#                                                                                                                                        {'dataset':0, 'name':"U_LV1", 'leg':'Voltage INV', 'offset':0.0, 'scale':1.0},], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
#                                                                                            'Active Power POC':                 {'channels':[{'dataset':0, 'name':"P_POC1", 'leg':'P - POC', 'offset':0.0, 'scale':1.0},
#                                                                                                                                        {'dataset':0, 'name':"P_LV1", 'leg':'P - INV', 'offset':0.0, 'scale':1.0},], 'unit':'MW', 'rank':2, 'yminspan':10},
#                                                                                            'Reactive Power POC':               {'channels':[{'dataset':0, 'name':"Q_POC1", 'leg':'Q - POC', 'offset':0.0, 'scale':1.0},
#                                                                                                                                        {'dataset':0, 'name':"Q_LV1", 'leg':'Q - INV', 'offset':0.0, 'scale':1.0},], 'unit':'MVAr', 'rank':3, 'yminspan':5},
##                                                                                            'Reactive Current POC':             {'channels':[{'dataset':0, 'name':"Iq_PLANT", 'leg':'Reactive Current HV', 'offset':0.0, 'scale':1.0},], 'unit':'p.u', 'rank':4, 'yminspan':0.1, 'ymaxlim':2.0, 'yminlim':-2.0},
##                                                                                            'Inverter State (FRT)':             {'channels':[{'dataset':0, 'name':"INV1_VDFLAG", 'leg':'INV1 state', 'offset':0.0, 'scale':1.0},], 'unit':'inverter code', 'rank':5},
#
#                                                                                            'Inverter State (FRT)':             {'channels':[{'dataset':0, 'name':"INV1_VDFLAG", 'leg':'INV1 state', 'offset':0.0, 'scale':1.0},], 'unit':'inverter code', 'rank':4},
#                                                                                            'Reactive Current INV':             {'channels':[{'dataset':0, 'name':"INV1_IQ", 'leg':'Reactive Current LV', 'offset':0.0, 'scale':1.0},], 'unit':'p.u', 'rank':5, 'yminspan':0.1, 'ymaxlim':2.0, 'yminlim':-2.0},
#
## 
#                                                                                            },
                                                                                                                                             
#                                                            'PSSE Network':                {
#                                                                                            'Voltage POC':                      {'channels':[{'dataset':0, 'name':"U_POC1", 'leg':'Voltage HV', 'offset':0.0, 'scale':1.0},], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
#                                                                                            'Active Power POC':                 {'channels':[{'dataset':0, 'name':"P_POC1", 'leg':'P - POC', 'offset':0.0, 'scale':1.0},], 'unit':'MW', 'rank':3, 'yminspan':10},
#                                                                                            'Reactive Power POC':               {'channels':[{'dataset':0, 'name':"Q_POC1", 'leg':'Q - POC', 'offset':0.0, 'scale':1.0},], 'unit':'MVAr', 'rank':5, 'yminspan':5},
##                                                                                            'Reactive Current POC':             {'channels':[{'dataset':0, 'name':"Iq_PLANT", 'leg':'Reactive Current HV', 'offset':0.0, 'scale':1.0},], 'unit':'p.u', 'rank':7, 'yminspan':0.1, 'ymaxlim':2.0, 'yminlim':-2.0},
##                                                                                            'Frequency':                        {'channels':[{'dataset':0, 'name':"F_POC1", 'leg':'Frequency', 'offset':50.0, 'scale':50.0},], 'unit':'Hz', 'rank':7, 'yminspan':5.0},
##                                                                                            'Angle':                           {'channels':[{'dataset':0, 'name':"ANG_POC1", 'leg':'Angle at POC', 'offset':0.0, 'scale':1.0},], 'unit':'MVAr', 'rank':7, 'yminspan':10.0},
#                                                                                            'Inverter State (FRT)':             {'channels':[{'dataset':0, 'name':"INV1_VDFLAG", 'leg':'INV1 state', 'offset':0.0, 'scale':1.0},], 'unit':'inverter code', 'rank':7},
#
#                                                                                                                                                
#                                                                                            'Voltage INV':                      {'channels':[{'dataset':0, 'name':"U_LV1", 'leg':'Voltage INV1', 'offset':0.0, 'scale':1.0},], 'unit':'p.u.', 'rank':2, 'yminspan':0.1},                                                                                                                                                
#                                                                                            'Active Power INV':                 {'channels':[{'dataset':0, 'name':"P_LV1", 'leg':'P_INV1', 'offset':0.0, 'scale':1.0},], 'unit':'MW', 'rank':4, 'yminspan':5, 'ymaxlim':200.0, 'yminlim':-200.0, },
#                                                                                            'Reactive Power INV':               {'channels':[{'dataset':0, 'name':"Q_LV1", 'leg':'Q_INV1', 'offset':0.0, 'scale':1.0},], 'unit':'MVAr', 'rank':6, 'yminspan':5, 'ymaxlim':200.0, 'yminlim':-200.0, },
##                                                                                            'Inverter Terminal Current':        {'channels':[{'dataset':0, 'name':"INV_ITOT", 'leg':'INV1 Itot', 'offset':0.0, 'scale':1.0},], 'unit':'p.u.', 'rank':8},
##                                                                                            'Inverter State (FRT)':             {'channels':[{'dataset':0, 'name':"INV1_VDFLAG", 'leg':'INV1 state', 'offset':0.0, 'scale':1.0},], 'unit':'inverter code', 'rank':8},
##                                                                                            'Frequency':                        {'channels':[{'dataset':0, 'name':"F_POC1", 'leg':'Frequency', 'offset':50.0, 'scale':50.0},], 'unit':'Hz', 'rank':8, 'yminspan':5.0},
#                                                                                            'Reactive Current INV':             {'channels':[{'dataset':0, 'name':"INV1_IQ", 'leg':'Reactive Current LV', 'offset':0.0, 'scale':1.0},], 'unit':'p.u', 'rank':8, 'yminspan':0.1, 'ymaxlim':2.0, 'yminlim':-2.0},
#                                                                                            
#                                                                                            },

#                                                            'Overlays(NEM)':             {
#                                                                                            'Lismore 330kV bus voltage':   {'channels':[{'dataset':1, 'name':"V_LSM_330", 'leg':'pre-SUMSF', 'offset':0.0, 'scale':1.0}, 
#                                                                                                                                     {'dataset':0, 'name':"V_LSM_330", 'leg':'post-SUMSF', 'offset':0.0, 'scale':1.0}], 'unit':'p.u.', 'rank':1, 'yminspan':0.1, 'ymaxlim':1.4, 'xmaxlim':3},#'GSMG':0, }, #first channel in array links to file in first result location
#                                                                                            'Coffs Harbour 330kV bus voltage':   {'channels':[{'dataset':1, 'name':"V_COFF_330", 'leg':'pre-LSF', 'offset':0.0, 'scale':1.0}, 
#                                                                                                                                     {'dataset':0, 'name':"V_COFF_330", 'leg':'post-LSF', 'offset':0.0, 'scale':1.0}], 'unit':'p.u.', 'rank':2, 'yminspan':0.1, 'ymaxlim':1.4},#'GSMG':0, }, #first channel in array links to file in first result location
#                                                                                            'Lismore 132kV bus voltage':   {'channels':[{'dataset':1, 'name':"V_LSM_132", 'leg':'pre-LSF', 'offset':0.0, 'scale':1.0}, 
#                                                                                                                                     {'dataset':0, 'name':"V_LSM_132", 'leg':'post-LSF', 'offset':0.0, 'scale':1.0}], 'unit':'p.u.', 'rank':3, 'yminspan':0.1, 'ymaxlim':1.4},#'GSMG':0, }, #first channel in array links to file in first result location
#                                                                                            'Koolkhan 132kV bus voltage':   {'channels':[{'dataset':1, 'name':"V_KOLK_132", 'leg':'pre-LSF', 'offset':0.0, 'scale':1.0}, 
#                                                                                                                                     {'dataset':0, 'name':"V_KOLK_132", 'leg':'post-LSF', 'offset':0.0, 'scale':1.0}], 'unit':'p.u.', 'rank':4, 'yminspan':0.1, 'ymaxlim':1.4},#'GSMG':0, }, #first channel in array links to file in first result location                                                                                                                                       
#                                                                                                                                       
#                                                                                         },
                                                                                                                                    


                                                            'Overlays(NEM)':             {
                                                                                            'Lismore 330kV bus voltage':   {'channels':[{'dataset':1, 'name':"V_LSM_330", 'leg':'pre-SUMSF', 'offset':0.0, 'scale':1.0}, 
                                                                                                                                     {'dataset':0, 'name':"V_LSM_330", 'leg':'post-SUMSF', 'offset':0.0, 'scale':1.0}], 'unit':'p.u.', 'rank':1, 'yminspan':0.1, 'ymaxlim':1.4, 'xmaxlim':3},#'GSMG':0, }, #first channel in array links to file in first result location
                                                                                            'Coffs Harbour 330kV bus voltage':   {'channels':[{'dataset':1, 'name':"V_COFF_330", 'leg':'pre-SUMSF', 'offset':0.0, 'scale':1.0}, 
                                                                                                                                     {'dataset':0, 'name':"V_COFF_330", 'leg':'post-SUMSF', 'offset':0.0, 'scale':1.0}], 'unit':'p.u.', 'rank':2, 'yminspan':0.1, 'ymaxlim':1.4, 'xmaxlim':3},#'GSMG':0, }, #first channel in array links to file in first result location
                                                                                            'Lismore 132kV bus voltage':   {'channels':[{'dataset':1, 'name':"V_LSM_132", 'leg':'pre-SUMSF', 'offset':0.0, 'scale':1.0}, 
                                                                                                                                     {'dataset':0, 'name':"V_LSM_132", 'leg':'post-SUMSF', 'offset':0.0, 'scale':1.0}], 'unit':'p.u.', 'rank':3, 'yminspan':0.1, 'ymaxlim':1.4},#'GSMG':0, }, #first channel in array links to file in first result location
                                                                                            'Koolkhan 132kV bus voltage':   {'channels':[{'dataset':1, 'name':"V_KOLK_132", 'leg':'pre-SUMSF', 'offset':0.0, 'scale':1.0}, 
                                                                                                                                     {'dataset':0, 'name':"V_KOLK_132", 'leg':'post-SUMSF', 'offset':0.0, 'scale':1.0}], 'unit':'p.u.', 'rank':4, 'yminspan':0.1, 'ymaxlim':1.4},#'GSMG':0, }, #first channel in array links to file in first result location                                                                                                                                       
#                                                                                            'SUM_POC 132kV bus voltage':   {'channels':[{'dataset':1, 'name':"V_SUM_POC", 'leg':'pre-SUMSF', 'offset':0.0, 'scale':1.0}, 
#                                                                                                                                     {'dataset':0, 'name':"V_SUM_POC", 'leg':'post-SUMSF', 'offset':0.0, 'scale':1.0}], 'unit':'p.u.', 'rank':4, 'yminspan':0.1, 'ymaxlim':1.4},#'GSMG':0, }, #first channel in array links to file in first result location                                                                                                                                       
                                                                                         },

#                                                                                                                                             
#                                                            'PSSE simple':                {
#                                                                                            'Voltage POC':                      {'channels':[{'dataset':0, 'name':"U_POC1", 'leg':'Voltage HV', 'offset':0.0, 'scale':1.0},
#                                                                                                                                        {'dataset':0, 'name':"VREF_POC", 'leg':'Voltage stp', 'offset':0.0, 'scale':1.0, 'colour':'grey', 'linestyle':'--', 'linewidth': 1.5},], 'unit':'p.u.', 'rank':1, 'yminspan':0.1},
#                                                                                            'Active Power POC':                 {'channels':[{'dataset':0, 'name':"P_POC1", 'leg':'P - POC', 'offset':0.0, 'scale':1.0},
#                                                                                                                                        {'dataset':0, 'name':"PREF_POC", 'leg':'Active Power stp', 'offset':0.0, 'scale':94.94, 'colour':'grey', 'linestyle':'--', 'linewidth': 1.5},], 'unit':'MW', 'rank':3, 'yminspan':10},
#                                                                                            'Reactive Power POC':               {'channels':[{'dataset':0, 'name':"Q_POC1", 'leg':'Q - POC', 'offset':0.0, 'scale':1.0},
#                                                                                                                                        {'dataset':0, 'name':"QREF_POC", 'leg':'Reactive Power stp', 'offset':0.0, 'scale':94.94, 'colour':'grey', 'linestyle':'--', 'linewidth': 1.5},], 'unit':'MVAr', 'rank':5, 'yminspan':5},
#                                                                                            'Reactive Current POC':             {'channels':[{'dataset':0, 'name':"Iq_PLANT", 'leg':'Reactive Current HV', 'offset':0.0, 'scale':1.0},], 'unit':'p.u', 'rank':7, 'yminspan':0.1, 'ymaxlim':2.0, 'yminlim':-2.0},
##                                                                                            'Frequency':                        {'channels':[{'dataset':0, 'name':"F_POC1", 'leg':'Frequency', 'offset':50.0, 'scale':50.0},], 'unit':'Hz', 'rank':7, 'yminspan':5.0},
##                                                                                            'Angle':                           {'channels':[{'dataset':0, 'name':"ANG_POC1", 'leg':'Angle at POC', 'offset':0.0, 'scale':1.0},], 'unit':'MVAr', 'rank':7, 'yminspan':10.0},
#
#                                                                                                                                                
#                                                                                            'Voltage INV':                      {'channels':[{'dataset':0, 'name':"U_LV1", 'leg':'Voltage INV1', 'offset':0.0, 'scale':1.0},], 'unit':'p.u.', 'rank':2, 'yminspan':0.1},                                                                                                                                                
#                                                                                            'Active Power INV':                 {'channels':[{'dataset':0, 'name':"P_LV1", 'leg':'P_INV1', 'offset':0.0, 'scale':1.0},], 'unit':'MW', 'rank':4, 'yminspan':5, 'ymaxlim':200.0, 'yminlim':-200.0, },
#                                                                                            'Reactive Power INV':               {'channels':[{'dataset':0, 'name':"Q_LV1", 'leg':'Q_INV1', 'offset':0.0, 'scale':1.0},], 'unit':'MVAr', 'rank':6, 'yminspan':5, 'ymaxlim':200.0, 'yminlim':-200.0, },
##                                                                                            'Inverter Terminal Current':        {'channels':[{'dataset':0, 'name':"INV_ITOT", 'leg':'INV1 Itot', 'offset':0.0, 'scale':1.0},], 'unit':'p.u.', 'rank':8},
#                                                                                            'Inverter State (FRT)':             {'channels':[{'dataset':0, 'name':"INV1_VDFLAG", 'leg':'INV1 state', 'offset':0.0, 'scale':1.0},], 'unit':'inverter code', 'rank':8},
##                                                                                            'Frequency':                        {'channels':[{'dataset':0, 'name':"F_POC1", 'leg':'Frequency', 'offset':50.0, 'scale':50.0},], 'unit':'Hz', 'rank':8, 'yminspan':5.0},
#                                                                                            
#                                                                                            },
                                                                                                                                             
#                                                            'PSSE simple':                {
#                                                                                            'Voltage POC':                      {'channels':[{'dataset':0, 'name':"V_POC", 'leg':'Voltage HV', 'offset':0.0, 'scale':1.0},
#                                                                                                                                        {'dataset':0, 'name':"VREF_POC", 'leg':'Voltage stp', 'offset':0.0, 'scale':1.0, 'colour':'grey', 'linestyle':'--', 'linewidth': 1.5},], 'unit':'p.u.', 'rank':1, 'yminspan':0.1, 'ymaxlim':1.3},
#                                                                                            'Active Power POC':                 {'channels':[{'dataset':0, 'name':"P_POC", 'leg':'P - POC', 'offset':0.0, 'scale':1.0},
#                                                                                                                                        {'dataset':0, 'name':"PREF_POC", 'leg':'Active Power stp', 'offset':0.0, 'scale':94.94, 'colour':'grey', 'linestyle':'--', 'linewidth': 1.5},], 'unit':'MW', 'rank':3, 'yminspan':10},
#                                                                                            'Reactive Power POC':               {'channels':[{'dataset':0, 'name':"Q_POC", 'leg':'Q - POC', 'offset':0.0, 'scale':1.0},
#                                                                                                                                        {'dataset':0, 'name':"QREF_POC", 'leg':'Reactive Power stp', 'offset':0.0, 'scale':94.94, 'colour':'grey', 'linestyle':'--', 'linewidth': 1.5},], 'unit':'MVAr', 'rank':5, 'yminspan':5},
#                                                                                            'Reactive Current POC':             {'channels':[{'dataset':0, 'name':"Iq_PLANT", 'leg':'Reactive Current HV', 'offset':0.0, 'scale':1.0},], 'unit':'p.u', 'rank':7, 'yminspan':0.1, 'ymaxlim':2.0, 'yminlim':-2.0},
##                                                                                            'Frequency':                        {'channels':[{'dataset':0, 'name':"F_POC1", 'leg':'Frequency', 'offset':50.0, 'scale':50.0},], 'unit':'Hz', 'rank':7, 'yminspan':5.0},
##                                                                                            'Angle':                           {'channels':[{'dataset':0, 'name':"ANG_POC1", 'leg':'Angle at POC', 'offset':0.0, 'scale':1.0},], 'unit':'MVAr', 'rank':7, 'yminspan':10.0},
#
#                                                                                                                                                
#                                                                                            'Voltage INV':                      {'channels':[{'dataset':0, 'name':"V_INV1", 'leg':'Voltage INV1', 'offset':0.0, 'scale':1.0},], 'unit':'p.u.', 'rank':2, 'yminspan':0.1},                                                                                                                                                
#                                                                                            'Active Power INV':                 {'channels':[{'dataset':0, 'name':"P_INV1", 'leg':'P_INV1', 'offset':0.0, 'scale':1.0},], 'unit':'MW', 'rank':4, 'yminspan':5, 'ymaxlim':200.0, 'yminlim':-200.0, },
#                                                                                            'Reactive Power INV':               {'channels':[{'dataset':0, 'name':"Q_INV1", 'leg':'Q_INV1', 'offset':0.0, 'scale':1.0},], 'unit':'MVAr', 'rank':6, 'yminspan':5, 'ymaxlim':200.0, 'yminlim':-200.0, },
#                                                                                            'Inverter State (FRT)':             {'channels':[{'dataset':0, 'name':"INV1_VDFLAG", 'leg':'INV1 state', 'offset':0.0, 'scale':1.0},], 'unit':'inverter code', 'rank':8},
#                                                                                            
#                                                                                            },


##                                                            
#                                                            'Overlays':             {
#                                                                                            'Voltage POC':                  {'channels':[{'dataset':4, 'name':"VPOC", 'leg':'PSCAD', 'offset':0.0, 'scale':1.0}, 
#                                                                                                                                     {'dataset':0, 'name':"U_POC1", 'leg':'PSSE', 'offset':0.0, 'scale':1.0}], 'unit':'p.u.', 'rank':1, 'yminspan':0.1, 'ymaxlim':1.3, 'tolerance_bands':{'trace': 1, 'percent':5.0, 'base':1.0}},#'GSMG':0, }, #first channel in array links to file in first result location
##                                                                                                                                     {'dataset':0, 'name':"U_POC1", 'leg':'PSSE', 'offset':0.0, 'scale':1.0}], 'unit':'p.u.', 'rank':1, 'yminspan':0.1, 'ymaxlim':1.3,'GSMG':0.1}, #first channel in array links to file in first result location
#                                                                                            'Active Power POC':             {'channels':[{'dataset':4, 'name':"PPOC", 'leg':'PSCAD', 'offset':0.0, 'scale':1.0}, 
#                                                                                                                                    {'dataset':0, 'name':"P_POC1", 'leg':'PSSE', 'offset':0.0, 'scale':1.0}], 'unit':'MW', 'rank':3, 'yminspan':50, 'yminlim':-200.0, 'ymaxlim':200.0, 'tolerance_bands':{'trace': 1, 'percent':5.0, 'leg':'error-bands', 'base':80.0}},#'GSMG':0},
#                                                                                            'Reactive Power POC':           {'channels':[{'dataset':4, 'name':"QPOC", 'leg':'PSCAD', 'offset':0.0, 'scale':1.0}, 
#                                                                                                                                    {'dataset':0, 'name':"Q_POC1", 'leg':'PSSE', 'offset':0.0, 'scale':1.0}], 'unit':'MVAr', 'rank':5, 'yminspan':10.0, 'yminlim':-200.0, 'ymaxlim':200.0, 'tolerance_bands':{'trace': 1, 'percent':5.0, 'leg':'error-bands', 'base':31.6}},  
#
#                                                                                            'Voltage INV':                  {'channels':[{'dataset':4, 'name':"Vrms_LV_pu", 'leg':'PSCAD', 'offset':0.0, 'scale':1.0}, 
#                                                                                                                                     {'dataset':0, 'name':"U_LV1", 'leg':'PSSE', 'offset':0.0, 'scale':1.0}], 'unit':'p.u.', 'rank':2, 'yminspan':0.1, 'ymaxlim':1.3, 'tolerance_bands':{'trace': 1, 'percent':5.0, 'base':1.0}},#'GSMG':0, }, #first channel in array links to file in first result location
#                                                                                            'Active Power INV':             {'channels':[{'dataset':4, 'name':"Pout_LV (MW)", 'leg':'PSCAD', 'offset':0.0, 'scale':1.0}, 
#                                                                                                                                    {'dataset':0, 'name':"P_LV1", 'leg':'PSSE', 'offset':0.0, 'scale':1.0}], 'unit':'MW', 'rank':4, 'yminspan':50, 'yminlim':-200.0, 'ymaxlim':200.0, 'tolerance_bands':{'trace': 1, 'percent':5.0, 'leg':'error-bands', 'base':80.0}},#'GSMG':0},
#                                                                                            'Reactive Power INV':           {'channels':[{'dataset':4, 'name':"Qout_LV (MVar)", 'leg':'PSCAD', 'offset':0.0, 'scale':1.0}, 
#                                                                                                                                    {'dataset':0, 'name':"Q_LV1", 'leg':'PSSE', 'offset':0.0, 'scale':1.0}], 'unit':'MVAr', 'rank':6, 'yminspan':10.0, 'yminlim':-200.0, 'ymaxlim':200.0, 'tolerance_bands':{'trace': 1, 'percent':5.0, 'leg':'error-bands', 'base':31.6}},  
#
#                                                                                         },

##                                                            
#                                                            'Overlays(NEM)':             {
#                                                                                            'SHTS 66 kV bus voltage':   {'channels':[{'dataset':1, 'name':"V_SHTS_66", 'leg':'pre-LSF', 'offset':0.0, 'scale':1.0}, 
#                                                                                                                                     {'dataset':0, 'name':"V_SHTS_66", 'leg':'post-LSF', 'offset':0.0, 'scale':1.0}], 'unit':'p.u.', 'rank':1, 'yminspan':0.1, 'ymaxlim':1.3},#'GSMG':0, }, #first channel in array links to file in first result location
#                                                                                            'ECA 66 kV bus voltage':   {'channels':[{'dataset':1, 'name':"V_ECA_66", 'leg':'pre-LSF', 'offset':0.0, 'scale':1.0}, 
#                                                                                                                                     {'dataset':0, 'name':"V_ECA_66", 'leg':'post-LSF', 'offset':0.0, 'scale':1.0}], 'unit':'p.u.', 'rank':2, 'yminspan':0.1, 'ymaxlim':1.3},#'GSMG':0, }, #first channel in array links to file in first result location
#                                                                                            'KYM 66 kV bus voltage':   {'channels':[{'dataset':1, 'name':"V_KYM_66", 'leg':'pre-LSF', 'offset':0.0, 'scale':1.0}, 
#                                                                                                                                     {'dataset':0, 'name':"V_KYM_66", 'leg':'post-LSF', 'offset':0.0, 'scale':1.0}], 'unit':'p.u.', 'rank':3, 'yminspan':0.1, 'ymaxlim':1.3},#'GSMG':0, }, #first channel in array links to file in first result location
#                                                                                            'SHP 66 kV bus voltage':   {'channels':[{'dataset':1, 'name':"V_SHP_66", 'leg':'pre-LSF', 'offset':0.0, 'scale':1.0}, 
#                                                                                                                                     {'dataset':0, 'name':"V_SHP_66", 'leg':'post-LSF', 'offset':0.0, 'scale':1.0}], 'unit':'p.u.', 'rank':4, 'yminspan':0.1, 'ymaxlim':1.3},#'GSMG':0, }, #first channel in array links to file in first result location                                                                                                                                       
#                                                                                                                                       
#                                                                                         },

#                                                            'Overlays(POC)':             {
#                                                                                            'Voltage':                  {'channels':[{'dataset':4, 'name':"VPOC", 'leg':'PSCAD', 'offset':0.0, 'scale':1.0}, 
#                                                                                                                                     {'dataset':0, 'name':"U_POC1", 'leg':'PSSE', 'offset':0.0, 'scale':1.0}], 'unit':'p.u.', 'rank':1, 'yminspan':0.1, 'ymaxlim':1.3, 'tolerance_bands':{'trace': 1, 'percent':5.0, 'base':1.0}},#'GSMG':0, }, #first channel in array links to file in first result location
##                                                                                                                                     {'dataset':0, 'name':"U_POC1", 'leg':'PSSE', 'offset':0.0, 'scale':1.0}], 'unit':'p.u.', 'rank':1, 'yminspan':0.1, 'ymaxlim':1.3,'GSMG':0.1}, #first channel in array links to file in first result location
#                                                                                            'Active Power':             {'channels':[{'dataset':4, 'name':"PPOC", 'leg':'PSCAD', 'offset':0.0, 'scale':1.0}, 
#                                                                                                                                    {'dataset':0, 'name':"P_POC1", 'leg':'PSSE', 'offset':0.0, 'scale':1.0}], 'unit':'MW', 'rank':2, 'yminspan':50, 'yminlim':-200.0, 'ymaxlim':200.0, 'tolerance_bands':{'trace': 1, 'percent':5.0, 'leg':'error-bands', 'base':80.0}},#'GSMG':0},
#                                                                                            'Reactive Power':           {'channels':[{'dataset':4, 'name':"QPOC", 'leg':'PSCAD', 'offset':0.0, 'scale':1.0}, 
#                                                                                                                                    {'dataset':0, 'name':"Q_POC1", 'leg':'PSSE', 'offset':0.0, 'scale':1.0}], 'unit':'MVAr', 'rank':3, 'yminspan':10.0, 'yminlim':-200.0, 'ymaxlim':200.0, 'tolerance_bands':{'trace': 1, 'percent':5.0, 'leg':'error-bands', 'base':31.6}},  
##                                                                                            'Voltage Low':              {'channels':[{'dataset':4, 'name':"Vrms_LV_inv1_pu", 'leg':'PSCAD', 'offset':0.0, 'scale':1.0}, 
##                                                                                                                                     {'dataset':0, 'name':"U_LV1", 'leg':'PSSE', 'offset':0.0, 'scale':1.0}], 'unit':'p.u.', 'rank':4, 'yminspan':0.1, 'ymaxlim':1.3},#'GSMG':0, }, #first channel in array links to file in first result location
#
#                                                                                         },

#                                                            'Overlays(INV)':             {
#                                                                                            'Voltage':                  {'channels':[{'dataset':4, 'name':"Vrms_LV_pu", 'leg':'PSCAD', 'offset':0.0, 'scale':1.0}, 
#                                                                                                                                     {'dataset':0, 'name':"U_LV1", 'leg':'PSSE', 'offset':0.0, 'scale':1.0}], 'unit':'p.u.', 'rank':1, 'yminspan':0.1, 'ymaxlim':1.3, 'tolerance_bands':{'trace': 1, 'percent':5.0, 'base':1.0}},#'GSMG':0, }, #first channel in array links to file in first result location
#                                                                                            'Active Power':             {'channels':[{'dataset':4, 'name':"Pout_LV (MW)", 'leg':'PSCAD', 'offset':0.0, 'scale':1.0}, 
#                                                                                                                                    {'dataset':0, 'name':"P_LV1", 'leg':'PSSE', 'offset':0.0, 'scale':1.0}], 'unit':'MW', 'rank':2, 'yminspan':50, 'yminlim':-200.0, 'ymaxlim':200.0, 'tolerance_bands':{'trace': 1, 'percent':5.0, 'leg':'error-bands', 'base':80.0}},#'GSMG':0},
#                                                                                            'Reactive Power':           {'channels':[{'dataset':4, 'name':"Qout_LV (MVar)", 'leg':'PSCAD', 'offset':0.0, 'scale':1.0}, 
#                                                                                                                                    {'dataset':0, 'name':"Q_LV1", 'leg':'PSSE', 'offset':0.0, 'scale':1.0}], 'unit':'MVAr', 'rank':3, 'yminspan':10.0, 'yminlim':-200.0, 'ymaxlim':200.0, 'tolerance_bands':{'trace': 1, 'percent':5.0, 'leg':'error-bands', 'base':31.6}},  
##                                                                                            'Voltage Low':              {'channels':[{'dataset':4, 'name':"Vrms_LV_inv1_pu", 'leg':'PSCAD', 'offset':0.0, 'scale':1.0}, 
##                                                                                                                                     {'dataset':0, 'name':"U_LV1", 'leg':'PSSE', 'offset':0.0, 'scale':1.0}], 'unit':'p.u.', 'rank':4, 'yminspan':0.1, 'ymaxlim':1.3},#'GSMG':0, }, #first channel in array links to file in first result location
#
#                                                                                         },
                                                                                                                                      
#                                                            'Overlays(INF)':             {
#                                                                                            'Voltage':                  {'channels':[{'dataset':4, 'name':"V_INF_pu", 'leg':'PSCAD', 'offset':0.0, 'scale':1.0}, 
#                                                                                                                                     {'dataset':0, 'name':"U_HV1", 'leg':'PSSE', 'offset':0.0, 'scale':1.0}], 'unit':'p.u.', 'rank':1, 'yminspan':0.1, 'ymaxlim':1.3, 'tolerance_bands':{'trace': 1, 'percent':5.0, 'base':1.0}},#'GSMG':0, }, #first channel in array links to file in first result location
#                                                                                            'Active Power':             {'channels':[{'dataset':4, 'name':"P_INF", 'leg':'PSCAD', 'offset':0.0, 'scale':1.0}, 
#                                                                                                                                    {'dataset':0, 'name':"P_HV1", 'leg':'PSSE', 'offset':0.0, 'scale':1.0}], 'unit':'MW', 'rank':2, 'yminspan':50, 'yminlim':-200.0, 'ymaxlim':200.0, 'tolerance_bands':{'trace': 1, 'percent':5.0, 'leg':'error-bands', 'base':80.0}},#'GSMG':0},
#                                                                                            'Reactive Power':           {'channels':[{'dataset':4, 'name':"Q_INF", 'leg':'PSCAD', 'offset':0.0, 'scale':1.0}, 
#                                                                                                                                    {'dataset':0, 'name':"Q_HV1", 'leg':'PSSE', 'offset':0.0, 'scale':1.0}], 'unit':'MVAr', 'rank':3, 'yminspan':10.0, 'yminlim':-200.0, 'ymaxlim':200.0, 'tolerance_bands':{'trace': 1, 'percent':5.0, 'leg':'error-bands', 'base':31.6}},  
##                                                                                            'Voltage Low':              {'channels':[{'dataset':4, 'name':"Vrms_LV_inv1_pu", 'leg':'PSCAD', 'offset':0.0, 'scale':1.0}, 
##                                                                                                                                     {'dataset':0, 'name':"U_LV1", 'leg':'PSSE', 'offset':0.0, 'scale':1.0}], 'unit':'p.u.', 'rank':4, 'yminspan':0.1, 'ymaxlim':1.3},#'GSMG':0, }, #first channel in array links to file in first result location
#
#                                                                                         },

#                                                            'PSCAD simple':                {
#                                                                                              'Voltage POC':                 {'channels':[{'dataset':4, 'name':"VPOC", 'leg':'Voltage HV', 'offset':0.0, 'scale':1.0}, 
#                                                                                                                                          {'dataset':4, 'name':"Vref_POC (pu)", 'leg':'Voltage stp', 'offset':0.0, 'scale':1.0, 'colour':'grey', 'linestyle':'--', 'linewidth': 1.5},  ], 'unit':'p.u.', 'rank':1, 'ymaxlim':1.3},
#                                                                                              'Active Power POC':                {'channels':[{'dataset':4, 'name':"PPOC", 'leg':'Active Power HV', 'offset':0.0, 'scale':1.0}, 
#                                                                                                                                          {'dataset':4, 'name':"Pref_POC (%)", 'leg':'Active Power stp', 'offset':0.0, 'scale':0.8, 'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}, ], 'unit':'MW', 'rank':3, 'yminlim':-200.0, 'ymaxlim':200.0}, #add active power setpoint here
#                                                                                              'Reactive Power POC':              {'channels':[{'dataset':4, 'name':"QPOC", 'leg':'Reactive Power HV', 'offset':0.0, 'scale':1.0}, 
#                                                                                                                                          {'dataset':4, 'name':"Qref_POC (%)", 'leg':'Reactive Power stp', 'offset':0.0, 'scale':0.8, 'colour':'grey', 'linestyle':'--', 'linewidth': 1.5}, ], 'unit':'MVAr', 'rank':5, 'yminlim':-200.0, 'ymaxlim':200.0}, #add reactive power setpoint here
#                                                                                              'Reactive Current POC':            {'channels':[{'dataset':4, 'name':"Iq_PLANT", 'leg':'Reactive Current HV', 'offset':0.0, 'scale':1.0},], 'unit':'p.u', 'rank':7, 'yminspan':0.1, 'ymaxlim':2.0, 'yminlim':-2.0 },
#                                                                                              'Frequency':                   {'channels':[{'dataset':4, 'name':"fpoc", 'leg':'Frequency', 'offset':0.0, 'scale':1.0},], 'unit':'Hz', 'rank':9, 'yminspan':5.0},
#                                                                                            # 'Voltage LV':                     {'channels':[{'dataset':0, 'name':"VPOC", 'leg':'Voltage LV', 'offset':0.0, 'scale':1.0},], 'unit':'p.u.', 'rank':2},# add voltage setpoint here
#
#                                                                                            'Voltage INV':                  {'channels':[{'dataset':4, 'name':"Vrms_LV_inv1_pu", 'leg':'Voltage INV1', 'offset':0.0, 'scale':1.0},
#                                                                                                                                        {'dataset':4, 'name':"Vrms_LV_inv2_pu", 'leg':'Voltage INV2', 'offset':0.0, 'scale':1.0},], 'unit':'p.u.', 'rank':2, 'yminspan':0.1},
#                                                                                            'Active Power INV':                 {'channels':[{'dataset':4, 'name':"Pout_LV (MW)", 'leg':'P_INV1', 'offset':0.0, 'scale':1.0},
#                                                                                                                                         {'dataset':4, 'name':"Pout_LV_x", 'leg':'P_INV2', 'offset':0.0, 'scale':1.0},], 'unit':'MW', 'rank':4, 'yminspan':5, 'ymaxlim':200.0, 'yminlim':-200.0, },
#                                                                                            'Reactive Power INV':                 {'channels':[{'dataset':4, 'name':"Qout_LV (MVar)", 'leg':'Q_INV1', 'offset':0.0, 'scale':1.0},
#                                                                                                                                         {'dataset':4, 'name':"Qout_LV_x", 'leg':'Q_INV2', 'offset':0.0, 'scale':1.0},], 'unit':'MVAr', 'rank':6, 'yminspan':5, 'ymaxlim':200.0, 'yminlim':-200.0, },
#                                                                                             'Inverter State (FRT)':        {'channels':[{'dataset':4, 'name':"LVRT_ON", 'leg':'Inverter LVRT state', 'offset':0.0, 'scale':1.0},], 'unit':'inverter code', 'rank':8},
#                                                                                             'Angle':                      {'channels':[{'dataset':4, 'name':"phAngSource", 'leg':'Angle at POC', 'offset':0.0, 'scale':1.0},], 'unit':'degrees', 'rank':10, 'yminspan':10.0},
#
#                                                                                             
#                                                                                              },
#                                                            'PSSE simple':                {
#                                                                                            'Voltage POC':                 {'channels':[{'dataset':0, 'name':"U_POC1", 'leg':'Voltage HV', 'offset':0.0, 'scale':1.0},
#                                                                                                                                        {'dataset':0, 'name':"VREF_POC", 'leg':'Voltage stp', 'offset':0.0, 'scale':1.0, 'colour':'grey', 'linestyle':'--', 'linewidth': 1.5},], 'unit':'p.u.', 'rank':1, 'yminspan':0.1, 'ymaxlim':1.3},
#                                                                                            'Active Power POC':                {'channels':[{'dataset':0, 'name':"P_POC1", 'leg':'P - POC', 'offset':0.0, 'scale':1.0},
#                                                                                                                                        {'dataset':0, 'name':"PREF_POC", 'leg':'Active Power stp', 'offset':0.0, 'scale':94.94, 'colour':'grey', 'linestyle':'--', 'linewidth': 1.5},], 'unit':'MW', 'rank':3, 'yminspan':10},
#                                                                                            'Reactive Power POC':              {'channels':[{'dataset':0, 'name':"Q_POC1", 'leg':'Q - POC', 'offset':0.0, 'scale':1.0},
#                                                                                                                                        {'dataset':0, 'name':"QREF_POC", 'leg':'Reactive Power stp', 'offset':0.0, 'scale':94.94, 'colour':'grey', 'linestyle':'--', 'linewidth': 1.5},], 'unit':'MVAr', 'rank':5, 'yminspan':5},
#                                                                                            'Reactive Current POC':            {'channels':[{'dataset':0, 'name':"Iq_PLANT", 'leg':'Reactive Current HV', 'offset':0.0, 'scale':1.0},], 'unit':'p.u', 'rank':7, 'yminspan':0.1, 'ymaxlim':2.0, 'yminlim':-2.0 },
##                                                                                            'Total Current':            {'channels':[{'dataset':0, 'name':"INV_ITOT", 'leg':'Current INV1', 'offset':0.0, 'scale':1.0},
##                                                                                                                                        {'dataset':0, 'name':"INV2_ITOT", 'leg':'Current INV2', 'offset':0.0, 'scale':1.0},], 'unit':'p.u', 'rank':4, 'yminspan':0.1, 'ymaxlim':2.0, 'yminlim':-2.0 },
#                                                                                            'Frequency':                   {'channels':[{'dataset':0, 'name':"F_POC1", 'leg':'Frequency', 'offset':50.0, 'scale':50.0},], 'unit':'Hz', 'rank':9, 'yminspan':5.0},
#
#
#                                                                                                                                                
#                                                                                            'Voltage INV':                  {'channels':[{'dataset':0, 'name':"U_LV1", 'leg':'Voltage INV1', 'offset':0.0, 'scale':1.0},
#                                                                                                                                        {'dataset':0, 'name':"U_LV2", 'leg':'Voltage INV2', 'offset':0.0, 'scale':1.0},], 'unit':'p.u.', 'rank':2, 'yminspan':0.1},                                                                                                                                                
#                                                                                            'Active Power INV':              {'channels':[{'dataset':0, 'name':"P_LV1", 'leg':'P_INV1', 'offset':0.0, 'scale':1.0},
#                                                                                                                                         {'dataset':0, 'name':"P_LV2", 'leg':'P_INV2', 'offset':0.0, 'scale':1.0},], 'unit':'MW', 'rank':4, 'yminspan':5, 'ymaxlim':200.0, 'yminlim':-200.0, },
#                                                                                            'Reactive Power INV':            {'channels':[{'dataset':0, 'name':"Q_LV1", 'leg':'Q_INV1', 'offset':0.0, 'scale':1.0},
#                                                                                                                                         {'dataset':0, 'name':"Q_LV2", 'leg':'Q_INV2', 'offset':0.0, 'scale':1.0},], 'unit':'MVAr', 'rank':6, 'yminspan':5, 'ymaxlim':200.0, 'yminlim':-200.0, },
#                                                                                            'Inverter State (FRT)':         {'channels':[{'dataset':0, 'name':"INV1_VDFLAG", 'leg':'INV1 state', 'offset':0.0, 'scale':1.0},
#                                                                                                                                        {'dataset':0, 'name':"INV2_VDFLAG", 'leg':'INV2 state', 'offset':0.0, 'scale':1.0},], 'unit':'inverter code', 'rank':8},
##                                                                                            'Angle':                       {'channels':[{'dataset':0, 'name':"ANG_POC1", 'leg':'Angle at POC', 'offset':0.0, 'scale':1.0},], 'unit':'MVAr', 'rank':10, 'yminspan':10.0},
#                                                                                            'Frequency':                   {'channels':[{'dataset':0, 'name':"F_POC1", 'leg':'Frequency', 'offset':50.0, 'scale':50.0},], 'unit':'Hz', 'rank':10, 'yminspan':5.0},
#                                                                                            
#                                                                                            },

                                                             
                                                             },

                                                    },
                                         ]
                    },
    

    
#            'GPS':{'batchname': 'first_batch', #in DMAT
#                    'report_definition':[          {'chapter':'Clause S5.2.5.1 SMIB tests', #Array of chapters. In this case there is only one chapter.
#                                                    'datasets': [
#                                                                    #{'label':'PSSE_34_2', 'path':r"PSSE_sim\result_data\20210317_BM_uptd87", 'ID': 0, 'timeID':'Time(s)', 'timeoffset':0.0}, #settling bands are added relating to the first dataset
#                                                                    {'label':'PSCAD', 'path':r"PSCAD_sim\result_data\20210421_DMAT", 'ID': 1, 'timeID':'time(s)', 'timeoffset':0.0, 'calcCurrents':[
#                                                                            {"P":"PLANT_P_HV", "Q":"PLANT_Q_HV", "I":"PV_Arms_HV", "nameLabel":"PLANT", "scaling":1.0, }, 
#                                                                            {"P":"PV_P_HV", "Q":"PV_Q_HV", "I":"PV_Arms_HV", "nameLabel":"SOLAR_HV", "scaling":1.0, },
#                                                                            {"P":"PCU2_P_LV", "Q":"PCU2_Q_LV", "I":"PCU2_Arms_LV", "nameLabel":"SOLAR_LV", "scaling":1.0, },
#                                                                            {"P":"SYNC_P_HV", "Q":"SYNC_Q_HV", "I":"SYNC_Arms_HV", "nameLabel":"SYNC_HV", "scaling":1.0, },
#                                                                            {"P":"SYNC_P_MV", "Q":"SYNC_Q_MV", "I":"SYNC_Arms_MV", "nameLabel":"SYNC_MV", "scaling":1.0, },
#                                                                            ]},
#                                                                ],
#                                                    'cases':[],
#                                                    'plots_for_report': ['POC - PSCAD', 'Solar_HV - PSCAD', 'Solar_LV - PSCAD', 'Solar_auxiliary - PSCAD'], #list pecifying which of the previous plots shoudl be be actually added to the report. Allows to omit plots without deleting the definition.
#                                                    'summary_items': {}, #Additional description/analysis of the result data can bebe specified here. For example amount of current injection for 5.2.5.5 analysis
#                                                    'report':True,
#                                                    'plots':{
#                                                            }
#                                                    },
#                                                    {'chapter':'Clauses S5.2.5.3, S5.2.5.2, S5.2.5.4, S5.2.5.7, S5.2.5.8, S5.2.5.11 and S5.2.5.14 SMIB tests',
#                                                    'datasets': [
#                                                                    #{'label':'PSSE_34_2', 'path':r"PSSE_sim\result_data\20210317_BM_uptd87", 'ID': 0, 'timeID':'Time(s)', 'timeoffset':0.0}, #settling bands are added relating to the first dataset
#                                                                    {'label':'PSCAD', 'path':r"PSCAD_sim\result_data\20210421_DMAT", 'ID': 1, 'timeID':'time(s)', 'timeoffset':0.0, 'calcCurrents':[
#                                                                            {"P":"PLANT_P_HV", "Q":"PLANT_Q_HV", "I":"PV_Arms_HV", "nameLabel":"PLANT", "scaling":1.0, }, 
#                                                                            {"P":"PV_P_HV", "Q":"PV_Q_HV", "I":"PV_Arms_HV", "nameLabel":"SOLAR_HV", "scaling":1.0, },
#                                                                            {"P":"PCU2_P_LV", "Q":"PCU2_Q_LV", "I":"PCU2_Arms_LV", "nameLabel":"SOLAR_LV", "scaling":1.0, },
#                                                                            {"P":"SYNC_P_HV", "Q":"SYNC_Q_HV", "I":"SYNC_Arms_HV", "nameLabel":"SYNC_HV", "scaling":1.0, },
#                                                                            {"P":"SYNC_P_MV", "Q":"SYNC_Q_MV", "I":"SYNC_Arms_MV", "nameLabel":"SYNC_MV", "scaling":1.0, },
#                                                                            ]},
#                                                                ],
#                                                    'cases':[],
#                                                    'plots_for_report': ['POC - PSCAD', 'Solar_HV - PSCAD', 'Solar_LV - PSCAD', 'Solar_auxiliary - PSCAD'], #list pecifying which of the previous plots shoudl be be actually added to the report. Allows to omit plots without deleting the definition.
#                                                    'summary_items': {}, #Additional description/analysis of the result data can bebe specified here. For example amount of current injection for 5.2.5.5 analysis
#                                                    'report':True,
#                                                    'plots':{'POC':{}, #Voltage + setpoint, Q+setpoint, P+setpoint, Frequency + Angle as per PSCAD Pll, 
#                                                             'SynCon - MV':{}, #Voltage, Q, P, protection block output
#                                                             'Solar - LV':{}, #Voltage, Q+setpoint, P+setpoint, Operating state+grid Error flag 
#                                                            }
#                                                    },
#                                                    {'chapter':'Clause S5.2.5.5 SMIB tests', 
#                                                    'datasets': [
#                                                                    #{'label':'PSSE_34_2', 'path':r"PSSE_sim\result_data\20210317_BM_uptd87", 'ID': 0, 'timeID':'Time(s)', 'timeoffset':0.0}, #settling bands are added relating to the first dataset
#                                                                    {'label':'PSCAD', 'path':r"PSCAD_sim\result_data\20210421_DMAT", 'ID': 1, 'timeID':'time(s)', 'timeoffset':0.0, 'calcCurrents':[
#                                                                            {"P":"PLANT_P_HV", "Q":"PLANT_Q_HV", "I":"PV_Arms_HV", "nameLabel":"PLANT", "scaling":1.0, }, 
#                                                                            {"P":"PV_P_HV", "Q":"PV_Q_HV", "I":"PV_Arms_HV", "nameLabel":"SOLAR_HV", "scaling":1.0, },
#                                                                            {"P":"PCU2_P_LV", "Q":"PCU2_Q_LV", "I":"PCU2_Arms_LV", "nameLabel":"SOLAR_LV", "scaling":1.0, },
#                                                                            {"P":"SYNC_P_HV", "Q":"SYNC_Q_HV", "I":"SYNC_Arms_HV", "nameLabel":"SYNC_HV", "scaling":1.0, },
#                                                                            {"P":"SYNC_P_MV", "Q":"SYNC_Q_MV", "I":"SYNC_Arms_MV", "nameLabel":"SYNC_MV", "scaling":1.0, },
#                                                                            ]},
#                                                                ],
#                                                    'cases':[],
#                                                    'plots_for_report': ['POC - PSCAD', 'Solar_HV - PSCAD', 'Solar_LV - PSCAD', 'Solar_auxiliary - PSCAD'], #list pecifying which of the previous plots shoudl be be actually added to the report. Allows to omit plots without deleting the definition.
#                                                    'summary_items': {}, #Additional description/analysis of the result data can bebe specified here. For example amount of current injection for 5.2.5.5 analysis
#                                                    'report':True,
#                                                    'plots':{'POC':{}, #Voltage + dV marker, Q, P+recovery time marker, Iq+dIq marker, Ip, Frequency + Angle as per PSCAD Pll, 
#                                                             'SynCon - MV':{}, #Voltage, Q, P, protection block output
#                                                             'Solar - LV':{}, #Voltage, Q+setpoint, P+setpoint, Operating state+grid Error flag 
#                                                            }
#                                                    },
#                                                    {'chapter':'Clause S5.2.5.13 SMIB tests', 
#                                                    'datasets': [
#                                                                    #{'label':'PSSE_34_2', 'path':r"PSSE_sim\result_data\20210317_BM_uptd87", 'ID': 0, 'timeID':'Time(s)', 'timeoffset':0.0}, #settling bands are added relating to the first dataset
#                                                                    {'label':'PSCAD', 'path':r"PSCAD_sim\result_data\20210421_DMAT", 'ID': 1, 'timeID':'time(s)', 'timeoffset':0.0, 'calcCurrents':[
#                                                                            {"P":"PLANT_P_HV", "Q":"PLANT_Q_HV", "I":"PV_Arms_HV", "nameLabel":"PLANT", "scaling":1.0, }, 
#                                                                            {"P":"PV_P_HV", "Q":"PV_Q_HV", "I":"PV_Arms_HV", "nameLabel":"SOLAR_HV", "scaling":1.0, },
#                                                                            {"P":"PCU2_P_LV", "Q":"PCU2_Q_LV", "I":"PCU2_Arms_LV", "nameLabel":"SOLAR_LV", "scaling":1.0, },
#                                                                            {"P":"SYNC_P_HV", "Q":"SYNC_Q_HV", "I":"SYNC_Arms_HV", "nameLabel":"SYNC_HV", "scaling":1.0, },
#                                                                            {"P":"SYNC_P_MV", "Q":"SYNC_Q_MV", "I":"SYNC_Arms_MV", "nameLabel":"SYNC_MV", "scaling":1.0, },
#                                                                            ]},
#                                                                ],
#                                                    'cases':[],
#                                                    'plots_for_report': ['POC - PSCAD', 'Solar_HV - PSCAD', 'Solar_LV - PSCAD', 'Solar_auxiliary - PSCAD'], #list pecifying which of the previous plots shoudl be be actually added to the report. Allows to omit plots without deleting the definition.
#                                                    'summary_items': {}, #Additional description/analysis of the result data can bebe specified here. For example amount of current injection for 5.2.5.5 analysis
#                                                    'report':True,
#                                                    'plots':{'POC':{}, #Voltage + settling time marker + settling band , V setpoint, Q+settling time maker + settling band, Q setpoint, P, Frequency + Angle as per PSCAD Pll, 
#                                                             'SynCon - MV':{}, #Voltage, Q, P, protection block output
#                                                             'Solar - LV':{}, #Voltage, Q+setpoint, P+setpoint, Operating state+grid Error flag 
#                                                            }
#                                                    }
#                                        ],
#                  },
            
            
            } #specify which type of output document(s) shall be generated

#channels={} #specified which channels to be used 

#datasets={} #specify locations in which to look for data sets
#------------------------------------------------------------------------------
# IMPORTS
#------------------------------------------------------------------------------
import os
import sys
main_folder_path=os.path.dirname(os.path.dirname(os.path.dirname(__file__)))
main_folder_path_out=os.path.dirname(os.path.dirname(os.path.dirname(__file__)))
#sys.path.append(r'C:\Users\Mervin Kall\OneDrive - ESCO Pacific\basics\ESCOPyTools')
sys.path.append(main_folder_path+"\\PSSE_sim\\scripts\\Libs")
sys.path.append(r"C:\Python27\Lib\site-packages")
import EscoPlot as ep
import shelve

import StringIO

import docx
from docx import Document, shape
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK
from docx import Document

from docx.enum.dml import MSO_THEME_COLOR_INDEX

import datetime
import readtestinfo


#------------------------------------------------------------------------------
# GLOBAL VARIABLES
#------------------------------------------------------------------------------
# [ProjectDetailsDict, PSCADmodelDict, PSSEmodelDict, SetpointsDict, TestTypesDict, ScenariosDict, ProfilesDict] =  readtestinfo.readTestdef(main_folder_path+"\\test_scenario_definitions\\"+TestDefinitionSheet)

return_dict =  readtestinfo.readTestdef(main_folder_path+"\\test_scenario_definitions\\"+TestDefinitionSheet, ['ProjectDetails', 'ModelDetailsPSSE', 'ModelDetailsPSCAD', 'Setpoints', 'ScenariosSMIB', 'Profiles'])
ProjectDetailsDict = return_dict['ProjectDetails']
# SimulationSettingsDict = return_dict['SimulationSettings']
PSSEmodelDict = return_dict['ModelDetailsPSSE']
PSCADmodelDict = return_dict['ModelDetailsPSCAD']
SetpointsDict = return_dict['Setpoints']
ScenariosDict = return_dict['ScenariosSMIB']
ProfilesDict = return_dict['Profiles']

landscape_flag=0
#------------------------------------------------------------------------------
# FUNCTIONS
#------------------------------------------------------------------------------
#sort cases by type and by ID (which is located at the end of the strin identifier)
def sort_cases(cases):
    sorted_cases=[]
    order = [ 'small', 'large', 'ort', 'tov', 'con', 'Case_']
    for testType in order:
        cases_temp={}
        for case in cases:
            if(testType in case): #will ignore any folders that are do not belong to one of thet est types
                number=float(case.replace(testType, ''))
                if(number>0):## additional criterion added for debugging purposes. Should be removed for final version
                    cases_temp[number]=case
        keys = cases_temp.keys()
        if(keys!=[]): 
            for id in range(0, int(max(keys))+1):
                if id in keys:
                    sorted_cases.append(cases_temp[id])    
    return sorted_cases

#check if string contains number
def is_number(s):
    try:
        float(s)
        return True
    except:
        return False
#function is given test profile (pointwise definition) and returns the start times and end times of the steps included in that profile. A step is considered to occur whenever the gradien is greater than 0, and to end whenever the gradient becomes 0 again
def detect_steps(offset, profile):
    steps=[] 
    magnitude=max(profile['y'])
    i=0
    prev_grad=0
    while (i < len(profile['x'])-1):
        gradient=float((profile['y'][i+1]-profile['y'][i]))/(magnitude*(profile['x'][i+1]-profile['x'][i]))
        if((abs(gradient)>0.01) and (abs(prev_grad)<=0.01)):
            stepStart=profile['x'][i]
        elif( (i>0) and ((abs(gradient)<=0.01) or ( (abs(gradient)>0.01) and (i+1==len(profile['x']-1)) ) ) ):
            stepEnd=profile['x'][i]
            steps.append([stepStart, stepEnd])
        i+=1        
    return steps

def add_bookmark(paragraph, bookmark_text, bookmark_name):
    run = paragraph.add_run()
    tag = run._r
    start = docx.oxml.shared.OxmlElement('w:bookmarkStart')
    start.set(docx.oxml.ns.qn('w:id'), '0')
    start.set(docx.oxml.ns.qn('w:name'), bookmark_name)
    tag.append(start)

    text = docx.oxml.OxmlElement('w:r')
    text.text = bookmark_text
    tag.append(text)

    end = docx.oxml.shared.OxmlElement('w:bookmarkEnd')
    end.set(docx.oxml.ns.qn('w:id'), '0')
    end.set(docx.oxml.ns.qn('w:name'), bookmark_name)
    tag.append(end)

def add_link(paragraph, link_to, text, tool_tip=None):
    # create hyperlink node
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')

    # set attribute for link to bookmark
    hyperlink.set(docx.oxml.shared.qn('w:anchor'), link_to,)

    if tool_tip is not None:
        # set attribute for link to bookmark
        hyperlink.set(docx.oxml.shared.qn('w:tooltip'), tool_tip,)

    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    r = paragraph.add_run()
    r._r.append(hyperlink)
    r.font.name = "Calibri"
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True
    

def initialise_report(report_type):
    #read report template 
    report=Document(main_folder_path+"\\Plots\\ReportTemplate.docx")
    return report

def add_report_intro(report, report_type, datasets, cases): #change it so that data location is an array of datasets. For overlay plots the inclusion of time step can then be decided based on if PSS/E results are present.
    #generate general description and intro based on ProjectDetailsDict
    p=report.add_paragraph('')
    run=p.add_run()
    run.add_break(WD_BREAK.PAGE)
    report.add_heading("Introduction", level=1 )
    if(report_type=="BENCHMARKING"):
        intro_text=str(ProjectDetailsDict['Dev'])+" is developing "+str(ProjectDetailsDict['Name'])+" in "+str(ProjectDetailsDict['Town'])+", "+str(ProjectDetailsDict['State'])+". "
        intro_text+="The project will connect into "+str(ProjectDetailsDict['Sub'])+" and feature a total active power rating of "+str(ProjectDetailsDict['PlantMW'])+" MW at the point of connection. " 
        pass
        p=report.add_paragraph(intro_text)
        
        software_types=[]
        for dataset in datasets:
            if (( ('PSS/E' in dataset['label']) or ("PSSE" in dataset['label']) ) and not ("PSS/E" in software_types)):
                software_types.append("PSS/E")
        if( ('PSCAD' in dataset['label']) and not ("PSCAD" in software_types)):
            software_types.append('PSCAD')
        intro_text="As part of the connection application, a PSS/E and a PSCAD model are submitted. These models are required to adequately represent the performance of the hardware proposed to be installed on site. As such, the two models need to show comparable behaviour and performance when subjected to the same test conditions in both PSS/E and PSCAD. "
        intro_text+="This report shows the results of benchmarking studies that have been carried out to demonstrate adequate alignment between the PSCAD and PSS/E model. "
        p=report.add_paragraph(intro_text)
        intro_text="The studies have been conducted in "+str(PSSEmodelDict['PSSEversion'])+" and "+str(PSCADmodelDict['pscad version'])+" using the compiler "+str(PSCADmodelDict['compiler'])+". "
        intro_text+="The tests included in this report are listed below." 
        summary_table=True

        
    elif(report_type=="DMAT"): 
        intro_text=str(ProjectDetailsDict['Dev'])+" is developing "+str(ProjectDetailsDict['Name'])+" in "+str(ProjectDetailsDict['Town'])+" ,"+str(ProjectDetailsDict['State'])+". "
        intro_text+="The project will connect into "+str(ProjectDetailsDict['Sub'])+" and feature a total active power rating of "+str(ProjectDetailsDict['PlantMW'])+" MW at the point of connection. " 
        pass
        p=report.add_paragraph(intro_text)
        
        software_types=[]
        for dataset in datasets:
            if (( ('PSS/E' in dataset['label']) or ("PSSE" in dataset['label']) ) and not ("PSS/E" in software_types)):
                software_types.append("PSS/E")
        if( ('PSCAD' in dataset['label']) and not ("PSCAD" in software_types)):
            software_types.append('PSCAD')
        
        intro_text="As part of the connection application "
        if(len(software_types)==1):
            intro_text+='a '+software_types[0]+" model is submitted. The model is required to adequately represent the performance of the hardware proposed to be installed on site, and the model is also expected to show acceptable performance in the test scenarios outlined in AEMO's Model Acceptance Test Guideline. "
            intro_text+="This report shows the results of Model Acceptance test studies that have been carried out to demonstrate the performance of the "+software_types[0]+" model. "
        elif(len(software_types)>1):
            intro_text+='models in '
            for software in range(0, len(software_types)-1):
                intro_text+=software_types[software]+", "
            intro_text+=' and '+software_types[-1]+" are submitted. The models are required to adequately represent the performance of the hardware proposed to be installed on site, and the models are also expected to show acceptable performance in the test scenarios outlined in AEMO's Model Acceptance Test Guideline. "
            intro_text+="This report shows the results of Model Acceptance test studies that have been carried out to demonstrate the performance of the models. "
        p=report.add_paragraph(intro_text)
        if('PSCAD' in software_types):
            intro_text="The PSCAD studies have been conducted in "+str(PSCADmodelDict['pscad version'])+" using the compiler "+str(PSCADmodelDict['compiler'])+". "
        if('PSS/E' in software_types):
            intro_text="The PSSE studies have been conducted in "+str(PSSEmodelDict['PSSEversion'])+". "
        intro_text+="The tests included in this report are listed below."   
        summary_table=True

    elif(report_type=="NetworkFault"): 
        intro_text=str(ProjectDetailsDict['Dev'])+" is developing "+str(ProjectDetailsDict['Name'])+" in "+str(ProjectDetailsDict['Town'])+" ,"+str(ProjectDetailsDict['State'])+". "
        intro_text+="The project will connect into "+str(ProjectDetailsDict['Sub'])+" and feature a total active power rating of "+str(ProjectDetailsDict['PlantMW'])+" MW at the point of connection. " 
        pass
        p=report.add_paragraph(intro_text)
        
        software_types=[]
        for dataset in datasets:
            if (( ('PSS/E' in dataset['label']) or ("PSSE" in dataset['label']) ) and not ("PSS/E" in software_types)):
                software_types.append("PSS/E")
        if( ('PSCAD' in dataset['label']) and not ("PSCAD" in software_types)):
            software_types.append('PSCAD')
        
        intro_text="As part of the connection application "
        if(len(software_types)==1):
            intro_text+='a '+software_types[0]+" model is submitted. The model is required to adequately represent the performance of the hardware proposed to be installed on site, and the model is also expected to show acceptable performance in the test scenarios outlined in AEMO's Model Acceptance Test Guideline. "
            intro_text+="This report shows the results of Model Acceptance test studies that have been carried out to demonstrate the performance of the "+software_types[0]+" model. "
        elif(len(software_types)>1):
            intro_text+='models in '
            for software in range(0, len(software_types)-1):
                intro_text+=software_types[software]+", "
            intro_text+=' and '+software_types[-1]+" are submitted. The models are required to adequately represent the performance of the hardware proposed to be installed on site, and the models are also expected to show acceptable performance in the test scenarios outlined in AEMO's Model Acceptance Test Guideline. "
            intro_text+="This report shows the results of Model Acceptance test studies that have been carried out to demonstrate the performance of the models. "
        p=report.add_paragraph(intro_text)
        if('PSCAD' in software_types):
            intro_text="The PSCAD studies have been conducted in "+str(PSCADmodelDict['pscad version'])+" using the compiler "+str(PSCADmodelDict['compiler'])+". "
        if('PSS/E' in software_types):
            intro_text="The PSSE studies have been conducted in "+str(PSSEmodelDict['PSSEversion'])+". "
        intro_text+="The tests included in this report are listed below."   
        summary_table=True
        
    if(summary_table):     
        p=report.add_paragraph(intro_text)        
        change_orientation(report)
        p=report.add_paragraph("")
        tableCnt=1
        if(any('small' in case for case in cases)):
            p.add_run('Table '+str(tableCnt)+': Scenario list - Small Disturbance tests').bold=True        

            if('PSS/E' in software_types):
#                headers=['Case Nr.', 'Test Type', 'Test profile', 'POC voltage', 'Grid MVA', 'Grid X/R', 'P at POC (MW)', 'Q at POC (MVAr)', 'PSSE time step', 'PSSE acc. factor', 'passed']
                headers=['Case Nr.', 'Test Type', 'Test profile', 'POC voltage', 'Grid SCR', 'Grid X/R', 'P at POC (MW)', 'Q at POC (MVAr)', 'PSSE time step (ms)', 'PSSE acc. factor', 'DMAT id', 'passed']
            else:
#                headers=['Case Nr.', 'Test Type', 'Test profile', 'POC voltage', 'Grid MVA', 'Grid X/R', 'P at POC (MW)', 'Q at POC (MVAr)', 'passed']
                headers=['Case Nr.', 'Test Type', 'Test profile', 'POC voltage', 'Grid SCR', 'Grid X/R', 'P at POC (MW)', 'Q at POC (MVAr)', 'DMAT id', 'passed']
            table=report.add_table(rows=1, cols=len(headers))
            table.style='ListTable3-Accent3'
            hdr_cells=table.rows[0].cells
            for header_id in range(0, len(headers)): hdr_cells[header_id].text=headers[header_id]
            # read test metadata to get test info
            for case_id in range(0, len(cases)):
                if ('small' in cases[case_id]):
                    dataset_number=0
                    test_details={}
                    while (dataset_number < len(datasets)) and (test_details=={}):
                        try:
                            test_details=shelve.open(main_folder_path_out+"\\"+datasets[dataset_number]['path']+"\\"+cases[case_id]+"\\testInfo\\"+cases[case_id])
                        except:
                            dataset_number+=1
                    row_cells=table.add_row().cells
                    cell_paragraph=row_cells[0].paragraphs[0]
                    add_link(paragraph=cell_paragraph, link_to=str(cases[case_id]), text=str(cases[case_id]), tool_tip="link to test results")
                    row_cells[1].text=str(test_details['scenario_params']['Test Type'])
                    row_cells[2].text=str(test_details['scenario_params']['Test profile'])
                    row_cells[3].text=str(test_details['setpoint']['V_POC'])
#                    row_cells[4].text=str(test_details['setpoint']['GridMVA'])
#                    row_cells[4].text=str(test_details['setpoint']['SCR'])
                    row_cells[4].text=str(round(test_details['setpoint']['SCR'],2))
                    row_cells[5].text=str(test_details['setpoint']['X_R'])
                    row_cells[6].text=str(test_details['setpoint']['P'])
                    row_cells[7].text=str(test_details['setpoint']['Q'])
                    if('PSS/E' in software_types):
                        row_cells[8].text=str(test_details['scenario_params']['TimeStep'])
                        row_cells[9].text=str(test_details['scenario_params']['AccFactor'])
#                    row_cells[10].text=str(test_details['scenario_params']['Comment/Corresponding DMAT case'])
                    row_cells[10].text=str(test_details['scenario_params']['comment'])
            for row in table.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size= Pt(9)
        
            tableCnt+=1
            #run.add_break()
            p=report.add_paragraph(" ")
            p=report.add_paragraph('')
        if(any('large' in case for case in cases)):
            p.add_run('Table '+str(tableCnt)+': Scenario list - Large Disturbance tests').bold=True  
            if('PSS/E' in software_types):
#                headers=['Case Nr.', 'Test Type', 'Fault Type', 'Fault time', 'Fault duration', 'Fault impedance', 'V residual', 'Fault X/R', 'POC voltage', 'Grid MVA', 'Grid X/R', 'Grid MVA post-fault','Grid X/R post-fault', 'P at POC (MW)', 'Q at POC (MVAr)', 'PSSE time step', 'PSSE acc. factor', 'passed']
                headers=['Case Nr.', 'Test Type', 'Fault Type', 'Fault time', 'Fault duration', 'Fault impedance', 'V residual', 'Fault X/R', 'POC voltage', 'Grid SCR', 'Grid X/R', 'Grid SCR post-fault','Grid X/R post-fault', 'P at POC (MW)', 'Q at POC (MVAr)', 'PSSE time step (ms)', 'PSSE acc. factor', 'DMAT id', 'passed']
            else:    
#                headers=['Case Nr.', 'Test Type', 'Fault Type', 'Fault time', 'Fault duration', 'Fault impedance', 'V residual', 'Fault X/R', 'POC voltage', 'Grid MVA', 'Grid X/R', 'Grid MVA post-fault','Grid X/R post-fault', 'P at POC (MW)', 'Q at POC (MVAr)', 'passed']
                headers=['Case Nr.', 'Test Type', 'Fault Type', 'Fault time', 'Fault duration', 'Fault impedance', 'V residual', 'Fault X/R', 'POC voltage', 'Grid SCR', 'Grid X/R', 'Grid SCR post-fault','Grid X/R post-fault', 'P at POC (MW)', 'Q at POC (MVAr)', 'DMAT id', 'passed']
            table1=report.add_table(rows=1, cols=len(headers))
            table1.style='ListTable3-Accent3'
            hdr_cells1=table1.rows[0].cells
            for header_id in range(0, len(headers)): hdr_cells1[header_id].text=headers[header_id]
            # read test metadata to get test info
            for case_id in range(0, len(cases)):
                if ('large' in cases[case_id]):
                    dataset_number=0
                    test_details={}
                    while (dataset_number < len(datasets)) and (test_details=={}):
                        try:
                            test_details=shelve.open(main_folder_path_out+"\\"+datasets[dataset_number]['path']+"\\"+cases[case_id]+"\\testInfo\\"+cases[case_id])
                        except:
                            dataset_number+=1
                    row_cells=table1.add_row().cells
                    cell_paragraph=row_cells[0].paragraphs[0]
                    add_link(paragraph=cell_paragraph, link_to=str(cases[case_id]), text=str(cases[case_id]), tool_tip="link to test results")
                    row_cells[1].text=str(test_details['scenario_params']['Test Type'])

                    if(not 'Multifault' in test_details['scenario_params']['Test Type'] ):#single fault test
                        row_cells[2].text=str(test_details['scenario_params']['Event_Type'])
                        row_cells[3].text=str(test_details['scenario_params']['Ftime'])
                        row_cells[4].text=str(test_details['scenario_params']['Fduration'])
                        row_cells[5].text=str(round(test_details['scenario_params']['F_Impedance'],2))
                        if(test_details['scenario_params']['Vresidual']!=''):
                            row_cells[6].text=str(test_details['scenario_params']['Vresidual'])
                        else:
                            row_cells[6].text='-'
                        row_cells[7].text=str(test_details['scenario_params']['Fault X_R'])
                    else:
                        row_cells[2].text='various'
                        row_cells[3].text='various'
                        row_cells[4].text='various'
                        row_cells[5].text='various'
                        row_cells[6].text='various'
                        row_cells[7].text='various'
                    row_cells[8].text=str(test_details['setpoint']['V_POC'])
#                    row_cells[9].text=str(test_details['setpoint']['GridMVA'])
                    row_cells[9].text=str(round(test_details['setpoint']['SCR'],2))
                    row_cells[10].text=str(test_details['setpoint']['X_R'])
                    if('SCL_post' in test_details['scenario_params']):
                        if(test_details['scenario_params']['SCL_post']!=''):
                            row_cells[11].text=str(test_details['scenario_params']['SCL_post'])
                        else:
#                            row_cells[11].text=str(test_details['setpoint']['GridMVA'])
                            row_cells[11].text=str(round(test_details['setpoint']['SCR'],2))
                    else:
#                        row_cells[11].text=str(test_details['setpoint']['GridMVA'])
                        row_cells[11].text=str(round(test_details['setpoint']['SCR'],2))
                    if('X_R_post' in test_details['scenario_params']):
                        if(test_details['scenario_params']['X_R_post']!=''):
                            row_cells[12].text=str(test_details['scenario_params']['X_R_post'])
                        else:
                            row_cells[12].text=str(test_details['setpoint']['X_R'])
                    else:
                        row_cells[12].text=str(test_details['setpoint']['X_R'])  
                    row_cells[13].text=str(test_details['setpoint']['P'])
                    row_cells[14].text=str(test_details['setpoint']['Q'])
                    if('PSS/E' in software_types):
                        row_cells[15].text=str(test_details['scenario_params']['TimeStep'])
                        row_cells[16].text=str(test_details['scenario_params']['AccFactor'])
                    row_cells[17].text=str(test_details['scenario_params']['Comment/Corresponding DMAT case'])
            for row in table1.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size= Pt(9)
                        
            tableCnt+=1
            #run.add_break()
            p=report.add_paragraph(" ")
            p=report.add_paragraph('')
        if(any('ort' in case for case in cases)):
            p.add_run('Table '+str(tableCnt)+': Scenario list - Oscillatory Rejection tests').bold=True        
            if('PSS/E' in software_types):
#                headers=['Case Nr.', 'Test Type', 'Time', 'Disturbance Frequency', 'Disturbance Magnitude', 'Phase Oscillation Magnitude', 'POC voltage', 'Grid MVA', 'Grid X/R',  'P at POC (MW)', 'Q at POC (MVAr)', 'PSSE time step', 'PSSE acc. factor', 'passed']
                headers=['Case Nr.', 'Test Type', 'Time', 'Disturbance Frequency', 'Disturbance Magnitude', 'Phase Oscillation Magnitude', 'POC voltage', 'Grid SCR', 'Grid X/R',  'P at POC (MW)', 'Q at POC (MVAr)', 'PSSE time step (ms)', 'PSSE acc. factor', 'DMAT id', 'passed']
            else:
#                headers=['Case Nr.', 'Test Type', 'Time', 'Disturbance Frequency', 'Disturbance Magnitude', 'Phase Oscillation Magnitude', 'POC voltage', 'Grid MVA', 'Grid X/R',  'P at POC (MW)', 'Q at POC (MVAr)', 'passed']
                headers=['Case Nr.', 'Test Type', 'Time', 'Disturbance Frequency', 'Disturbance Magnitude', 'Phase Oscillation Magnitude', 'POC voltage', 'Grid SCR', 'Grid X/R',  'P at POC (MW)', 'Q at POC (MVAr)', 'DMAT id', 'passed']
            table1=report.add_table(rows=1, cols=len(headers))
            table1.style='ListTable3-Accent3'
            hdr_cells1=table1.rows[0].cells
            for header_id in range(0, len(headers)): hdr_cells1[header_id].text=headers[header_id]
            # read test metadata to get test info
            for case_id in range(0, len(cases)):
                if ('ort' in cases[case_id]):
                    dataset_number=0
                    test_details={}
                    while (dataset_number < len(datasets)) and (test_details=={}):
                        try:
                            test_details=shelve.open(main_folder_path_out+"\\"+datasets[dataset_number]['path']+"\\"+cases[case_id]+"\\testInfo\\"+cases[case_id])
                        except:
                            dataset_number+=1
                    row_cells=table1.add_row().cells
                    cell_paragraph=row_cells[0].paragraphs[0]
                    add_link(paragraph=cell_paragraph, link_to=str(cases[case_id]), text=str(cases[case_id]), tool_tip="link to test results")
                    row_cells[1].text=str(test_details['scenario_params']['Test Type'])
                    if('time' in test_details['scenario_params'].keys()):
                        row_cells[2].text=str(test_details['scenario_params']['time'])
                    else:
                        row_cells[3].text='0.0'
                    row_cells[3].text=str(test_details['scenario_params']['Disturbance Frequency'])
                    row_cells[4].text=str(round(test_details['scenario_params']['Disturbance Magnitude'],2))
                    row_cells[5].text=str(test_details['scenario_params']['PhaseOsc Magnitude'])
                    row_cells[6].text=str(test_details['setpoint']['V_POC'])
#                    row_cells[7].text=str(test_details['setpoint']['GridMVA'])
                    row_cells[7].text=str(round(test_details['setpoint']['SCR'],2))
                    row_cells[8].text=str(test_details['setpoint']['X_R'])
                    row_cells[9].text=str(test_details['setpoint']['P'])
                    row_cells[10].text=str(test_details['setpoint']['Q'])
                    if('PSS/E' in software_types):
                        if ('TimeStep' in test_details['scenario_params'].keys()):
                            row_cells[11].text=str(test_details['scenario_params']['TimeStep'])
                        elif('time step' in test_details['scenario_params'].keys()):
                            row_cells[11].text=str(test_details['scenario_params']['time step'])
                            row_cells[12].text=str(test_details['scenario_params']['AccFactor'])
                    row_cells[13].text=str(test_details['scenario_params']['Comment/Corresponding DMAT case'])
            for row in table1.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size= Pt(9)
                        
            tableCnt+=1
            #run.add_break()
            p=report.add_paragraph(" ")
            p=report.add_paragraph('')
        if(any('tov' in case for case in cases)):
            p.add_run('Table '+str(tableCnt)+': Scenario list - Temporary Over-Voltage tests').bold=True        
            if('PSS/E' in software_types):
#                headers=['Case Nr.', 'Test Type', 'Time', 'TOV duration', 'Capacity(F)', 'V residual', 'POC voltage', 'Grid MVA', 'Grid X/R', 'Grid MVA post-TOV','Grid X/R post-TOV', 'P at POC (MW)', 'Q at POC (MVAr)', 'PSSE time step', 'PSSE acc. factor', 'passed']
                headers=['Case Nr.', 'Test Type', 'Time', 'TOV duration', 'Capacity(F)', 'V residual', 'POC voltage', 'Grid SCR', 'Grid X/R', 'Grid SCR post-TOV','Grid X/R post-TOV', 'P at POC (MW)', 'Q at POC (MVAr)', 'PSSE time step (ms)', 'PSSE acc. factor', 'DMAT id', 'passed']
            else:
#                headers=['Case Nr.', 'Test Type', 'Time', 'TOV duration', 'Capacity(F)', 'V residual', 'POC voltage', 'Grid MVA', 'Grid X/R', 'Grid MVA post-TOV','Grid X/R post-TOV', 'P at POC (MW)', 'Q at POC (MVAr)', 'passed']
                headers=['Case Nr.', 'Test Type', 'Time', 'TOV duration', 'Capacity(F)', 'V residual', 'POC voltage', 'Grid SCR', 'Grid X/R', 'Grid SCR post-TOV','Grid X/R post-TOV', 'P at POC (MW)', 'Q at POC (MVAr)', 'DMAT id', 'passed']
            table1=report.add_table(rows=1, cols=len(headers))
            table1.style='ListTable3-Accent3'
            hdr_cells1=table1.rows[0].cells
            for header_id in range(0, len(headers)): hdr_cells1[header_id].text=headers[header_id]
            # read test metadata to get test info
            for case_id in range(0, len(cases)):
                if ('tov' in cases[case_id]):
                    dataset_number=0
                    test_details={}
                    while (dataset_number < len(datasets)) and (test_details=={}):
                        try:
                            test_details=shelve.open(main_folder_path_out+"\\"+datasets[dataset_number]['path']+"\\"+cases[case_id]+"\\testInfo\\"+cases[case_id])
                        except:
                            dataset_number+=1
                    row_cells=table1.add_row().cells
                    cell_paragraph=row_cells[0].paragraphs[0]
                    add_link(paragraph=cell_paragraph, link_to=str(cases[case_id]), text=str(cases[case_id]), tool_tip="link to test results")
                    row_cells[1].text=str(test_details['scenario_params']['Test Type'])
                    row_cells[2].text=str(test_details['scenario_params']['time'])
                    row_cells[3].text=str(test_details['scenario_params']['Fduration'])
                    if('Capacity(F)' in test_details['scenario_params'].keys()):
                        if(test_details['scenario_params']['Capacity(F)']!=''):
                            row_cells[4].text=str(round(float(test_details['scenario_params']['Capacity(F)']),2))
                        else:row_cells[4].text='-'
                    else:row_cells[4].text='-'
                    row_cells[5].text=str(test_details['scenario_params']['Vresidual'])
                    row_cells[6].text=str(test_details['setpoint']['V_POC'])
#                    row_cells[7].text=str(test_details['setpoint']['GridMVA'])
                    row_cells[7].text=str(round(test_details['setpoint']['SCR'],2))
                    row_cells[8].text=str(test_details['setpoint']['X_R'])
                    if('SCL_post' in test_details['scenario_params']):
                        if(test_details['scenario_params']['SCL_post']!=''):
                            row_cells[9].text=str(test_details['scenario_params']['SCL_post'])
                        else:
#                            row_cells[9].text=str(test_details['setpoint']['GridMVA'])
                            row_cells[9].text=str(round(test_details['setpoint']['SCR'],2))
                    else:
#                        row_cells[9].text=str(test_details['setpoint']['GridMVA'])
                        row_cells[9].text=str(round(test_details['setpoint']['SCR'],2))
                    if('X_R_post' in test_details['scenario_params']):
                        if(test_details['scenario_params']['X_R_post']!=''):
                            row_cells[10].text=str(test_details['scenario_params']['X_R_post'])
                        else:
                            row_cells[10].text=str(test_details['setpoint']['X_R'])
                    else:
                        row_cells[10].text=str(test_details['setpoint']['X_R'])            
                    row_cells[11].text=str(test_details['setpoint']['P'])
                    row_cells[12].text=str(test_details['setpoint']['Q'])
                    if('PSS/E' in software_types):
                        row_cells[13].text=str(test_details['scenario_params']['TimeStep'])
                        row_cells[14].text=str(test_details['scenario_params']['AccFactor'])
                    row_cells[15].text=str(test_details['scenario_params']['Comment/Corresponding DMAT case'])
            for row in table1.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size= Pt(9)
                            
            tableCnt+=1    
            #run.add_break()
            p=report.add_paragraph(" ")
            p=report.add_paragraph('') 


        if(any('con' in case for case in cases)):
            p.add_run('Table '+str(tableCnt)+': Scenario list - Network Contingency tests').bold=True        

            if('PSS/E' in software_types):
#                headers=['Case Nr.', 'Test Type', 'Test profile', 'POC voltage', 'Grid MVA', 'Grid X/R', 'P at POC (MW)', 'Q at POC (MVAr)', 'PSSE time step', 'PSSE acc. factor', 'passed']
                headers=['Case Nr.', 'Test Type', 'Test profile', 'POC voltage', 'Grid SCR', 'Grid X/R', 'P at POC (MW)', 'Q at POC (MVAr)', 'PSSE time step (ms)', 'PSSE acc. factor', 'DMAT id', 'passed']
            else:
#                headers=['Case Nr.', 'Test Type', 'Test profile', 'POC voltage', 'Grid MVA', 'Grid X/R', 'P at POC (MW)', 'Q at POC (MVAr)', 'passed']
                headers=['Case Nr.', 'Test Type', 'Test profile', 'POC voltage', 'Grid SCR', 'Grid X/R', 'P at POC (MW)', 'Q at POC (MVAr)', 'DMAT id', 'passed']
            table=report.add_table(rows=1, cols=len(headers))
            table.style='ListTable3-Accent3'
            hdr_cells=table.rows[0].cells
            for header_id in range(0, len(headers)): hdr_cells[header_id].text=headers[header_id]
            # read test metadata to get test info
            for case_id in range(0, len(cases)):
                if ('small' in cases[case_id]):
                    dataset_number=0
                    test_details={}
                    while (dataset_number < len(datasets)) and (test_details=={}):
                        try:
                            test_details=shelve.open(main_folder_path_out+"\\"+datasets[dataset_number]['path']+"\\"+cases[case_id]+"\\testInfo\\"+cases[case_id])
                        except:
                            dataset_number+=1
                    row_cells=table.add_row().cells
                    cell_paragraph=row_cells[0].paragraphs[0]
                    add_link(paragraph=cell_paragraph, link_to=str(cases[case_id]), text=str(cases[case_id]), tool_tip="link to test results")
                    row_cells[1].text=str(test_details['scenario_params']['Test Type'])
                    row_cells[2].text=str(test_details['scenario_params']['Test profile'])
                    row_cells[3].text=str(test_details['setpoint']['V_POC'])
#                    row_cells[4].text=str(test_details['setpoint']['GridMVA'])
#                    row_cells[4].text=str(test_details['setpoint']['SCR'])
                    row_cells[4].text=str(round(test_details['setpoint']['SCR'],2))
                    row_cells[5].text=str(test_details['setpoint']['X_R'])
                    row_cells[6].text=str(test_details['setpoint']['P'])
                    row_cells[7].text=str(test_details['setpoint']['Q'])
                    if('PSS/E' in software_types):
                        row_cells[8].text=str(test_details['scenario_params']['TimeStep'])
                        row_cells[9].text=str(test_details['scenario_params']['AccFactor'])
#                    row_cells[10].text=str(test_details['scenario_params']['Comment/Corresponding DMAT case'])
                    row_cells[10].text=str(test_details['scenario_params']['comment'])
            for row in table.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size= Pt(9)
        

        p=report.add_paragraph('')        
        run=p.add_run()
        run.add_break(WD_BREAK.PAGE)
        change_orientation(report)
            
    report.add_heading("Simulation Results", level=1 )       
               
                    
        #generate table with scenarios and scenario details based on list of cases and metadata saved in 'data_location'
    
    return 0

def add_plots_to_report(case, report, datasets, plots, plot_list, assessment):
    software_type=''
    if(assessment['PSSE_flag']>0):
        software_type='PSS/E'
    #retrieve test info
    dataset_number=0
    test_details={}
    while (dataset_number < len(datasets)) and (test_details=={}):
        try:
            test_details=shelve.open(main_folder_path_out+"\\"+datasets[dataset_number]['path']+"\\"+case+"\\testInfo\\"+case)
        except:
            dataset_number+=1
    if("Comment/Corresponding DMAT case" in test_details['scenario_params'].keys()):
        if(test_details['scenario_params']["Comment/Corresponding DMAT case"]!=''):
            paragraph_temp=report.add_heading('Case '+case+' (DMAT '+str(test_details['scenario_params']["Comment/Corresponding DMAT case"])+')', level=2 )
        else:
            paragraph_temp=report.add_heading('Case '+case, level=2 ) 
    elif("Comment/corresponding DMAT case" in test_details['scenario_params'].keys()):
        if(test_details['scenario_params']["Comment/corresponding DMAT case"]!=''):
            paragraph_temp=report.add_heading('Case '+case+' (DMAT '+str(test_details['scenario_params']["Comment/corresponding DMAT case"])+')', level=2 )
        else:
            paragraph_temp=report.add_heading('Case '+case, level=2 ) 
    elif("comment/corresponding DMAT case" in test_details['scenario_params'].keys()):
        if(test_details['scenario_params']["comment/corresponding DMAT case"]!=''):
            paragraph_temp=report.add_heading('Case '+case+' (DMAT '+str(test_details['scenario_params']["comment/corresponding DMAT case"])+')', level=2 )
        else:
            paragraph_temp=report.add_heading('Case '+case, level=2 ) 
    elif("comment" in test_details['scenario_params'].keys()):
        if(test_details['scenario_params']["comment"]!=''):
            paragraph_temp=report.add_heading('Case '+case+' (DMAT '+str(test_details['scenario_params']["comment"])+')', level=2 )
        else:
            paragraph_temp=report.add_heading('Case '+case, level=2 ) 
    else:        
    #add level 2 headline with test name
        paragraph_temp=report.add_heading('Case '+case, level=2 ) 
    add_bookmark(paragraph=paragraph_temp, bookmark_text='', bookmark_name=case )#add bookmark that link in table at start of document points to
    #add summary table with test details
    table=report.add_table(rows=1, cols=2)
    table.style='ListTable3-Accent3'   
    headers=['Parameter', 'Value']
    hdr_cells=table.rows[0].cells
    for header_id in range(0, len(headers)): hdr_cells[header_id].text=headers[header_id]    
    if('small') in case:
        dataset_number=0
        test_details={}
        while (dataset_number < len(datasets)) and (test_details=={}):
            try:
                test_details=shelve.open(main_folder_path_out+"\\"+datasets[dataset_number]['path']+"\\"+case+"\\testInfo\\"+case)
            except:
                dataset_number+=1
        row_cells=table.add_row().cells
        row_cells[0].text='Test Type'
        row_cells[1].text=str(test_details['scenario_params']['Test Type'])
        row_cells=table.add_row().cells
        row_cells[0].text='Test profile'
        row_cells[1].text=str(test_details['scenario_params']['Test profile'])
        row_cells=table.add_row().cells
        row_cells[0].text='POC voltage (p.u.)'
        row_cells[1].text=str(test_details['setpoint']['V_POC'])
        row_cells=table.add_row().cells
#        row_cells[0].text='Grid Fault level (MVA)'
        row_cells[0].text='Grid short circuit ratio'
#        row_cells[1].text=str(test_details['setpoint']['GridMVA'])
        row_cells[1].text=str(round(test_details['setpoint']['SCR'],2))
        row_cells=table.add_row().cells
        row_cells[0].text='Grid X/R-ratio'
        row_cells[1].text=str(test_details['setpoint']['X_R'])
        row_cells=table.add_row().cells
        row_cells[0].text='P at POC (MW)'
        row_cells[1].text=str(test_details['setpoint']['P'])
        row_cells=table.add_row().cells
        row_cells[0].text='Q at POC (MVAr)'
        row_cells[1].text=str(test_details['setpoint']['Q'])
        if(software_type=='PSS/E'):
            row_cells=table.add_row().cells
            row_cells[0].text='PSSE time step (ms)'
            row_cells[1].text=str(test_details['scenario_params']['TimeStep'])
            row_cells=table.add_row().cells
            row_cells[0].text='PSSE acc. factor'
            row_cells[1].text=str(test_details['scenario_params']['AccFactor'])
    if('large') in case:
#        headers=['Case Nr.', 'Test Type', 'Fault Type', 'Fault time', ' Fault duration', 'Fault impedance', 'V residual', 'Fault X/R', 'POC voltage', 'Grid MVA', 'Grid X/R',  'P at POC (MW)', 'Q at POC (MVAr)', 'PSSE time step', 'PSSE acc. factor']
#        headers=['Case Nr.', 'Test Type', 'Fault Type', 'Fault time', ' Fault duration', 'Fault impedance', 'V residual', 'Fault X/R', 'POC voltage', 'Grid SCR', 'Grid X/R',  'P at POC (MW)', 'Q at POC (MVAr)', 'PSSE time step', 'PSSE acc. factor']
        headers=['Case Nr.', 'Test Type', 'Fault Type', 'Fault time', ' Fault duration', 'Fault impedance', 'V residual', 'Fault X/R', 'POC voltage', 'Grid SCR', 'Grid X/R',  'P at POC (MW)', 'Q at POC (MVAr)', 'PSSE time step (ms)', 'PSSE acc. factor']
        dataset_number=0
        test_details={}
        while (dataset_number < len(datasets)) and (test_details=={}):
            try:
                test_details=shelve.open(main_folder_path_out+"\\"+datasets[dataset_number]['path']+"\\"+case+"\\testInfo\\"+case)
            except:
                dataset_number+=1
        row_cells=table.add_row().cells
        row_cells[0].text='Test Type'
        row_cells[1].text=str(test_details['scenario_params']['Test Type'])

        if(not 'Multifault' in test_details['scenario_params']['Test Type'] ):#single fault test
            row_cells=table.add_row().cells
            row_cells[0].text='Fault Type'
            row_cells[1].text=str(test_details['scenario_params']['Event_Type'])
            row_cells=table.add_row().cells
            row_cells[0].text='Fault time (s)'
            row_cells[1].text=str(test_details['scenario_params']['Ftime'])
            row_cells=table.add_row().cells
            row_cells[0].text='Fault duration (s)'
            row_cells[1].text=str(test_details['scenario_params']['Fduration'])
            row_cells=table.add_row().cells
            row_cells[0].text='Fault impedance (Ohm)'
            row_cells[1].text=str(round(test_details['scenario_params']['F_Impedance'],2))
            row_cells=table.add_row().cells
            row_cells[0].text='V residual (p.u.)'
            if(test_details['scenario_params']['Vresidual']!=''):
                row_cells[1].text=str(test_details['scenario_params']['Vresidual'])
            else:
                row_cells[1].text='-'
            row_cells=table.add_row().cells
            row_cells[0].text='Fault X/R-ratio'
            row_cells[1].text=str(test_details['scenario_params']['Fault X_R'])
        else: #Multifault case
            if(len(test_details['scenario_params']['Event_Type'])>0):
                tmp=str(test_details['scenario_params']['Event_Type'][0])
#                for fault_id in range(0, len(test_details['scenario_params']['Ftype'])-1) :
                for fault_id in range(1, len(test_details['scenario_params']['Event_Type'])) : # #11/8/2022: correcting the fault list
                    tmp+=', '+str(test_details['scenario_params']['Event_Type'][fault_id])
                row_cells=table.add_row().cells
                row_cells[0].text='Fault Type'
                row_cells[1].text=tmp
            if(len(test_details['scenario_params']['Ftime'])>0):
                tmp=str(round(test_details['scenario_params']['Ftime'][0],3))
#                for fault_id in range(0, len(test_details['scenario_params']['Ftime'])-1) :
                for fault_id in range(1, len(test_details['scenario_params']['Ftime'])) :
                    tmp+=', '+str(round(test_details['scenario_params']['Ftime'][fault_id],3))
                row_cells=table.add_row().cells
                row_cells[0].text='Fault time (s)'
                row_cells[1].text=tmp
            if(len(test_details['scenario_params']['Fduration'])>0):
                tmp=str(test_details['scenario_params']['Fduration'][0])
#                for fault_id in range(0, len(test_details['scenario_params']['Fduration'])-1) :
                for fault_id in range(1, len(test_details['scenario_params']['Fduration'])) :
                    tmp+=', '+str(test_details['scenario_params']['Fduration'][fault_id])
                row_cells=table.add_row().cells
                row_cells[0].text='Fault duration (s)'
                row_cells[1].text=tmp
            if(len(test_details['scenario_params']['F_Impedance'])>0):
                tmp=str(round(test_details['scenario_params']['F_Impedance'][0],2))
                for fault_id in range(0, len(test_details['scenario_params']['F_Impedance'])-1) :
#                for fault_id in range(1, len(test_details['scenario_params']['F_Impedance'])) :
                    tmp+=', '+str(round(test_details['scenario_params']['F_Impedance'][fault_id],2))
                row_cells=table.add_row().cells
                row_cells[0].text='Fault impedance (Ohm)'
                row_cells[1].text=tmp
            if(len(test_details['scenario_params']['Vresidual'])>0):
                tmp=str(test_details['scenario_params']['Vresidual'][0])
#                for fault_id in range(0, len(test_details['scenario_params']['Vresidual'])-1) :
                for fault_id in range(1, len(test_details['scenario_params']['Vresidual'])) :
                    tmp+=', '+str(test_details['scenario_params']['Vresidual'][fault_id])
                row_cells=table.add_row().cells
                row_cells[0].text='V residual (p.u.)'
                row_cells[1].text=tmp
    
        row_cells=table.add_row().cells
        row_cells[0].text='POC voltage (p.u.)'
        row_cells[1].text=str(test_details['setpoint']['V_POC'])
        row_cells=table.add_row().cells
#        row_cells[0].text='Grid Fault Level (MVA)'
        row_cells[0].text='Grid short circuit ratio'
#        row_cells[1].text=str(test_details['setpoint']['GridMVA'])
        row_cells[1].text=str(round(test_details['setpoint']['SCR'],2))
        row_cells=table.add_row().cells
        row_cells[0].text='Grid X/R-ratio'
        row_cells[1].text=str(test_details['setpoint']['X_R'])
        row_cells=table.add_row().cells
#        row_cells[0].text='Grid Fault Level post-fault (MVA)'
        row_cells[0].text='Grid short circuit ratio post-fault'
        if('SCL_post' in test_details['scenario_params']):
            if(test_details['scenario_params']['SCL_post']!=''):
                row_cells[1].text=str(test_details['scenario_params']['SCL_post'])
#            else: row_cells[1].text=str(test_details['setpoint']['GridMVA'])
            else: row_cells[1].text=str(round(test_details['setpoint']['SCR'],2))
#        else: row_cells[1].text=str(test_details['setpoint']['GridMVA'])
        else: row_cells[1].text=str(round(test_details['setpoint']['SCR'],2))
        row_cells=table.add_row().cells
        row_cells[0].text='Grid X/R-ratio post-fault'
        if('X_R_post' in test_details['scenario_params']):
            if(test_details['scenario_params']['X_R_post']!=''): 
                row_cells[1].text=str(test_details['scenario_params']['X_R_post'])
            else:row_cells[1].text=str(test_details['setpoint']['X_R'])
        else:row_cells[1].text=str(test_details['setpoint']['X_R'])

        row_cells=table.add_row().cells
        row_cells[0].text='P at POC (MW)'
        row_cells[1].text=str(test_details['setpoint']['P'])
        row_cells=table.add_row().cells
        row_cells[0].text='Q at POC (MVAr)'
        row_cells[1].text=str(test_details['setpoint']['Q'])
        if(software_type=='PSS/E'):
            row_cells=table.add_row().cells
            row_cells[0].text='PSSE time step'
            row_cells[1].text=str(test_details['scenario_params']['TimeStep'])
            row_cells=table.add_row().cells
            row_cells[0].text='PSSE acc. factor'
            row_cells[1].text=str(test_details['scenario_params']['AccFactor'])
            
    if('ort') in case:
#        headers=['Case Nr.', 'Test Type', 'Time', 'Disturbance Frequency', 'Disturbance Magnitude', 'Phase Oscillation Magnitude', 'POC voltage', 'Grid MVA', 'Grid X/R',  'P at POC (MW)', 'Q at POC (MVAr)', 'PSSE time step', 'PSSE acc. factor']
        headers=['Case Nr.', 'Test Type', 'Time', 'Disturbance Frequency', 'Disturbance Magnitude', 'Phase Oscillation Magnitude', 'POC voltage', 'Grid SCR', 'Grid X/R',  'P at POC (MW)', 'Q at POC (MVAr)', 'PSSE time step (ms)', 'PSSE acc. factor']
        dataset_number=0
        test_details={}
        while (dataset_number < len(datasets)) and (test_details=={}):
            try:
                test_details=shelve.open(main_folder_path_out+"\\"+datasets[dataset_number]['path']+"\\"+case+"\\testInfo\\"+case)
            except:
                dataset_number+=1
        row_cells=table.add_row().cells
        row_cells[0].text='Test Type'
        row_cells[1].text=str(test_details['scenario_params']['Test Type'])
        if('time' in test_details['scenario_params'].keys()):
            row_cells=table.add_row().cells
            row_cells[0].text='Fault time (s)'
            row_cells[1].text=str(test_details['scenario_params']['time'])
        row_cells=table.add_row().cells
        row_cells[0].text='Disturbance Frequency'
        row_cells[1].text=str(test_details['scenario_params']['Disturbance Frequency'])
        row_cells=table.add_row().cells
        row_cells[0].text='Disturbance Magnitude'
        row_cells[1].text=str(round(test_details['scenario_params']['Disturbance Magnitude'],2))
        row_cells=table.add_row().cells
        row_cells[0].text='Phase Oscillation Magnitude'
        row_cells[1].text=str(test_details['scenario_params']['PhaseOsc Magnitude'])
        row_cells=table.add_row().cells
        row_cells[0].text='POC voltage (p.u.)'
        row_cells[1].text=str(test_details['setpoint']['V_POC'])
        row_cells=table.add_row().cells
#        row_cells[0].text='Grid Fault Level (MVA)'
        row_cells[0].text='Grid short circuit ratio'
#        row_cells[1].text=str(test_details['setpoint']['GridMVA'])
        row_cells[1].text=str(round(test_details['setpoint']['SCR'],2))
        row_cells=table.add_row().cells
        row_cells[0].text='Grid X/R-ratio'
        row_cells[1].text=str(test_details['setpoint']['X_R'])
        row_cells=table.add_row().cells
        row_cells[0].text='P at POC (MW)'
        row_cells[1].text=str(test_details['setpoint']['P'])
        row_cells=table.add_row().cells
        row_cells[0].text='Q at POC (MVAr)'
        row_cells[1].text=str(test_details['setpoint']['Q'])
        if(software_type=='PSS/E'):
            row_cells=table.add_row().cells
            row_cells[0].text='PSSE time step'
            if ('TimeStep' in test_details['scenario_params'].keys()):
                row_cells[1].text=str(test_details['scenario_params']['TimeStep'])
            elif ('time step' in test_details['scenario_params'].keys()):
                row_cells[1].text=str(test_details['scenario_params']['time step'])
            row_cells=table.add_row().cells
            row_cells[0].text='PSSE acc. factor'
            row_cells[1].text=str(test_details['scenario_params']['AccFactor'])
            
    if('tov') in case:
#        headers=['Case Nr.', 'Test Type', 'Time', 'TOV duration', 'Capacity(F)', 'V residual', 'POC voltage', 'Grid MVA', 'Grid X/R', 'Grid MVA post-TOV','Grid X/R post-TOV', 'P at POC (MW)', 'Q at POC (MVAr)', 'PSSE time step', 'PSSE acc. factor']
        headers=['Case Nr.', 'Test Type', 'Time', 'TOV duration', 'Capacity(F)', 'V residual', 'POC voltage', 'Grid SCR', 'Grid X/R', 'Grid SCR post-TOV','Grid X/R post-TOV', 'P at POC (MW)', 'Q at POC (MVAr)', 'PSSE time step (ms)', 'PSSE acc. factor']
        dataset_number=0
        test_details={}
        while (dataset_number < len(datasets)) and (test_details=={}):
            try:
                test_details=shelve.open(main_folder_path_out+"\\"+datasets[dataset_number]['path']+"\\"+case+"\\testInfo\\"+case)
            except:
                dataset_number+=1
        row_cells=table.add_row().cells
        row_cells[0].text='Test Type'
        row_cells[1].text=str(test_details['scenario_params']['Test Type'])
        row_cells=table.add_row().cells
        row_cells[0].text='Time (s)'
        row_cells[1].text=str(test_details['scenario_params']['time'])
        row_cells=table.add_row().cells
        row_cells[0].text='TOV duration (s)'
        row_cells[1].text=str(test_details['scenario_params']['Fduration'])
        if('Capacity(F)' in test_details['scenario_params'].keys()):
            if(test_details['scenario_params']['Capacity(F)']!=''):
                row_cells=table.add_row().cells
                row_cells[0].text='Capacity (uF)'
                row_cells[1].text=str(round(1000000*test_details['scenario_params']['Capacity(F)'],2))
        row_cells=table.add_row().cells
        row_cells[0].text='V residual (p.u.)'
        row_cells[1].text=str(test_details['scenario_params']['Vresidual'])
        row_cells=table.add_row().cells
        row_cells[0].text='POC voltage (p.u.)'
        row_cells[1].text=str(test_details['setpoint']['V_POC'])
        row_cells=table.add_row().cells
#        row_cells[0].text='Grid Fault Level (MVA)'
        row_cells[0].text='Grid short circuit ratio'
#        row_cells[1].text=str(test_details['setpoint']['GridMVA'])
        row_cells[1].text=str(round(test_details['setpoint']['SCR'],2))
        row_cells=table.add_row().cells
        row_cells[0].text='Grid X/R-ratio'
        row_cells[1].text=str(test_details['setpoint']['X_R'])
        row_cells=table.add_row().cells
        row_cells[0].text='Grid Fault Level post-TOV (MVA)'
        if('SCL_post' in test_details['scenario_params']):
            if(test_details['scenario_params']['SCL_post']!=''):
                row_cells[1].text=str(test_details['scenario_params']['SCL_post'])
#            else: row_cells[1].text=str(test_details['setpoint']['GridMVA'])
            else: row_cells[1].text=str(round(test_details['setpoint']['SCR'],2))
#        else: row_cells[1].text=str(test_details['setpoint']['GridMVA'])
        else: row_cells[1].text=str(round(test_details['setpoint']['SCR'],2))
        row_cells=table.add_row().cells
        row_cells[0].text='Grid X/R-ratio post-TOV'
        if('X_R_post' in test_details['scenario_params']):
            if(test_details['scenario_params']['X_R_post']!=''): 
                row_cells[1].text=str(test_details['scenario_params']['X_R_post'])
            else:row_cells[1].text=str(test_details['setpoint']['X_R'])
        else:row_cells[1].text=str(test_details['setpoint']['X_R'])
        row_cells=table.add_row().cells
        row_cells[0].text='P at POC (MW)'
        row_cells[1].text=str(test_details['setpoint']['P'])
        row_cells=table.add_row().cells
        row_cells[0].text='Q at POC (MVAr)'
        row_cells[1].text=str(test_details['setpoint']['Q'])
        if(software_type=='PSS/E'):
            row_cells=table.add_row().cells
            row_cells[0].text='PSSE time step'
            row_cells[1].text=str(test_details['scenario_params']['TimeStep'])
            row_cells=table.add_row().cells
            row_cells[0].text='PSSE acc. factor'
            row_cells[1].text=str(test_details['scenario_params']['AccFactor'])
            
    #add comment generated from automated assessment of the results (optional)
    #add Graphs (with titles?)
    for plot_id in range(0, len(plot_list)): 
        if plot_list[plot_id] in plots.keys():
            report.add_heading(plot_list[plot_id], level=3 )
            report.add_picture(plots[plot_list[plot_id]], Inches(6)) #normal graphs smaller
    # add page break
    p=report.add_paragraph('')
    run=p.add_run()
    run.add_break(WD_BREAK.PAGE)     
    return 0

def add_plots_to_report_nw(case, report, datasets, plots, plot_list, assessment):
    software_type=''
    if(assessment['PSSE_flag']>0):
        software_type='PSS/E'
    #retrieve test info
    dataset_number=0
    test_details={}
    while (dataset_number < len(datasets)) and (test_details=={}):
        try:
            test_details=shelve.open(main_folder_path_out+"\\"+datasets[dataset_number]['path']+"\\"+case+"\\testInfo\\"+case)
        except:
            dataset_number+=1

  
    #add level 2 headline with test name
    paragraph_temp=report.add_heading('Case '+case, level=2 ) 
    add_bookmark(paragraph=paragraph_temp, bookmark_text='', bookmark_name=case )#add bookmark that link in table at start of document points to
    #add summary table with test details
    table=report.add_table(rows=1, cols=2)
    table.style='ListTable3-Accent3'   
    headers=['Parameter', 'Value']
    hdr_cells=table.rows[0].cells
    for header_id in range(0, len(headers)): hdr_cells[header_id].text=headers[header_id]    
    if('con') in case:
        dataset_number=0
        test_details={}
        while (dataset_number < len(datasets)) and (test_details=={}):
            try:
                test_details=shelve.open(main_folder_path_out+"\\"+datasets[dataset_number]['path']+"\\"+case+"\\testInfo\\"+case)
            except:
                dataset_number+=1
        row_cells=table.add_row().cells
        row_cells[0].text='Test Type'
        row_cells[1].text=str(test_details['scenario_params'][0]['Test Type'])
        row_cells=table.add_row().cells
        row_cells[0].text='Fault Type'
        row_cells[1].text=str(test_details['scenario_params'][0]['Event_Type'])
        row_cells=table.add_row().cells
        row_cells[0].text='Fault Element'
        row_cells[1].text=str(test_details['scenario_params'][0]['Event_Element'])
        row_cells=table.add_row().cells
        row_cells[0].text='Fault Description'
        row_cells[1].text=str(test_details['scenario_params'][0]['CaseDescription'])
        row_cells=table.add_row().cells
        row_cells[0].text='Fault From Bus'
        row_cells[1].text=str(test_details['scenario_params'][0]['i_bus'])           
        if test_details['scenario_params'][0]['j_bus'] != '':
            row_cells=table.add_row().cells
            row_cells[0].text='Fault To Bus'
            row_cells[1].text=str(test_details['scenario_params'][0]['j_bus'])   
#        row_cells=table.add_row().cells
#        row_cells[0].text='Fault Time'
#        row_cells[1].text=str(test_details['scenario_params'][0]['Ftime'])  
        row_cells=table.add_row().cells
        row_cells[0].text='Local Clearing Time (s)'
        row_cells[1].text=str(test_details['scenario_params'][0]['trip_near']) 
        if test_details['scenario_params'][0]['trip_far'] != '':
            row_cells=table.add_row().cells
            row_cells[0].text='Remote Clearing Time (s)'
            row_cells[1].text=str(test_details['scenario_params'][0]['trip_far'])         
        if test_details['scenario_params'][0]['arc_success'] != '' and test_details['scenario_params'][0]['arc_time'] > 0:
            row_cells=table.add_row().cells
            row_cells[0].text='Auto-reclose Time (s)'
            row_cells[1].text=str(test_details['scenario_params'][0]['arc_time'])     
        if len(test_details['scenario_params']) > 1:
            for rb_id in range(1,len(test_details['scenario_params'])):
                row_cells=table.add_row().cells
                row_cells[0].text='Runback Action'
                if test_details['scenario_params'][rb_id]['j_bus'] != '':
                    row_cells[1].text=str(test_details['scenario_params'][rb_id]['CaseDescription']) +' ('+ str(test_details['scenario_params'][rb_id]['i_bus']) +'-'+ str(test_details['scenario_params'][rb_id]['j_bus'])+')'
                else:
                    row_cells[1].text=str(test_details['scenario_params'][rb_id]['CaseDescription']) +' ('+ str(test_details['scenario_params'][rb_id]['i_bus'])+')'
                    

#        if len(test_details['scenario_params']) > 1:
#            for rb_id in range(1,len(test_details['scenario_params'])):
#                row_cells=table.add_row().cells
#                row_cells[0].text='Runback Action'
#                row_cells[1].text=str(test_details['scenario_params'][rb_id]['CaseDescription']) +' at '+ str(test_details['scenario_params'][rb_id]['Ftime'])      
#                row_cells=table.add_row().cells
#                row_cells[0].text='From Bus'
#                row_cells[1].text=str(test_details['scenario_params'][rb_id]['i_bus'])           
#                if test_details['scenario_params'][rb_id]['j_bus'] != '':
#                    row_cells=table.add_row().cells
#                    row_cells[0].text='To Bus'
#                    row_cells[1].text=str(test_details['scenario_params'][rb_id]['j_bus'])   
                    
    #add comment generated from automated assessment of the results (optional)
    #add Graphs (with titles?)
    for plot_id in range(0, len(plot_list)): 
        if plot_list[plot_id] in plots.keys():
            report.add_heading(plot_list[plot_id], level=3 )
            report.add_picture(plots[plot_list[plot_id]], Inches(6)) #normal graphs smaller
    # add page break
    p=report.add_paragraph('')
    run=p.add_run()
    run.add_break(WD_BREAK.PAGE)     
    return 0

    
def change_orientation(document):
    global landscape_flag
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    
    if(landscape_flag==0):
        landscape_flag = 1
    else:
        landscape_flag=0

    return new_section    

        
def DMAT_report(info, batchname):
    #checks which data is available and generates all defined plots for which all data is available. If 'cases' is not empty, the plots (and report) are generated only for the cases listed in 'cases'. All Plots for which only partial data is availabel are left out (e.g. an overlay plot for which only the PSS/E results data but no PSCAD results data is available would not be generated)
    #create output folder
    output_loc=main_folder_path+"\\plots\\DMAT"
    for chapter_cnt in range(0,len(info)):
        chapter_info=info[chapter_cnt]
        datasets=chapter_info['datasets']
        relevant_cases=chapter_info['cases']# make this more sopphisticated to allow for ranges of cases to be defined in a more elegant manner
        all_cases=[] #all cases for which any data is available (may just be from a single dataset and not across multiple)
        for dataset_id in range(0, len(datasets)):#create list of results included in each of the listed datasets. for each plot check if data specified in the plot is available in the datasets. Only produce plots for which all data is available, otherwise produce an error message.
            cases_temp=next(os.walk(main_folder_path_out+"\\"+datasets[dataset_id]['path']))
            cases_temp=cases_temp[1]
            datasets[dataset_id]['cases']=sort_cases(cases_temp)
            for case in datasets[dataset_id]['cases']:
                if not case in all_cases:
                    all_cases.append(case)
        
        cases=[]
        if(relevant_cases==[]):
            cases=sort_cases(all_cases)
        else:
            for case in all_cases:
                if case in relevant_cases:
                    cases.append(case)
            cases=sort_cases(cases)
            
        if(chapter_info['report']==True):
            report=initialise_report('DMAT')
            add_report_intro(report, 'DMAT', datasets , cases) #This should include summary table of the test cases the pass/no pass column can be left empty and filled in manually while reviewing the results. 

        for case in cases:
            plots, assessment = generate_plots(case, output_loc, chapter_info, batchname) #"info" contains dataset information as well
            if(plots==-1):
                print(assessment)
            if(chapter_info['report']==True):
                add_plots_to_report(case, report, datasets, plots, chapter_info['plots_for_report'], assessment)
        if( chapter_info['report']==True):
           reportname= str(datetime.datetime.now().strftime("%Y%m%d-%H%M"))+"-"+str(ProjectDetailsDict['NameShrt'])+"_"+str(batchname)+"_report.docx"
           report.save(main_folder_path+"\\Plots\\DMAT\\"+reportname)
            
            
#Generate Benchmarking plots and report
def benchmarking_report(info, batchname):

    for chapter_cnt in range(0,len(info)):
        chapter_info=info[chapter_cnt]
        datasets=chapter_info['datasets']
        relevant_cases=chapter_info['cases']# make this more sopphisticated to allow for ranges of cases to be defined in a more elegant manner
        all_cases=[] #all cases for which any data is available (may just be from a single dataset and not across multiple)
        overlap_cases=[]
        for dataset_id in range(0, len(datasets)):#create list of results included in each of the listed datasets. for each plot check if data specified in the plot is available in the datasets. Only produce plots for which all data is available, otherwise produce an error message.
            datasets[dataset_id]['cases']=sort_cases(next(os.walk(main_folder_path_out+"\\"+datasets[dataset_id]['path']))[1])
            for case in datasets[dataset_id]['cases']:
                if not case in all_cases:
                    all_cases.append(case)
        for case in all_cases:
            n_datasets_including_case=0
            for dataset_id in range(0, len(datasets)):
                if case in datasets[dataset_id]['cases']:
                   n_datasets_including_case+=1
            if(n_datasets_including_case>=1):
                overlap_cases.append(case)
                
        if(relevant_cases==[]):
            benchmark_cases=sort_cases(overlap_cases)
        else:
            benchmark_cases=[]
            for case in all_cases:
                if case in relevant_cases:
                    benchmark_cases.append(case)
            benchmark_cases=sort_cases(benchmark_cases)
        

        #create output folder
        output_loc=main_folder_path+"\\plots\\Overlays"
        if(chapter_info['report']==True):
            report=initialise_report('BENCHMARKING')
            add_report_intro(report, 'BENCHMARKING', datasets , benchmark_cases) #This should include summary table of the test cases the pass/no pass column can be left empty and filled in manually while reviewing the results. 
        for case in benchmark_cases:
            plots, assessment = generate_plots(case, output_loc, chapter_info, batchname, x_range='common') #"info" contains dataset information as well
            if(plots==-1):
                print(assessment)
            if(chapter_info['report']==True):
                add_plots_to_report(case, report, datasets, plots, chapter_info['plots_for_report'], assessment)
    if(chapter_info['report']==True):       
        reportname= str(datetime.datetime.now().strftime("%Y%m%d-%H%M"))+"-"+str(ProjectDetailsDict['NameShrt'])+"_benchmarking_report.docx" 
        report.save(main_folder_path+"\\Plots\\Overlays\\"+reportname)

#Generates the partial GPS report, based on SMIB test data
def GPS_report(info):
    pass

def NetworkFault_report(info, batchname):
    #checks which data is available and generates all defined plots for which all data is available. If 'cases' is not empty, the plots (and report) are generated only for the cases listed in 'cases'. All Plots for which only partial data is availabel are left out (e.g. an overlay plot for which only the PSS/E results data but no PSCAD results data is available would not be generated)
    #create output folder
    output_loc=main_folder_path+"\\plots\\NetworkFault"
    for chapter_cnt in range(0,len(info)):
        chapter_info=info[chapter_cnt]
        datasets=chapter_info['datasets']
        relevant_cases=chapter_info['cases']# make this more sopphisticated to allow for ranges of cases to be defined in a more elegant manner
        all_cases=[] #all cases for which any data is available (may just be from a single dataset and not across multiple)
        for dataset_id in range(0, len(datasets)):#create list of results included in each of the listed datasets. for each plot check if data specified in the plot is available in the datasets. Only produce plots for which all data is available, otherwise produce an error message.
            cases_temp=next(os.walk(main_folder_path_out+"\\"+datasets[dataset_id]['path']))
            cases_temp=cases_temp[1]
            datasets[dataset_id]['cases']=sort_cases(cases_temp)
            for case in datasets[dataset_id]['cases']:
                if not case in all_cases:
                    all_cases.append(case)
        
        cases=[]
        if(relevant_cases==[]):
            cases=sort_cases(all_cases)
        else:
            for case in all_cases:
                if case in relevant_cases:
                    cases.append(case)
            cases=sort_cases(cases)
            
        if(chapter_info['report']==True):
            report=initialise_report('NetworkFault')
            add_report_intro(report, 'NetworkFault', datasets , cases) #This should include summary table of the test cases the pass/no pass column can be left empty and filled in manually while reviewing the results. 

        for case in cases:
            plots, assessment = generate_plots(case, output_loc, chapter_info, batchname) #"info" contains dataset information as well
            if(plots==-1):
                print(assessment)
            if(chapter_info['report']==True):
                add_plots_to_report_nw(case, report, datasets, plots, chapter_info['plots_for_report'], assessment)
        if( chapter_info['report']==True):
           reportname= str(datetime.datetime.now().strftime("%Y%m%d-%H%M"))+"-"+str(ProjectDetailsDict['NameShrt'])+"_"+str(batchname)+"_report.docx"
           report.save(main_folder_path+"\\Plots\\NetworkFault\\"+reportname)
            
def generate_plots(case, output_loc, info, batchname, x_range='max'): #xrange can be 'max', 'common' or can be user-defined as a range
    #read all relevant datasets (possibly cross-check which ones are needed for specified plots, to reduce loading time
    assessment={}
    assessment['PSSE_flag']=0
    plots={}
    plot = ep.ESCOPlot()
    dataset_names={}
    dataset_pos={} #this stores the position of the datasets within the report declaration in the header of this script (addressable by ID)
    dataset_pos_in_mem={} #this stores the position of the datasets within the plot object (addressable by ID)
    dataset_in_mem_cnt=0
    dataset_aux_cnt=0
    for dataset_id in range(0, len(info['datasets'])):
        dataset=info['datasets'][dataset_id]
        # check if data exists for the given case and if not, don't load.
        try:           
            plot.read_data(main_folder_path_out+"\\"+dataset['path']+"\\"+case+"\\"+case+'_results.csv', timeID=dataset['timeID'])
            dataset_pos_in_mem[info['datasets'][dataset_id]['ID']]=dataset_in_mem_cnt
            if('PSSE' in dataset['path']):
                assessment['PSSE_flag']=1
            dataset_in_mem_cnt+=1
            dataset_names[info['datasets'][dataset_id]['ID']]=dataset['label']
            dataset_pos[info['datasets'][dataset_id]['ID']]=dataset_id
            
    #        dataset_names[dataset_id]=dataset['label']
            plot.timeoffset[dataset_id]=dataset['timeoffset']
            
            if 'calcCurrents' in dataset.keys():
                    calcCurrents=dataset['calcCurrents']
                    for calcCurrentCnt in range (0, len(calcCurrents)):
                        calcCurrent=calcCurrents[calcCurrentCnt]
                        if('I' in calcCurrent.keys()):
                            currents=plot.calcCurrents(dataset_aux_cnt, P=calcCurrent['P'], Q=calcCurrent['Q'], I=calcCurrent['I'], nameLabel=calcCurrent['nameLabel'], scaling=calcCurrent['scaling'])
                        else:
                            currents=plot.calcCurrents(dataset_aux_cnt, P=calcCurrent['P'], Q=calcCurrent['Q'], V=calcCurrent['V'], nameLabel=calcCurrent['nameLabel'], scaling=calcCurrent['scaling'])
            dataset_aux_cnt+=1
        except IOError:
            print(str(main_folder_path_out+"\\"+dataset['path']+"\\"+case+"\\"+case+'_results.csv')+' is not available.')
                        
        
    time_range=plot.check_min_max_time() #returns ---> by default limit x-axis of plots to this value. 
    common_range=[max(time_range[0]), min(time_range[1])]
    max_range=[min(time_range[0]), max(time_range[1])]
#    settling_time = {}
    filename = 'Rise, settling and recovery times.csv'
    csvfile = output_loc + '\\' + filename
    f = open(csvfile, 'w')
#    f.write('Filename,Vsettle,Qsettle,Qrise,Vrise\n')
    f.write('Record rise, settling and recovery time\n')
    
    for plot_name in info['plots'].keys():
        #determine datasets required for the plot and traces required in the datasets. Then send request to plot script so see if the signals are available in the dataset 
        #--> if not, generate error message and don't plot, Otherwise proceed to plot
        proceed=True
        for subplot in info['plots'][plot_name].keys():
            for channel in info['plots'][plot_name][subplot]['channels']:
                if( (channel['dataset'] not in dataset_names.keys()) or not (case in info['datasets'][dataset_pos[channel['dataset']]]['cases']) ):
                    proceed=False # "Data missing - Plots could not be generated."
        if(proceed):
            print(plot_name)
            test=1
            plot_xrange=[] #this is 
            for subplot_ID in range(0, len(info['plots'][plot_name].keys())):
                for subplot in info['plots'][plot_name].keys():
                    if(info['plots'][plot_name][subplot]['rank']-1==subplot_ID) :
                        subplot_info=info['plots'][plot_name][subplot]
                        print(subplot)
                        dataset_number=0
                        test_details={}
                        while (dataset_number < len(info['datasets'])) and (test_details=={}):
                            try:
                                test_details=shelve.open(main_folder_path_out+"\\"+info['datasets'][dataset_number]['path']+"\\"+case+"\\testInfo\\"+case)#take metadata from case of first dataset (shoudl be the same across all datasets for a given case)
                            except:
                                dataset_number+=1
                        traces=info['plots'][plot_name][subplot]['channels']                        
                        subplot_legend=[]
                        for trace_ID in range(0, len(traces)):
                            print(subplot_info['channels'][trace_ID]['name'])
                            
                            if('linestyle' in subplot_info['channels'][trace_ID].keys() ):
                                linestyle=subplot_info['channels'][trace_ID]['linestyle']
                            else:
                                linestyle='-'
                            if('colour' in subplot_info['channels'][trace_ID].keys() ):
                                colour=subplot_info['channels'][trace_ID]['colour']
                            else:
                                colour=''
                            if('linewidth' in subplot_info['channels'][trace_ID].keys() ):
                                linewidth=subplot_info['channels'][trace_ID]['linewidth']
                            else:
                                linewidth=2.5
                            #check on which axis trace shouls be plotted
                            if('twinX' in subplot_info['channels'][trace_ID].keys()):
                                plot.subplot_spec(subplot_ID, (dataset_pos_in_mem[subplot_info['channels'][trace_ID]['dataset']], subplot_info['channels'][trace_ID]['name']), title=subplot,y2label=subplot_info['unit2'], scale=subplot_info['channels'][trace_ID]['scale'], offset=subplot_info['channels'][trace_ID]['offset'], colour=colour, linestyle=linestyle, linewidth=linewidth, twinX=subplot_info['channels'][trace_ID]['twinX'])
                            else:
                                if('GSMG' in subplot_info.keys() ): #if marker is set and trace belongs to the dataset 
                                    plot.subplot_spec(subplot_ID, (dataset_pos_in_mem[subplot_info['channels'][trace_ID]['dataset']], subplot_info['channels'][trace_ID]['name']), title=subplot,ylabel=subplot_info['unit'], scale=subplot_info['channels'][trace_ID]['scale'], offset=subplot_info['channels'][trace_ID]['offset'], colour=colour, linestyle=linestyle, linewidth=linewidth, markers=['GSMG'],)
                                elif('tolerance_bands' in subplot_info.keys()):
                                    if (trace_ID==subplot_info['tolerance_bands']['trace']):
                                        plot.subplot_spec(subplot_ID, (dataset_pos_in_mem[subplot_info['channels'][trace_ID]['dataset']], subplot_info['channels'][trace_ID]['name']), title=subplot,ylabel=subplot_info['unit'], scale=subplot_info['channels'][trace_ID]['scale'], offset=subplot_info['channels'][trace_ID]['offset'], colour=colour, linestyle=linestyle, linewidth=linewidth, markers=['tolerance_bands'],tolerance_band_offset=subplot_info['tolerance_bands']['percent']/100.0, tolerance_band_base=subplot_info['tolerance_bands']['base'])
                                    else:
                                        plot.subplot_spec(subplot_ID, (dataset_pos_in_mem[subplot_info['channels'][trace_ID]['dataset']], subplot_info['channels'][trace_ID]['name']), title=subplot,ylabel=subplot_info['unit'], scale=subplot_info['channels'][trace_ID]['scale'], offset=subplot_info['channels'][trace_ID]['offset'], colour=colour, linestyle=linestyle, linewidth=linewidth,)                                
                                else:
                                    plot.subplot_spec(subplot_ID, (dataset_pos_in_mem[subplot_info['channels'][trace_ID]['dataset']], subplot_info['channels'][trace_ID]['name']), title=subplot,ylabel=subplot_info['unit'], scale=subplot_info['channels'][trace_ID]['scale'], offset=subplot_info['channels'][trace_ID]['offset'], colour=colour, linestyle=linestyle, linewidth=linewidth,)                                

                            #add Legend: If not specified: only add legend to first subplot only, using the label of the respective dataset (in case there are multiple datasets). Otherwise add no legend.
    #                        if( (trace_ID==0) and not ('leg' in (subplot_info['channels'][trace_ID].keys())) ):
    #                            subplot_legend.append(info['datasets'][subplot_info['channels'][trace_ID]['dataset']]['label']) #this should return label of dataset to which trace belongs and add it as lagend for the first trace.
                            #if legend is explicitly specified for a trace, add legend for that trace.
                            if('leg' in subplot_info['channels'][trace_ID].keys() ): 
                                if(subplot_info['channels'][trace_ID]['leg']!=''):
                                    subplot_legend.append(subplot_info['channels'][trace_ID]['leg'])
                        #include settling bands if marker is set                   
                        if('GSMG' in subplot_info.keys() ): #only if marker is set and there is either a profile defined or a a fault test applied
                            for dataset_pos in range (0, len(info['datasets'])):
                                if (info['datasets'][dataset_pos]['ID']==subplot_info['channels'][subplot_info['GSMG']]['dataset']):
                                    timeoffset_0=dataset['timeoffset']
                                    GSMG_dataset_pos=dataset_pos #determine at which position in memory the plot object has stored the dataset for which GSMG shall be included.
                            #timeoffset_0=info['datasets'][subplot_info['channels'][subplot_info['GSMG']]['dataset']]['timeoffset'] #extract time offset at dataset to which settling band should be applied
                            if ('Test profile' in test_details['scenario_params'].keys()): #if scenario contains a test profile
                                test_profile=ProfilesDict[test_details['scenario_params']['Test profile']]
                                #test_profile=test_details['Test profile']
                                steps=detect_steps(subplot_info['channels'][subplot_info['GSMG']]['offset'], {'x':test_profile['x_data'], 'y':test_profile['y_data']})#whatever test profile is given 
                                if(len(steps)==1):
    #                                    endtime=#last entry in profile -0.105
                                    starttime=steps[0][0]-0.105+timeoffset_0
                                    endtime=test_profile['x_data'][-1]-0.105+timeoffset_0 #endtime is end of scenario -105 ms
                                    plot.GSMG_bands(GSMG_dataset_pos, subplot_info['channels'][subplot_info['GSMG']]['name'], starttime=starttime, endtime=endtime, startWindow=0.1, endWindow=0.1)
                                elif(len(steps)>1):
                                    starttime=steps[0][0]-0.105+timeoffset_0
                                    endtime=steps[1][0]-0.105+timeoffset_0 #endtime right before next step - 105 ms
                                    plot.GSMG_bands(subplot_info['GSMG'], subplot_info['channels'][subplot_info['GSMG']]['name'], starttime=starttime, endtime=endtime, startWindow=0.1, endWindow=0.1)
                            elif (test_details['scenario_params']['Test Type']=='Fault'):
                                starttime=test_details['scenario_params']['Ftime']-0.105+timeoffset_0
#                                endtime=starttime+0.8*test_details['scenario_params']['Fduration']-0.05+timeoffset_0 #leave 5 ms buffer. use as "endTime" the last 20% of the time where the fault is applied
                                endtime=starttime+0.8*test_details['scenario_params']['Fduration']-0.05 #leave 5 ms buffer. use as "endTime" the last 20% of the time where the fault is applied; timeoffset_0 is not needed to include into endtime as already included in starttime.
                                startWindow=0.1
                                endWindow=0.2*test_details['scenario_params']['Fduration']
                                plot.GSMG_bands(subplot_info['GSMG'], subplot_info['channels'][subplot_info['GSMG']]['name'], starttime=starttime, endtime=endtime, startWindow=startWindow, endWindow=endWindow)
                                pass


                        if('set_t' in subplot_info.keys() ): #only if marker is set and there is either a profile defined or a a fault test applied
                            for dataset_pos in range (0, len(info['datasets'])):
                                if (info['datasets'][dataset_pos]['ID']==subplot_info['channels'][subplot_info['set_t']]['dataset']):
                                    timeoffset_0=dataset['timeoffset']
#                                    GSMG_dataset_pos=dataset_pos #determine at which position in memory the plot object has stored the dataset for which GSMG shall be included.
                            if (test_details['scenario_params']['Test Type']=='Fault') or 'fault' in test_details['scenario_params']: # Designed to estimate the settling time of fault current Iq
                                starttime=test_details['scenario_params']['Ftime']+test_details['scenario_params']['Fduration']+timeoffset_0 # starting at end of fault
#                                endtime=starttime+0.8*test_details['scenario_params']['Fduration']-0.05+timeoffset_0 #leave 5 ms buffer. use as "endTime" the last 20% of the time where the fault is applied
                                endtime=test_details['scenario_params']['Ftime']+test_details['scenario_params']['Fduration']+timeoffset_0 + 1 #1second after fault exit
                                startWindow=0.1
                                endWindow=1
                                tempsettime = plot.settleTime(subplot_info['set_t'], subplot_info['channels'][subplot_info['set_t']]['name'], starttime=starttime, endtime=endtime, startWindow=startWindow, endWindow=endWindow)
                                settletime = tempsettime - starttime
                                print "settingling period = {:2.4f}".format(settletime)
#                                settling_time{"settingling period"} = settletime
                                f.write('Case ID: {},Test Type: {},Channel: {},Settling Time: {},{:1.3f}\n'.format(case, test_details['scenario_params']['Test Type'], subplot_info['channels'][subplot_info['set_t']]['name'], settletime, settletime))
                                pass

                        if('rise_t' in subplot_info.keys() ): #only if marker is set and there is either a profile defined or a a fault test applied
                            for dataset_pos in range (0, len(info['datasets'])):
                                if (info['datasets'][dataset_pos]['ID']==subplot_info['channels'][subplot_info['rise_t']]['dataset']):
                                    timeoffset_0=dataset['timeoffset']
#                                    GSMG_dataset_pos=dataset_pos #determine at which position in memory the plot object has stored the dataset for which GSMG shall be included.
                            if (test_details['scenario_params']['Test Type']=='Fault') or 'fault' in test_details['scenario_params']: # Designed to estimate the settling time of fault current Iq
                                starttime=test_details['scenario_params']['Ftime']+timeoffset_0 # starting at begining of fault
                                endtime=starttime+0.8*test_details['scenario_params']['Fduration'] #use as "endTime" the last 20% of the time where the fault is applied
#                                startWindow=0.1
#                                endWindow=1
                                risetime = plot.qrise(subplot_info['rise_t'], subplot_info['channels'][subplot_info['rise_t']]['name'], starttime=starttime, endtime=endtime)
#                                settletime = tempsettime - starttime
                                print "rising period = {:2.4f}".format(risetime)
#                                settling_time{"settingling period"} = settletime
                                f.write('Case ID: {},Test Type: {},Channel: {},Rising Time: {},{:1.3f}\n'.format(case, test_details['scenario_params']['Test Type'], subplot_info['channels'][subplot_info['rise_t']]['name'], risetime, risetime))
                                pass
                            
                        if('rec_t' in subplot_info.keys() ): #only if marker is set and there is either a profile defined or a a fault test applied
                            for dataset_pos in range (0, len(info['datasets'])):
                                if (info['datasets'][dataset_pos]['ID']==subplot_info['channels'][subplot_info['rec_t']]['dataset']):
                                    timeoffset_0=dataset['timeoffset']
#                                    GSMG_dataset_pos=dataset_pos #determine at which position in memory the plot object has stored the dataset for which GSMG shall be included.
                            if (test_details['scenario_params']['Test Type']=='Fault') or 'fault' in test_details['scenario_params']: # Designed to estimate the settling time of fault current Iq
                                starttime=test_details['scenario_params']['Ftime']+timeoffset_0 # starting at end of fault
#                                endtime=starttime+0.8*test_details['scenario_params']['Fduration']-0.05+timeoffset_0 #leave 5 ms buffer. use as "endTime" the last 20% of the time where the fault is applied
                                endtime=test_details['scenario_params']['Ftime']+test_details['scenario_params']['Fduration']+timeoffset_0 #Disturbance ended
#                                startWindow=0.1
#                                endWindow=1
                                p_recovery = plot.prise(subplot_info['rec_t'], subplot_info['channels'][subplot_info['rec_t']]['name'],100, 'U_POC1',1, distStartTime=starttime, distEndTime=endtime)
                                p_recovery = p_recovery - endtime
                                print "p_recovery period = {:2.4f}".format(p_recovery)
#                                settling_time{"settingling period"} = settletime
                                f.write('Case ID: {},Test Type: {},Channel: {},p_recovery Time: {},{:1.3f}\n'.format(case, test_details['scenario_params']['Test Type'], subplot_info['channels'][subplot_info['rec_t']]['name'], p_recovery, p_recovery))

                                pass

                            
#                            plot.prise(self, entry, pchan, pmax=100, vchan=-1, vbase = 1.0, distStartTime=0, distEndTime=-1)
#                            GSMG_bands(self, entry, channel, starttime = 0.0, endtime = 10.0, startWindow = 1.0, endWindow = 5.0):
#                                plot.prise(subplot_info['rec_t'], subplot_info['channels'][subplot_info['rec_t']]['name'],100, subplot_info['channels']['U_POC1'],1, distStartTime=starttime, distEndTime=endtime)
#                                plot.GSMG_bands(subplot_info['GSMG'], subplot_info['channels'][subplot_info['GSMG']]['name'], starttime=starttime, endtime=endtime, startWindow=startWindow, endWindow=endWindow)
                            
#                        if('set_band' in subplot_info.keys() ): #only if marker is set and there is either a profile defined or a a fault test applied
                            
    #                    if('dV' in subplot_info.keys() ):
    #                        
    #                    if('dIq' in subplot_info.keys() ):
                                    
                        # if plot contains overlays, include labels of datasets --> include legend using the label of the datasets for that subplot
                        legendAll=''
                        for i in range(0, len(subplot_legend)  ):
                            legendAll+=subplot_legend[i]
                        if(subplot_ID==0):
                            if (legendAll=='') and len(subplot_info['channels'])>1:
                                if(subplot_info['channels'][0]['dataset'] != subplot_info['channels'][1]['dataset']): #if first two traces belong to different datasets it is likely that it is an overlay plot
                                    for i in range(0, len(subplot_info['channels'])):
                                        subplot_legend.append(dataset_names[subplot_info['channels'][i]['dataset']])
                                
                        plot.legends[subplot_ID] = subplot_legend#try and set legend per subplot if possible. If not possible, modify EscoPlot lib to be able to do that.
                        if('xrange' in subplot_info.keys()):
                            plot_xrange=subplot_info['xrange']
                        if('yrange' in subplot_info.keys()):
                            plot.ylimit[subplot_ID]=subplot_info['yrange']
                        elif('yminspan' in subplot_info.keys()):
                            plot.yspan[subplot_ID]=subplot_info['yminspan']
                        if('y2range' in subplot_info.keys()):
                            plot.y2limit[subplot_ID]=subplot_info['y2range'] 
                        elif('y2minspan' in subplot_info.keys()):
                            plot.y2span[subplot_ID]=subplot_info['y2minspan']
                        if('ymaxlim' in subplot_info.keys()):
                            plot.ymaxlim[subplot_ID]=subplot_info['ymaxlim']
                        if('yminlim' in subplot_info.keys()):
                            plot.yminlim[subplot_ID]=subplot_info['yminlim']
            output_folder=output_loc+"\\"+batchname
            try:
                os.mkdir(output_folder)
            except:
                print("plot folder already exists")
            else:
                print("plot directory created")
                
            if(plot_xrange!=[]):
                plot.xlimit=plot_xrange
            else:
                if(x_range=='max'):
                    plot.xlimit=max_range
                elif(x_range=='common'):
                    plot.xlimit=common_range
            
            imgdata=plot.plot(figname = output_folder+"\\"+case+'_'+plot_name, show = 0, legloc = 'best')#, markers={'settleTime':['Qout_MV (MW)']} ) #legloc likely =legend location
            plots[plot_name]=imgdata
            plot.clear_subplot_spec()
            plot.clear_ylabels()
            plot.clear_ylimits()
    f.close()

    return plots, assessment
            
        
        
    #iterate over all plot definitions
        #iterale over all subplot definitions
        #generate plots and metadata, add lables and legen, etc.
        #reset subplot definition
    
    

def generate_plot(case, output_loc, result_data, info, plot_type):
    if(plot_type=='BENCHMARKING'): #expecting result_data to contain two file paths for the two overlays
        assessment={}        
        plots={}        
        for plot_name in info['overlay_plots'].keys():
            print(plot_name)
            plot = ep.ESCOPlot()            
            dataset_names={}
            for i in [0,1]:
                for dataset in info['datasets'].keys():
                    if(info['datasets'][dataset]['ID']==i):
                        dataset_names[i]=dataset
                        plot.read_data(result_data[i], timeID=info['datasets'][dataset]['timeID'])    
                        plot.timeoffset[i]=info['datasets'][dataset]['timeoffset']
            test=1
            for subplot_id in range(0, len(info['overlay_plots'][plot_name].keys()) ) : 
                for subplot in info['overlay_plots'][plot_name].keys():
                    if(info['overlay_plots'][plot_name][subplot]['rank']-1==subplot_id):
                        subplot_info=info['overlay_plots'][plot_name][subplot]
                        print(subplot)                
                        duration_0=plot.subplot_spec(subplot_id, (0, subplot_info['channels'][0]['name']), title = subplot,  ylabel = subplot_info['unit'], scale=subplot_info['channels'][0]['scale'], offset = subplot_info['channels'][0]['offset'], markers=['GSMG'])
                        duration_1=plot.subplot_spec(subplot_id, (1, subplot_info['channels'][1]['name']), title = subplot,  ylabel = subplot_info['unit'], scale=subplot_info['channels'][1]['scale'], offset = subplot_info['channels'][1]['offset'], markers=['GSMG'])
                        #read test scenario metadata
                        test_details=shelve.open(os.path.dirname(result_data[0])+"\\testInfo\\"+case)
                        timeoffset_0=info['datasets'][dataset]['timeoffset']
                        if('GSMG' in subplot_info.keys() ): #only if marker is set and there is either a profile defined or a a fault test applied
                            if ('Test profile' in test_details['scenario_params'].keys()): #if scenario contains a test profile
                                test_profile=ProfilesDict[test_details['scenario_params']['Test profile']]
                                #test_profile=test_details['Test profile']
                                steps=detect_steps(subplot_info['channels'][subplot_info['GSMG']]['offset'], {'x':test_profile['x_data'], 'y':test_profile['y_data']})#whatever test profile is given 
                                if(len(steps)==1):
#                                    endtime=#last entry in profile -0.105
                                    starttime=steps[0][0]-0.105+timeoffset_0
                                    endtime=test_profile['x_data'][-1]-0.105+timeoffset_0 #endtime is end of scenario -105 ms
                                    plot.GSMG_bands(subplot_info['GSMG'], subplot_info['channels'][subplot_info['GSMG']]['name'], starttime=starttime, endtime=endtime, startWindow=0.1, endWindow=0.1)
                                elif(len(steps)>1):
                                    starttime=steps[0][0]-0.105+timeoffset_0
                                    endtime=steps[1][0]-0.105+timeoffset_0 #endtime right before next step - 105 ms
                                    plot.GSMG_bands(subplot_info['GSMG'], subplot_info['channels'][subplot_info['GSMG']]['name'], starttime=starttime, endtime=endtime, startWindow=0.1, endWindow=0.1)
                            elif (test_details['scenario_params']['Test Type']=='Fault'):
                                starttime=test_details['scenario_params']['Ftime']-0.105+timeoffset_0
                                endtime=starttime+0.8*test_details['scenario_params']['Fduration']-0.05+timeoffset_0 #leave 5 ms buffer. use as "endTime" the last 20% of the time where the fault is applied
                                startWindow=0.1
                                endWindow=0.2*test_details['scenario_params']['Fduration']
                                plot.GSMG_bands(subplot_info['GSMG'], subplot_info['channels'][subplot_info['GSMG']]['name'], starttime=starttime, endtime=endtime, startWindow=startWindow, endWindow=endWindow)
                                pass
                        common_range=max([duration_0, duration_1])
                        plot.xlimit=[0,common_range]        
                                
                            
            plot.legends[0] = [dataset_names[0], dataset_names[1]]  
                
            #plot.xlimit=[0,10] #cut x-axis at the point where the shorter dataset ends.
            
            #plot.ylimit = [[0.8,1.2],[], [],[]] #scale y-axis automatically, with option to scale manually in plot_info
            #modify plot routine to return hander that can be used to copy plot into word report.
            # create output folder
            if('batchname' in info.keys()):
                if(info['batchname']!=''):
                    output_folder=output_loc+"\\"+info['batchname']
            else:
                output_folder=output_loc+"\\"+dataset_names[0]+'_vs_'+dataset_names[1]
            try:
                os.mkdir(output_folder)
            except:
                print("plot folder already exists")
            else:
                print("plot directory created")
                        
            imgdata=plot.plot(figname = output_folder+"\\"+case+'_'+plot_name, show = 1, legloc = 'best')#, markers={'settleTime':['Qout_MV (MW)']} ) #legloc likely =legend location
            plots[plot_name]=imgdata        
        return plots, assessment
    
    if(plot_type=='MAT'):
        assessment={}
        plots={}     
        for plot_name in info['plots']:
            print(plot_name)
            plot = ep.ESCOPlot()
            dataset_names={}
            dataset=info['datasets'].keys()[0]
            plot.read_data(result_data, timeID=info['datasets'][dataset]['timeID'])
            plot.timeoffset[0]=info['datasets'][dataset]['timeoffset']
        
            plot_xrange=[]
            subplots=info['plots'][plot_name].keys()
            for subplot_ID in range(0, len(subplots) ):
                for subplot in subplots:
                    if(info['plots'][plot_name][subplot]['rank']==subplot_ID+1):
                        subplot_info=info['plots'][plot_name][subplot]
                        traces=info['plots'][plot_name][subplot]['channels']
                        
                        subplot_legend=[]
                        for trace_ID in range(0, len(traces)):
                            print(subplot_info['channels'][trace_ID]['name'])
                            plot.subplot_spec(subplot_ID, (0, subplot_info['channels'][trace_ID]['name']), title=subplot,ylabel=subplot_info['unit'], scale=subplot_info['channels'][trace_ID]['scale'], offset=subplot_info['channels'][trace_ID]['offset'])
                            subplot_legend.append(subplot_info['channels'][trace_ID]['leg'])
                        plot.legends[subplot_ID] = subplot_legend#try and set legend per subplot if possible. If not possible, modify EscoPlot lib to be able to do that.
                        if('xrange' in subplot_info.keys()):
                            plot_xrange=subplot_info['xrange']
            if('batchname' in info.keys()):
                if(info['batchname']!=''):
                    output_folder=output_loc+"\\"+info['batchname']
                else:
                    output_folder=output_loc+"\\"+info['datasets'].keys()[0] #just one dataset included for MAT
            try:
                os.mkdir(output_folder)
            except:
                print("plot folder already exists")
            else:
                print("plot directory created")
                
            if(plot_xrange!=[]):
                plot.xlimit=plot_xrange
                            
            imgdata=plot.plot(figname = output_folder+"\\"+case+'_'+plot_name, show = 1, legloc = 'best')#, markers={'settleTime':['Qout_MV (MW)']} ) #legloc likely =legend location
            plots[plot_name]=imgdata
        return plots, assessment
                 

def main():
#    if('BENCH' in reports.keys()):
#        benchmarking_report(reports['BENCH']['report_definition'], reports['BENCH']['batchname'])
#    if('MAT' in reports.keys()):
#        for batch in reports['MAT'].keys():
#            reports['MAT'][batch]['batchname']=batch
#            MAT_report(reports['MAT'][batch], batch)
    if('DMAT' in reports.keys()):
        DMAT_report(reports['DMAT']['report_definition'], reports['DMAT']['batchname'])

    elif('NetworkFault' in reports.keys()):
        NetworkFault_report(reports['NetworkFault']['report_definition'], reports['NetworkFault']['batchname'])
        
    pass        
        
        
if __name__ == "__main__":
    main()
        






