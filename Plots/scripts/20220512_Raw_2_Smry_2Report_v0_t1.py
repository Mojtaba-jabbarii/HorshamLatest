# -*- coding: utf-8 -*-
"""
Created on Fri Mar 11 10:04:13 2022

@author: Mani Aulakh

The script is WIP. The return dataframes and excel sheet output cane be used for reporting purposes. Functions and naming convetions requires update at this stage.
"""
import os
import sys
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
sys.path.append(r"C:\ProgramData\Anaconda2\Lib\site-packages")
sys.path.append(r"C:\Python27\Lib\site-packages")
import docx
import openpyxl
import re
#from io 
#import StringIO
try:
    from StringIO import StringIO
except ImportError:
    from io import StringIO
from io import BytesIO

from docx import Document, shape
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK
from docx import Document

from docx.enum.dml import MSO_THEME_COLOR_INDEX

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

import datetime
main_folder_path=os.path.dirname(os.path.dirname(os.path.dirname(__file__)))
#main_folder_path=C:\Users\Mervin Kall\Documents\GitHub\PowerSystemStudyTool\20220203_APE\
sys.path.append(main_folder_path+"\\PSSE_sim\\scripts\\Libs")
import readtestinfo

#==============================================================================
#USER INPUTS
#Allign these inputs with the PSSE_NEM_testing_SteadyState script
#==============================================================================
TestDefinitionSheet=r'20230828_SUM_TESTINFO_V1.xlsx'
raw_SS_result_folder = '20230919-123914_SS'
simulation_batch_label = 'SS' # same as SS analysis script
#cases = ['_HighLoad','_LowLoad'] # use _ in addittion to the input model names
  
summary_dfs={'HighLoad':{}, 'LowLoad':{},} # use input model names
probs={'voltage_levels':0, 'line_ldngs':0, 'vltg_fluc_gen_chng':0, 'vltg_fluc_lol':0} 
cases=[]
for case in summary_dfs.keys():
    cases.append('_'+case)


# Create result folder
def make_dir(dir_path, dir_name=""):
    dir_make = os.path.join(dir_path, dir_name)
    try:
        os.mkdir(dir_make)
    except OSError:
        pass
    
dir_path =  main_folder_path +"\\Plots\\steady_state"
make_dir(dir_path)


#==============================================================================
#Global variables
#==============================================================================
return_dict =  readtestinfo.readTestdef(main_folder_path+"\\test_scenario_definitions\\"+TestDefinitionSheet, ['ProjectDetails','ModelDetailsPSSE'])
ProjectDetailsDict = return_dict['ProjectDetails']
PSSEmodelDict = return_dict['ModelDetailsPSSE']


#two_up = os.path.dirname(os.path.dirname(main_folder_path)) # this may need to change to match when finalising the script it looks like this will be three up
result_sheet_path = main_folder_path +"\\PSSE_sim\\result_data\\steady_state\\" + raw_SS_result_folder + "\\Steady State Analysis Results.xlsx"



# =============================================================================
# Functions below are for conditional formatting i.e. to check the violation of voltage levels, line loadings and voltage fluctuations.
# =============================================================================

def hl_vltg_lvls_violation(val):
   color = 'red' if (val < 0.9 or val >1.1) else 'black'
   return 'color: %s' % color

def hl_line_ldng_violation(val):
    color = 'red' if (val >99.99) else 'black'
    return 'color: %s' % color

def hl_vltg_fluc_violation(val):
    color = 'red' if (val > 0.0299) else 'black'
    return 'color: %s' % color


# read inputs
# Iterate through the raw steady state analysis results file and create dataframes require for analysis based on high,low,genon and genoff scenarios.

genoff = '_off' 
genon = '_on'
def read_result_sheet():
    sheets_dict  = pd.read_excel(result_sheet_path, sheet_name = None)
    result_dfs={}
    for case in cases:
        for name,sheet in sheets_dict.items():
            if name == 'Voltage Level'+ case + genoff:
                result_dfs['df_vltg_lvl'+case+genoff] = pd.DataFrame(data=sheet, columns=['bus_name','PU']).round(3)
                result_dfs['df_vltg_lvl'+case+genoff].rename(columns = {'bus_name':'Bus Name','PU':'Voltage Level(pu) GenOFF'}, inplace = True)
            elif name == 'Voltage Level' + case + genon:
                result_dfs['df_vltg_lvl'+case+genon] = pd.DataFrame(data=sheet, columns=['bus_name','PU']).round(3)
                result_dfs['df_vltg_lvl'+case+genon].rename(columns = {'bus_name':'Bus Name','PU':'Voltage Level(pu) GenON'}, inplace = True)
            elif name == 'Line Loadings'+case+genoff:
                result_dfs['df_line_ldng' + case+genoff] = pd.DataFrame(data=sheet, columns = ['CaseNr','CaseRef','brch_name','Loading (%)']).round(3)
                result_dfs['df_line_ldng' + case+genoff].rename(columns = {'CaseNr':'Case Number','CaseRef': 'Case Reference','brch_name':'Branch Name','Loading (%)': 'Loading(%) GenOFF'}, inplace = True)
            elif name == 'Line Loadings'+case+genon:
                 result_dfs['df_line_ldng' + case+genon] = pd.DataFrame(data=sheet, columns = ['CaseNr','CaseRef','brch_name','Loading (%)']).round(3)
                 result_dfs['df_line_ldng' + case+genon].rename(columns = {'CaseNr':'Case Number','CaseRef': 'Case Reference','brch_name':'Branch Name','Loading (%)': 'Loading(%) GenON'}, inplace = True)
            elif name == 'Volt Fluc GenChg'+case+genon:
                result_dfs['df_vltfluc_genchg'+case+genon] = pd.DataFrame(data=sheet, columns = ['CaseNr','CaseRef','bus_name','PU','PU_final','VolDev (%)']).round(3)
                result_dfs['df_vltfluc_genchg'+case+genon].rename(columns = {'CaseNr':'Case Number','CaseRef': 'Case Reference','bus_name':'Bus Name','PU':'Voltage Level(pu) Prior','PU_final': 'Voltage Level(pu) After','VolDev (%)': 'Volt Fluc(%)'}, inplace = True)
            elif name == 'Volt Fluc Lol'+case+genon:
                result_dfs['df_vltfluc_lol'+case+genon] = pd.DataFrame(data=sheet, columns = ['CaseNr','CaseRef','bus_name','PU','VolDev (%)']).round(3)
                result_dfs['df_vltfluc_lol'+case+genon].rename(columns = {'CaseNr':'Case Number','CaseRef': 'Case Reference','bus_name':'Bus Name','PU':'Voltage Level(pu) GenON','VolDev (%)': 'Volt Fluc(%) GenON'}, inplace = True)
            elif name == 'Volt Fluc Lol'+case+genoff:
                result_dfs['df_vltfluc_lol'+case+genoff] = pd.DataFrame(data=sheet, columns = ['CaseNr','CaseRef','bus_name','PU','VolDev (%)']).round(3)
                result_dfs['df_vltfluc_lol'+case+genoff].rename(columns = {'CaseNr':'Case Number','CaseRef': 'Case Reference','bus_name':'Bus Name','PU':'Voltage Level(pu) GenOFF','VolDev (%)': 'Volt Fluc(%) GenOFF'}, inplace = True)
            elif name == 'Fault Level' + case + genon:
                result_dfs['df_fault_lvl'+case+genon] = pd.DataFrame(data=sheet, columns=['bus_name','bus_no','fault_0_genoff','fault_0_genon', 'fault_1_genoff','fault_1_genon']).round(3)
                result_dfs['df_fault_lvl'+case+genon].rename(columns = {'bus_name':'Bus Name','bus_no':'Bus No.','fault_0_genoff':'1ph Fault (MVA) GenOFF','fault_0_genon':'1ph Fault (MVA) GenON', 'fault_1_genoff':'3ph Fault (MVA) GenOFF','fault_1_genon':'3ph Fault (MVA) GenON'}, inplace = True)
            else:
                pass
        

    return result_dfs
   

# =============================================================================
# This function uses the steady state analysis results excel file and summarise the voltage level results with and without plant.
# This will return violated voltage level table, voltage level summary table and will save the voltage level figure.
# =============================================================================

def voltage_levels(result_dfs, summary_dfs):
    for case in summary_dfs.keys():
        summary_dfs[case]['voltage_levels']={}
        for key,value in result_dfs.items():
            if 'df_vltg_lvl_'+case+'_off' in key:
                #Summary table
                df_vlt_lvl = pd.merge(result_dfs['df_vltg_lvl_'+case+'_off'],result_dfs['df_vltg_lvl_'+case+'_on'], how = 'outer',on = 'Bus Name')
                pvt_df = pd.pivot_table(data = df_vlt_lvl,index = ['Bus Name'],values = ['Voltage Level(pu) GenON','Voltage Level(pu) GenOFF'] )
                pvt_df.style.applymap(hl_vltg_lvls_violation,subset = ['Voltage Level(pu) GenOFF','Voltage Level(pu) GenON']).format({'Voltage Level(pu) GenOFF':'{0:,.3f}','Voltage Level(pu) GenON':'{0:,.3f}'})
                summary_dfs[case]['voltage_levels']['summary'] = []
                summary_dfs[case]['voltage_levels']['summary'].append(df_vlt_lvl)
                
                # Voltage levels plot
                fig = plt.figure(figsize=(7,5))
                fig.add_axes([0.1,0.1,0.8,0.8])
                plt.scatter(df_vlt_lvl['Bus Name'], df_vlt_lvl['Voltage Level(pu) GenON'], label = 'VL(pu) GenON', color = 'k', marker = '*', s = 120)
                plt.scatter(df_vlt_lvl['Bus Name'], df_vlt_lvl['Voltage Level(pu) GenOFF'], label = 'VL(pu) GenOFF')
                plt.title('Voltage Levels', fontsize =12, color = 'black' )
                plt.ylabel('Voltage(pu)', fontsize = 10)
                plt.xlabel('Bus Names', fontsize = 10)
                plt.legend()
                plt.xticks(rotation = 45)
                plt.margins( tight = True)
                plt.grid()
                plt.minorticks_off()
                #plt.savefig(case +'voltage_levels'+ 'plot.png', bbox_inches = 'tight')
                #imgdata= StringIO.StringIO()
                #imgdata = StringIO()
                imgdata= BytesIO() # version issues
                plt.savefig(imgdata, bbox_inches = 'tight', dpi =200)
                #plt.savefig(imgdata,format = 'svg', dpi =200) # version issues
                summary_dfs[case]['voltage_levels']['plot'] = []
                summary_dfs[case]['voltage_levels']['plot'].append(imgdata)
                
                #Violation voltage level table
    
                vl_min_violation = df_vlt_lvl.loc[df_vlt_lvl['Voltage Level(pu) GenOFF']<0.9]
                vl_max_violation = df_vlt_lvl.loc[df_vlt_lvl['Voltage Level(pu) GenOFF']>1.1]
                vl_violation = pd.concat([vl_max_violation,vl_min_violation], axis =0)
                
                # function to check pass fail criterion
                def vl_check(row):
                    if row['Voltage Level(pu) GenOFF'] == row['Voltage Level(pu) GenON']:
                        val = 'yes'
                    elif (row['Voltage Level(pu) GenOFF'] <0.9) and (row['Voltage Level(pu) GenON'] > row['Voltage Level(pu) GenOFF']):
                        val = 'yes'
                    elif (row['Voltage Level(pu) GenOFF'] >1.1) and (row['Voltage Level(pu) GenON'] < row['Voltage Level(pu) GenOFF']):
                        val = 'yes'
                    else:
                        val = 'no'
                    return val
                
                if vl_violation.empty:
                    pass
                else:
                    vl_violation['Pass'] = vl_violation.apply(vl_check, axis = 1)
                #vl_violation.style.applymap(hl_vltg_lvls_violation,subset = ['Voltage Level(pu) GenOFF','Voltage Level(pu) GenON']).format({'Voltage Level(pu) GenOFF':'{0:,.3f}','Voltage Level(pu) GenON':'{0:,.3f}'})
                summary_dfs[case]['voltage_levels']['violations'] = []
                summary_dfs[case]['voltage_levels']['violations'].append(vl_violation)
                
                #Empty apendix
                summary_dfs[case]['voltage_levels']['appendix'] = []
                
            else:
                pass          

# =============================================================================
# This function summarise the line loading results for network normal and contingencies. This function requires results with and without the plant. 
# It will return the violation of lines table, line loading summary table and will the save the line loading image. 
# =============================================================================
      
def line_ldngs(result_dfs, summary_dfs):
    
    for case in summary_dfs.keys():
        summary_dfs[case]['line_loadings']={}
        for key,value in result_dfs.items():
            if 'df_line_ldng_'+case+'_off' in key:
                # line loading summary table
                
                df_high = pd.concat([result_dfs['df_line_ldng_'+case+'_off'],result_dfs['df_line_ldng_'+case+'_on']],axis = 1)
                df_high = df_high.T.drop_duplicates().T
                pvt_df_high = pd.pivot_table(data=df_high, index=['Case Reference', 'Branch Name'], values = ['Loading(%) GenON','Loading(%) GenOFF'], aggfunc = ['max'])
                #pvt_df_high.style.applymap(hl_line_ldng_violation,subset = [('max','Loading(%) GenON'),('max','Loading(%) GenOFF')]).format({'Loading(%) GenON':'{0:,.3f}','Loading(%) GenOFF':'{0:,.3f}'})
                summary_dfs[case]['line_loadings']['appendix'] = []
                summary_dfs[case]['line_loadings']['appendix'].append(pvt_df_high)
                
                # line loading plot
                fig = plt.figure(figsize=(7,5))
                fig.add_axes([0.1,0.1,0.8,0.8])
                l_profile_plot = df_high.loc[(df_high['Case Number'] == 'case0 (base)')]
                plt.scatter(l_profile_plot['Branch Name'], l_profile_plot['Loading(%) GenON'], label = 'Loading(%) GenON', color = 'k', marker = '*', s = 120)
                plt.scatter(l_profile_plot['Branch Name'], l_profile_plot['Loading(%) GenOFF'], label = 'Loading(%) GenOFF')
                plt.xticks(rotation = 45, size = 8, visible = True)
                plt.title('Line Loadings(%)', fontsize = 12)
                plt.ylabel('Loading(%)', fontsize = 10)
                plt.xlabel('Branch Name', fontsize = 10)
                plt.legend()
                plt.grid()
                plt.minorticks_off()
                #plt.savefig(case + 'line_loadings'+ 'plot.png', bbox_inches = 'tight')
                #imgdata= StringIO.StringIO()
                #imgdata= StringIO() # version issues
                imgdata= BytesIO() # version issues
                plt.savefig(imgdata, bbox_inches = 'tight', dpi =200)
                #plt.savefig(imgdata,format = 'svg', dpi =200) # version issues
                summary_dfs[case]['line_loadings']['plot'] = []
                summary_dfs[case]['line_loadings']['plot'].append(imgdata)
                
                # line loading violation table
                line_max_violation = df_high.loc[(df_high['Loading(%) GenOFF'] > 99.99) | (df_high['Loading(%) GenON'] > 99.99)]
                
                
                 # function to check pass fail criterion
                def ll_check(row):
                    if (row['Loading(%) GenOFF'] == row['Loading(%) GenON']):
                        val = 'yes'
                    elif (row['Loading(%) GenOFF'] > row['Loading(%) GenON']):
                        val = 'yes'
                    else:
                        val = 'no'
                    return val
                
                if line_max_violation.empty:
                    pass
                else:
                    line_max_violation['Pass'] = line_max_violation.apply(ll_check, axis = 1)
                #line_max_violation.style.applymap(hl_line_ldng_violation,subset = ['Loading(%) GenON','Loading(%) GenOFF']).format({'Loading(%) GenON':'{0:,.3f}','Loading(%) GenOFF':'{0:,.3f}'})
                summary_dfs[case]['line_loadings']['violations'] = []
                summary_dfs[case]['line_loadings']['violations'].append(line_max_violation)
                
                #Line Loading Summary
                df_high.pop('Case Number')
                df_high['Loading(%) GenON'] = pd.to_numeric(df_high['Loading(%) GenON'])
                line_ldng_max = df_high.loc[df_high.groupby('Case Reference')['Loading(%) GenON'].idxmax()]
                line_ldng_min = df_high.loc[df_high.groupby('Case Reference')['Loading(%) GenON'].idxmin()]
                line_ldng_smry = pd.merge(line_ldng_max,line_ldng_min, how = 'outer', on = ['Case Reference'])
                #vltg_fluc_gen_chng_smry = vltg_fluc_gen_chng_smry.T.drop_duplicates().T
                summary_dfs[case]['line_loadings']['summary'] = []
                summary_dfs[case]['line_loadings']['summary'].append(line_ldng_smry)
                
                #doc = docx.Document('test.docx')
                #t = doc.add_table(summary_dfs[case]['line_loadings']['appendix'][0].shape[0]+1,summary_dfs[case]['line_loadings']['appendix'][0].shape[1])
                # add the header rows.
                #for j in range(summary_dfs[case]['line_loadings']['appendix'][0].shape[-1]):
                 #   t.cell(0,j).text = summary_dfs[case]['line_loadings']['appendix'][0].columns[j]
                
                # add the rest of the data frame
                #for i in range(summary_dfs[case]['line_loadings']['appendix'][0].shape[0]):
                 #   for j in range(summary_dfs[case]['line_loadings']['appendix'][0].shape[-1]):
                  #      t.cell(i+1,j).text = str(summary_dfs[case]['line_loadings']['appendix'][0].values[i,j])
                #doc.save('test.docx')
            else:
                pass
    

    
# =============================================================================
# This function returns summary tables for voltage fluctuation due to change in generation (cloud cover and gen trip). Thw tables will be a summary table, violated voltage fluctuation table, detailed table for appendix.
# Image to show voltage levels for intial and final condition. Only one set of input sheet is required.  
# =============================================================================

def vltg_fluc_gen_chng(result_dfs, summary_dfs):
    
    for case in summary_dfs.keys():
        summary_dfs[case]['vltg_fluc_gen_change']={}
        for key,value in result_dfs.items():
            if 'df_vltfluc_genchg_'+case+'_on' in key:
                vltg_fluc_gen_chng = result_dfs['df_vltfluc_genchg_'+case+'_on']
                
                # adppendix of voltage flcutuation due to gen chnage
                vltg_fluc_gen_chng.pop('Case Number')
                vltg_fluc_gen_chng_pivot = pd.pivot_table(data = vltg_fluc_gen_chng, index = ['Case Reference','Bus Name'], aggfunc = ['max'])
                vltg_fluc_gen_chng_pivot.style.applymap( hl_vltg_fluc_violation,subset = [('max','Volt Fluc(%)')]).format({'Volt Fluc(%)':'{0:,.3f}'})
                summary_dfs[case]['vltg_fluc_gen_change']['appendix'] = []
                summary_dfs[case]['vltg_fluc_gen_change']['appendix'].append(vltg_fluc_gen_chng_pivot)
                
                # Plot for voltage flcutuation due to gen change
                fig = plt.figure(figsize=(7,5))
                fig.add_axes([0.1,0.1,0.8,0.8])
                plt.scatter(vltg_fluc_gen_chng['Bus Name'], vltg_fluc_gen_chng['Voltage Level(pu) Prior'], label = 'Voltage Level(pu) Prior', color = 'k', marker = '*', s = 120)
                plt.scatter(vltg_fluc_gen_chng['Bus Name'], vltg_fluc_gen_chng['Voltage Level(pu) After'], label = 'Voltage Level(pu) After')
                plt.xticks(rotation = 45, visible = True)
                plt.title('Voltage Levels Gen Change', fontsize =12, color = 'black' )
                plt.ylabel('Voltage(pu)', fontsize = 10)
                plt.xlabel('Bus Names', fontsize = 10)
                plt.legend()
                plt.grid()
                #plt.savefig(case+'vltg_fluc_gen_change'+'plot.png', bbox_inches = 'tight')
                #imgdata= StringIO.StringIO()
                #imgdata= StringIO() # version issues
                imgdata= BytesIO() # version issues
                plt.savefig(imgdata, bbox_inches = 'tight', dpi =200)
                #plt.savefig(imgdata,format = 'svg', dpi =200) # version issues
                summary_dfs[case]['vltg_fluc_gen_change']['plot'] = []
                summary_dfs[case]['vltg_fluc_gen_change']['plot'].append(imgdata)
                
                #Violation results
                vltg_fluc_gen_chng_violation = vltg_fluc_gen_chng.loc[vltg_fluc_gen_chng['Volt Fluc(%)']>3.00]
                
                 # function to check pass fail criterion
                def vf_gc_check(row):
                    if (abs(row['Volt Fluc(%)']) < 3.00):
                        val = 'yes'
                    else:
                        val = 'no'
                    return val
                
                if vltg_fluc_gen_chng_violation.empty:
                    pass
                else:
                    vltg_fluc_gen_chng_violation['Pass'] = vltg_fluc_gen_chng_violation.apply( vf_gc_check, axis = 1)
                    
                summary_dfs[case]['vltg_fluc_gen_change']['violations'] = []
                summary_dfs[case]['vltg_fluc_gen_change']['violations'].append(vltg_fluc_gen_chng_violation)
                
                
                #Summary results for voltage fluctuation due to gen change
                vltg_fluc_gen_chng.pop('Voltage Level(pu) Prior')
                vltg_fluc_gen_chng.pop('Voltage Level(pu) After')
                vltg_fluc_gen_chng_max = vltg_fluc_gen_chng.loc[vltg_fluc_gen_chng.groupby('Case Reference')['Volt Fluc(%)'].idxmax()]
                vltg_fluc_gen_chng_min = vltg_fluc_gen_chng.loc[vltg_fluc_gen_chng.groupby('Case Reference')['Volt Fluc(%)'].idxmin()]
                vltg_fluc_gen_chng_smry = pd.merge(vltg_fluc_gen_chng_max,vltg_fluc_gen_chng_min, how = 'outer', on = ['Case Reference'])
                #vltg_fluc_gen_chng_smry = vltg_fluc_gen_chng_smry.T.drop_duplicates().T
                summary_dfs[case]['vltg_fluc_gen_change']['summary'] = []
                summary_dfs[case]['vltg_fluc_gen_change']['summary'].append(vltg_fluc_gen_chng_smry)
                
####
####This functions need to be updated, but the result provided by this can be used for reporting purpopse at this stage.###
####    

##Voltage fluctuation results due to loss of line
#def highlight_violation_vfluc_lol(val):
#    color = 'red' if (val > 0.0299) else 'black'
#    return 'color: %s' % color

    
def vltg_fluc_lol(result_dfs, summary_dfs):
    for case in summary_dfs.keys():
        summary_dfs[case]['vltg_fluc_lol']={}
        for key,value in result_dfs.items():
            if 'df_vltfluc_lol_'+case+'_on' in key:
                vltfluc_lol = pd.concat([result_dfs['df_vltfluc_lol_'+case+'_on'],result_dfs['df_vltfluc_lol_'+case+'_off']],axis = 1)
                #vltfluc_lol = vltfluc_lol.T.drop_duplicates().T not sure why not working took lot of my time
                vltfluc_lol = vltfluc_lol.loc[:,~vltfluc_lol.columns.duplicated()]
                
                # adppendix of voltage flcutuation due to loss of line
                vltfluc_lol.pop('Case Number')
                vltfluc_lol_pivot = pd.pivot_table(data =  vltfluc_lol, index = ['Case Reference','Bus Name'], aggfunc = ['max'])
                vltfluc_lol_pivot.style.applymap( hl_vltg_fluc_violation,subset = [('max','Volt Fluc(%) GenON'),('max','Volt Fluc(%) GenOFF')]).format({'Volt Fluc(%) GenOFF':'{0:,.3f}','Volt Fluc(%) GenON':'{0:,.3f}'})
                summary_dfs[case]['vltg_fluc_lol']['appendix'] = []
                summary_dfs[case]['vltg_fluc_lol']['appendix'].append(vltfluc_lol_pivot)
                
                # Plot for voltage flcutuation due to lol
                fig = plt.figure(figsize=(7,5))
                fig.add_axes([0.1,0.1,0.8,0.8])
                plt.scatter(vltfluc_lol['Bus Name'], vltfluc_lol['Volt Fluc(%) GenON'], label = 'Volt Fluc(%) GenON', color = 'k', marker = '*', s = 120)
                plt.scatter(vltfluc_lol['Bus Name'], vltfluc_lol['Volt Fluc(%) GenOFF'], label = 'Volt Fluc(%) GenOFF')
                plt.xticks(rotation = 45, visible = True)
                plt.title('Voltage Fluctuation Loss of Line', fontsize =12, color = 'black' )
                plt.ylabel('Volt Fluc(%)', fontsize = 10)
                plt.xlabel('Bus Names', fontsize = 10)
                plt.grid()
                #plt.savefig(case + 'vltg_fluc_lol'+'plot.png', bbox_inches = 'tight')
                #imgdata= StringIO.StringIO()
                #imgdata= StringIO() # version issues
                imgdata= BytesIO() # version issues
                plt.savefig(imgdata, bbox_inches = 'tight', dpi =200)
                #plt.savefig(imgdata,format = 'svg', dpi =200) # version issues
                summary_dfs[case]['vltg_fluc_lol']['plot'] = []
                summary_dfs[case]['vltg_fluc_lol']['plot'].append(imgdata)
                
                
                #Violation results
                vltfluc_lol_violation = vltfluc_lol.loc[abs(vltfluc_lol['Volt Fluc(%) GenON'])>3.00]
                
                # function to check pass fail criterion
                def vl_lol_check(row):
                    if (abs(row['Volt Fluc(%) GenOFF']) > 3.00) and (abs(row['Volt Fluc(%) GenOFF']) > abs(row['Volt Fluc(%) GenON'])):
                        val = 'yes'
                    else:
                        val = 'no'
                    return val
                
                if vltfluc_lol_violation.empty:
                    pass
                else:
                    vltfluc_lol_violation['Pass'] = vltfluc_lol_violation.apply(vl_lol_check, axis = 1)
                summary_dfs[case]['vltg_fluc_lol']['violations'] = []
                summary_dfs[case]['vltg_fluc_lol']['violations'].append(vltfluc_lol_violation)
                
                #Summary results for voltage fluctuation due to gen change
                vltfluc_lol.pop('Voltage Level(pu) GenOFF')
                vltfluc_lol.pop('Voltage Level(pu) GenON')
                vltfluc_lol_max = vltfluc_lol.loc[vltfluc_lol.groupby('Case Reference')['Volt Fluc(%) GenON'].idxmax()]
                vltfluc_lol_min = vltfluc_lol.loc[vltfluc_lol.groupby('Case Reference')['Volt Fluc(%) GenON'].idxmin()]
                vltfluc_lol_smry = pd.merge( vltfluc_lol_max,vltfluc_lol_min, how = 'outer', on = ['Case Reference'])
                #vltfluc_lol_smry = vltfluc_lol_smry.T.drop_duplicates().T
                summary_dfs[case]['vltg_fluc_lol']['summary'] = []
                summary_dfs[case]['vltg_fluc_lol']['summary'].append(vltfluc_lol_smry)

# This function populate the summary for fault level results
def fault_levels(result_dfs,summary_dfs):
    for case in summary_dfs.keys():
        summary_dfs[case]['Fault_levels'] = {}
        for key,value in result_dfs.items():
            if 'df_fault_lvl_' + case + '_on' in key:
                #summary table
                df_fault_lvl = pd.DataFrame(result_dfs['df_fault_lvl_'+case+'_on'])
                #df_fault_lvl = df_fault_lvl.T.drop_duplicates().T
                summary_dfs[case]['Fault_levels']['summary'] = []
                summary_dfs[case]['Fault_levels']['summary'].append(df_fault_lvl)
                
                #appendix
                summary_dfs[case]['Fault_levels']['appendix'] = []
                
                #plots
                summary_dfs[case]['Fault_levels']['plot'] = []
                
                #violations
                summary_dfs[case]['Fault_levels']['violations'] = []
  

def initialise_report():
    #read report template 
    report=Document(main_folder_path+"\\Plots\\SSReportTemplate1.docx")
    return report


def replace_placeholders(report):
    replace_dict = {'[Project Name]':str(ProjectDetailsDict['Name']), '[Project Name Short]':str(ProjectDetailsDict['NameShrt']), '[Total Plant MW at POC]': str(ProjectDetailsDict['PlantMW']), 
                    '[Developer]': str(ProjectDetailsDict['Dev']), '[Network Service Provider]':str(ProjectDetailsDict['NSP']), '[Town]': str(ProjectDetailsDict['Town']), 
                    '[State]': str(ProjectDetailsDict['State']), '[Connection type]': str(ProjectDetailsDict['contyp']), '[POC Feeder]': str(ProjectDetailsDict['poc_fdr']),
                    '[Nominal POC voltage (kV)]': str(ProjectDetailsDict['VPOCkv']), '[PSSEversion]': str(PSSEmodelDict['PSSEversion']), '[Lot/DP]': str(ProjectDetailsDict['lot_dp']),
                    '[Address]': str(ProjectDetailsDict['addrs']), '[LGA]': str(ProjectDetailsDict['lga']), '[POC Substation]': str(ProjectDetailsDict['Sub']),
                    '[Plant Model]': str(ProjectDetailsDict['plnt_mdl'])
                    }
    for key,value in replace_dict.items():
        for p in report.paragraphs:
            if key in p.text:
                p.text = p.text.replace(key,value)
    #for p in report.paragraphs:
     #   inline = p.runs
      #  for j in range(0,len(inline)):
       #     for k,v in replace_dict.items():
        #        if k in inline[j].text:
         #           inline[j].text = inline[j].text.replace(k,v)
    for key,value in replace_dict.items():
        for table in report.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if key in p.text:
                            p.text = p.text.replace(key,value)
                        
                    #inline = p.runs
                    #for j in range(0,len(inline)):
                    #    for k,v in replace_dict.items():
                    #       if k in inline[j].text:
                    #           inline[j].text = inline[j].text.replace(k,v)
    
    return report    

def add_report_intro(report):
    plant_rating=ProjectDetailsDict['PlantMW']
    POC_name=ProjectDetailsDict['Sub']
    location=ProjectDetailsDict['Town']+", "+ProjectDetailsDict['State']
    #generate general description and intro based on ProjectDetailsDict
    #Headline
    '''
    p=report.add_paragraph('')
    run=p.add_run()
    run.add_break(WD_BREAK.PAGE)
    report.add_heading("Introduction", level=1 )
    intro_text="This report summarises the findings of a steady state study carried out for a "+str(plant_rating)+" MW generator connected to "+str(POC_name)+" in "+str(location)+"."
    intro_text+="\n\nThe Report consists of five separate sections, each of which illustrate the findings of the respective part of the study:"
    p=report.add_paragraph(intro_text)
    #add description of each subsection of the report    
    temp_text="1) The first part of the steady state analysis looks at Bus voltages under system normal conditions (i.e. no line outages or other contingencies. The results of the analysis are presented in section 2. Adding generation to a bus or a line may shift the voltage levels at that bus or surrounding buses due to the impact it has on on power flows in the area. "
    temp_text+="\nThe normal operating band of the NEM is between 0.9 p.u. voltage and 1.1 p.u. voltage. The results presented in that section will highlight the changes in relevant bus voltages due to the addition of the new generator."
    p=report.add_paragraph('')
    for style in report.styles:
        print("style.name == %s" % style.name)
    report.add_paragraph(temp_text, style='List Paragraph')
    report.add_paragraph('', style='List Paragraph')
    temp_text="2) The second part of the Steady State analysis investigates line loading and transformer loading under N-1 conditions, with and without the proposed generator. This reveals pre-existing issues as well as issues caused by the inclusion of the proposed generator. "
    report.add_paragraph(temp_text, style='List Paragraph')
    report.add_paragraph('', style='List Paragraph')
    temp_text="3) The third part of the Steady State analysis explores voltage stability for a change in generation output. If there is a sudden disconnect the plant output can drop from 100% to 0%, which will instantly change the voltage at surrounding buses. This is quantified and compared against applicable thresholds. "
    report.add_paragraph(temp_text, style='List Paragraph')
    report.add_paragraph('', style='List Paragraph')
    temp_text="4) The fourth part of the Steady State analysis focusses on voltage stability under credible contingencies. The voltage magnitude of the voltage fluctuations "
    report.add_paragraph(temp_text, style='List Paragraph')
    report.add_paragraph('', style='List Paragraph')
    temp_text="5) The last part of the Steady State analysis investigates Fault current at buses of interest. These fault levels are required to not exceed planning levels. This item is unlikely to be problematic in most instances as the contribution from inverter-based resources behind transformer and cable impedances is generally normally small compared against existing headroom."
    report.add_paragraph(temp_text, style='List Paragraph')
    report.add_paragraph('', style='List Paragraph')
    '''
    p=report.add_paragraph('')
    run=p.add_run()
    run.add_break(WD_BREAK.PAGE)
    report.add_heading("Simulation Results", level=1 )
    return 0

#adds descritpion of the voltage level checks to the summary report and will add the table for high and low case, or just one if only one included.
#It will also add a graph showing the voltage levels.
def volt_lvl_report(report, summary_dfs, probs):
    report.add_heading("Voltage Levels", level=2 )
    temp_text="The voltage levels at buses of interest have been analysed for system normal conditions (no contingencies) with "+str(ProjectDetailsDict['Name'])+" in service and compared to the voltage levels prior to adding the new generator. "
    temp_text+="The normal operating band is defined as 0.9 p.u. to 1.1 p.u. voltage. If the plant pushes the voltage outside these boundaries this indicates a problem. Normally the voltage would be expected to remain within +/-5% of 1 p.u. "
    temp_text+="Any cases where the voltage goes outside this range should be carefully looked at and may indicate a need for further reactive support or voltage control."
    report.add_paragraph(temp_text)
    report.add_paragraph('')
    #Add Summary tables
    report.add_heading("Summary of findings", level=3 )
    temp_text="The results for each bus before and after addition of  "+str(ProjectDetailsDict['Name'])+" are listed in the tables below."
    report.add_paragraph(temp_text)
    for case in summary_dfs.keys():
        if(summary_dfs[case]['voltage_levels']!={}):
            report.add_heading(ident_case_name(case), level=4)#maybe make this a label of the table instead of making it a  heading.
            for frame_id in range(0, len(summary_dfs[case]['voltage_levels']['summary'])):#add summary tables of results to word doc.
                data_frame_to_docx_table(summary_dfs[case]['voltage_levels']['summary'][frame_id], report)
                report.add_paragraph('')
    #Add plots
    plots_present=False
    for case in summary_dfs.keys():
        if('plot' in summary_dfs[case]['voltage_levels'].keys()):
            if(summary_dfs[case]['voltage_levels']['plot']!=[]):
                plots_present=True
    if(plots_present):
        report.add_heading("Plots", level=3)
        temp_text="The scatter plots show the absolute voltage level in the reference case(s) analysed in this study."
        report.add_paragraph(temp_text)
        for case in summary_dfs.keys():
            if(summary_dfs[case]['voltage_levels']!={}):
                report.add_heading(ident_case_name(case), level=4)#maybe make this a label of the table instead of making it a  heading.
                for frame_id in range(0, len(summary_dfs[case]['voltage_levels']['plot'])):#add summary tables of results to word doc.
                    report.add_picture(summary_dfs[case]['voltage_levels']['plot'][frame_id], Inches(6))
                    report.add_paragraph('')
        
    #Add overview of violations
        #add check whether any violations exist
    report.add_heading("Violations", level=3 )
    temp_text="A voltage violation exists wherever the voltage is outside 0.9 to 1.1 p.u. If a given voltage violation is not exacerbated by the addition of the plant, it is considered unproblematic. "
    temp_text+="Where a violation occurs at dedicated buses, such as within other plant (SVCs etc.) these can often easily be mitigated by adjusting local transformer taps or control settings and may not be prohibitive."
    report.add_paragraph(temp_text)
    causer_flag=0
    for case in summary_dfs.keys():
        case_results_present=False
        viol_in_act_case=False
        if(summary_dfs[case]['voltage_levels']!={}):
            if('violations' in summary_dfs[case]['voltage_levels'].keys()):
                case_results_present=True
                if(summary_dfs[case]['voltage_levels']['violations']!=[]):
                    viol_in_act_case=True
                    probs['voltage_levels']=1
                    report.add_heading(ident_case_name(case), level=4)#maybe make this a label of the table instead of making it a  heading.
                    for frame_id in range(0, len(summary_dfs[case]['voltage_levels']['violations'])):#add summary tables of results to word doc.
                        problem_flag=data_frame_to_docx_table(summary_dfs[case]['voltage_levels']['violations'][frame_id], report)
                        if(problem_flag>0):
                            causer_flag=1#indicates whether the project is creating or exacerbating the issue.
                        report.add_paragraph('')
        if(not viol_in_act_case and case_results_present):
            report.add_paragraph('No voltage violations observed for '+ident_case_name(case)+'.') 
            
    if (causer_flag>0):
       probs['voltage_levels']=2 #Add conclusion based on whether violations exist and provide some generic advice.
    return 0
#Adds description of line loading analysis along with summary table (comparing with and withou gen). It will also include a graph showing the highest loading of each line (this may occur for a different contingency per line. this should be marked in the graph)
#The function should be able - based on the result data - to point out pre-existing overloadings and new overloadings and distinguish between them.
def line_loadings_report(report, summary_dfs, probs):
    p=report.add_paragraph('')
    run=p.add_run()
    run.add_break(WD_BREAK.PAGE)
    report.add_heading("Line Loading", level=2 )
    temp_text="Line and transformer loading in the area around "+ProjectDetailsDict['Town']+" has been analysed under various conditions, including relevant N-1 scenarios. "
    temp_text+="Where overloading occurs, it is important to understand whether the overloading is caused and/or exacerbated by the proposed generator. Relevant information is listed under 'violations' in below. "
    report.add_paragraph(temp_text)
    report.add_paragraph('')
    #Add Summary tables
    report.add_heading("Summary of findings", level=3 )
    temp_text="The results for each bus before and after addition of  "+str(ProjectDetailsDict['Name'])+" are listed in the tables below."
    report.add_paragraph(temp_text)
    for case in summary_dfs.keys():
        if(summary_dfs[case]['line_loadings']!={}):
            report.add_heading(ident_case_name(case), level=4)#maybe make this a label of the table instead of making it a  heading.
            for frame_id in range(0, len(summary_dfs[case]['line_loadings']['summary'])):#add summary tables of results to word doc.
                data_frame_to_docx_table(summary_dfs[case]['line_loadings']['summary'][frame_id], report)
                report.add_paragraph('')
    #Add plots
    plots_present=False
    for case in summary_dfs.keys():
        if('plot' in summary_dfs[case]['line_loadings'].keys()):
            if(summary_dfs[case]['line_loadings']['plot']!=[]):
                plots_present=True
    if(plots_present):
        report.add_heading("Plots", level=3 )
        temp_text="The scatter plots show the result for the worst case loading of each element (line or transformer). "
        temp_text+="The worst case scenario may differ between different network elements. If one particular outage lead to overloading of line X, it does not mean that the same outage constitutes the worst case for line Y. (line X and Y serving as a generic example in this explanation)."
        report.add_paragraph(temp_text)
        for case in summary_dfs.keys():
            if(summary_dfs[case]['line_loadings']!={}):
                report.add_heading(ident_case_name(case), level=4)#maybe make this a label of the table instead of making it a  heading.
                for frame_id in range(0, len(summary_dfs[case]['line_loadings']['plot'])):#add summary tables of results to word doc.
                    report.add_picture(summary_dfs[case]['line_loadings']['plot'][frame_id], Inches(6))
                    report.add_paragraph('')
        
    #Add overview of violations
        #add check whether any violations exist
    report.add_heading("Violations", level=3 )
    temp_text="An overloading exists wherever the line or transformer loading exceeds 100% for any credible N-1 scenario. If a given overloading is not exacerbated by the addition of the plant, it is not generally a threat to the project, but may me an indication of emerging problems in the area and needs to be given careful consideration. "
    temp_text+="Overloading that is caused or exacerbated by the propsoed project may be addressed with a runback scheme. If the overloading occurs under System normal conditions, the "
    report.add_paragraph(temp_text)
    causer_flag=0
    for case in summary_dfs.keys():
        case_results_present=False
        viol_in_act_case=False
        if(summary_dfs[case]['line_loadings']!={}):
            if('violations' in summary_dfs[case]['line_loadings'].keys()):
                case_results_present=True
                if(summary_dfs[case]['voltage_levels']['violations']!=[]):
                    viol_in_act_case=True
                    probs['line_ldngs']=1
                    report.add_heading(ident_case_name(case), level=4)#maybe make this a label of the table instead of making it a  heading.
                    for frame_id in range(0, len(summary_dfs[case]['line_loadings']['violations'])):#add summary tables of results to word doc.
                        problem_flag=data_frame_to_docx_table(summary_dfs[case]['line_loadings']['violations'][frame_id], report)
                        if(problem_flag>0):
                            causer_flag=1#indicates whether the project is creating or exacerbating the issue.
                        report.add_paragraph('')
        if(not viol_in_act_case and case_results_present):
            report.add_paragraph('No overloading observed for '+ident_case_name(case)+'.') 
    if (causer_flag>0):
       probs['line_ldngs']=2             #overloadings caused by proposed generator       
    return 0

#Add description of test type and context. include summary table of voltage fluctuations and critical fluctuations. 
#Additional comment on the table with fluctuations that excceed the limits and conclude whether or not there are problematic cases.
def bus_volt_fluct_gen_chng_report(report, summary_dfs, probs):
    p=report.add_paragraph('')
    run=p.add_run()
    run.add_break(WD_BREAK.PAGE)
    report.add_heading("Voltage Fluctuations for Change of Generation", level=2 )
    temp_text="This section explores the impact of a sudden loss of generation on the bus voltages at the buses of interest. The output of "+str(ProjectDetailsDict['Name'])+" can suddenly change due to unexpected tripping of the plant or weather-related circumstances such as a change in cloud cover. "
    temp_text+="A sudden change in the output of the plant must not cause a voltage disturbance larger than 3% due to loss of generation and larger than 5% due to trip of the plant. This voltage flcutation is acceptable within 5% at distribution level. If a project causes steps greater than the allowed margin at surrounding buses, mitigation measures can be considered, such as a Static Var Compensator (SVC) located at a nearby bus. "
    report.add_paragraph(temp_text)
    report.add_paragraph('')
    #Add Summary tables
    report.add_heading("Summary of findings", level=3 )
    temp_text="The maximum and minimum voltage fluctuation for each bus under the given scenarios are listed in the tables below. "
    report.add_paragraph(temp_text)
    for case in summary_dfs.keys():
        if(summary_dfs[case]['vltg_fluc_gen_change']!={}):
            report.add_heading(ident_case_name(case), level=4)#maybe make this a label of the table instead of making it a  heading.
            for frame_id in range(0, len(summary_dfs[case]['vltg_fluc_gen_change']['summary'])):#add summary tables of results to word doc.
                data_frame_to_docx_table(summary_dfs[case]['vltg_fluc_gen_change']['summary'][frame_id], report)
                report.add_paragraph('')
    #Add plots
    plots_present=False
    for case in summary_dfs.keys():
        if('plot' in summary_dfs[case]['vltg_fluc_gen_change'].keys()):
            if(summary_dfs[case]['vltg_fluc_gen_change']['plot']!=[]):
                plots_present=True
    if(plots_present):
        report.add_heading("Plots", level=3 )
        temp_text="The scatter plots shows the voltage profile of the network during loss of generation events prior and after connecting the proposed plant. "
        report.add_paragraph(temp_text)
        for case in summary_dfs.keys():
            if(summary_dfs[case]['vltg_fluc_gen_change']!={}):
                report.add_heading(ident_case_name(case), level=4)#maybe make this a label of the table instead of making it a  heading.
                for frame_id in range(0, len(summary_dfs[case]['vltg_fluc_gen_change']['plot'])):#add summary tables of results to word doc.
                    report.add_picture(summary_dfs[case]['vltg_fluc_gen_change']['plot'][frame_id], Inches(6))
                    report.add_paragraph('')
        
    #Add overview of violations
        #add check whether any violations exist
    report.add_heading("Violations", level=3 )
    temp_text="All voltage variations exceeding 3% are listed in this section. Where the variations occur at transmission level, up to 5% may be acceptable. "
    report.add_paragraph(temp_text)
    causer_flag=0 #flag contains information whether or not a given violation of limits is due to the addition of the proposed generator
    for case in summary_dfs.keys():
        case_results_present=False
        viol_in_act_case=False
        if(summary_dfs[case]['vltg_fluc_gen_change']!={}):
            if('violations' in summary_dfs[case]['vltg_fluc_gen_change'].keys()):
                case_results_present=True
                if(summary_dfs[case]['vltg_fluc_gen_change']['violations']!=[]):
                    viol_in_act_case=True
                    probs['vltg_fluc_gen_change']=1
                    report.add_heading(ident_case_name(case), level=4)#maybe make this a label of the table instead of making it a  heading.
                    for frame_id in range(0, len(summary_dfs[case]['vltg_fluc_gen_change']['violations'])):#add summary tables of results to word doc.
                        problem_flag=data_frame_to_docx_table(summary_dfs[case]['vltg_fluc_gen_change']['violations'][frame_id], report)
                        if(problem_flag>0):
                            causer_flag=1#indicates whether the project is creating or exacerbating the issue.
                        report.add_paragraph('')
        if(not viol_in_act_case and case_results_present):
            report.add_paragraph('No critical voltage fluctuatation observed for '+ident_case_name(case)+'.')
    if (causer_flag>0):
       probs['vltg_fluc_gen_change']=2 
    return 0
#Add description of test type and context. include summary table of voltage fluctuations and critical fluctuations. 
#Additional comment on the table with fluctuations and state whether or not the proposes project is causing the fluctuations to exceed limits.
def bus_volt_fluct_cont_report(report, summary_dfs, probs):
    p=report.add_paragraph('')
    run=p.add_run()
    run.add_break(WD_BREAK.PAGE)
    report.add_heading("Voltage fluctuations due to loss of line contingencies", level=2 )
    temp_text="This section explores voltage fluctuations at the buses of interest due to loss of line contingencies."
    temp_text+=" These voltage fluctions are than compared with the existing voltage fluctuations i.e. prior connecting the proposed plant. The voltage fluctuations should not be worse than the existing voltage fluctuations. "
    report.add_paragraph(temp_text)
    report.add_paragraph('')
    #Add Summary tables
    report.add_heading("Summary of findings", level=3 )
    temp_text="The maximum and minimum voltage fluctuation for each bus under the given scenarios are listed in the tables below. "
    report.add_paragraph(temp_text)
    for case in summary_dfs.keys():
        if(summary_dfs[case]['vltg_fluc_lol']!={}):
            report.add_heading(ident_case_name(case), level=4)#maybe make this a label of the table instead of making it a  heading.
            for frame_id in range(0, len(summary_dfs[case]['vltg_fluc_lol']['summary'])):#add summary tables of results to word doc.
                data_frame_to_docx_table(summary_dfs[case]['vltg_fluc_lol']['summary'][frame_id], report)
                report.add_paragraph('')
    #Add plots
    plots_present=False
    for case in summary_dfs.keys():
        if('plot' in summary_dfs[case]['vltg_fluc_lol'].keys()):
            if(summary_dfs[case]['vltg_fluc_lol']['plot']!=[]):
                plots_present=True
    if(plots_present):
        report.add_heading("Plots", level=3 )
        temp_text="The scatter plots shows the voltage fluctuations due to loss of line contingencies. "
        report.add_paragraph(temp_text)
        for case in summary_dfs.keys():
            if(summary_dfs[case]['vltg_fluc_lol']!={}):
                report.add_heading(ident_case_name(case), level=4)#maybe make this a label of the table instead of making it a  heading.
                for frame_id in range(0, len(summary_dfs[case]['vltg_fluc_lol']['plot'])):#add summary tables of results to word doc.
                    report.add_picture(summary_dfs[case]['vltg_fluc_lol']['plot'][frame_id], Inches(6))
                    report.add_paragraph('')
        
    #Add overview of violations
        #add check whether any violations exist
    report.add_heading("Violations", level=3 )
    temp_text="All voltage fluctuations in this assessment are "
    report.add_paragraph(temp_text)
    causer_flag=0 #flag contains information whether or not a given violation of limits is due to the addition of the proposed generator
    for case in summary_dfs.keys():
        case_results_present=False
        viol_in_act_case=False
        if(summary_dfs[case]['vltg_fluc_lol']!={}):
            if('violations' in summary_dfs[case]['vltg_fluc_lol'].keys()):
                case_results_present=True
                if(summary_dfs[case]['vltg_fluc_lol']['violations']!=[]):
                    viol_in_act_case=True
                    probs['vltg_fluc_lol']=1
                    report.add_heading(ident_case_name(case), level=4)#maybe make this a label of the table instead of making it a  heading.
                    for frame_id in range(0, len(summary_dfs[case]['vltg_fluc_lol']['violations'])):#add summary tables of results to word doc.
                        problem_flag=data_frame_to_docx_table(summary_dfs[case]['vltg_fluc_lol']['violations'][frame_id], report)
                        if(problem_flag>0):
                            causer_flag=1#indicates whether the project is creating or exacerbating the issue.
                        report.add_paragraph('')
        if(not viol_in_act_case and case_results_present):
            report.add_paragraph('No critical voltage fluctuatation observed for '+ident_case_name(case)+'.')
    if (causer_flag>0):
       probs['vltg_fluc_lol']=2 
    return 0
#Descritpion of the test type and context. Result table with fault levels. This shoudl include not only IEEE methodology but also N-something methodology.
#Maybe add logic  to check against some pre-defined levels.
def fautl_levels_report(report, summary_dfs, probs):
    p=report.add_paragraph('')
    run=p.add_run()
    run.add_break(WD_BREAK.PAGE)
    report.add_heading("Fault Level Analysis", level=2 )
    temp_text="The fault levels were calculated using the NCSFCC (for current source generators) and ASCC function in PSS®E) for both 3 phase and phase to ground faults. The short circuit current studies have been performed on the maximum load case."
    temp_text+=" Following table shows the fault levels on the monitored buses. "
    report.add_paragraph(temp_text)
    report.add_paragraph('')
    #Add Summary tables
    #report.add_heading("Summary of findings", level=3 )
    #temp_text="The maximum and minimum voltage fluctuation for each bus under the given scenarios are listed in the tables below. "
    #report.add_paragraph(temp_text)
    for case in summary_dfs.keys():
        if(summary_dfs[case]['Fault_levels']!={}):
            report.add_heading(ident_case_name(case), level=3)#maybe make this a label of the table instead of making it a  heading.
            for frame_id in range(0, len(summary_dfs[case]['Fault_levels']['summary'])):#add summary tables of results to word doc.
                data_frame_to_docx_table(summary_dfs[case]['Fault_levels']['summary'][frame_id], report)
                report.add_paragraph('')
    return 0

def add_conclusion(report, probs):
    p=report.add_paragraph('')
    run=p.add_run()
    run.add_break(WD_BREAK.PAGE)
    report.add_heading("Conclusion", level=1 )   
    temp_text=''
    flawless=True
    for test_type in probs.keys():
        if(probs[test_type])>0:
            flawless=False
    if not flawless:
        temp_text+="Some potential issues have been detected during the Steady State Analysis. "
        if(probs['voltage_levels']>0):
            temp_text+="\nThe studies have shown issues with the normal voltage levels at one or more buses. "
            if(probs['voltage_levels']>1):
                temp_text+="These issues are caused or exacerbated by the proposed project, indicating that there is a need for additional reactive power support like capacitor banks or an SVC, or a downsizing of the project. Please refer to section 2.1 for details. "
            else:
                temp_text+="These issues are pre-existing and are not exacerbated by the plant, hence this is not a reason against connecting the project, however it may indicate an issue with the network in the area and further investigation is recommended. "
        if(probs['line_ldngs']>0):
            temp_text+="\nOne or more network elements exhibit overloading, the details are provided in section 2.2."
            if(probs['line_ldngs']>1):
                temp_text+="The overloading is caused or exacerbated by "+str(ProjectDetailsDict['Name'])+" and must be addressed using runback schemes, a network augmentation or downsizing the project. "
            else:
                temp_text+="The overloading is pre-existing and not exacerbated by the project and is not a reason against the development, but an indicator for a nearby network issue. "
        if(probs['vltg_fluc_gen_chng']>0):            
            temp_text+="\nVoltage violations have been observed for a change in generation output. Similar to voltage violations under system normal conditions this can be addressed using an SVC in a nearby location that will not trip for a trip of the proposed generator. The violations are detailed in section 2.3. "            
        if(probs['vltg_fluc_lol']>0):
            temp_text+="\nSome contingency events have been observed to lead to unacceptable voltage fluctuations on the network. "
            if(probs['vltg_fluc_lol']>1):
                temp_text+="These fluctuations are caused or exacerbated by the proposed generator and need to be mitigated in order for the project to procees.  Similar to voltage violations under system normal conditions this can also be addressed with the help of an SVC. Further details about the relevant cases are provided in the tables in section 2.4. "
            else:
                temp_text+="The observed voltage fluctuations under N-1 are not caused by the project and shoudl not negatively impact the development. Further details about the relevant cases are provided in the tables in section 2.4. "       
    else:
        temp_text+="No issues have been detected in the Steady State analysis, which may indicate that the proposed location on the network is suitable to host a" + str(ProjectDetailsDict['PlantMW']) + "MW generator without any further limitations from a grid perspective. It should however be kept in mind that the the network is subject to constant change and new incoming generation can change this situation. "
        temp_text+="The steady state study should be repeated from time to time as the project progresses and new information becomes available."
    
    p=report.add_paragraph(temp_text)
    #report.add_paragraph("test", color='red')

def add_appendices(report,summary_dfs):
    p=report.add_paragraph('')
    run=p.add_run()
    run.add_break(WD_BREAK.PAGE)
    report.add_heading("Appendices", level=1)
    temp_text ="This section provide the detailed results where applicable for the assessments performed in this steady state analysis."
    report.add_paragraph(temp_text)
    report.add_paragraph('')
    #Add Appendix tables
    #Voltage levels
    report.add_heading("Appendix 1", level=3 )
    temp_text="The voltage levels for all the monitored buses as per section[]"
    report.add_paragraph(temp_text)
    for case in summary_dfs.keys():
        if(summary_dfs[case]['voltage_levels']!={}):
            report.add_heading(ident_case_name(case), level=4)#maybe make this a label of the table instead of making it a  heading.
            for frame_id in range(0, len(summary_dfs[case]['voltage_levels']['appendix'])):#add summary tables of results to word doc.
                data_frame_to_docx_table(summary_dfs[case]['voltage_levels']['appendix'][frame_id], report)
                report.add_paragraph('')
    
    #Line Loadings
    report.add_heading("Appendix 2", level=3 )
    temp_text="The thermal loadings for all the monitored branches in network normal and N-1 contingencies as per section[]"
    report.add_paragraph(temp_text)
    for case in summary_dfs.keys():
        if(summary_dfs[case]['line_loadings']!={}):
            report.add_heading(ident_case_name(case), level=4)#maybe make this a label of the table instead of making it a  heading.
            for frame_id in range(0, len(summary_dfs[case]['line_loadings']['appendix'])):#add summary tables of results to word doc.
                data_frame_to_docx_table(summary_dfs[case]['line_loadings']['appendix'][frame_id], report)
                report.add_paragraph('')
    
    #Voltage fluctuations GenChange
    report.add_heading("Appendix 3", level=3 )
    temp_text="The voltage fluctuations due to change in generation output on all the monitored buses as per section[]"
    report.add_paragraph(temp_text)
    for case in summary_dfs.keys():
        if(summary_dfs[case]['vltg_fluc_gen_change']!={}):
            report.add_heading(ident_case_name(case), level=4)#maybe make this a label of the table instead of making it a  heading.
            for frame_id in range(0, len(summary_dfs[case]['vltg_fluc_gen_change']['appendix'])):#add summary tables of results to word doc.
                data_frame_to_docx_table(summary_dfs[case]['vltg_fluc_gen_change']['appendix'][frame_id], report)
                report.add_paragraph('')
                
    #Line Loadings
    report.add_heading("Appendix4", level=3 )
    temp_text="The voltage fluctuations due to loss of line contingencies on all the monitored buses as per section[]"
    report.add_paragraph(temp_text)
    for case in summary_dfs.keys():
        if(summary_dfs[case]['vltg_fluc_lol']!={}):
            report.add_heading(ident_case_name(case), level=4)#maybe make this a label of the table instead of making it a  heading.
            for frame_id in range(0, len(summary_dfs[case]['vltg_fluc_lol']['appendix'])):#add summary tables of results to word doc.
                data_frame_to_docx_table(summary_dfs[case]['vltg_fluc_lol']['appendix'][frame_id], report)
                report.add_paragraph('')
    
    return 0
                
            
    
            
#maps the shrot case names against a more detailed version and returns it as a string. This can be expanded to cover more network cases in the future.    
def ident_case_name(case):
    if('high' in case):
        return 'High demand Network Scenario'
    elif('low' in case):
        return "Low demand Network Scenario"
    else:
        return case

def data_frame_to_docx_table(df, report, skiprows=0, skipcolumns=0 ):
    from docx.shared import RGBColor
    grey='A5A5A5'
    white="FFFFFF"
    # add a table to the end and create a reference variable
    # extra row is so we can add the header row
    problem_flag=0
    # add the header rows
    #identify whether pivot table or regular table
    #if pivot table
    if(df.index.names[0]!=None):
        #add header rows
        t = report.add_table(df.shape[0]+2, df.shape[1]+len(df.index.names))
        #t.style='ESCO Data'
        t.style='ListTable3-Accent3'
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="A5A5A5"/>'.format(nsdecls('w')))
        t.rows[1].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)

        for lvl_id in range(0, len(df.columns.levels[0])): #create first header line, skipping cells on the left equal to lengt of df.index.names, then label remaining cells with df.columns.levels[0]. merging cells equal to amount of length of df.columns.levels[1]
            col_id= len(df.index.names)+lvl_id*len(df.columns.levels[1])
            add_cell_text(t.cell(0, col_id), str(df.columns.levels[0][lvl_id]), text_color=white, cell_color=grey, bolt=True)
            #merge cell(s) right of the entry, equivalent to the amount of len(df.columns.levels[1])
        for col_id in range(0, len(df.index.names)):#add first part of second header row
            add_cell_text(t.cell(1, col_id), str(df.index.names[col_id]), text_color=white, cell_color=grey, bolt=True)
        for col_id in range (0, len(df.columns.levels[1])): #add secodn part of second header row
            add_cell_text(t.cell(1, col_id+len(df.index.names)), str(df.columns.levels[1][col_id]), text_color=white, cell_color=grey, bolt=True)
        #add body of table
        for row_id in range(0, df.shape[0]):
            #first couple of cells from index table
            for index_id in range(0, len(df.index.names)):
                if(hasattr(df.index, 'levels')):
                    entry=df.index.levels[index_id][df.index.codes[index_id][row_id]]
                    if(row_id>0):
                        prev_entry=df.index.levels[index_id][df.index.codes[index_id][row_id-1]]
                    else:
                        prev_entry=''
                    if(prev_entry!=entry):
                        t.cell(row_id+2, index_id).text=str(entry)
                    #elif(entry == prev_entry):
                     #   t.cell(row_id-1, index_id).merge(t.cell(row_id+2, index_id))
                        
                else:
                    entry=df.index[row_id] #only single index avaialble
                    if(row_id>0):
                        prev_entry=df.index[row_id-1] #only single index available
                    else:
                        prev_entry=''
                    if(prev_entry!=entry):
                        t.cell(row_id+2, index_id).text=str(entry)
            
                    
                    
                #if entry same as previous entry: merge cells and only make one entry
                
            #remaining cells from values table
            for col_id in range(0, len(df.values[0])):
                t.cell(row_id+2, col_id+len(df.index.names)).text=str(df.values[row_id][col_id])
                
            
        
    else: #regular table
        t = report.add_table(df.shape[0]+1, df.shape[1])
        #t.style='ESCO Data'
        t.style='ListTable3-Accent3'
        for j in range(df.shape[-1]):
            t.cell(0,j).text = df.columns[j]
        
        # add the rest of the data frame
        
        for i in range(0, df.shape[0]):
            for j in range(df.shape[-1]):
                if(str(df.values[i,j])=='no'):
                    cell=t.cell(i+1,j)
                    paragraph=cell.paragraphs[0]
                    run1=paragraph.add_run('no')
                    red = RGBColor(255, 0, 0)
                    run1.font.color.rgb = red                
                    #cell.add_paragraph(str(df.values[i,j]), color="red")
                    problem_flag=1
                else:
                    t.cell(i+1,j).text = str(df.values[i,j])
    
    #df.values
   
    return problem_flag

def add_cell_text(cell, content, text_color, cell_color, bolt):
    from docx.shared import RGBColor
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="A5A5A5"/>'.format(nsdecls('w'))) #shading element for table cell
    cell._tc.get_or_add_tcPr().append(shading_elm_1)
        
    paragraph=cell.paragraphs[0]
    run1=paragraph.add_run(content)
    color = RGBColor(int('0x'+str(text_color[0:2]),0), int('0x'+str(text_color[2:4]),0), int('0x'+str(text_color[4:6]),0))
    run1.font.color.rgb = color    


#def set_table_header_bg_color(cell):
#    """
#    set background shading for Header Rows
#    """
#    tblCell = cell._tc
#    tblCellProperties = tc.get_or_add_tcPr()
#    clShading = OxmlElement('w:shd')
#    clShading.set(qn('w:fill'), "FF0000") #Hex of Dark Blue Shade {R:0x00, G:0x51, B:0x9E}
#    tblCellProperties.append(clShading)
#    return cell

# =============================================================================
# MAIN BODY OF THE SCRIPT
# =============================================================================   
    

#Calling functions as required

   
result_dfs=read_result_sheet()
reportname_prefix= str(datetime.datetime.now().strftime("%Y%m%d-%H%M"))+"-"+str(ProjectDetailsDict['NameShrt']+ str(simulation_batch_label))


voltage_levels(result_dfs, summary_dfs)
line_ldngs(result_dfs, summary_dfs)
vltg_fluc_gen_chng(result_dfs, summary_dfs)
vltg_fluc_lol(result_dfs, summary_dfs)
fault_levels(result_dfs,summary_dfs)


'''    
vp_violation,vp_tbl2 = voltage_profile(result_dfs, summary_dfs)
line_max_violation,tbl2 = line_ldng_profile(result_dfs, summary_dfs)
vltg_fluc_gen_chng_smry,vltg_fluc_gen_chng_violation,vltg_fluc_gen_chng_pivot = vltg_fluc_gen_chng(result_dfs, summary_dfs)
vltg_fluc_lol_smry,vltg_fluc_lol_violation,vltg_fluc_lol_pivot = vltg_fluc_lol(result_dfs, summary_dfs)
'''

# Styling and writing the return dataframes to excel file for overall summary. Currently 'xlsxwriter' append function was not working.
writer = pd.ExcelWriter(main_folder_path+"\\Plots\\steady_state\\"+reportname_prefix+"_SteadyStateResultsSummaryTable.xlsx",engine = 'xlsxwriter')


    
# Exporting voltage level results to excel file
for key1 in summary_dfs.keys():
    for key2 in summary_dfs[key1].keys():
        if summary_dfs[key1][key2] == {}:
            pass
        else:
            for result in summary_dfs[key1][key2]['summary']:
                result.to_excel(writer,sheet_name = str(key1)+str(key2),startrow=0)
            for result in summary_dfs[key1][key2]['appendix']:
                result.to_excel(writer,sheet_name = str(key1)+str(key2),startrow=20)
            for result in summary_dfs[key1][key2]['violations']:
                result.to_excel(writer,sheet_name = str(key1)+str(key2),startcol=12)
            for result in summary_dfs[key1][key2]['plot']:
                worksheet = writer.sheets[str(key1)+str(key2)]
                worksheet.insert_image('AA1', str(key1)+str(key2) + 'plot.png' )
                

writer.save()  
#Generate summary Word doc.
report=initialise_report()
replace_placeholders(report)
add_report_intro(report)
volt_lvl_report(report, summary_dfs, probs) #add description of the results along with table(s) and plots(s)
line_loadings_report(report, summary_dfs, probs) #add description of the results along with table(s) and plots(s)
bus_volt_fluct_gen_chng_report(report, summary_dfs, probs)
bus_volt_fluct_cont_report(report, summary_dfs, probs)
fautl_levels_report(report, summary_dfs, probs)
add_conclusion(report, probs)
add_appendices(report,summary_dfs)


report.save(main_folder_path+"\\Plots\\steady_state\\"+reportname_prefix+"_SteadyStateResultsSummaryReport.docx")

